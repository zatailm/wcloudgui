import sys
from collections import Counter
import os
import asyncio
from qasync import QEventLoop
from deep_translator import GoogleTranslator
from PySide6.QtWidgets import (
    QApplication,
    QMainWindow,
    QFileDialog,
    QMessageBox,
    QLabel,
    QPushButton,
    QVBoxLayout,
    QGridLayout,
    QWidget,
    QLineEdit,
    QComboBox,
    QSpinBox,
    QDialog,
    QTextEdit,
    QProgressBar,
    QFrame,
    QColorDialog,
    QInputDialog,
    QTextBrowser,
)
from PySide6.QtCore import QThread, Signal, Qt, QTimer, QMutex
from PySide6.QtGui import QIcon
import matplotlib

matplotlib.use("QtAgg")
from wordcloud import WordCloud, STOPWORDS
import numpy as np
from PIL import Image
import socket

class StartupThread(QThread):
    def run(self):
        try:
            import matplotlib.pyplot as plt
            from matplotlib.font_manager import FontProperties
        except Exception as e:
            pass

class FileLoaderThread(QThread):
    file_loaded = Signal(str, str)
    file_error = Signal(str)

    def __init__(self, file_path):
        super().__init__()
        self.file_path = file_path

    def run(self):
        try:
            if not os.path.exists(self.file_path):
                raise FileNotFoundError(f"File not found: {self.file_path}")

            if self.file_path.endswith(".txt"):
                with open(self.file_path, "r", encoding="utf-8") as file:
                    text_data = file.read()
            elif self.file_path.endswith(".pdf"):
                text_data = self.extract_text_from_pdf(self.file_path)
            elif self.file_path.endswith((".doc", ".docx")):
                text_data = self.extract_text_from_word(self.file_path)
            elif self.file_path.endswith((".xlsx", ".xls")):
                text_data = self.extract_text_from_excel(self.file_path)
            elif self.file_path.endswith(".csv"):
                text_data = self.extract_text_from_csv(self.file_path)
            else:
                raise ValueError("Unsupported file format")

            if not text_data.strip():
                raise ValueError("File is empty or contains no extractable text")

            self.file_loaded.emit(self.file_path, text_data)

        except Exception as e:
            self.file_error.emit(
                f"Error loading {os.path.basename(self.file_path)}: {str(e)}"
            )

    def extract_text_from_pdf(self, pdf_path):
        import pypdf

        try:
            text = ""
            with open(pdf_path, "rb") as file:
                reader = pypdf.PdfReader(file)
                if not reader.pages:
                    raise ValueError("PDF contains no readable pages")
                for page in reader.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text += page_text + "\n"
            return text
        except Exception as e:
            raise RuntimeError(f"PDF processing failed: {str(e)}") from e

    def extract_text_from_word(self, word_path):
        import docx

        try:
            doc = docx.Document(word_path)
            if not doc.paragraphs:
                raise ValueError("Word document contains no readable text")
            return "\n".join(
                [para.text for para in doc.paragraphs if para.text.strip()]
            )
        except Exception as e:
            raise RuntimeError(f"Word processing failed: {str(e)}") from e

    def extract_text_from_excel(self, excel_path):
        import pandas as pd

        try:
            df = pd.read_excel(excel_path, sheet_name=None)
            if not df:
                raise ValueError("Excel file is empty or unreadable")
            text_data = ""
            for sheet_name, sheet_df in df.items():
                text_data += f"\n--- {sheet_name} ---\n"
                text_data += sheet_df.to_string(index=False, header=True)
            return text_data
        except Exception as e:
            raise RuntimeError(f"Excel processing failed: {str(e)}") from e

    def extract_text_from_csv(self, csv_path):
        import pandas as pd

        try:
            df = pd.read_csv(csv_path)
            if df.empty:
                raise ValueError("CSV file is empty")
            return df.to_string(index=False, header=True)
        except Exception as e:
            raise RuntimeError(f"CSV processing failed: {str(e)}") from e

class CustomFileLoaderThread(QThread):
    file_loaded = Signal(object, bool)

    def __init__(self, file_path, file_type):
        super().__init__()
        self.file_path = file_path
        self.file_type = file_type

    def run(self):
        try:
            if self.file_type == "lexicon":
                from vaderSentiment.vaderSentiment import SentimentIntensityAnalyzer

                self.vader_analyzer = CustomVaderSentimentIntensityAnalyzer(
                    custom_lexicon_file=self.file_path
                )
                self.file_loaded.emit(self.file_path, True)

            elif self.file_type == "TextBlob (Custom Lexicon)":
                from textblob import TextBlob

                self.textblob_analyzer = CustomTextBlobSentimentAnalyzer(
                    self.file_path
                    )
                self.file_loaded.emit(self.file_path, True)

            elif self.file_type == "model":
                from flair.models import TextClassifier

                model = TextClassifier.load(self.file_path)
                self.file_loaded.emit(model, True)
        except Exception as e:
            self.file_loaded.emit(str(e), False)

class CustomVaderSentimentIntensityAnalyzer:
    def __init__(self, lexicon_file="vader_lexicon.txt", custom_lexicon_file=None):
        from vaderSentiment.vaderSentiment import SentimentIntensityAnalyzer

        self.analyzer = SentimentIntensityAnalyzer()
        if custom_lexicon_file:
            self.load_custom_lexicon(custom_lexicon_file)

    def load_custom_lexicon(self, custom_lexicon_file):
        try:
            with open(custom_lexicon_file, "r", encoding="utf-8") as file:
                for line in file:
                    parts = line.strip().split("\t")
                    if len(parts) == 2:
                        word, measure = parts
                        try:
                            self.analyzer.lexicon[word] = float(measure)
                        except ValueError:
                            pass
        except Exception:
            pass

    def polarity_scores(self, text):
        return self.analyzer.polarity_scores(text)

class CustomTextBlobSentimentAnalyzer:
    def __init__(self, lexicon_file=None):
        self.lexicon = {}
        self.negations = {"not", "no", "never", "none"}  # Default negations (expandable)
        self.intensifiers = {"very": 1.3, "extremely": 1.5}  # Default intensifiers (word: multiplier)
        if lexicon_file and os.path.exists(lexicon_file):
            self.load_custom_lexicon(lexicon_file)

    def load_custom_lexicon(self, custom_lexicon_file):
        try:
            with open(custom_lexicon_file, "r", encoding="utf-8") as file:
                for line in file:
                    parts = line.strip().split("\t")
                    if len(parts) >= 2:
                        word, measure = parts[0], parts[1]
                        # Handle special markers for negations/intensifiers
                        if measure.lower() == "negation":
                            self.negations.add(word.lower())
                        elif measure.startswith("intensifier:"):
                            try:
                                multiplier = float(measure.split(":")[1])
                                self.intensifiers[word.lower()] = multiplier
                            except ValueError:
                                pass
                        else:
                            try:
                                self.lexicon[word.lower()] = float(measure)
                            except ValueError:
                                pass
        except Exception as e:
            print(f"Failed to load TextBlob custom lexicon: {e}")

    def analyze(self, text):
        from textblob import TextBlob
        
        blob = TextBlob(text)
        total_polarity = 0.0
        words_with_context = []

        # Preprocess with context markers
        for sentence in blob.sentences:
            words = sentence.words
            for i, word in enumerate(words):
                word_lower = word.lower()
                current_score = self.lexicon.get(word_lower, 0.0)
                
                # Check for negations (previous word)
                if i > 0 and words[i-1].lower() in self.negations:
                    current_score *= -1
                
                # Check for intensifiers (previous word)
                if i > 0 and words[i-1].lower() in self.intensifiers:
                    multiplier = self.intensifiers[words[i-1].lower()]
                    current_score *= multiplier
                
                words_with_context.append(current_score)

        # Calculate adjusted polarity
        valid_scores = [s for s in words_with_context if s != 0]
        if valid_scores:
            avg_polarity = sum(valid_scores) / len(valid_scores)
        else:
            avg_polarity = 0.0

        # Use native TextBlob subjectivity
        return {
            'polarity': avg_polarity,
            'subjectivity': blob.sentiment.subjectivity,
            'processed_words': len(valid_scores)
        }

class SentimentAnalysisThread(QThread):
    sentiment_analyzed = Signal(dict)
    offline_warning = Signal(str)

    def __init__(
        self,
        text_data,
        sentiment_mode,
        vader_analyzer,
        flair_classifier,
        flair_classifier_cuslang,
        textblob_analyzer=None,
    ):
        super().__init__()
        self.text_data = text_data
        self.sentiment_mode = sentiment_mode
        self.vader_analyzer = vader_analyzer
        self.flair_classifier = flair_classifier
        self.flair_classifier_cuslang = flair_classifier_cuslang
        self.textblob_analyzer = textblob_analyzer

    async def _async_translate(self, text):
        try:
            translated = GoogleTranslator(source="auto", target="en").translate(text)
            return translated
        except Exception as e:
            print(f"Translation error: {str(e)}")
            return None

    def translate_text(self, text):
        try:
            loop = asyncio.get_event_loop()
        except RuntimeError:
            loop = asyncio.new_event_loop()
            asyncio.set_event_loop(loop)

        sentences = [s.strip() for s in text.split(".") if s.strip()]
        translated = []

        async def run_translations():
            for sentence in sentences:
                try:
                    result = await self._async_translate(sentence)
                    if result:
                        translated.append(result)
                except Exception as e:
                    print(f"Failed to translate sentence: {str(e)}")
            return ". ".join(translated)

        try:
            return loop.run_until_complete(run_translations())
        except Exception as e:
            print(f"Translation failed: {str(e)}")
            return None
        finally:
            if loop.is_running():
                loop.close()

    def analyze_textblob(self, text, use_custom=False):
        if use_custom and self.textblob_analyzer:
            analysis = self.textblob_analyzer.analyze(text)
            polarity = analysis['polarity']
            subjectivity = analysis['subjectivity']
        else:
            from textblob import TextBlob
            blob = TextBlob(text)
            polarity = blob.sentiment.polarity
            subjectivity = blob.sentiment.subjectivity

        return {
            "positive_score": max(polarity, 0),
            "negative_score": max(-polarity, 0),
            "neutral_score": 1 - abs(polarity),
            "compound_score": polarity,
            "sentiment_label": "POSITIVE" if polarity > 0 else "NEGATIVE" if polarity < 0 else "NEUTRAL",
            "subjectivity": subjectivity
        }

    def run(self):
        result = {
            "positive_score": 0,
            "neutral_score": 0,
            "negative_score": 0,
            "compound_score": 0,
            "sentiment_label": "N/A",
            "subjectivity": "N/A (only available in TextBlob mode)",
        }

        try:
            text_to_analyze = self.text_data
            needs_translation = False

            if self.sentiment_mode in ["VADER", "Flair", "TextBlob"]:
                try:
                    from langdetect import detect

                    detected_lang = detect(self.text_data)
                    needs_translation = detected_lang != "en"
                except:
                    needs_translation = False

            if needs_translation:
                if not is_connected():
                    self.offline_warning.emit(
                        "Translation required but no internet connection"
                    )
                    return

                translated_text = self.translate_text(self.text_data)
                if translated_text:
                    text_to_analyze = translated_text

            if self.sentiment_mode == "TextBlob":
                result.update(self.analyze_textblob(text_to_analyze))
            elif self.sentiment_mode == "TextBlob (Custom Lexicon)":
                result.update(self.analyze_textblob(self.text_data, True))

            elif self.sentiment_mode == "VADER":
                if self.vader_analyzer:
                    sentiment = self.vader_analyzer.polarity_scores(text_to_analyze)
                    result.update(
                        {
                            "positive_score": sentiment["pos"],
                            "negative_score": sentiment["neg"],
                            "neutral_score": sentiment["neu"],
                            "compound_score": sentiment["compound"],
                            "sentiment_label": "POSITIVE"
                            if sentiment["compound"] >= 0.05
                            else "NEGATIVE"
                            if sentiment["compound"] <= -0.05
                            else "NEUTRAL",
                        }
                    )
                else:
                    result["sentiment_label"] = "Error: VADER not initialized"

            elif self.sentiment_mode == "VADER (Custom Lexicon)":
                if self.vader_analyzer:
                    sentiment = self.vader_analyzer.polarity_scores(self.text_data)
                    result.update(
                        {
                            "positive_score": sentiment["pos"],
                            "negative_score": sentiment["neg"],
                            "neutral_score": sentiment["neu"],
                            "compound_score": sentiment["compound"],
                            "sentiment_label": "POSITIVE"
                            if sentiment["compound"] >= 0.05
                            else "NEGATIVE"
                            if sentiment["compound"] <= -0.05
                            else "NEUTRAL",
                        }
                    )
                else:
                    result["sentiment_label"] = "Error: Custom VADER not initialized"

            elif self.sentiment_mode == "Flair":
                from flair.data import Sentence

                sentence = Sentence(text_to_analyze)
                self.flair_classifier.predict(sentence)
                sentiment = sentence.labels[0]
                result.update(
                    {
                        "compound_score": sentiment.score,
                        "sentiment_label": sentiment.value,
                        "positive_score": sentiment.score
                        if sentiment.value == "POSITIVE"
                        else 0,
                        "negative_score": sentiment.score
                        if sentiment.value == "NEGATIVE"
                        else 0,
                        "neutral_score": 1 - sentiment.score,
                    }
                )

            elif self.sentiment_mode == "Flair (Custom Model)":
                from flair.data import Sentence

                sentence = Sentence(self.text_data)
                self.flair_classifier_cuslang.predict(sentence)
                sentiment = sentence.labels[0]
                result.update(
                    {
                        "compound_score": sentiment.score,
                        "sentiment_label": sentiment.value,
                        "positive_score": sentiment.score
                        if sentiment.value == "POSITIVE"
                        else 0,
                        "negative_score": sentiment.score
                        if sentiment.value == "NEGATIVE"
                        else 0,
                        "neutral_score": 1 - sentiment.score,
                    }
                )

            self.sentiment_analyzed.emit(result)
        except Exception as e:
            result["sentiment_label"] = f"Error: {str(e)}"
            self.sentiment_analyzed.emit(result)

class FlairModelLoaderThread(QThread):
    model_loaded = Signal(object)
    error_occurred = Signal(str)

    _cached_model = None

    def __init__(self, model_path="sentiment"):
        super().__init__()
        self.model_path = model_path

    def run(self):
        try:
            import time
            from flair.models import TextClassifier

            if FlairModelLoaderThread._cached_model is None:
                time.sleep(0.1)
                model = TextClassifier.load(self.model_path)
                FlairModelLoaderThread._cached_model = model
            else:
                model = FlairModelLoaderThread._cached_model
            self.model_loaded.emit(model)
        except Exception as e:
            self.error_occurred.emit(str(e))

class WordCloudGenerator(QMainWindow):
    def __init__(self):
        super().__init__()
        self.file_path = ""
        self.text_data = ""
        self.mask_path = ""
        self.additional_stopwords = set()
        self.custom_color_palettes = {}
        self.current_figure = None

        from vaderSentiment.vaderSentiment import SentimentIntensityAnalyzer

        self.vader_analyzer = SentimentIntensityAnalyzer()
        self.sentiment_mode = "TextBlob"
        self.custom_lexicon_path = None

        self.custom_textblob_lexicon_path = None
        self.textblob_analyzer = None

        self.custom_model_path = None
        self.flair_classifier = None
        self.flair_classifier_cuslang = None

        self.active_threads = []
        self.threads_mutex = QMutex()

        self.file_loader_thread = None
        self.startup_thread = StartupThread()
        self.startup_thread.start()

        self.initUI()

    def initUI(self):
        self.setWindowTitle("WCGen")
        self.setFixedSize(550, 750)
        self.setWindowIcon(QIcon("D:/python_proj/wcloudgui/res/gs.ico"))

        layout = QGridLayout()

        self.file_label = QLineEdit(self)
        self.file_label.setReadOnly(True)
        self.file_label.setFixedHeight(30)
        self.file_label.setPlaceholderText("No file selected")
        layout.addWidget(self.file_label, 0, 0, 1, 6)

        self.upload_button = QPushButton("Load Text File", self)
        self.upload_button.clicked.connect(self.pilih_file)
        self.upload_button.setFixedHeight(30)
        layout.addWidget(self.upload_button, 1, 0, 1, 4)

        self.view_text_button = QPushButton("View Full Text", self)
        self.view_text_button.setFixedHeight(30)
        self.view_text_button.setToolTip("View the full text in a separate window")
        self.view_text_button.clicked.connect(self.view_full_text)
        self.view_text_button.setEnabled(False)
        layout.addWidget(self.view_text_button, 1, 4, 1, 2)

        self.progress_bar = QProgressBar(self)
        self.progress_bar.setRange(0, 0)
        self.progress_bar.setFixedHeight(4)
        self.progress_bar.setStyleSheet(
            """
            QProgressBar {
                border-radius: 1px;
                background-color: white;
                text-align: center;
                margin-left: 5px;  
                margin-right: 5px;                        
            }
            QProgressBar::chunk {
                background-color: red;
                width: 1px;  /* Ukuran kecil agar animasi terlihat lebih baik */
                margin: 0px; /* Pastikan tidak ada margin */
            }
        """
        )
        self.progress_bar.setContentsMargins(10, 0, 10, 0)
        self.progress_bar.setVisible(False)
        layout.addWidget(self.progress_bar, 2, 0, 1, 6)

        self.stopword_entry = QTextEdit(self)
        self.stopword_entry.setFixedHeight(75)
        self.stopword_entry.setPlaceholderText(
            "Enter stopwords, separated by spaces or new lines (optional)"
        )
        self.stopword_entry.setToolTip("Stopwords")
        layout.addWidget(self.stopword_entry, 3, 0, 1, 6)

        hline = QFrame()
        hline.setFrameShape(QFrame.HLine)
        hline.setFrameShadow(QFrame.Sunken)
        layout.addWidget(hline, 4, 0, 1, 6)

        self.color_theme_label = QLabel("Color Theme:", self)
        layout.addWidget(self.color_theme_label, 5, 0, 1, 2)

        self.color_theme = QComboBox(self)
        self.color_theme.setFixedHeight(30)
        QTimer.singleShot(100, self.load_colormaps)
        self.color_theme.setToolTip("Select a color theme for the word cloud")
        layout.addWidget(self.color_theme, 5, 2, 1, 3)

        self.custom_palette_button = QPushButton("Custom", self)
        self.custom_palette_button.setFixedHeight(30)
        self.custom_palette_button.setToolTip("Create a custom color palette")
        self.custom_palette_button.clicked.connect(self.create_custom_palette)
        layout.addWidget(self.custom_palette_button, 5, 5, 1, 1)

        self.bg_color_label = QLabel("Background Color:", self)
        layout.addWidget(self.bg_color_label, 7, 0, 1, 2)

        self.bg_color = QComboBox(self)
        self.bg_color.setFixedHeight(30)
        self.bg_color.addItems(["white", "black", "gray", "blue", "red", "yellow"])
        self.bg_color.setToolTip("Select a background color for the word cloud")
        layout.addWidget(self.bg_color, 7, 2, 1, 3)

        self.custom_bg_color_button = QPushButton("Custom", self)
        self.custom_bg_color_button.setFixedHeight(30)
        self.custom_bg_color_button.setToolTip("Select a custom background color")
        self.custom_bg_color_button.clicked.connect(self.select_custom_bg_color)
        layout.addWidget(self.custom_bg_color_button, 7, 5, 1, 1)

        hline = QFrame()
        hline.setFrameShape(QFrame.HLine)
        hline.setFrameShadow(QFrame.Sunken)
        layout.addWidget(hline, 8, 0, 1, 6)

        self.title_label = QLabel("WordCloud Title:", self)
        layout.addWidget(self.title_label, 9, 0, 1, 2)

        self.title_entry = QLineEdit(self)
        self.title_entry.setFixedHeight(30)
        self.title_entry.setPlaceholderText("Enter title (optional)")
        layout.addWidget(self.title_entry, 9, 2, 1, 2)

        self.title_font_size = QSpinBox(self)
        self.title_font_size.setRange(8, 72)
        self.title_font_size.setValue(14)
        self.title_font_size.setFixedHeight(30)
        layout.addWidget(self.title_font_size, 9, 4, 1, 1)

        self.title_position = QComboBox(self)
        self.title_position.addItems(["Left", "Center", "Right"])
        self.title_position.setCurrentText("Center")
        self.title_position.setFixedHeight(30)
        layout.addWidget(self.title_position, 9, 5, 1, 1)

        self.font_choice_label = QLabel("Font Choice:", self)
        layout.addWidget(self.font_choice_label, 10, 0, 1, 2)

        self.font_choice = QComboBox(self)
        self.font_choice.setFixedHeight(30)
        self.font_choice.addItem("Default")
        self.font_choice.setToolTip("Select a font family for the word cloud")
        layout.addWidget(self.font_choice, 10, 2, 1, 4)
        QTimer.singleShot(100, self.load_matplotlib_fonts)

        hline = QFrame()
        hline.setFrameShape(QFrame.HLine)
        hline.setFrameShadow(QFrame.Sunken)
        layout.addWidget(hline, 11, 0, 1, 6)

        self.min_font_size_label = QLabel("Minimum Font Size:", self)
        layout.addWidget(self.min_font_size_label, 12, 0, 1, 2)

        self.min_font_size_entry = QSpinBox(self)
        self.min_font_size_entry.setFixedHeight(30)
        self.min_font_size_entry.setValue(11)
        layout.addWidget(self.min_font_size_entry, 12, 2, 1, 1)

        self.max_words_label = QLabel("Maximum Words:", self)
        layout.addWidget(self.max_words_label, 12, 3, 1, 2, Qt.AlignRight)

        self.max_words_entry = QSpinBox(self)
        self.max_words_entry.setFixedHeight(30)
        self.max_words_entry.setMaximum(10000)
        self.max_words_entry.setValue(200)
        layout.addWidget(self.max_words_entry, 12, 5, 1, 1)

        hline = QFrame()
        hline.setFrameShape(QFrame.HLine)
        hline.setFrameShadow(QFrame.Sunken)
        layout.addWidget(hline, 13, 0, 1, 6)

        self.mask_label = QLabel("Mask Image:", self)
        layout.addWidget(self.mask_label, 14, 0, 1, 2)

        self.mask_path_label = QLineEdit("default (rectangle)", self)
        self.mask_path_label.setReadOnly(True)
        self.mask_path_label.setFixedHeight(30)
        layout.addWidget(self.mask_path_label, 14, 2, 1, 4)

        self.mask_button = QPushButton("Load Mask Image", self)
        self.mask_button.setFixedHeight(30)
        self.mask_button.setToolTip(
            "Select an image to use as a mask for the word cloud"
        )
        self.mask_button.clicked.connect(self.pilih_mask)
        layout.addWidget(self.mask_button, 15, 2, 1, 2)

        self.reset_mask_button = QPushButton("Remove Mask Image", self)
        self.reset_mask_button.setFixedHeight(30)
        self.reset_mask_button.setToolTip("Remove the selected mask image")
        self.reset_mask_button.clicked.connect(self.reset_mask)
        layout.addWidget(self.reset_mask_button, 15, 4, 1, 2)

        hline = QFrame()
        hline.setFrameShape(QFrame.HLine)
        hline.setFrameShadow(QFrame.Sunken)
        layout.addWidget(hline, 16, 0, 1, 6)

        self.buat_wordcloud_button = QPushButton("Generate WordCloud", self)
        self.buat_wordcloud_button.setFixedHeight(50)
        self.buat_wordcloud_button.setToolTip("Generate the word cloud")
        self.buat_wordcloud_button.clicked.connect(self.buat_wordcloud)
        self.buat_wordcloud_button.setEnabled(False)
        layout.addWidget(self.buat_wordcloud_button, 17, 0, 1, 6)

        self.wordcloud_progress_bar = QProgressBar(self)
        self.wordcloud_progress_bar.setRange(0, 0)
        self.wordcloud_progress_bar.setFixedHeight(4)
        self.wordcloud_progress_bar.setStyleSheet(
            """
            QProgressBar {
                border-radius: 1px;
                background-color: white;
                text-align: center;
                margin-left: 5px;  
                margin-right: 5px;
            }
            QProgressBar::chunk {
                background-color: blue;
                width: 1px;  /* Small size for better animation */
                margin: 0px; /* Ensure no margin */
            }
        """
        )
        self.wordcloud_progress_bar.setVisible(False)
        layout.addWidget(self.wordcloud_progress_bar, 18, 0, 1, 6)

        self.statistik_button = QPushButton("Word Count", self)
        self.statistik_button.setFixedHeight(30)
        self.statistik_button.setToolTip("Show word frequency statistics")
        self.statistik_button.clicked.connect(self.tampilkan_statistik)
        self.statistik_button.setEnabled(False)
        layout.addWidget(self.statistik_button, 19, 0, 1, 3)

        self.simpan_button = QPushButton("Save WordCloud", self)
        self.simpan_button.setFixedHeight(30)
        self.simpan_button.setToolTip("Save the generated word cloud")
        self.simpan_button.clicked.connect(self.simpan_wordcloud)
        self.simpan_button.setEnabled(False)
        layout.addWidget(self.simpan_button, 19, 3, 1, 3)

        hline = QFrame()
        hline.setFrameShape(QFrame.HLine)
        hline.setFrameShadow(QFrame.Sunken)
        layout.addWidget(hline, 20, 0, 1, 6)

        self.sentiment_mode_label = QLabel("Sentiment Analysis Mode:", self)
        layout.addWidget(self.sentiment_mode_label, 21, 0, 1, 2)

        self.sentiment_mode_combo = QComboBox(self)
        self.sentiment_mode_combo.setFixedHeight(30)
        self.sentiment_mode_combo.addItems(
            [
                "TextBlob",
                "TextBlob (Custom Lexicon)",
                "VADER",
                "VADER (Custom Lexicon)",
                "Flair",
                "Flair (Custom Model)",
            ]
        )
        self.sentiment_mode_combo.setToolTip("Select sentiment analysis mode")
        self.sentiment_mode_combo.setCurrentText("TextBlob")
        self.sentiment_mode_combo.currentTextChanged.connect(self.change_sentiment_mode)
        layout.addWidget(self.sentiment_mode_combo, 21, 2, 1, 3)

        self.sentiment_mode_info_button = QPushButton("Info", self)
        self.sentiment_mode_info_button.setFixedHeight(30)
        self.sentiment_mode_info_button.setToolTip(
            "Show description for each sentiment analysis mode"
        )
        self.sentiment_mode_info_button.clicked.connect(self.show_sentiment_mode_info)
        layout.addWidget(self.sentiment_mode_info_button, 21, 5, 1, 1)

        self.custom_lexicon_button = QPushButton("Load Lexicon", self)
        self.custom_lexicon_button.setFixedHeight(30)
        self.custom_lexicon_button.setToolTip("Load a custom lexicon file for VADER")
        self.custom_lexicon_button.clicked.connect(self.load_custom_lexicon)
        self.custom_lexicon_button.setEnabled(False)
        layout.addWidget(self.custom_lexicon_button, 22, 2, 1, 2)

        self.custom_model_button = QPushButton("Load Model", self)
        self.custom_model_button.setFixedHeight(30)
        self.custom_model_button.setToolTip("Load a custom model file for Flair")
        self.custom_model_button.clicked.connect(self.load_custom_model)
        self.custom_model_button.setEnabled(False)
        layout.addWidget(self.custom_model_button, 22, 4, 1, 2)

        self.model_progress_bar = QProgressBar(self)
        self.model_progress_bar.setRange(0, 0)
        self.model_progress_bar.setFixedHeight(4)
        self.model_progress_bar.setStyleSheet(
            """
            QProgressBar {
                border-radius: 1px;
                background-color: white;
                text-align: center;
                margin-left: 5px;  
                margin-right: 5px;
            }
            QProgressBar::chunk {
                background-color: orange;
                width: 1px; 
                margin: 0px; 
            }
        """
        )
        self.model_progress_bar.setVisible(False)
        layout.addWidget(self.model_progress_bar, 23, 2, 1, 4)

        self.sentiment_button = QPushButton("Analyze Sentiment", self)
        self.sentiment_button.setFixedHeight(50)
        self.sentiment_button.setToolTip("Analyze sentiment using current mode")
        self.sentiment_button.clicked.connect(self.analyze_sentiment)
        self.sentiment_button.setEnabled(False)
        layout.addWidget(self.sentiment_button, 24, 0, 1, 6)

        hline = QFrame()
        hline.setFrameShape(QFrame.HLine)
        hline.setFrameShadow(QFrame.Sunken)
        layout.addWidget(hline, 25, 0, 1, 6)

        self.about_button = QPushButton("About", self)
        self.about_button.setFixedHeight(30)
        self.about_button.setToolTip("Show information about the application")
        self.about_button.clicked.connect(self.show_about)
        layout.addWidget(self.about_button, 26, 0, 1, 2)

        self.panic_button = QPushButton("STOP", self)
        self.panic_button.setFixedHeight(30)
        self.panic_button.setToolTip("Stop all ongoing processes")
        self.panic_button.clicked.connect(self.stop_all_processes)
        layout.addWidget(self.panic_button, 26, 2, 1, 2)

        self.quit_button = QPushButton("Quit", self)
        self.quit_button.setFixedHeight(30)
        self.quit_button.setToolTip("Quit WCGen :(")
        self.quit_button.clicked.connect(self.close)
        layout.addWidget(self.quit_button, 26, 4, 1, 2)

        container = QWidget()
        container.setLayout(layout)
        self.setCentralWidget(container)

    def load_matplotlib_fonts(self):
        try:
            from matplotlib import font_manager

            self.font_map = {}

            weight_conversion = {
                100: "Thin",
                200: "Extra Light",
                300: "Light",
                400: "Regular",
                500: "Medium",
                600: "Semi Bold",
                700: "Bold",
                800: "Extra Bold",
                900: "Black",
            }

            for font in font_manager.fontManager.ttflist:
                try:
                    family = (
                        font.family_name if hasattr(font, "family_name") else font.name
                    )
                    style = (
                        font.style_name.lower()
                        if hasattr(font, "style_name")
                        else font.style.lower()
                    )

                    weight = weight_conversion.get(font.weight, str(font.weight))

                    display_parts = [family]

                    if "italic" in style or "oblique" in style:
                        display_parts.append("Italic")
                    elif "normal" not in style:
                        display_parts.append(style.title())

                    if weight != "Regular":
                        display_parts.append(weight)

                    display_name = " ".join(display_parts)

                    if display_name not in self.font_map:
                        self.font_map[display_name] = font.fname

                except Exception as e:
                    continue

            self.font_choice.clear()
            self.font_choice.addItem("Default")
            self.font_choice.addItems(sorted(self.font_map.keys()))

        except Exception as e:
            QMessageBox.warning(self, "Font Error", f"Failed to load fonts: {str(e)}")
            self.font_choice.addItems(["Arial", "Times New Roman", "Verdana"])

    def load_colormaps(self):
        import matplotlib.pyplot as plt

        try:
            import matplotlib.pyplot as plt

            self.color_theme.addItems(plt.colormaps())
        except ImportError as e:
            self.color_theme.clear()
            self.color_theme.addItem("Default")
            QMessageBox.warning(
                self, "Dependency Error", f"Failed to load color maps: {str(e)}"
            )

    def analyze_sentiment(self):
        if not self.text_data:
            QMessageBox.critical(
                self,
                "Error",
                "No text data available for sentiment analysis.\nPlease load a text file first.",
            )
            return

        if self.sentiment_mode in ["Flair", "TextBlob"]:
            try:
                from langdetect import detect
                detected_lang = detect(self.text_data)
                if detected_lang == "en":
                    self.sentiment_mode = "Flair"
            except:
                pass  # If detection fails, continue with the selected mode

        if self.sentiment_mode == "Flair" and not self.flair_classifier:
            QMessageBox.warning(
                self,
                "Model Loading",
                "Default Flair model is still loading. Please wait...",
            )
            return

        if self.sentiment_mode == "Flair (Custom Model)":
            if not self.flair_classifier_cuslang:
                QMessageBox.warning(
                    self, "Model Required", "Please load a custom Flair model first!"
                )
                return

        classifier = (
            self.flair_classifier_cuslang
            if self.sentiment_mode == "Flair (Custom Model)"
            else self.flair_classifier
        )

        self.progress_bar.show()
        self.sentiment_thread = SentimentAnalysisThread(
            self.text_data,
            self.sentiment_mode,
            self.vader_analyzer,
            classifier,
            self.flair_classifier_cuslang,
            self.textblob_analyzer,
        )
        self.sentiment_thread.offline_warning.connect(self.handle_offline_warning)
        self.sentiment_thread.sentiment_analyzed.connect(self.on_sentiment_analyzed)
        self.active_threads.append(self.sentiment_thread)
        self.sentiment_thread.start()

    def stop_all_processes(self):
        self.threads_mutex.lock()
        threads = self.active_threads.copy()
        self.threads_mutex.unlock()

        for thread in threads:
            if thread.isRunning():
                thread.requestInterruption()
                thread.quit()
                thread.wait(2000)

        self.threads_mutex.lock()
        self.active_threads.clear()
        self.threads_mutex.unlock()

        self.progress_bar.setVisible(False)
        self.model_progress_bar.setVisible(False)
        self.model_progress_bar.setVisible(False)
        self.active_threads.clear()
        QMessageBox.information(
            self, "Processes Stopped", "All ongoing processes have been stopped."
        )

    def closeEvent(self, event):
        import matplotlib.pyplot as plt

        plt.close("all")

        reply = QMessageBox.question(
            self,
            ":(",
            "Are you sure you want to quit?",
            QMessageBox.Yes | QMessageBox.No,
            QMessageBox.No,
        )
        if reply == QMessageBox.Yes:
            self.threads_mutex.lock()
            threads = self.active_threads.copy()
            self.threads_mutex.unlock()

            for thread in threads:
                if thread.isRunning():
                    thread.requestInterruption()
                    thread.quit()
                    thread.wait(2000)

            self.threads_mutex.lock()
            self.active_threads.clear()
            self.threads_mutex.unlock()

            self.progress_bar.setVisible(False)
            self.model_progress_bar.setVisible(False)
            self.model_progress_bar.setVisible(False)
            self.active_threads.clear()

            for widget in QApplication.topLevelWidgets():
                if widget is not self:
                    widget.close()
            event.accept()
        else:
            event.ignore()

    def show_about(self):
        about_text = (
            "<h2>🌟 WCGen - WordCloud Generator</h2>"
            "<p><b>Version:</b> 1.5</p>"
            "<p>&copy; 2025 MAZ Ilmam</p>"
            '<p><a href="https://github.com/zatailm/wcloudgui">📌 GitHub Repository</a></p>'

            "<h3>🔍 What is WCGen?</h3>"
            "<p>WCGen is an advanced yet easy-to-use application designed for creating visually appealing word clouds from text data. "
            "It enables users to quickly analyze word frequency, making it a valuable tool for researchers, students, writers, and data enthusiasts.</p>"

            "<h3>🚀 Key Features</h3>"
            "<ul>"
            "<li><b>Supports multiple file formats:</b> TXT, PDF, DOC/DOCX, CSV, XLSX</li>"
            "<li><b>Fully customizable word clouds:</b> Choose colors, fonts, shapes, and themes</li>"
            "<li><b>Stopword filtering:</b> Remove common words for clearer insights</li>"
            "<li><b>Smart text processing:</b> Handles large datasets efficiently</li>"
            "<li><b>Export & save options:</b> High-resolution image saving in multiple formats</li>"
            "<li><b>Sentiment analysis:</b> Optional analysis using TextBlob, VADER, and Flair</li>"
            "</ul>"

            "<h3>📖 How WCGen Helps You</h3>"
            "<p>Whether you’re analyzing customer feedback, conducting academic research, or visualizing text-based insights, "
            "WCGen simplifies the process and enhances your workflow with its intuitive design and powerful features.</p>"

            "<h3>📜 License</h3>"
            "<p>WCGen is free for personal and educational use. For commercial applications, please refer to the licensing terms.</p>"
        )

        dialog = QDialog(self)
        dialog.setWindowModality(Qt.NonModal)
        dialog.setWindowTitle("About WCGen")
        dialog.setMinimumSize(500, 400)
        dialog.setSizeGripEnabled(True)
        layout = QVBoxLayout()

        text_browser = QTextBrowser()
        text_browser.setHtml(about_text)
        text_browser.setOpenExternalLinks(True)
        text_browser.setReadOnly(True)

        layout.addWidget(text_browser)

        close_button = QPushButton("Close")
        close_button.clicked.connect(dialog.accept)
        layout.addWidget(close_button)

        dialog.setLayout(layout)
        dialog.show()

    def pilih_file(self):
        options = QFileDialog.Options()
        file_path, _ = QFileDialog.getOpenFileName(
            self,
            "Load Text File",
            "",
            "Supported Files (*.txt *.pdf *.doc *.docx *.csv *.xlsx *.xls);;All Files (*)",
            options=options,
        )
        if not file_path:
            return

        self.progress_bar.setVisible(True)

        try:
            self.file_loader_thread = FileLoaderThread(file_path)
            self.file_loader_thread.file_loaded.connect(self.on_file_loaded)
            self.file_loader_thread.file_error.connect(self.handle_file_error)

            self.threads_mutex.lock()
            self.active_threads.append(self.file_loader_thread)
            self.threads_mutex.unlock()

            self.file_loader_thread.start()

        except Exception as e:
            self.progress_bar.setVisible(False)
            QMessageBox.critical(self, "Error", f"Failed to load file: {e}")

    def handle_file_error(self, error_message):
        self.progress_bar.setVisible(False)
        QMessageBox.critical(self, "File Load Error", error_message)
        self._reset_file_state()

    def on_file_loaded(self, file_path, text_data):
        self.progress_bar.setVisible(False)

        try:
            if not text_data.strip():
                raise ValueError("File appears empty after loading")

            self.text_data = text_data
            self.file_label.setText(os.path.basename(file_path))

            self.buat_wordcloud_button.setEnabled(True)
            self.simpan_button.setEnabled(True)
            self.statistik_button.setEnabled(True)
            self.view_text_button.setEnabled(True)

            self.change_sentiment_mode(self.sentiment_mode)

        except ValueError as e:
            self.handle_file_error(str(e))
            self._reset_file_state()

    def _reset_file_state(self):
        self.text_data = ""
        self.file_label.setText("")
        self.buat_wordcloud_button.setEnabled(False)
        self.simpan_button.setEnabled(False)
        self.statistik_button.setEnabled(False)
        self.view_text_button.setEnabled(False)
        self.sentiment_button.setEnabled(False)

    def handle_file_error(self, error_message):
        QMessageBox.critical(self, "File Load Error", error_message)

    def pilih_mask(self):
        options = QFileDialog.Options()
        mask_path, _ = QFileDialog.getOpenFileName(
            self,
            "Select Mask Image",
            "",
            "Image Files (*.png *.jpg *.bmp)",
            options=options,
        )
        if mask_path:
            self.mask_path = mask_path
            self.mask_path_label.setText(f"{mask_path}")

    def reset_mask(self):
        self.mask_path = ""
        self.mask_path_label.setText("default (rectangle)")

    def ambil_stopwords(self):
        custom_words = self.stopword_entry.toPlainText().strip().lower()
        if custom_words:
            self.additional_stopwords = set(custom_words.split())
        return STOPWORDS.union(self.additional_stopwords)

    def tampilkan_statistik(self):
        if not self.text_data:
            QMessageBox.critical(self, "Error", "Load text file first!")
            return

        stopwords = self.ambil_stopwords()
        words = [word.lower() for word in self.text_data.split() if word.lower() not in stopwords]
        word_counts = Counter(words)

        sorted_word_counts = sorted(word_counts.items(), key=lambda item: item[1], reverse=True)

        if hasattr(self, "stats_dialog") and self.stats_dialog is not None:
            self.stats_dialog.close()

        self.stats_dialog = QDialog(self)
        self.stats_dialog.setWindowModality(Qt.NonModal)
        self.stats_dialog.setWindowTitle("Word Frequency Statistics")
        self.stats_dialog.setMinimumSize(500, 400)

        layout = QVBoxLayout()

        text_browser = QTextBrowser()

        html_content = """
        <h2>Word Frequency Statistics</h2>
        <table border="1" cellspacing="0" cellpadding="5" width="100%">
            <tr><th align="left">Word</th><th align="left">Count</th></tr>
        """
        for word, count in sorted_word_counts:
            html_content += f"<tr><td>{word}</td><td>{count}</td></tr>"

        html_content += "</table>"

        text_browser.setHtml(html_content)
        text_browser.setReadOnly(True)

        layout.addWidget(text_browser)

        close_button = QPushButton("Close")
        close_button.clicked.connect(self.stats_dialog.accept)
        layout.addWidget(close_button)

        self.stats_dialog.setLayout(layout)
        self.stats_dialog.show()

    def buat_wordcloud(self):
        import matplotlib.pyplot as plt

        if not self.text_data:
            QMessageBox.critical(self, "Error", "Load text file first!")
            return

        font_path = None
        selected_font = self.font_choice.currentText()
        if selected_font != "Default":
            font_path = self.font_map.get(selected_font)
            if not font_path or not os.path.exists(font_path):
                QMessageBox.warning(
                    self, "Font Error", f"Font file not found: {font_path}"
                )
                font_path = None
                return

        stopwords = self.ambil_stopwords()
        mask = None
        if self.mask_path:
            try:
                mask = np.array(Image.open(self.mask_path))
            except Exception as e:
                QMessageBox.critical(self, "Error", f"Failed to load mask image: {e}")
                return

        try:
            if self.current_figure:
                plt.close(self.current_figure)
                self.current_figure = None

            self.wordcloud_progress_bar.setVisible(True)
            QApplication.processEvents()

            wc = WordCloud(
                width=800,
                height=400,
                background_color=self.bg_color.currentText(),
                stopwords=stopwords,
                colormap=self.color_theme.currentText(),
                max_words=self.max_words_entry.value(),
                min_font_size=self.min_font_size_entry.value(),
                mask=mask,
                font_path=font_path,
            ).generate(self.text_data)

            plt.ion()
            self.current_figure = plt.figure(figsize=(10, 5))
            plt.imshow(wc, interpolation="bilinear")

            title_text = self.title_entry.text().strip()
            if title_text:
                title_font = None
                if selected_font != "Default" and font_path:
                    title_font = FontProperties(
                        fname=font_path, size=self.title_font_size.value()
                    )
                plt.title(
                    title_text,
                    loc=self.title_position.currentText().lower(),
                    fontproperties=title_font,
                )

            plt.axis("off")
            plt.show(block=False)
        except Exception as e:
            if self.current_figure:
                plt.close(self.current_figure)
                self.current_figure = None
            QMessageBox.critical(self, "Error", f"Failed to generate word cloud: {e}")
        finally:
            self.wordcloud_progress_bar.setVisible(False)

    def simpan_wordcloud(self):
        options = QFileDialog.Options()
        save_path, _ = QFileDialog.getSaveFileName(
            self,
            "Save WordCloud",
            "",
            "PNG file (*.png);;JPG file (*.jpg)",
            options=options,
        )
        if not save_path:
            return
        stopwords = self.ambil_stopwords()
        mask = None

        font_path = None
        if self.font_choice.currentText() != "Default":
            from matplotlib import font_manager

            selected_font = self.font_choice.currentText()
            font_path = self.font_map.get(selected_font)

            try:
                font_manager.findfont(self.font_choice.currentText())
                font_path = self.font_choice.currentText()
            except:
                QMessageBox.warning(
                    self, "Font Error", "Selected font not found, using default"
                )

        if self.mask_path:
            try:
                mask = np.array(Image.open(self.mask_path))
            except Exception as e:
                QMessageBox.critical(self, "Error", f"Failed to load mask image: {e}")
                return
        try:
            wc = WordCloud(
                width=800,
                height=400,
                background_color=self.bg_color.currentText(),
                stopwords=stopwords,
                colormap=self.color_theme.currentText(),
                max_words=self.max_words_entry.value(),
                min_font_size=self.min_font_size_entry.value(),
                mask=mask,
                font_path=None
                if self.font_choice.currentText() == "Default"
                else self.font_choice.currentText(),
            ).generate(self.text_data)
            wc.to_file(save_path)
            QMessageBox.information(self, "Succeed", "WordCloud saved!")
        except Exception as e:
            QMessageBox.critical(self, "Error", f"Failed to save word cloud: {e}")

    def create_custom_palette(self):
        dialog = QDialog(self)
        dialog.setWindowTitle("Create Custom Palette")
        dialog.setFixedSize(400, 150)

        layout = QVBoxLayout()
        color_list = []

        def add_color():
            color = QColorDialog.getColor()
            if color.isValid():
                color_list.append(color.name())
                color_label.setText(", ".join(color_list))

        color_label = QLabel("", self)
        layout.addWidget(color_label)

        add_color_button = QPushButton("Add Color", self)
        add_color_button.clicked.connect(add_color)
        layout.addWidget(add_color_button)

        save_palette_button = QPushButton("Save Palette", self)
        save_palette_button.clicked.connect(
            lambda: self.save_custom_palette(color_list, dialog)
        )
        layout.addWidget(save_palette_button)

        dialog.setLayout(layout)
        dialog.exec()

    def save_custom_palette(self, color_list, dialog):
        palette_name, ok = QInputDialog.getText(
            self, "Save Palette", "Enter palette name:"
        )
        if ok and palette_name:
            self.custom_color_palettes[palette_name] = color_list
            self.color_theme.addItem(palette_name)
            dialog.accept()

    def select_custom_bg_color(self):
        color = QColorDialog.getColor()
        if color.isValid():
            self.bg_color.addItem(color.name())
            self.bg_color.setCurrentText(color.name())

    def view_full_text(self):
        if not self.text_data:
            QMessageBox.critical(self, "Error", "No text data available to display.")
            return

        if hasattr(self, "text_dialog") and self.text_dialog is not None:
            self.text_dialog.close()

        self.text_dialog = QDialog(self)
        self.text_dialog.setWindowModality(Qt.NonModal)
        self.text_dialog.setWindowTitle("Full Text")
        self.text_dialog.setMinimumSize(600, 400)
        self.text_dialog.setSizeGripEnabled(True)

        layout = QVBoxLayout()

        text_browser = QTextBrowser()
        text_browser.setPlainText(self.text_data)
        text_browser.setOpenExternalLinks(True)
        text_browser.setReadOnly(True)

        layout.addWidget(text_browser)

        close_button = QPushButton("Close")
        close_button.clicked.connect(self.text_dialog.accept)
        layout.addWidget(close_button)

        self.text_dialog.setLayout(layout)
        self.text_dialog.show()

    def handle_offline_warning(self, msg):
        self.progress_bar.setVisible(False)
        QMessageBox.warning(self, "Offline Mode", msg)

    def on_sentiment_analyzed(self, result):
        self.progress_bar.setVisible(False)
        text_length = len(self.text_data)
        word_count = len(self.text_data.split())
        char_count_excl_spaces = len(self.text_data.replace(" ", ""))
        avg_word_length = char_count_excl_spaces / word_count if word_count > 0 else 0
        most_frequent_words = self.get_most_frequent_words(self.text_data, 5)

        self.show_sentiment_analysis(
            self.sentiment_mode,
            result["positive_score"],
            result["neutral_score"],
            result["negative_score"],
            result["compound_score"],
            result["sentiment_label"],
            text_length,
            result["subjectivity"],
            word_count,
            char_count_excl_spaces,
            avg_word_length,
            most_frequent_words,
        )

    def show_sentiment_analysis(
        self,
        analysis_mode,
        positive_score,
        neutral_score,
        negative_score,
        compound_score,
        sentiment_label,
        text_length,
        subjectivity,
        word_count,
        char_count_excl_spaces,
        avg_word_length,
        most_frequent_words,
    ):
        dialog = QDialog(self)
        dialog.setWindowModality(Qt.NonModal)
        dialog.setWindowTitle("Sentiment Analysis")
        dialog.setMinimumSize(500, 400)
        dialog.setSizeGripEnabled(True)

        layout = QVBoxLayout()

        text_browser = QTextBrowser()

        sentiment_result = f"""
        <h2>Sentiment Analysis Results</h2>
        <table border="1" cellspacing="0" cellpadding="5" width="100%">
            <tr><th align="left">Metric</th><th align="left">Value</th></tr>
            <tr><td>Analysis Mode</td><td>{analysis_mode}</td></tr>
            <tr><td>Sentiment Label</td><td>{sentiment_label}</td></tr>
            <tr><td>Positive Sentiment</td><td>{positive_score:.2f}</td></tr>
            <tr><td>Neutral Sentiment</td><td>{neutral_score:.2f}</td></tr>
            <tr><td>Negative Sentiment</td><td>{negative_score:.2f}</td></tr>
            <tr><td>Compound Score</td><td>{compound_score:.2f}</td></tr>
            <tr><td>Subjectivity</td><td>{subjectivity if isinstance(subjectivity, str) else f"{subjectivity:.2f}"}</td></tr>
            <tr><td>Text Length</td><td>{text_length} characters</td></tr>
            <tr><td>Word Count</td><td>{word_count}</td></tr>
            <tr><td>Character Count (excluding spaces)</td><td>{char_count_excl_spaces}</td></tr>
            <tr><td>Average Word Length</td><td>{avg_word_length:.2f}</td></tr>
            <tr><td>Most Frequent Words</td><td>{most_frequent_words}</td></tr>
        </table>
        """

        text_browser.setHtml(sentiment_result)
        text_browser.setOpenExternalLinks(True)
        text_browser.setReadOnly(True)

        layout.addWidget(text_browser)

        close_button = QPushButton("Close")
        close_button.clicked.connect(dialog.accept)
        layout.addWidget(close_button)

        dialog.setLayout(layout)
        dialog.show()

    def get_most_frequent_words(self, text, n):
        stopwords = self.ambil_stopwords().union(STOPWORDS)
        words = [word.lower() for word in text.split() if word.lower() not in stopwords]
        word_counts = Counter(words)
        most_common = word_counts.most_common(n)
        return [word for word, count in most_common]

    def change_sentiment_mode(self, mode):
        self.sentiment_mode = mode
        status_text = ""
        has_text = bool(self.text_data.strip())

        self.custom_lexicon_button.setEnabled(mode == "VADER (Custom Lexicon)")
        self.custom_model_button.setEnabled(mode == "Flair (Custom Model)")

        if mode == "Flair":
            status_text = (
                "Flair: Loading default model..."
                if not self.flair_classifier
                else "Ready"
            )
            self.sentiment_button.setEnabled(bool(self.flair_classifier) and has_text)
            self.custom_model_button.setEnabled(False)
            if not self.flair_classifier:
                self.load_flair_model()

        elif mode == "Flair (Custom Model)":
            if not self.flair_classifier:
                self.custom_model_button.setEnabled(False)
                self.load_flair_model()
                status_text = "Initializing Flair environment..."
            else:
                if self.flair_classifier_cuslang:
                    status_text = "Ready"
                    self.sentiment_button.setEnabled(has_text)
                else:
                    status_text = "Load custom model first!"
                    self.sentiment_button.setEnabled(False)
                self.custom_model_button.setEnabled(True)

        elif mode == "VADER (Custom Lexicon)":
            status_text = (
                "Load custom lexicon first!"
                if not self.custom_lexicon_path
                else "Ready"
            )
            self.custom_lexicon_button.setEnabled(True)
            self.sentiment_button.setEnabled(
                has_text and bool(self.custom_lexicon_path)
            )
            if not self.custom_lexicon_path:
                QMessageBox.warning(
                    self,
                    "Lexicon Required",
                    "You have selected 'VADER (Custom Lexicon)'. "
                    "Please load a custom lexicon before analyzing sentiment.",
                )

        elif mode == "TextBlob (Custom Lexicon)":
            status_text = (
                "Load custom lexicon first!"
                if not self.custom_textblob_lexicon_path
                else "Ready"
            )
            self.custom_lexicon_button.setEnabled(True)
            self.sentiment_button.setEnabled(
                has_text and bool(self.custom_textblob_lexicon_path)
            )
            if not self.custom_textblob_lexicon_path:
                QMessageBox.warning(
                    self,
                    "Lexicon Required",
                    "You have selected 'TextBlob (Custom Lexicon)'. "
                    "Please load a custom lexicon before analyzing sentiment.",
                )            

        else:
            self.sentiment_button.setEnabled(has_text)
            status_text = "Ready" if has_text else "Load text first!"

        self.sentiment_button.setToolTip(
            f"{mode} - {status_text}\n"
            "Click to analyze sentiment\n"
            "Subjectivity scores only available in TextBlob modes"
        )

    def load_custom_lexicon(self):
        options = QFileDialog.Options()
        lexicon_path, _ = QFileDialog.getOpenFileName(
            self,
            "Select Custom Lexicon File",
            "",
            "Text Files (*.txt)",
            options=options,
        )
        if lexicon_path:
            self.model_progress_bar.setVisible(True)
            self.model_progress_bar.setRange(0, 0)
            self.custom_lexicon_button.setEnabled(False)

            self.lexicon_loader_thread = CustomFileLoaderThread(
                lexicon_path, "TextBlob (Custom Lexicon)" if self.sentiment_mode == "TextBlob (Custom Lexicon)" else "lexicon"
            )
            self.lexicon_loader_thread.file_loaded.connect(self.on_lexicon_loaded)

            self.threads_mutex.lock()
            self.active_threads.append(self.lexicon_loader_thread)
            self.threads_mutex.unlock()

            self.lexicon_loader_thread.start()

    def on_lexicon_loaded(self, result, success):
        self.model_progress_bar.setVisible(False)
        self.custom_lexicon_button.setEnabled(True)
        self.sentiment_button.setEnabled(False)

        if success:
            if self.sentiment_mode == "TextBlob (Custom Lexicon)":
                self.custom_textblob_lexicon_path = result
                self.textblob_analyzer = CustomTextBlobSentimentAnalyzer(
                    self.custom_textblob_lexicon_path
                )
                QMessageBox.information(
                    self,
                    "Success",
                    "Custom TextBlob lexicon loaded successfully!"
                )
                self.change_sentiment_mode(self.sentiment_mode)
            else:
                self.custom_lexicon_path = result
                self.vader_analyzer = CustomVaderSentimentIntensityAnalyzer(
                    custom_lexicon_file=self.custom_lexicon_path
                )
                QMessageBox.information(
                    self,
                    "Success",
                    "Custom lexicon loaded successfully! VADER will now use this lexicon.",
                )
                self.change_sentiment_mode(self.sentiment_mode)
        else:
            QMessageBox.critical(
                self, "Error", f"Failed to load custom lexicon: {result}"
            )

    def load_custom_model(self):
        options = QFileDialog.Options()
        model_path, _ = QFileDialog.getOpenFileName(
            self, "Select Custom Model File", "", "Model Files (*.pt)", options=options
        )
        if model_path:
            self.model_progress_bar.setVisible(True)
            self.custom_model_button.setEnabled(False)

            self.model_loader_thread = CustomFileLoaderThread(model_path, "model")
            self.model_loader_thread.file_loaded.connect(self.on_model_loaded)
            self.active_threads.append(self.model_loader_thread)

            self.threads_mutex.lock()
            self.active_threads.append(self.model_loader_thread)
            self.threads_mutex.unlock()

            self.model_loader_thread.start()

    def load_flair_model(self):
        if (
            hasattr(self, "flair_loader_thread")
            and self.flair_loader_thread
            and self.flair_loader_thread.isRunning()
        ):
            self.flair_loader_thread.quit()
            self.flair_loader_thread.wait()

        self.model_progress_bar.setVisible(True)
        self.model_progress_bar.setRange(0, 0)
        self.custom_model_button.setEnabled(False)
        self.sentiment_button.setEnabled(False)
        QApplication.processEvents()

        self.flair_loader_thread = FlairModelLoaderThread()
        self.flair_loader_thread.model_loaded.connect(self.on_flair_model_loaded)
        self.flair_loader_thread.error_occurred.connect(self.on_flair_model_error)
        self.flair_loader_thread.finished.connect(self.cleanup_flair_thread)

        self.threads_mutex.lock()
        self.active_threads.append(self.flair_loader_thread)
        self.threads_mutex.unlock()
        self.flair_loader_thread.start()

    def on_flair_model_loaded(self, model):
        if model:
            self.flair_classifier = model
            self.model_progress_bar.setVisible(False)

            if self.sentiment_mode == "Flair (Custom Model)":
                self.custom_model_button.setEnabled(True)
                if not self.flair_classifier_cuslang:
                    QMessageBox.warning(
                        self,
                        "Custom Model Required",
                        "Flair library loaded successfully!\n"
                        "Please load your custom model using the 'Load Model' button",
                    )

            if self.flair_loader_thread is not None:
                self.flair_loader_thread.quit()
                self.flair_loader_thread.wait()
                self.flair_loader_thread = None

            self.change_sentiment_mode(self.sentiment_mode)
        else:
            QMessageBox.critical(
                self, "Error", "Flair model failed to load. Please try again."
            )

    def on_flair_model_error(self, error):
        self.model_progress_bar.setVisible(False)
        QMessageBox.critical(
            self, "Loading Error", f"Failed to load Flair model: {error}"
        )

    def on_model_loaded(self, result, success):
        self.model_progress_bar.setVisible(False)
        self.custom_model_button.setEnabled(True)

        if success:
            try:
                from flair.models import TextClassifier
                from flair.data import Sentence

                if not isinstance(result, TextClassifier):
                    QMessageBox.critical(
                        self,
                        "Error",
                        "Invalid model type. Please load a valid Flair TextClassifier model.",
                    )
                    return

                test_sentence = Sentence("This is a test sentence")
                result.predict(test_sentence)
                if not test_sentence.labels:
                    raise ValueError("Model didn't produce any labels")

                self.flair_classifier_cuslang = result
                QMessageBox.information(
                    self,
                    "Success",
                    "Custom model loaded successfully! Flair will now use this model.",
                )

            except Exception as e:
                QMessageBox.critical(
                    self,
                    "Model Test Failed",
                    f"Model validation failed: {str(e)}\nPlease ensure this is a valid sentiment analysis model.",
                )
        else:
            QMessageBox.critical(
                self, "Error", f"Failed to load custom model: {result}"
            )

    def cleanup_flair_thread(self):
        self.threads_mutex.lock()
        if self.flair_loader_thread in self.active_threads:
            self.active_threads.remove(self.flair_loader_thread)
        self.threads_mutex.unlock()
        QTimer.singleShot(100, lambda: setattr(self, "flair_loader_thread", None))

    def show_sentiment_mode_info(self):
        dialog = QDialog(self)
        dialog.setWindowModality(Qt.NonModal)
        dialog.setWindowTitle("Sentiment Analysis Modes")
        dialog.setMinimumSize(500, 400)
        dialog.setSizeGripEnabled(True)

        layout = QVBoxLayout()

        text_browser = QTextBrowser()

        mode_info = (
            "<h2>Sentiment Analysis Modes</h2>"
            "<p>Select the most suitable sentiment analysis method based on your text type and analysis needs.</p>"

            "<h3>📝 TextBlob</h3>"
            "<p><b>Best for:</b> Formal texts, well-structured documents, news articles, research papers, and reports.</p>"
            "<p>TextBlob is a lexicon-based sentiment analysis tool that provides a simple yet effective approach for evaluating "
            "the sentiment of structured text. It assigns a polarity score (positive, negative, or neutral) and can also analyze "
            "subjectivity levels.</p>"

            "<h3>💬 VADER (Valence Aware Dictionary and sEntiment Reasoner)</h3>"
            "<p><b>Best for:</b> Social media posts, tweets, short comments, chat messages, and informal reviews.</p>"
            "<p>VADER is specifically designed for analyzing short, informal texts that often contain slang, emojis, and punctuation-based emotions. "
            "It is a rule-based sentiment analysis model that efficiently determines sentiment intensity and works exceptionally well for real-time applications.</p>"

            "<h3>🔬 Flair</h3>"
            "<p><b>Best for:</b> Long-form content, product reviews, professional documents, and AI-based deep sentiment analysis.</p>"
            "<p>Flair utilizes deep learning techniques for sentiment analysis, making it highly accurate for complex texts. "
            "It is ideal for analyzing large-scale textual data, capturing context more effectively than traditional rule-based models. "
            "However, it requires more computational resources compared to TextBlob and VADER.</p>"

            "<h3>🌐 Important Note for Language Support</h3>"
            "<p>While this application supports non-English text through automatic translation, it is <b>highly recommended</b> to use <b>manually translated and "
            "refined English text</b> for the most accurate sentiment analysis. The built-in automatic translation feature may not always function correctly, "
            "leading to potential misinterpretations or inaccurate sentiment results.</p>"
            "<p>For the best performance, ensure that non-English text is properly reviewed and adjusted before sentiment analysis. 🚀</p>"

            "<h3>📌 Custom Lexicon Format Example</h3>"
            "<p>Below is an example of a custom lexicon format for sentiment analysis:</p>"
            "<pre style='background-color:#f4f4f4; padding:10px; border-radius:5px;'>"
            "excellent   1.5\n"
            "awful      -1.5\n"
            "not        negation         # Mark as negation word\n"
            "intensely  intensifier:1.7  # Custom intensifier with multiplier"
            "</pre>"
            "<p>This custom lexicon allows fine-tuning of sentiment scores by adding custom words, negations, and intensifiers to improve sentiment analysis accuracy.</p>"
        )

        text_browser.setHtml(mode_info)
        text_browser.setOpenExternalLinks(True)
        text_browser.setReadOnly(True)

        layout.addWidget(text_browser)

        close_button = QPushButton("Close")
        close_button.clicked.connect(dialog.accept)
        layout.addWidget(close_button)

        dialog.setLayout(layout)
        dialog.show()

def is_connected():
    try:
        socket.create_connection(("8.8.8.8", 53), timeout=3)
        return True
    except OSError:
        return False

if __name__ == "__main__":
    app = QApplication(sys.argv)
    loop = QEventLoop(app)
    asyncio.set_event_loop(loop)

    ex = WordCloudGenerator()
    ex.show()

    with loop:
        sys.exit(loop.run_forever())
