import sys
import os
import re
import gc
from pathlib import Path
from functools import lru_cache, cached_property
from contextlib import contextmanager
import asyncio
from collections import Counter
import socket
import base64
import importlib
import hashlib
import tempfile
import shutil
import logging
import logging.handlers
import time
import json
from typing import Optional, Union, Dict, Any
import torch
import psutil
os.environ["QT_API"] = "pyside6"
from wordcloud import WordCloud, STOPWORDS
from sklearn.feature_extraction.text import CountVectorizer, TfidfVectorizer
from sklearn.decomposition import LatentDirichletAllocation, NMF
import yake
import rake_nltk
from qasync import QEventLoop
from deep_translator import GoogleTranslator
from PySide6.QtWidgets import (
    QApplication, QMainWindow, QFileDialog, QMessageBox, QLabel, QPushButton,
    QVBoxLayout, QGridLayout, QWidget, QLineEdit, QComboBox, QSpinBox, QDialog,
    QTextEdit, QProgressBar, QFrame, QColorDialog, QInputDialog, QTextBrowser,
    QHBoxLayout, QGroupBox, QTabWidget, QTableWidget, QTableWidgetItem, QHeaderView, QStyle,
    QDialogButtonBox, QSizePolicy, QTreeWidget, QTreeWidgetItem, QScrollArea
)
from PySide6.QtCore import QThread, Signal, Qt, QTimer, QMutex, QThreadPool, QObject, QMutexLocker
from PySide6.QtGui import QIcon, QGuiApplication, QPixmap, QPainter, QColor, QLinearGradient
import matplotlib
matplotlib.use("QtAgg")
import numpy as np
from PIL import Image
from matplotlib.font_manager import FontProperties
from matplotlib.colors import LinearSegmentedColormap
from sumy.parsers.plaintext import PlaintextParser
from sumy.nlp.tokenizers import Tokenizer
from sumy.summarizers.lsa import LsaSummarizer
os.environ["QT_SCALE_FACTOR"] = "1"
"""FUNGSI SETUP & UTILITAS GLOBAL"""
def setup_logging():
    """Setup application logging with enhanced formatting"""
    log_dir = Path.home() / ".textplora" / "logs"
    log_dir.mkdir(parents=True, exist_ok=True)
    log_file = log_dir / "textplora.log"
    fmt = logging.Formatter(
        '%(asctime)s - %(name)s - [%(levelname)s] - %(message)s'
    )
    file_handler = logging.handlers.RotatingFileHandler(
        log_file, 
        maxBytes=2*1024*1024,
        backupCount=3
    )
    file_handler.setFormatter(fmt)
    console_handler = logging.StreamHandler()
    console_handler.setFormatter(fmt)
    root_logger = logging.getLogger()
    root_logger.setLevel(logging.INFO)
    root_logger.addHandler(file_handler)
    root_logger.addHandler(console_handler)
    return root_logger
logger = setup_logging()
def setup_dll_path():
    """Mengatur jalur DLL"""
    if getattr(sys, 'frozen', False):
        application_path = os.path.dirname(sys.executable)
        dll_path = os.path.join(application_path, 'lib')
        if dll_path not in os.environ['PATH']:
            os.environ['PATH'] = dll_path + os.pathsep + os.environ['PATH']
setup_dll_path()
if getattr(sys, 'frozen', False):
    APP_DIR = Path(sys.executable).parent
else:
    APP_DIR = Path(__file__).parent
ICON_PATH = APP_DIR / "icon.ico"
logo_path = APP_DIR / "logo.png"
def is_connected():
    """Cek koneksi internet"""
    try:
        socket.create_connection(("8.8.8.8", 53), timeout=3)
        return True
    except OSError:
        return False
@lru_cache(maxsize=128)
def load_stopwords():
    """Memuat stopwords"""
    if getattr(sys, 'frozen', False):
        stopwords_path = os.path.join(os.path.dirname(sys.executable), "lib", "wordcloud", "stopwords")
    else:
        from wordcloud import STOPWORDS
        return set(STOPWORDS)
    try:
        with open(stopwords_path, 'r', encoding='utf-8') as f:
            return set(word.strip() for word in f)
    except Exception as e:
        logger.warning(f"Could not load stopwords from {stopwords_path}: {e}")
        return set()        
def safe_open(file_path, mode='r', encoding=None, **kwargs):
    """
    Membuka file dengan aman setelah memvalidasi path
    Args:
        file_path (str): Path file yang akan dibuka
        mode (str): Mode untuk membuka file ('r', 'w', dll)
        encoding (str, optional): Encoding yang digunakan
        **kwargs: Argumen tambahan untuk fungsi open()
    Returns:
        file: File object yang telah dibuka
    Raises:
        ValueError: Jika path tidak valid
        IOError: Jika file tidak dapat dibuka
    """
    validated_path = PathValidator.sanitize_path(file_path)
    return open(validated_path, mode=mode, encoding=encoding, **kwargs)
def safe_read_file(file_path, encoding='utf-8'):
    """
    Membaca isi file dengan aman dengan penanganan kesalahan yang lebih baik
    Args:
        file_path (str): Path file yang akan dibaca
        encoding (str): Encoding yang digunakan
    Returns:
        str: Isi file
    Raises:
        FileNotFoundError: Jika file tidak ditemukan
        PermissionError: Jika tidak ada izin untuk membaca file
        UnicodeDecodeError: Jika file tidak dapat didekode dengan encoding yang diberikan
        IOError: Jika file tidak dapat dibaca karena alasan lain
    """
    try:
        validated_path = PathValidator.sanitize_path(file_path)
        if not os.path.exists(validated_path):
            raise FileNotFoundError(f"File not found: {validated_path}")
        if not os.access(validated_path, os.R_OK):
            raise PermissionError(f"Permission denied: {validated_path}")
        file_size = os.path.getsize(validated_path)
        if file_size > 100 * 1024 * 1024:
            raise IOError(f"File too large: {file_size} bytes")
        with open(validated_path, 'r', encoding=encoding) as f:
            return f.read()
    except UnicodeDecodeError:
        raise
    except Exception as e:
        logger.error(f"Error reading file {file_path}: {e}")
        raise FileLoadError(f"Tidak dapat membaca file: {e}")
def safe_write_file(file_path, content, encoding='utf-8'):
    """
    Menulis ke file dengan aman
    Args:
        file_path (str): Path file yang akan ditulis
        content (str): Konten yang akan ditulis
        encoding (str): Encoding yang digunakan
    Raises:
        ValueError: Jika path tidak valid
        IOError: Jika file tidak dapat ditulis
    """
    try:
        with safe_open(file_path, 'w', encoding=encoding) as f:
            f.write(content)
    except Exception as e:
        logger.error(f"Error writing to file {file_path}: {e}")
        raise IOError(f"Tidak dapat menulis ke file: {e}")
def is_qobject_valid(obj):
    """
    Memeriksa apakah objek QObject masih valid (belum dihapus)
    Lebih aman daripada menggunakan sip.isdeleted()
    """
    if obj is None:
        return False
    try:
        obj.objectName()
        return True
    except RuntimeError:
        return False        
"""CLASS DASAR & UTILITAS"""
class AppConstants:
    """Application constants and configuration values with improved organization"""
    WINDOW_SIZE = (550, 870)
    DEFAULT_FONT_SIZE = 14
    DEFAULT_MAX_WORDS = 200
    MAX_FILE_SIZE = 100 * 1024 * 1024
    SUPPORTED_EXTENSIONS = {
        'text': ['.txt', '.csv'],
        'document': ['.doc', '.docx', '.pdf'],
        'spreadsheet': ['.xlsx', '.xls'],
    }
    CACHE_SIZE = 128
    THREAD_TIMEOUT = {
        "SOFT": 1000,
        "FORCE": 500,
        "TERMINATION": 2000
    }
    SENTIMENT_MODES = [
        "TextBlob", "TextBlob (Custom Lexicon)",
        "VADER", "VADER (Custom Lexicon)",
        "Flair", "Flair (Custom Model)"
    ]
    SCORE_TYPES = {
        "TextBlob": "Polarity Score", 
        "TextBlob (Custom Lexicon)": "Polarity Score",
        "VADER": "Compound Score",
        "VADER (Custom Lexicon)": "Compound Score",        
        "Flair": "Confidence Score",
        "Flair (Custom Model)": "Confidence Score"
    }
    ABOUT_TEXT = "PGh0bWw+Cjxib2R5IHN0eWxlPSJmb250LWZhbWlseTogQXJpYWwsIHNhbnMtc2VyaWY7IHRleHQtYWxpZ246IGNlbnRlcjsgcGFkZGluZzogMTBweDsiPgogICAgPGltZyBzcmM9Intsb2dvX3BhdGh9IiB3aWR0aD0iMTAwIiBoZWlnaHQ9IjEwMCIgYWx0PSJUZXh0cGxvcmEgTG9nbyI+CiAgICA8aDI+VGV4dHBsb3JhPC9oMj4KICAgIDxwPjxiPlZlcnNpb246PC9iPiAxLjY8L3A+CiAgICA8cD4mY29weTsgMjAyNSBNLiBBZGliIFphdGEgSWxtYW08L3A+CiAgICA8aHI+CiAgICA8cD5UZXh0cGxvcmEgKGZvcm1lcmx5IFdDR2VuKSBpcyBhIFB5dGhvbi1iYXNlZCBhcHBsaWNhdGlvbiB0aGF0IGhlbHBzIHlvdSBhbmFseXplIGZlZWRiYWNrLCBjb25kdWN0IHJlc2VhcmNoLCBhbmQgZXh0cmFjdCBtZWFuaW5nZnVsIGluc2lnaHRzIHdpdGggZWFzZS48L3A+CiAgICA8aHI+CiAgICA8cD48Yj5NYWluIGxpYnJhcmllczo8L2I+IFB5U2lkZTYsIG1hdHBsb3RsaWIsIHdvcmRjbG91ZCwgc2tsZWFybiwgdmFkZXJTZW50aW1lbnQsIHRleHRibG9iLCBmbGFpciwgc3VteSwgeWFrZSwgcmFrZV9ubHRrLCBQaWxsb3cuPC9wPgogICAgPGhyPgogICAgPHA+PGI+R2l0SHViIFJlcG9zaXRvcnk6PC9iPjwvcD4KICAgIDxwPjxhIGhyZWY9Imh0dHBzOi8vZ2l0aHViLmNvbS96YXRhaWxtL3djbG91ZGd1aSI+aHR0cHM6Ly9naXRodWIuY29tL3phdGFpbG0vd2Nsb3VkZ3VpPC9hPjwvcD4KICAgIDxocj4KICAgIDxwPjxiPkxpY2Vuc2U6PC9iPjwvcD4KICAgIDxwPkZyZWUgZm9yIHBlcnNvbmFsICYgZWR1Y2F0aW9uYWwgdXNlIChDQyBCWS1OQyA0LjApLjwvcD4KICAgIDxwPkxlYXJuIG1vcmU6IDxhIGhyZWY9Imh0dHBzOi8vY3JlYXRpdmVjb21tb25zLm9yZy9saWNlbnNlcy9ieS1uYy80LjAvIj5DQyBCWS1OQyA0LjA8L2E+PC9wPgo8L2JvZHk+CjwvaHRtbD4="
    MODE_INFO = "PGgyPlNlbnRpbWVudCBBbmFseXNpcyBNb2RlczwvaDI+CjxwPlNlbGVjdCB0aGUgbW9zdCBzdWl0YWJsZSBzZW50aW1lbnQgYW5hbHlzaXMgbWV0aG9kIGJhc2VkIG9uIHlvdXIgdGV4dCB0eXBlIGFuZCBhbmFseXNpcyBuZWVkcy48L3A+CjxoMz5UZXh0QmxvYjwvaDM+CjxwPjxiPkJlc3QgZm9yOjwvYj4gRm9ybWFsIHRleHRzLCB3ZWxsLXN0cnVjdHVyZWQgZG9jdW1lbnRzLCBuZXdzIGFydGljbGVzLCByZXNlYXJjaCBwYXBlcnMsIGFuZCByZXBvcnRzLjwvcD4KPHA+VGV4dEJsb2IgaXMgYSBsZXhpY29uLWJhc2VkIHNlbnRpbWVudCBhbmFseXNpcyB0b29sIHRoYXQgcHJvdmlkZXMgYSBzaW1wbGUgeWV0IGVmZmVjdGl2ZSBhcHByb2FjaCBmb3IgZXZhbHVhdGluZyB0aGUgc2VudGltZW50IG9mIHN0cnVjdHVyZWQgdGV4dC4gSXQgYXNzaWducyBhIHBvbGFyaXR5IHNjb3JlIChwb3NpdGl2ZSwgbmVnYXRpdmUsIG9yIG5ldXRyYWwpIGFuZCBjYW4gYWxzbyBhbmFseXplIHN1YmplY3Rpdml0eSBsZXZlbHMuPC9wPgo8aDM+VkFERVIgKFZhbGVuY2UgQXdhcmUgRGljdGlvbmFyeSBhbmQgc0VudGltZW50IFJlYXNvbmVyKTwvaDM+CjxwPjxiPkJlc3QgZm9yOjwvYj4gU29jaWFsIG1lZGlhIHBvc3RzLCB0d2VldHMsIHNob3J0IGNvbW1lbnRzLCBjaGF0IG1lc3NhZ2VzLCBhbmQgaW5mb3JtYWwgcmV2aWV3cy48L3A+CjxwPlZBREVSIGlzIHNwZWNpZmljYWxseSBkZXNpZ25lZCBmb3IgYW5hbHl6aW5nIHNob3J0LCBpbmZvcm1hbCB0ZXh0cyB0aGF0IG9mdGVuIGNvbnRhaW4gc2xhbmcsIGVtb2ppcywgYW5kIHB1bmN0dWF0aW9uLWJhc2VkIGVtb3Rpb25zLiBJdCBpcyBhIHJ1bGUtYmFzZWQgc2VudGltZW50IGFuYWx5c2lzIG1vZGVsIHRoYXQgZWZmaWNpZW50bHkgZGV0ZXJtaW5lcyBzZW50aW1lbnQgaW50ZW5zaXR5IGFuZCB3b3JrcyBleGNlcHRpb25hbGx5IHdlbGwgZm9yIHJlYWwtdGltZSBhcHBsaWNhdGlvbnMuPC9wPgo8aDM+RmxhaXI8L2gzPgo8cD48Yj5CZXN0IGZvcjo8L2I+IExvbmctZm9ybSBjb250ZW50LCBwcm9kdWN0IHJldmlld3MsIHByb2Zlc3Npb25hbCBkb2N1bWVudHMsIGFuZCBBSS1iYXNlZCBkZWVwIHNlbnRpbWVudCBhbmFseXNpcy48L3A+CjxwPkZsYWlyIHV0aWxpemVzIGRlZXAgbGVhcm5pbmcgdGVjaG5pcXVlcyBmb3Igc2VudGltZW50IGFuYWx5c2lzLCBtYWtpbmcgaXQgaGlnaGx5IGFjY3VyYXRlIGZvciBjb21wbGV4IHRleHRzLiBJdCBpcyBpZGVhbCBmb3IgYW5hbHl6aW5nIGxhcmdlLXNjYWxlIHRleHR1YWwgZGF0YSwgY2FwdHVyaW5nIGNvbnRleHQgbW9yZSBlZmZlY3RpdmVseSB0aGFuIHRyYWRpdGlvbmFsIHJ1bGUtYmFzZWQgbW9kZWxzLiBIb3dldmVyLCBpdCByZXF1aXJlcyBtb3JlIGNvbXB1dGF0aW9uYWwgcmVzb3VyY2VzIGNvbXBhcmVkIHRvIFRleHRCbG9iIGFuZCBWQURFUi48L3A+CjxoMj5JbXBvcnRhbnQgTm90ZSBmb3IgTGFuZ3VhZ2UgU3VwcG9ydDwvaDI+CjxwPldoaWxlIHRoaXMgYXBwbGljYXRpb24gc3VwcG9ydHMgbm9uLUVuZ2xpc2ggdGV4dCB0aHJvdWdoIGF1dG9tYXRpYyB0cmFuc2xhdGlvbiwgaXQgaXMgPGI+aGlnaGx5IHJlY29tbWVuZGVkPC9iPiB0byB1c2UgPGI+bWFudWFsbHkgdHJhbnNsYXRlZCBhbmQgcmVmaW5lZCBFbmdsaXNoIHRleHQ8L2I+IGZvciB0aGUgbW9zdCBhY2N1cmF0ZSBzZW50aW1lbnQgYW5hbHlzaXMuIFRoZSBidWlsdC1pbiBhdXRvbWF0aWMgdHJhbnNsYXRpb24gZmVhdHVyZSBtYXkgbm90IGFsd2F5cyBmdW5jdGlvbiBjb3JyZWN0bHksIGxlYWRpbmcgdG8gcG90ZW50aWFsIG1pc2ludGVycHJldGF0aW9ucyBvciBpbmFjY3VyYXRlIHNlbnRpbWVudCByZXN1bHRzLiBFbnN1cmUgeW91ciBjdXN0b20gbGV4aWNvbiBhbmQgbW9kZWwgaXMgdHJhaW5lZCBvbiBzaW1pbGFyIHRleHQgZG9tYWlucyBhcyB5b3VyIGFuYWx5c2lzIGRhdGEgZm9yIG9wdGltYWwgcGVyZm9ybWFuY2UuPC9wPgo8cD5Gb3IgdGhlIGJlc3QgcGVyZm9ybWFuY2UsIGVuc3VyZSB0aGF0IG5vbi1FbmdsaXNoIHRleHQgaXMgcHJvcGVybHkgcmV2aWV3ZWQgYW5kIGFkanVzdGVkIGJlZm9yZSBzZW50aW1lbnQgYW5hbHlzaXMuPC9wPgo8aDI+Q3VzdG9tIExleGljb24gRm9ybWF0IEV4YW1wbGU8L2gyPgo8cD5CZWxvdyBpcyBhbiBleGFtcGxlIG9mIGEgY3VzdG9tIGxleGljb24gZm9ybWF0IGZvciBzZW50aW1lbnQgYW5hbHlzaXM6PC9wPgo8cHJlIHN0eWxlPSdiYWNrZ3JvdW5kLWNvbG9yOiNmNGY0ZjQ7IGNvbG9yOiBibGFjazsgcGFkZGluZzoxMHB4OyBib3JkZXItcmFkaXVzOjVweDsnPgpleGNlbGxlbnQgICAxLjUKYXdmdWwgICAgICAtMS41Cm5vdCAgICAgICAgbmVnYXRpb24gICAgICAgICAjIE1hcmsgYXMgbmVnYXRpb24gd29yZAppbnRlbnNlbHkgIGludGVuc2lmaWVyOjEuNyAgIyBDdXN0b20gaW50ZW5zaWZpZXIgd2l0aCBtdWx0aXBsaWVyCjwvcHJlPgo8cD5UaGlzIGN1c3RvbSBsZXhpY29uIGFsbG93cyBmaW5lLXR1bmluZyBvZiBzZW50aW1lbnQgc2NvcmVzIGJ5IGFkZGluZyBjdXN0b20gd29yZHMsIG5lZ2F0aW9ucywgYW5kIGludGVuc2lmaWVycyB0byBpbXByb3ZlIHNlbnRpbWVudCBhbmFseXNpcyBhY2N1cmFjeS48L3A+CjxoMj5DdXN0b20gRmxhaXIgTW9kZWwgUmVxdWlyZW1lbnRzPC9oMj4KPHA+Rm9yIGN1c3RvbSBGbGFpciBtb2RlbHMgKDxiPi5wdDwvYj4gZmlsZXMpLCBlbnN1cmU6PC9wPgo8cHJlIHN0eWxlPSdiYWNrZ3JvdW5kLWNvbG9yOiNmNGY0ZjQ7IGNvbG9yOiBibGFjazsgcGFkZGluZzoxMHB4OyBib3JkZXItcmFkaXVzOjVweDsnPgoxLiBMYWJlbHMgbXVzdCBiZTogUE9TSVRJVkUsIE5FR0FUSVZFLCBvciBORVVUUkFMCjIuIE1vZGVsIG91dHB1dHMgY29uZmlkZW5jZSBzY29yZXMgKDAtMSkKMy4gVGVzdGVkIHdpdGggRmxhaXIgdjAuMTIrIGNvbXBhdGliaWxpdHkKPC9wcmU+CjxwPlRoZSBhcHBsaWNhdGlvbiB3aWxsIGF1dG9tYXRpY2FsbHkgdmFsaWRhdGUgbW9kZWwgbGFiZWxzIGR1cmluZyBsb2FkaW5nLjwvcD4="
class PathValidator:
    """Sanitasi jalur file untuk mencegah serangan path traversal"""
    @staticmethod
    def get_allowed_dirs():
        return [
            os.path.abspath(os.path.expanduser("~")),
            os.path.abspath(APP_DIR),
            os.path.abspath(tempfile.gettempdir())
        ]
    @staticmethod
    def sanitize_path(path: str) -> str:
        if not path:
            return path
        path = os.path.normpath(path)
        if os.path.isabs(path):
            abs_path = os.path.abspath(path)
            allowed_dirs = PathValidator.get_allowed_dirs()
            if not any(abs_path.startswith(allowed_dir) for allowed_dir in allowed_dirs):
                logger.warning(f"Path traversal attempt blocked: {path}")
                raise ValueError(f"Access denied to path: {path}")
        return path
class LazyLoader:
    """
    Optimized lazy loading system for managing heavy dependencies.
    This class implements a singleton pattern with caching to ensure
    efficient loading and reuse of module imports. It helps reduce
    memory usage and startup time by loading modules only when needed.
    Attributes:
        _instances (dict): Singleton instances storage
        _cache (dict): Module cache storage
    Methods:
        load(module_name, class_name=None): 
            Lazily loads a module or class and caches it
    Example:
        >>> loader = LazyLoader()
        >>> TextBlob = loader.load('textblob', 'TextBlob')
        >>> vader = loader.load('vaderSentiment.vaderSentiment')
    """
    _instances = {}
    _cache = {}
    @classmethod
    def load(cls, module_name: str, class_name: str = None) -> object:
        """
        Lazily load and cache a module or specific class.
        Args:
            module_name (str): Name of the module to import
            class_name (str, optional): Specific class to import from module
        Returns:
            object: Loaded module or class
        Raises:
            ImportError: If module/class cannot be loaded
        """
        key = f"{module_name}.{class_name}" if class_name else module_name
        if key not in cls._cache:
            try:
                module = __import__(module_name, fromlist=[class_name] if class_name else [])
                cls._cache[key] = getattr(module, class_name) if class_name else module
            except ImportError as e:
                logger.error(f"Failed to load {key}: {e}")
                return None
        return cls._cache[key]
class ButtonStateManager:
    """Manages button states during processing operations"""
    def __init__(self, main_window):
        self.main_window = main_window
        self.button_states = {}
        self.process_buttons = [
            'generate_wordcloud_button',
            'sentiment_button',
            'analyze_topics_btn',
            'extract_keywords_btn',
            'view_fulltext_button',
            'text_stats_button',
            'summarize_button',
            'load_file_button',
        ]
    def save_states(self):
        """Save current state of all process buttons"""
        for btn_name in self.process_buttons:
            if hasattr(self.main_window, btn_name):
                btn = getattr(self.main_window, btn_name)
                self.button_states[btn_name] = btn.isEnabled()
            elif hasattr(self.main_window.topic_tab, btn_name):
                btn = getattr(self.main_window.topic_tab, btn_name)
                self.button_states[btn_name] = btn.isEnabled()
    def disable_other_buttons(self, active_button):
        """Disable all process buttons except the active one"""
        self.save_states()
        for btn_name in self.process_buttons:
            if btn_name != active_button:
                if hasattr(self.main_window, btn_name):
                    getattr(self.main_window, btn_name).setEnabled(False)
                elif hasattr(self.main_window.topic_tab, btn_name):
                    getattr(self.main_window.topic_tab, btn_name).setEnabled(False)
    def restore_states(self):
        """Restore buttons to their previous states"""
        for btn_name, was_enabled in self.button_states.items():
            if hasattr(self.main_window, btn_name):
                getattr(self.main_window, btn_name).setEnabled(was_enabled)
            elif hasattr(self.main_window.topic_tab, btn_name):
                getattr(self.main_window.topic_tab, btn_name).setEnabled(was_enabled)
        self.button_states.clear()
"""CLASS RESOURCE & CACHE MANAGEMENT"""
class ResourceController(QObject):
    """Unified resource management system for the application"""
    font_loaded = Signal(str)
    palette_updated = Signal(str)
    cache_cleared = Signal(int)
    resource_error = Signal(str, str)
    colormap_loading_started = Signal()
    colormap_loading_progress = Signal(str)
    colormap_loading_finished = Signal(list)
    def __init__(self, parent=None):
        super().__init__(parent)
        if hasattr(parent, 'cache_manager'):
            self.cache_manager = parent.cache_manager
        else:
            self.cache_manager = EnhancedCacheManager()        
        self.parent = parent
        self._mutex = QMutex()
        self._resources = {}
        self._cleanup_hooks = {}
        self._cached_fonts = {}
        self._cached_colormaps = {}
        self._cached_models = {}
        self.model_cache = {}
        self.custom_color_palettes = {}
        self.current_figures = []
        self.colormap_loader_thread = None
        self.font_loader_thread = None
    def register_resource(self, resource_id, resource, cleanup_hook=None):
        """Register a resource with optional cleanup hook"""
        with QMutexLocker(self._mutex):
            self._resources[resource_id] = resource
            if cleanup_hook and callable(cleanup_hook):
                self._cleanup_hooks[resource_id] = cleanup_hook
            logger.debug(f"Resource registered: {resource_id}")
    def get_resource(self, resource_id, creator_func=None):
        """Get a resource, creating it if necessary"""
        with QMutexLocker(self._mutex):
            if resource_id not in self._resources and creator_func:
                try:
                    self._resources[resource_id] = creator_func()
                except Exception as e:
                    logger.error(f"Error creating resource {resource_id}: {e}")
                    self.resource_error.emit("resource", f"Failed to create {resource_id}: {e}")
                    return None
            return self._resources.get(resource_id)
    def cleanup_resource(self, resource_id):
        """Cleanup a specific resource"""
        with QMutexLocker(self._mutex):
            if resource_id in self._cleanup_hooks:
                try:
                    self._cleanup_hooks[resource_id](self._resources.get(resource_id))
                    del self._cleanup_hooks[resource_id]
                    del self._resources[resource_id]
                    logger.debug(f"Resource cleaned up: {resource_id}")
                    return True
                except Exception as e:
                    logger.error(f"Error cleaning up resource {resource_id}: {e}")
                    self.resource_error.emit("cleanup", f"Failed to cleanup {resource_id}: {e}")
                    return False
            return False
    def cleanup_all_resources(self):
        """Cleanup all registered resources"""
        success_count = 0
        error_count = 0
        with QMutexLocker(self._mutex):
            if self.cleanup_all_graphics():
                success_count += 1
            else:
                error_count += 1
            if self.res_clear_model_cache():
                success_count += 1
            else:
                error_count += 1
            for resource_id in list(self._cleanup_hooks.keys()):
                if self.cleanup_resource(resource_id):
                    success_count += 1
                else:
                    error_count += 1
            self._resources.clear()
            self._cleanup_hooks.clear()
        logger.info(f"Resource cleanup completed: {success_count} successful, {error_count} failed")
        return success_count > 0
    def cleanup_all_graphics(self):
        """Cleanup all graphics resources"""
        try:
            self.res_cleanup_figure()
            self._cached_colormaps.clear()
            self._cached_fonts.clear()
            return True
        except Exception as e:
            logger.error(f"Error cleaning up graphics: {e}")
            self.resource_error.emit("graphics", f"Error cleaning up graphics: {e}")
            return False
    def initialize(self):
        """Inisialisasi semua resource"""
        self.res_load_matplotlib_fonts()
        self.res_load_colormaps_async()
    def res_load_colormaps_async(self):
        """Load colormap matplotlib secara asynchronous"""
        try:
            if self.colormap_loader_thread and self.colormap_loader_thread.isRunning():
                self.colormap_loader_thread.stop()
                self.colormap_loader_thread.wait(500)
            self.colormap_loader_thread = ColormapLoaderThread(self)
            self.colormap_loader_thread.loading_complete.connect(self._on_colormaps_loaded)
            self.colormap_loader_thread.loading_progress.connect(self.colormap_loading_progress.emit)
            self.colormap_loader_thread.error_occurred.connect(
                lambda msg: self.resource_error.emit("colormap", msg)
            )
            self.colormap_loading_started.emit()
            self.colormap_loader_thread.start()
            return True
        except Exception as e:
            logger.error(f"Error starting colormap loader thread: {e}")
            self.resource_error.emit("colormap", f"Error loading colormaps: {e}")
            return False
    def _on_colormaps_loaded(self, colormaps):
        """Handle completion of colormap loading"""
        self._mutex.lock()
        try:
            self._cached_colormaps = colormaps
            for name, colors in self.custom_color_palettes.items():
                try:
                    from matplotlib.colors import LinearSegmentedColormap
                    self._cached_colormaps[name] = LinearSegmentedColormap.from_list(name, colors)
                except Exception as e:
                    logger.debug(f"Failed to load custom palette {name}: {e}")
            self.colormap_loading_finished.emit(list(self._cached_colormaps.keys()))
            logger.info(f"Loaded {len(self._cached_colormaps)} colormaps")
        finally:
            self._mutex.unlock()
    def res_load_matplotlib_fonts(self):
        """Adaptive font loading - delegate ke AppUI jika tersedia"""
        parent_app = self.parent
        if hasattr(parent_app, 'load_matplotlib_fonts_optimized'):
            logger.debug("ResourceController: Delegating font loading to AppUI optimized method")
            parent_app.load_matplotlib_fonts_optimized()
            return True
        try:
            self.font_loader_thread = FontLoaderThread(self)
            self.font_loader_thread.loading_complete.connect(self._on_fonts_loaded)
            self.font_loader_thread.loading_progress.connect(self.font_loaded.emit)
            self.font_loader_thread.error_occurred.connect(
                lambda msg: self.resource_error.emit("font", msg)
            )
            self.font_loader_thread.start()
            return True
        except Exception as e:
            logger.error(f"ResourceController: Error loading fonts: {e}")
            self.resource_error.emit("font", f"Error loading fonts: {e}")
            return False
    def res_get_available_fonts(self):
        """Get list of available fonts dengan delegate ke AppUI jika perlu"""
        if not self._cached_fonts or len(self._cached_fonts) == 0:
            parent_app = self.parent
            if hasattr(parent_app, '_cached_fonts') and parent_app._cached_fonts:
                self._cached_fonts = parent_app._cached_fonts.copy()
                logger.debug("ResourceController: Using fonts from AppUI cache")
                return self._cached_fonts
            logger.debug("ResourceController: No cached fonts, loading fonts via traditional method")
            self.res_load_matplotlib_fonts()
        return self._cached_fonts
    def _on_fonts_loaded(self, fonts):
        """Handle completion of font loading"""
        self._mutex.lock()
        try:
            if not fonts or len(fonts) == 0:
                logger.warning("Font loader returned empty font dictionary")
                return
            if not self._cached_fonts or self._cached_fonts != fonts:
                self._cached_fonts = fonts
                if not hasattr(self, '_fonts_logged') or not self._fonts_logged:
                    logger.info(f"Loaded {len(self._cached_fonts)} fonts")
                    self._fonts_logged = True
        finally:
            self._mutex.unlock()      
    def res_load_colormaps(self):
        """Load colormap matplotlib"""
        try:
            import matplotlib.pyplot as plt
            self._cached_colormaps = {}
            for name in plt.colormaps():
                try:
                    self._cached_colormaps[name] = plt.get_cmap(name)
                except Exception as e:
                    logger.debug(f"Failed to load colormap {name}: {e}")
            for name, colors in self.custom_color_palettes.items():
                try:
                    from matplotlib.colors import LinearSegmentedColormap
                    self._cached_colormaps[name] = LinearSegmentedColormap.from_list(name, colors)
                except Exception as e:
                    logger.debug(f"Failed to load custom palette {name}: {e}")
            logger.info(f"Loaded {len(self._cached_colormaps)} colormaps")
            return self._cached_colormaps
        except Exception as e:
            logger.error(f"Error loading colormaps: {e}")
            self.resource_error.emit("colormap", f"Error loading colormaps: {e}")
            return {}
    def res_create_custom_palette(self, color_list, palette_name=None):
        """Buat custom palette dari daftar warna"""
        try:
            if not color_list:
                logger.error("Empty color list provided for custom palette")
                self.resource_error.emit("palette", "Empty color list provided")
                return None
            if not palette_name:
                palette_name = f"custom_palette_{len(self.custom_color_palettes) + 1}"
            if not palette_name.startswith('custom_'):
                palette_name = f"custom_{palette_name}"
            if not hasattr(self, 'custom_color_palettes'):
                self.custom_color_palettes = {}
            color_list_copy = color_list.copy()
            self.custom_color_palettes[palette_name] = color_list_copy
            from matplotlib.colors import LinearSegmentedColormap
            self._cached_colormaps[palette_name] = LinearSegmentedColormap.from_list(palette_name, color_list_copy)
            self.palette_updated.emit(palette_name)
            logger.info(f"Created custom palette: {palette_name} with {len(color_list_copy)} colors: {color_list_copy}")
            return palette_name
        except Exception as e:
            logger.error(f"Error creating custom palette: {e}")
            self.resource_error.emit("palette", f"Error creating custom palette: {e}")
            return None
    def res_get_palette_names(self):
        """Dapatkan daftar nama palette yang tersedia"""
        if not self._cached_colormaps:
            self.res_load_colormaps()
        return list(self._cached_colormaps.keys())
    def res_load_flair_model(self, model_path=None, custom=False):
        """Load model Flair untuk analisis sentimen"""
        try:
            cache_key = f"flair_model_{model_path}_{custom}"
            if cache_key in self.model_cache:
                return self.model_cache[cache_key]
            from flair.models import TextClassifier
            if custom and model_path:
                model = TextClassifier.load(model_path)
            else:
                model = TextClassifier.load('en-sentiment')
            self.model_cache[cache_key] = model
            logger.info(f"Loaded Flair model: {model_path if custom else 'en-sentiment'}")
            return model
        except Exception as e:
            logger.error(f"Error loading Flair model: {e}")
            self.resource_error.emit("model", f"Error loading Flair model: {e}")
            return None
    def res_unload_ml_models(self):
        """Unload semua model ML"""
        try:
            self.model_cache.clear()
            import gc
            gc.collect()
            logger.info("ML models unloaded")
            return True
        except Exception as e:
            logger.error(f"Error unloading ML models: {e}")
            self.resource_error.emit("model", f"Error unloading ML models: {e}")
            return False
    def res_cleanup_figure(self, figure=None):
        """Cleanup matplotlib figure"""
        try:
            import matplotlib.pyplot as plt
            if figure:
                try:
                    plt.close(figure)
                    if figure in self.current_figures:
                        self.current_figures.remove(figure)
                except Exception as e:
                    logger.debug(f"Error closing specific figure: {e}")
            else:
                try:
                    plt.close('all')
                    self.current_figures.clear()
                except Exception as e:
                    logger.debug(f"Error closing all figures: {e}")
            return True
        except Exception as e:
            logger.error(f"Error cleaning up figure: {e}")
            self.resource_error.emit("figure", f"Error cleaning up figure: {e}")
            return False
    def register_figure(self, figure):
        """Register figure untuk di-track"""
        if figure not in self.current_figures:
            self.current_figures.append(figure)
    def cleanup_all_graphics(self):
        """Cleanup semua resource grafis"""
        try:
            self.res_cleanup_figure()
            self._cached_colormaps.clear()
            self._cached_fonts.clear()
            return True
        except Exception as e:
            logger.error(f"Error cleaning up graphics: {e}")
            self.resource_error.emit("graphics", f"Error cleaning up graphics: {e}")
            return False
class EnhancedCacheManager(QObject):
    """Manages application caches and ensures proper cleanup during shutdown"""
    def __init__(self):
        super().__init__()
        self._regular_caches = {}
        self._graphics_caches = set()
        self._model_caches = set()
        self._mutex = QMutex()
    def register_graphics_cache(self, cache_obj):
        """Register a graphics-related cache object"""
        with QMutexLocker(self._mutex):
            self._graphics_caches.add(cache_obj)
    def register_model_cache(self, cache_obj):
        """Register an ML model cache object"""
        with QMutexLocker(self._mutex):
            self._model_caches.add(cache_obj)
    def register(self, cache_id, cache_obj=None, clear_method=None):
        """Register a regular cache with optional custom clear method"""
        with QMutexLocker(self._mutex):
            self._regular_caches[cache_id] = {
                'object': cache_obj,
                'clear_method': clear_method
            }
    def register_cached_function(self, func):
        """Register a function decorated with @lru_cache"""
        if hasattr(func, 'cache_clear'):
            with QMutexLocker(self._mutex):
                self._regular_caches[func.__name__] = {
                    'object': func,
                    'clear_method': func.cache_clear
                }
    def register_cached_property(self, instance, property_name):
        """Register a cached property"""
        if hasattr(instance, property_name):
            with QMutexLocker(self._mutex):
                self._regular_caches[f"{instance.__class__.__name__}.{property_name}"] = {
                    'object': instance,
                    'property': property_name
                }
    def _clear_lazy_loader_cache(self):
        """Clear LazyLoader module cache"""
        try:
            if LazyLoader._cache:
                LazyLoader._cache.clear()
                logger.info("LazyLoader cache cleared")
        except Exception as e:
            logger.error(f"Error clearing LazyLoader cache: {e}")
    def _clear_matplotlib_cache(self):
        """Clear matplotlib internal caches"""
        try:
            import matplotlib.pyplot as plt
            plt.rcParams.update(plt.rcParamsDefault)
            plt.close('all')
            logger.info("Matplotlib caches cleared")
        except Exception as e:
            logger.error(f"Error clearing matplotlib cache: {e}")
    def cleanup_on_exit(self):
        """Perform complete cache cleanup during application shutdown"""
        logger.info("Starting cache cleanup...")
        cleared_count = 0
        error_count = 0
        with QMutexLocker(self._mutex):
            for cache_id, cache_info in self._regular_caches.items():
                try:
                    if 'clear_method' in cache_info and cache_info['clear_method']:
                        cache_info['clear_method']()
                    elif 'property' in cache_info:
                        instance = cache_info['object']
                        prop_name = cache_info['property']
                        if hasattr(instance.__class__, prop_name):
                            delattr(instance.__class__, prop_name)
                    elif hasattr(cache_info['object'], 'clear'):
                        cache_info['object'].clear()
                    cleared_count += 1
                    logger.debug(f"Cleared cache: {cache_id}")
                except Exception as e:
                    error_count += 1
                    logger.error(f"Error clearing cache {cache_id}: {e}")
            for cache in self._graphics_caches:
                try:
                    if hasattr(cache, 'clear'):
                        cache.clear()
                    cleared_count += 1
                except Exception as e:
                    error_count += 1
                    logger.error(f"Error clearing graphics cache: {e}")
            for cache in self._model_caches:
                try:
                    if hasattr(cache, 'cleanup'):
                        cache.cleanup()
                    cleared_count += 1
                except Exception as e:
                    error_count += 1
                    logger.error(f"Error clearing model cache: {e}")
            try:
                self._clear_lazy_loader_cache()
                self._clear_matplotlib_cache()
                cleared_count += 2
            except Exception as e:
                error_count += 1
                logger.error(f"Error clearing system caches: {e}")
        logger.info(f"Cache cleanup completed. Cleared: {cleared_count}, Errors: {error_count}")
"""CLASS THREAD MANAGEMENT"""
class ThreadManager(QObject):
    """Manages application threads to ensure proper cleanup during shutdown"""
    thread_added = Signal(object)
    thread_removed = Signal(object)
    all_threads_stopped = Signal()
    thread_error = Signal(str)
    error_occurred = Signal(str)
    _is_stopping = False
    _is_shutting_down = False
    _threads = {}
    _lock = QMutex()
    def __init__(self, parent=None):
        super().__init__(parent)
        self._threads = {}
        self._mutex = QMutex()
        self._is_shutting_down = False
    @property
    def active_threads(self):
        """Return the number of active threads"""
        with QMutexLocker(self._lock):
            return len(self._threads)
    def add_thread(self, thread, thread_id=None):
        """Register a thread for management
        Args:
            thread: QThread instance
            thread_id: Optional identifier for the thread
        """
        if not thread:
            return
        self._mutex.lock()
        try:
            if thread_id is None:
                thread_id = id(thread)
            self._threads[thread_id] = thread
            thread.finished.connect(lambda: self._on_thread_finished(thread_id))
            logger.debug(f"Thread registered: {thread_id} ({thread.objectName() or 'unnamed'})")
        finally:
            self._mutex.unlock()
    def remove_thread(self, thread_or_id):
        """Remove a thread from management
        Args:
            thread_or_id: Thread instance or thread_id
        """
        self._mutex.lock()
        try:
            thread_id = thread_or_id if isinstance(thread_or_id, (int, str)) else id(thread_or_id)
            if thread_id in self._threads:
                thread = self._threads[thread_id]
                try:
                    thread.finished.disconnect()
                except:
                    pass
                del self._threads[thread_id]
                logger.debug(f"Thread removed: {thread_id}")
        finally:
            self._mutex.unlock()
    def _on_thread_finished(self, thread_id):
        """Handle thread completion"""
        self._mutex.lock()
        try:
            if thread_id in self._threads:
                thread = self._threads[thread_id]
                logger.debug(f"Thread finished: {thread_id} ({thread.objectName() or 'unnamed'})")
                if self._is_shutting_down:
                    thread.deleteLater()
                    del self._threads[thread_id]
        finally:
            self._mutex.unlock()
    def cleanup_finished_threads(self):
        """Remove finished threads from registry"""
        self._mutex.lock()
        try:
            finished_threads = []
            for thread_id, thread in self._threads.items():
                if not thread.isRunning():
                    finished_threads.append(thread_id)
            for thread_id in finished_threads:
                thread = self._threads[thread_id]
                try:
                    thread.finished.disconnect()
                except:
                    pass
                thread.deleteLater()
                del self._threads[thread_id]
            logger.debug(f"Cleaned up {len(finished_threads)} finished threads")
        finally:
            self._mutex.unlock()
    def stop_all_threads(self, wait_timeout=1000):
        """Stop all managed threads with proper error handling for deleted objects"""
        if self._is_stopping:
            return
        self._is_stopping = True
        logger.info("Stopping all threads...")
        thread_ids = list(self._threads.keys())
        stopped_threads = 0
        failed_threads = 0
        for thread_id in thread_ids:
            try:
                thread = self._threads.get(thread_id)
                if thread is None:
                    logger.debug(f"Thread {thread_id} already removed")
                    continue
                if is_qobject_valid(thread) and thread.isRunning():
                    logger.debug(f"Stopping thread: {thread_id}")
                    if hasattr(thread, 'cancel') and callable(thread.cancel):
                        try:
                            thread.cancel()
                        except Exception as e:
                            logger.debug(f"Error calling cancel on thread {thread_id}: {e}")
                    if hasattr(thread, 'requestInterruption') and callable(thread.requestInterruption):
                        try:
                            thread.requestInterruption()
                        except Exception as e:
                            logger.debug(f"Error requesting interruption on thread {thread_id}: {e}")
                    if not thread.wait(wait_timeout):
                        logger.warning(f"Thread {thread_id} did not stop gracefully, forcing termination")
                        if hasattr(thread, 'cleanup_resources') and callable(thread.cleanup_resources):
                            try:
                                thread.cleanup_resources()
                            except Exception as e:
                                logger.debug(f"Error cleaning up thread {thread_id}: {e}")
                        try:
                            if is_qobject_valid(thread) and thread.isRunning():
                                thread.terminate()
                                if thread.wait(500):
                                    stopped_threads += 1
                                else:
                                    failed_threads += 1
                            else:
                                stopped_threads += 1
                        except Exception as e:
                            logger.debug(f"Error terminating thread {thread_id}: {e}")
                            failed_threads += 1
                    else:
                        stopped_threads += 1
                else:
                    with QMutexLocker(self._lock):
                        if thread_id in self._threads:
                            logger.debug(f"Removing invalid thread: {thread_id}")
                            del self._threads[thread_id]
            except RuntimeError as e:
                logger.debug(f"C++ object for thread {thread_id} already deleted: {e}")
                with QMutexLocker(self._lock):
                    if thread_id in self._threads:
                        del self._threads[thread_id]
            except Exception as e:
                logger.error(f"Unexpected error stopping thread {thread_id}: {e}")
                failed_threads += 1
        self.cleanup_finished_threads()
        self._is_stopping = False
        self.all_threads_stopped.emit()
        logger.info(f"All threads stopped: {stopped_threads} successful, {failed_threads} failed")
        return {
            "total": len(thread_ids),
            "stopped": stopped_threads,
            "failed": failed_threads,
            "remaining": len(self._threads)
        }
class ThreadFactory:
    """Factory class untuk membuat dan mengonfigurasi thread"""
    @staticmethod
    def create_worker_thread(worker, on_finished=None, on_progress=None, on_error=None, thread_name=None):
        """
        Membuat dan mengonfigurasi worker thread
        Args:
            worker: Worker object yang akan dijalankan di thread
            on_finished: Callback ketika thread selesai
            on_progress: Callback untuk update progress
            on_error: Callback ketika terjadi error
            thread_name: Nama thread untuk debugging
        Returns:
            QThread: Thread yang sudah dikonfigurasi
        """
        thread = QThread()
        if thread_name:
            thread.setObjectName(thread_name)
        worker.moveToThread(thread)
        thread.started.connect(worker.run)
        worker.finished.connect(thread.quit)
        worker.finished.connect(worker.deleteLater)
        thread.finished.connect(thread.deleteLater)
        if on_finished:
            worker.finished.connect(on_finished)
        if on_progress and hasattr(worker, 'progress'):
            worker.progress.connect(on_progress)
        if on_error and hasattr(worker, 'error'):
            worker.error.connect(on_error)
        return thread        
class ImportThread(QThread):
    """Dedicated thread for import initialization"""
    finished = Signal()
    def __init__(self):
        super().__init__()
        self._is_running = True
        self.setObjectName("ImportThread")
    def cleanup_resources(self):
        """Clean up resources used by this thread"""
        logger.debug(f"Cleaning up resources for {self.objectName()}")        
    def run(self):
        try:
            if not self._is_running or self.isInterruptionRequested():
                return
            import matplotlib.pyplot as plt
            if not self._is_running or self.isInterruptionRequested():
                return
            import nltk
            if not self._is_running or self.isInterruptionRequested():
                return
            from textblob import TextBlob
        except Exception as e:
            logger.error(f"Import thread error: {e}")
        finally:
            if self._is_running and not self.isInterruptionRequested():
                self.finished.emit()
            self.cleanup_resources()
    def stop(self):
        """Method to safely stop thread"""
        self._is_running = False
        self.requestInterruption()
        self.wait(100)
class FileLoaderThread(QThread):
    file_loaded = Signal(str, str)
    file_error = Signal(str)
    def __init__(self, file_path):
        super().__init__()
        self.file_path = file_path
        self._interrupt_requested = False
        self.setObjectName("FileLoaderThread")
        self._opened_resources = []
    def cleanup_resources(self):
        """Clean up resources used by this thread"""
        logger.debug(f"Cleaning up resources for {self.objectName()}")
        try:
            for resource in self._opened_resources:
                try:
                    if hasattr(resource, 'close'):
                        resource.close()
                except Exception as e:
                    logger.error(f"Error closing resource: {e}")
            self._opened_resources.clear()
        except Exception as e:
            logger.error(f"Error cleaning up file loader thread: {e}")        
    def run(self):
        try:
            if self.isInterruptionRequested():
                self._interrupt_requested = True
                return
            if not os.path.exists(self.file_path):
                raise FileNotFoundError(f"File not found: {self.file_path}")
            file_size = os.path.getsize(self.file_path)
            if file_size > AppConstants.MAX_FILE_SIZE:
                raise ValueError(f"File too large: {file_size} bytes")
            else:
                self.process_normal_file()
        except Exception as e:
            if not self._interrupt_requested and not self.isInterruptionRequested():
                self.file_error.emit(f"Error loading {os.path.basename(self.file_path)}: {str(e)}")
        finally:
            self.cleanup_resources()
    def extract_text_from_pdf(self, pdf_path: str) -> str:
        """Extract text from PDF with proper error handling"""
        try:
            text = ""
            if self.isInterruptionRequested():
                self._interrupt_requested = True
                return ""
            try:
                import pypdf
                with open(pdf_path, "rb") as file:
                    self._opened_resources.append(file)
                    reader = pypdf.PdfReader(file)
                    if not reader.pages:
                        raise ValueError("PDF contains no readable pages")
                    for i, page in enumerate(reader.pages):
                        if self.isInterruptionRequested():
                            self._interrupt_requested = True
                            return ""
                        page_text = page.extract_text()
                        if page_text:
                            text += page_text + "\n\n"
                return text
            except Exception as e:
                logger.error(f"PDF read error: {e}")
                raise ValueError(f"Failed to read PDF: {e}")
        except Exception as e:
            logger.error(f"Unexpected error processing PDF: {e}")
            raise ValueError(f"PDF processing failed: {e}")
    def process_normal_file(self):
        """Process normal sized files with improved error handling"""
        try:
            if self.isInterruptionRequested():
                self._interrupt_requested = True
                return
            text_data = ""
            file_ext = os.path.splitext(self.file_path)[1].lower()
            try:
                if file_ext == ".txt":
                    if self.isInterruptionRequested():
                        self._interrupt_requested = True
                        return
                    try:
                        with open(self.file_path, "r", encoding="utf-8") as file:
                            self._opened_resources.append(file)
                            text_data = file.read()
                    except UnicodeDecodeError:
                        with open(self.file_path, "r", encoding="latin-1") as file:
                            self._opened_resources.append(file)
                            text_data = file.read()
                elif file_ext == ".pdf":
                    if self.isInterruptionRequested():
                        self._interrupt_requested = True
                        return
                    text_data = self.extract_text_from_pdf(self.file_path)
                elif file_ext in (".doc", ".docx"):
                    if self.isInterruptionRequested():
                        self._interrupt_requested = True
                        return
                    text_data = self.extract_text_from_word(self.file_path)
                elif file_ext in (".xlsx", ".xls"):
                    if self.isInterruptionRequested():
                        self._interrupt_requested = True
                        return
                    text_data = self.extract_text_from_excel(self.file_path)
                elif file_ext == ".csv":
                    if self.isInterruptionRequested():
                        self._interrupt_requested = True
                        return
                    text_data = self.extract_text_from_csv(self.file_path)
                else:
                    raise ValueError(f"Unsupported file format: {file_ext}")
            except (PermissionError, OSError) as e:
                logger.error(f"File access error: {e}")
                raise ValueError(f"Cannot access file: {e}")
            except UnicodeError as e:
                logger.error(f"Text encoding error: {e}")
                raise ValueError("File contains invalid text encoding")
            if self.isInterruptionRequested():
                self._interrupt_requested = True
                return
            if not text_data.strip():
                raise ValueError("File is empty or contains no extractable text")
            self.file_loaded.emit(self.file_path, text_data)
        except Exception as e:
            logger.error(f"File load error: {e}")
            if not self._interrupt_requested and not self.isInterruptionRequested():
                self.file_error.emit(str(e))
    def extract_text_from_word(self, word_path):
        try:
            import docx
            doc = docx.Document(word_path)
            if not doc.paragraphs:
                raise FileLoadError("Word document contains no readable text")
            text = ""
            for p in doc.paragraphs:
                if self.isInterruptionRequested():
                    return ""
                if p.text.strip():
                    text += p.text.strip() + "\n"
            for table in doc.tables:
                if self.isInterruptionRequested():
                    return ""
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            text += cell.text.strip() + "\n"
            if not text.strip():
                raise FileLoadError("No readable text found in document")
            text = re.sub(r'\s+', ' ', text)
            text = text.replace('\n\n', '\n').strip()
            return text
        except Exception as e:
            logger.error(f"Word processing failed: {str(e)}")
            raise FileLoadError(f"Word processing failed: {str(e)}") from e
    def extract_text_from_excel(self, excel_path):
        try:
            import pandas as pd
            text_data = ""
            df = pd.read_excel(excel_path, sheet_name=None)
            if not df:
                raise FileLoadError("Excel file is empty or unreadable")
            for sheet_name, sheet_df in df.items():
                if self.isInterruptionRequested():
                    return ""
                text_data += f"\n--- {sheet_name} ---\n"
                text_data += sheet_df.to_string(index=False, header=True)
            return text_data
        except Exception as e:
            logger.error(f"Excel processing failed: {str(e)}")
            raise FileLoadError(f"Excel processing failed: {str(e)}") from e
    def extract_text_from_csv(self, csv_path):
        try:
            import pandas as pd
            df = pd.read_csv(csv_path)
            if df.empty:
                raise FileLoadError("CSV file is empty")
            if self.isInterruptionRequested():
                return ""
            return df.to_string(index=False, header=True)
        except Exception as e:
            logger.error(f"CSV processing failed: {str(e)}")
            raise FileLoadError(f"CSV processing failed: {str(e)}") from e       
class CustomFileLoaderThread(QThread):
    file_loaded = Signal(object, bool)
    def __init__(self, file_path, file_type):
        super().__init__()
        self.file_path = file_path
        self.file_type = file_type
    def run(self):
        try:
            if self.file_type == "lexicon":
                from vaderSentiment.vaderSentiment import SentimentIntensityAnalyzer
                self.vader_analyzer = CustomVaderSentimentIntensityAnalyzer(custom_lexicon_file=self.file_path)
                self.file_loaded.emit(self.file_path, True)
            elif self.file_type == "TextBlob (Custom Lexicon)":
                from textblob import TextBlob
                self.textblob_analyzer = CustomTextBlobSentimentAnalyzer(self.file_path)
                self.file_loaded.emit(self.file_path, True)
            elif self.file_type == "model":
                from flair.models import TextClassifier
                from flair.data import Sentence
                try:
                    model = TextClassifier.load(self.file_path)
                    test = Sentence("test")
                    model.predict(test)
                    if not test.labels or test.labels[0].value not in ['POSITIVE', 'NEGATIVE', 'NEUTRAL']:
                        raise ValueError("Model is not a valid sentiment classifier")
                    self.file_loaded.emit(model, True)
                except Exception as e:
                    self.file_loaded.emit(f"Invalid sentiment model: {str(e)}", False)
        except Exception as e:
            self.file_loaded.emit(str(e), False)
"""CLASS SERVICE"""    
class WordCloudService(QObject):
    """Service untuk generate dan manipulasi word cloud"""
    generation_progress = Signal(int)
    generation_complete = Signal(object)
    error_occurred = Signal(str)
    save_complete = Signal(bool, str)
    def __init__(self, parent=None):
        super().__init__(parent)
        if hasattr(parent, 'cache_manager'):
            self.cache_manager = parent.cache_manager
        else:
            self.cache_manager = EnhancedCacheManager()        
        self.parent = parent
        self.resource_controller = parent.resource_controller if parent else None
        self.generator_thread = None
        self.current_figure = None
        self.mask_path = ""
        self.custom_color_palettes = {}
        self._cached_colormaps = {}
        self.font_map = {}
    def set_font_map(self, font_map):
        """Set font map dari UI"""
        if not font_map:
            logger.warning("Received empty font_map")
            return
        self.font_map = font_map.copy()
        logger.debug(f"Updated font_map with {len(font_map)} fonts")
    def set_custom_color_palettes(self, palettes):
        """Set custom color palettes dari UI"""
        self.custom_color_palettes = palettes
    def set_cached_colormaps(self, colormaps):
        """Set cached colormaps dari UI"""
        self._cached_colormaps = colormaps
    def generate_wordcloud(self, text_data, params):
        """Generate word cloud dengan parameter yang diberikan"""
        try:
            if self.generator_thread:
                if self.generator_thread.isRunning():
                    self.stop_generation()
                else:
                    self.generator_thread.cleanup_resources()
                    self.generator_thread.deleteLater()
                    self.generator_thread = None
            if self.current_figure:
                self.resource_controller.res_cleanup_figure(self.current_figure)
                self.current_figure = None
            params['font_map'] = self.font_map
            self.generator_thread = WordCloudGeneratorThread(text_data, params, self.mask_path)
            self.generator_thread.setParent(self)
            self.generator_thread.progress_updated.connect(self.generation_progress.emit)
            self.generator_thread.generation_complete.connect(self._on_generation_complete)
            self.generator_thread.error_occurred.connect(self.error_occurred.emit)
            self.generator_thread.start()
        except Exception as e:
            logger.error(f"Failed to start word cloud generation: {e}")
    def stop_generation(self):
        """Stop word cloud generation process"""
        if self.generator_thread and self.generator_thread.isRunning():
            try:
                self.generator_thread.cancel()
                self.generator_thread.wait(500)
                if self.current_figure:
                    self.resource_controller.res_cleanup_figure(self.current_figure)
                    self.current_figure = None
                self.generation_progress.emit(0)
                self.error_occurred.emit("Word cloud generation cancelled by user")
                self.generator_thread.cleanup_resources()
                self.generator_thread.deleteLater()
                self.generator_thread = None
            except Exception as e:
                logger.error(f"Error stopping word cloud generation: {e}")
                self.generation_progress.emit(0)
    def _on_generation_complete(self, wordcloud):
        try:
            import matplotlib.pyplot as plt
            from matplotlib.font_manager import FontProperties
            from matplotlib.backends.backend_qt5 import NavigationToolbar2QT
            from PySide6.QtCore import QTimer
            import numpy as np
            from PIL import Image, ImageDraw, ImageFont
            plt.ion()
            figure = plt.figure(figsize=(10, 6))
            ax = figure.add_subplot(111)
            title_text = self.generator_thread.params.get('title_text', '').strip()
            title_position = self.generator_thread.params.get('title_position', 'center').lower()
            title_font_size = self.generator_thread.params.get('title_font_size', 14)
            if title_text:
                wc_array = wordcloud.to_array()
                wc_image = Image.fromarray(wc_array)
                """
                Buat image baru dengan ruang untuk title
                Ini menentukan total tinggi keseluruhan gambar akhir
                Faktor 1.15 berarti gambar akhir akan 15% lebih tinggi dari word cloud asli
                Ini membuat "kanvas" atau ruang total yang tersedia untuk menampung judul dan word cloud
                Seperti menentukan tinggi kertas keseluruhan
                """
                total_height = int(wc_image.height * 1.20) 
                combined_image = Image.new('RGBA', (wc_image.width, total_height), 
                                        color=self.generator_thread.params.get('bg_color', 'white'))
                try:
                    selected_font = self.generator_thread.params.get('font_choice')
                    font_map = self.generator_thread.params.get('font_map', {})
                    dpi_scale = wc_image.width / 800
                    scaled_font_size = int(title_font_size * dpi_scale)
                    if selected_font and selected_font != "Default" and selected_font in font_map:
                        font_path = font_map.get(selected_font)
                        if font_path and os.path.exists(font_path):
                            title_font = ImageFont.truetype(font_path, scaled_font_size)
                        else:
                            from matplotlib.font_manager import findfont, FontProperties
                            font_path = findfont(FontProperties())
                            title_font = ImageFont.truetype(font_path, scaled_font_size)
                    else:
                        from matplotlib.font_manager import findfont, FontProperties
                        font_path = findfont(FontProperties())
                        title_font = ImageFont.truetype(font_path, scaled_font_size)
                except Exception as e:
                    logger.warning(f"Failed to load font, using default: {e}")
                    try:
                        from matplotlib.font_manager import findfont, FontProperties
                        font_path = findfont(FontProperties())
                        title_font = ImageFont.truetype(font_path, scaled_font_size)
                    except:
                        title_font = ImageFont.load_default()
                draw = ImageDraw.Draw(combined_image)
                title_bbox = draw.textbbox((0, 0), title_text, font=title_font)
                title_width = title_bbox[2] - title_bbox[0]
                if title_position == 'center':
                    x = (wc_image.width - title_width) // 2
                elif title_position == 'left':
                    x = 10
                else:
                    x = wc_image.width - title_width - 10
                """ jarak judul dari tepi atas. default = 0.02 """
                y = int(wc_image.height * 0) 
                draw.text((x, y), title_text, font=title_font, 
                        fill=self.generator_thread.params.get('title_color', 'black'))
                """ jarak wordcloud dari judul. default = 0.15 """
                combined_image.paste(wc_image, (0, int(wc_image.height * 0.15)))
                combined_array = np.array(combined_image)
                ax.imshow(combined_array)
            else:
                ax.imshow(wordcloud)
            ax.axis("off")
            manager = plt.get_current_fig_manager()
            if hasattr(manager, "window"):
                main_window = manager.window
                toolbar = main_window.findChild(NavigationToolbar2QT)
                if toolbar:
                    def keep_only_actions(toolbar, keep_texts):
                        for action in toolbar.actions():
                            if action.text() not in keep_texts:
                                toolbar.removeAction(action)
                    keep_only_actions(toolbar, ['Save', 'Subplots'])
                    def rename_subplots_window():
                        app = QApplication.instance()
                        for widget in app.topLevelWidgets():
                            if isinstance(widget, QDialog):
                                title = widget.windowTitle().strip().lower()
                                if title in ["", "python", "configure subplots"]:
                                    widget.setWindowTitle("Word Cloud Layout Configuration")
                                    break
                    def delayed_rename():
                        QTimer.singleShot(500, rename_subplots_window)
                    for action in toolbar.actions():
                        if action.text() == 'Subplots':
                            action.triggered.connect(delayed_rename)
            main_window.setWindowTitle("Word Cloud Result")
            self.current_figure = figure
            self.resource_controller.register_figure(figure)
            self.generation_complete.emit(figure)
            plt.show()
        except Exception as e:
            logger.error(f"Error creating figure: {e}")
            self.error_occurred.emit(f"Error creating figure: {e}")
    def choose_mask(self, mask_path):
        """Set mask image path"""
        self.mask_path = mask_path
    def reset_mask(self):
        """Reset mask image path"""
        self.mask_path = ""
    def res_create_custom_palette(self, color_list, palette_name=None):
        """Buat custom palette dari daftar warna"""
        try:
            if not color_list:
                logger.error("Empty color list provided for custom palette")
                self.resource_error.emit("palette", "Empty color list provided")
                return None
            if not palette_name:
                palette_name = f"custom_palette_{len(self.custom_color_palettes) + 1}"
            if not palette_name.startswith('custom_'):
                palette_name = f"custom_{palette_name}"
            color_list_copy = list(color_list)
            self.custom_color_palettes[palette_name] = color_list_copy
            from matplotlib.colors import LinearSegmentedColormap
            self._cached_colormaps[palette_name] = LinearSegmentedColormap.from_list(palette_name, color_list_copy)
            self.palette_updated.emit(palette_name)
            logger.info(f"Created custom palette: {palette_name} with {len(color_list_copy)} colors: {color_list_copy}")
            return palette_name
        except Exception as e:
            logger.error(f"Error creating custom palette: {e}")
            self.resource_error.emit("palette", f"Error creating custom palette: {e}")
            return None
    def cleanup(self):
        """Cleanup resources"""
        if self.current_figure:
            self.resource_controller.res_cleanup_figure(self.current_figure)
            self.current_figure = None
        if self.generator_thread and self.generator_thread.isRunning():
            try:
                self.generator_thread.cancel()
                self.generator_thread.wait(500)
            except:
                pass        
class SentimentAnalysisService(QObject):
    """Service for handling sentiment analysis operations asynchronously"""
    analysis_started = Signal()
    analysis_progress = Signal(int)
    analysis_complete = Signal(dict)
    error_occurred = Signal(str)
    model_loading_started = Signal(str)
    model_loading_progress = Signal(int)
    model_loaded = Signal(str, object)
    model_loading_failed = Signal(str, str)
    sentiment_timeline_ready = Signal(object)
    def __init__(self, parent=None):
        QObject.__init__(self, parent)
        self.parent = parent
        self.analysis_thread = None
        self.model_loader_thread = None
        self._cached_results = {}
        self._cached_models = {}
        self._mutex = QMutex()
    def load_model(self, model_type, model_path=None, custom=False):
        """Load ML model asynchronously
        Args:
            model_type (str): Type of model ('flair', 'vader', 'textblob')
            model_path (str, optional): Path to custom model
            custom (bool): Whether this is a custom model
        """
        try:
            if self.model_loader_thread and self.model_loader_thread.isRunning():
                self.model_loader_thread.requestInterruption()
                self.model_loader_thread.wait(500)
            cache_key = f"{model_type}_{model_path if model_path else 'default'}"
            if cache_key in self._cached_models:
                self.model_loaded.emit(model_type, self._cached_models[cache_key])
                return True
            self.model_loading_started.emit(model_type)
            if model_type.lower() == 'flair':
                self.model_loader_thread = FlairModelLoaderThread(model_path if model_path else "sentiment")
            else:
                self.error_occurred.emit(f"Unsupported model type: {model_type}")
                return False
            self.model_loader_thread.model_loaded.connect(
                lambda model: self._on_model_loaded(model_type, model, cache_key))
            self.model_loader_thread.error_occurred.connect(
                lambda error: self.model_loading_failed.emit(model_type, error))
            self.model_loader_thread.start()
            return True
        except Exception as e:
            self.error_occurred.emit(f"Failed to start model loading: {str(e)}")
            return False
    def _on_model_loaded(self, model_type, model, cache_key):
        """Handle completion of model loading"""
        self._mutex.lock()
        try:
            self._cached_models[cache_key] = model
            if len(self._cached_models) > 5:
                oldest_key = next(iter(self._cached_models))
                del self._cached_models[oldest_key]
        finally:
            self._mutex.unlock()
        self.model_loaded.emit(model_type, model)
    def analyze_text(self, text_data, sentiment_mode, analyzers, translation_handler=None):
        """Analyze text sentiment asynchronously"""
        try:
            if self.analysis_thread and self.analysis_thread.isRunning():
                self.analysis_thread.requestInterruption()
                self.analysis_thread.wait(500)
            cache_key = f"{hashlib.md5(text_data.encode()).hexdigest()}_{sentiment_mode}"
            if cache_key in self._cached_results:
                cached_result = self._cached_results[cache_key]
                self.analysis_complete.emit(cached_result)
                if 'sentiment_timeline' in cached_result:
                    self.sentiment_timeline_ready.emit(cached_result['sentiment_timeline'])
                return
            self.analysis_started.emit()
            self.analysis_thread = SentimentAnalysisThread(
                text_data, 
                sentiment_mode,
                analyzers.get('vader_analyzer'),
                analyzers.get('flair_classifier'),
                analyzers.get('flair_classifier_cuslang'),
                analyzers.get('textblob_analyzer')
            )
            self.analysis_thread.sentiment_analyzed.connect(self._on_analysis_complete)
            self.analysis_thread.translation_failed.connect(
                lambda msg: self.error_occurred.emit(f"Translation error: {msg}")
            )
            self.analysis_thread.offline_warning.connect(
                lambda msg: self.error_occurred.emit(f"Offline warning: {msg}")
            )
            self.analysis_thread.sentiment_timeline_ready.connect(
                lambda timeline: self.sentiment_timeline_ready.emit(timeline)
            )
            if translation_handler:
                self.analysis_thread.translation_progress.connect(translation_handler)
            self.analysis_thread.start()
        except Exception as e:
            self.error_occurred.emit(f"Failed to start sentiment analysis: {str(e)}")
    def _on_analysis_complete(self, result):
        """Handle completion of sentiment analysis"""
        if result and "sentiment_label" in result:
            cache_key = f"{hashlib.md5(self.analysis_thread.text_data.encode()).hexdigest()}_{self.analysis_thread.sentiment_mode}"
            self._cached_results[cache_key] = result
            if len(self._cached_results) > 20:
                oldest_key = next(iter(self._cached_results))
                del self._cached_results[oldest_key]
        self.analysis_complete.emit(result)
    def get_cached_timeline(self, text_data, sentiment_mode):
        """Get cached sentiment timeline if available"""
        cache_key = f"{hashlib.md5(text_data.encode()).hexdigest()}_{sentiment_mode}"
        if cache_key in self._cached_results:
            result = self._cached_results[cache_key]
            if 'sentiment_timeline' in result:
                return result['sentiment_timeline']
        return None
    def cancel_analysis(self):
        """Cancel ongoing analysis"""
        if self.analysis_thread and self.analysis_thread.isRunning():
            self.analysis_thread.requestInterruption()
            self.analysis_thread.wait(500)
            return True
        return False
    def cancel_model_loading(self):
        """Cancel ongoing model loading"""
        if self.model_loader_thread and self.model_loader_thread.isRunning():
            self.model_loader_thread.requestInterruption()
            self.model_loader_thread.wait(500)
            return True
        return False
    def cleanup(self):
        """Clean up resources"""
        self.cancel_analysis()
        self.cancel_model_loading()
        self._cached_results.clear()
        self._cached_models.clear()
class TextAnalysisService(QObject):
    """Service untuk text analysis dan statistik"""
    analysis_started = Signal()
    analysis_progress = Signal(int)
    analysis_complete = Signal(dict)
    error_occurred = Signal(str)
    dialog_shown = Signal()
    dialog_closed = Signal()
    def __init__(self, parent=None):
        super().__init__(parent)
        self.current_thread = None
        self.stats_data = None
        self._parent = parent
    def analyze_text(self, text_data, include_stopwords=False, custom_stopwords=None):
        """Memulai thread untuk menganalisis teks"""
        try:
            self.analysis_started.emit()
            if self.current_thread and self.current_thread.isRunning():
                self.current_thread.quit()
                self.current_thread.wait()
            self.current_thread = TextStatsThread(text_data)
            self.current_thread.include_stopwords = include_stopwords
            self.current_thread.custom_stopwords = custom_stopwords
            self.current_thread.stats_ready.connect(self._on_thread_complete)
            self.current_thread.error_occurred.connect(self._on_thread_error)
            self.current_thread.progress_updated.connect(
                lambda value: self.analysis_progress.emit(value)
            )
            logger.info(f"Starting text analysis with stopwords {'included' if include_stopwords else 'excluded'}")            
            self.current_thread.start()            
        except Exception as e:
            logger.error(f"Error starting text analysis: {str(e)}")
            self.error_occurred.emit(str(e))
    def show_stats_dialog(self, parent_widget, stats):
        """Menampilkan dialog statistik teks"""
        try:
            dialog = QDialog(parent_widget)
            dialog.setWindowTitle("Text Statistics")
            dialog.setMinimumSize(800, 750)
            layout = QVBoxLayout()
            stopwords_info = QLabel()
            if stats.get('stopwords_removed', False):
                stopwords_info.setText("Statistics calculated with stopwords removed")
                stopwords_info.setStyleSheet("color: #0066cc; font-weight: bold;")
            else:
                stopwords_info.setText("Statistics calculated with stopwords included")
                stopwords_info.setStyleSheet("color: #009933; font-weight: bold;")
            layout.addWidget(stopwords_info)
            tab_widget = QTabWidget()
            self._populate_stats_tabs(tab_widget, stats)
            layout.addWidget(tab_widget)
            close_button = QPushButton("Close")
            close_button.clicked.connect(dialog.accept)
            button_layout = QHBoxLayout()
            button_layout.addStretch()
            button_layout.addWidget(close_button)
            layout.addLayout(button_layout)
            dialog.setLayout(layout)
            self.dialog_shown.emit()
            dialog.show()
            self.dialog_closed.emit()
        except Exception as e:
            logger.error(f"Error showing stats dialog: {str(e)}")
            self.error_occurred.emit(f"Failed to display text statistics: {str(e)}")
    def _populate_stats_tabs(self, tab_widget, stats):
        """Create and populate tabs for different statistics categories"""
        try:
            basic_tab = self._create_basic_stats_tab(stats)
            tab_widget.addTab(basic_tab, "Basic Statistics")
            word_freq_tab = self._create_word_frequency_tab(stats)
            tab_widget.addTab(word_freq_tab, "Word Frequency")
            if 'readability' in stats and isinstance(stats['readability'], dict):
                readability_tab = self._create_readability_tab(stats)
                tab_widget.addTab(readability_tab, "Readability")
            if 'pos_categories' in stats:
                pos_tab = QWidget()
                pos_layout = QVBoxLayout()
                pos_frame = QFrame()
                pos_chart_layout = QVBoxLayout()
                pos_chart_view = self._create_pos_chart(stats['pos_categories'])
                pos_chart_layout.addWidget(pos_chart_view)
                pos_frame.setLayout(pos_chart_layout)
                pos_layout.addWidget(pos_frame)
                if 'pos_counts' in stats:
                    pos_table = QTableWidget()
                    pos_table.setColumnCount(2)
                    pos_table.setHorizontalHeaderLabels(["POS Tag", "Count"])
                    pos_table.horizontalHeader().setSectionResizeMode(0, QHeaderView.Stretch)
                    sorted_pos = sorted(stats['pos_counts'].items(), key=lambda x: x[1], reverse=True)
                    pos_table.setRowCount(len(sorted_pos))
                    for i, (tag, count) in enumerate(sorted_pos):
                        pos_table.setItem(i, 0, QTableWidgetItem(tag))
                        pos_table.setItem(i, 1, QTableWidgetItem(str(count)))
                    pos_layout.addWidget(QLabel("Detailed POS Counts:"))
                    pos_layout.addWidget(pos_table)
                pos_tab.setLayout(pos_layout)
                tab_widget.addTab(pos_tab, "Part of Speech Analysis")
        except Exception as e:
            import traceback
            logger.error(f"Error populating stats tabs: {str(e)}\n{traceback.format_exc()}")
    def _create_basic_stats_tab(self, stats):
        """Create the basic statistics tab"""
        tab = QWidget()
        layout = QVBoxLayout()
        basic_stats_items = [
            ("Character count", "{:,}".format(stats.get('char_count', 0))),
            ("Word count", "{:,}".format(stats.get('word_count', 0))),
            ("Sentence count", "{:,}".format(stats.get('sentence_count', 0))),
            ("Paragraph count", "{:,}".format(stats.get('paragraph_count', 0))),
            ("Average word length", "{:.2f} characters".format(stats.get('avg_word_length', 0))),
            ("Average sentence length", "{:.2f} words".format(stats.get('avg_sentence_length', 0))),
            ("Unique words", "{:,}".format(stats.get('unique_word_count', 0))),
            ("Lexical diversity", "{:.4f}".format(stats.get('lexical_diversity', 0)))
        ]
        if 'stopwords_removed' in stats and stats['stopwords_removed']:
            basic_stats_items.extend([
                ("Words after stopwords removal", "{:,}".format(stats.get('filtered_word_count', 0))),
                ("Unique words after stopwords removal", "{:,}".format(stats.get('filtered_unique_words', 0)))
            ])
        table = QTableWidget()
        table.setColumnCount(2)
        table.setRowCount(len(basic_stats_items))
        table.setHorizontalHeaderLabels(["Statistic", "Value"])
        table.horizontalHeader().setSectionResizeMode(0, QHeaderView.Stretch)
        table.horizontalHeader().setSectionResizeMode(1, QHeaderView.Stretch)
        table.verticalHeader().setVisible(False)
        for i, (name, value) in enumerate(basic_stats_items):
            table.setItem(i, 0, QTableWidgetItem(name))
            table.setItem(i, 1, QTableWidgetItem(value))
        layout.addWidget(table)
        tab.setLayout(layout)
        return tab
    def _create_word_frequency_tab(self, stats):
        """Create the word frequency tab"""
        tab = QWidget()
        layout = QVBoxLayout()
        word_freq = stats.get('word_freq', {})
        most_common = stats.get('most_common_words', [])
        
        if most_common:
            try:
                # Create visualization container with strict size control
                viz_container = QWidget()
                viz_container.setSizePolicy(QSizePolicy.Expanding, QSizePolicy.Fixed)
                viz_container.setFixedHeight(350)
                viz_layout = QVBoxLayout(viz_container)
                viz_layout.setContentsMargins(0, 0, 0, 0)
                
                # Import and setup matplotlib
                from matplotlib.figure import Figure
                from matplotlib.backends.backend_qt5agg import FigureCanvasQTAgg
                import matplotlib.ticker as ticker
                
                # Get top 15 words
                top_words = most_common[:15]
                
                # Create figure with fixed DPI
                figure = Figure(dpi=100)
                canvas = FigureCanvasQTAgg(figure)
                canvas.setSizePolicy(QSizePolicy.Expanding, QSizePolicy.Fixed)
                canvas.setFixedHeight(350)
                
                def update_plot():
                    figure.clear()
                    figure.subplots_adjust(left=0.15, right=0.95, bottom=0.15, top=0.95)
                    ax = figure.add_subplot(111)
                    
                    words = [word for word, count in top_words]
                    counts = [count for word, count in top_words]
                    
                    # Create horizontal bars with improved style
                    bar_color = '#4285F4'
                    bars = ax.barh(range(len(words)), counts, align='center', 
                                color=bar_color, alpha=0.7)
                    
                    # Customize axes with smaller fonts
                    ax.set_yticks(range(len(words)))
                    ax.set_yticklabels(words, fontsize=8)
                    ax.tick_params(axis='x', labelsize=8, pad=5)
                    
                    ax.invert_yaxis()  # Invert to show highest frequency at top
                    
                    # Set labels with smaller fonts
                    ax.set_xlabel('Frequency', fontsize=9, labelpad=10)
                    
                    # Set axis limits and grid
                    max_count = max(counts) if counts else 0
                    ax.set_xlim(0, max_count * 1.05)
                    ax.xaxis.set_major_locator(ticker.MaxNLocator(integer=True))
                    ax.xaxis.grid(True, linestyle='--', alpha=0.3)
                    
                    # Remove unnecessary spines
                    ax.spines['top'].set_visible(False)
                    ax.spines['right'].set_visible(False)
                    
                    canvas.draw()
                
                # Add canvas to container
                viz_layout.addWidget(canvas)
                layout.addWidget(viz_container)
                
                # Connect resize event
                canvas.mpl_connect('resize_event', lambda evt: update_plot())
                update_plot()
                
                if hasattr(self._parent, 'cleanup_figure'):
                    self._parent.cleanup_figure(figure)
                
                # Create content widget that will contain both table and stopwords
                content_widget = QWidget()
                content_widget.setSizePolicy(QSizePolicy.Expanding, QSizePolicy.Expanding)
                content_layout = QVBoxLayout(content_widget)
                
                # Word frequency table in an expandable container
                table_container = QWidget()
                table_container.setSizePolicy(QSizePolicy.Expanding, QSizePolicy.Expanding)
                table_layout = QVBoxLayout(table_container)
                
                table_group = QGroupBox("Word Frequency Details")
                table_group.setSizePolicy(QSizePolicy.Expanding, QSizePolicy.Expanding)
                group_layout = QVBoxLayout()
                
                table = QTableWidget()
                table.setSizePolicy(QSizePolicy.Expanding, QSizePolicy.Expanding)
                table.setColumnCount(2)
                table.setHorizontalHeaderLabels(["Word", "Frequency"])
                table.horizontalHeader().setSectionResizeMode(0, QHeaderView.Stretch)
                table.setRowCount(len(most_common))
                
                for i, (word, count) in enumerate(most_common):
                    table.setItem(i, 0, QTableWidgetItem(word))
                    table.setItem(i, 1, QTableWidgetItem(str(count)))
                
                group_layout.addWidget(table)
                table_group.setLayout(group_layout)
                table_layout.addWidget(table_group)
                content_layout.addWidget(table_container, stretch=1)
                
                # Add stopwords section with fixed size if available
                if stats.get('stopwords_removed', False) and 'stopwords_used' in stats:
                    stopwords_container = QWidget()
                    stopwords_container.setSizePolicy(QSizePolicy.Expanding, QSizePolicy.Fixed)
                    stopwords_container.setFixedHeight(120)  # Fixed height for stopwords section
                    stopwords_layout = QVBoxLayout(stopwords_container)
                    
                    stopwords_group = QGroupBox("Stopwords Used")
                    stopwords_group.setCheckable(True)
                    stopwords_group.setChecked(False)
                    group_layout = QVBoxLayout()
                    
                    stopwords_text = QTextEdit()
                    stopwords_text.setReadOnly(True)
                    stopwords_text.setText(", ".join(sorted(stats['stopwords_used'])))
                    stopwords_text.setFixedHeight(60)
                    
                    group_layout.addWidget(stopwords_text)
                    stopwords_group.setLayout(group_layout)
                    stopwords_layout.addWidget(stopwords_group)
                    
                    content_layout.addWidget(stopwords_container)
                
                # Add the content widget to main layout with stretch
                layout.addWidget(content_widget, stretch=1)
                
            except Exception as e:
                logger.error(f"Error creating word frequency chart: {str(e)}")
                layout.addWidget(QLabel("Failed to create word frequency chart"))
        else:
            layout.addWidget(QLabel("No word frequency data available."))
        
        tab.setLayout(layout)
        return tab
    def _create_readability_tab(self, stats):
        """Create the readability tab"""
        tab = QWidget()
        layout = QVBoxLayout()
        readability = stats.get('readability', {})
        if isinstance(readability, dict) and not 'error' in readability:
            if 'readability_consensus' in readability:
                consensus_label = QLabel()
                consensus_label.setText(f"<b>Overall Readability: {readability.get('readability_consensus', 'Unknown')}</b>")
                consensus_label.setStyleSheet("font-size: 14px; margin-bottom: 10px;")
                layout.addWidget(consensus_label)
            table = QTableWidget()
            table.setColumnCount(2)
            table.setHorizontalHeaderLabels(["Readability Metric", "Score"])
            table.horizontalHeader().setSectionResizeMode(0, QHeaderView.Stretch)
            metrics = [
                ("flesch_reading_ease", "Flesch Reading Ease", "0-100 scale where higher scores indicate easier reading. 60-70 is considered plain English."),
                ("flesch_kincaid_grade", "Flesch-Kincaid Grade Level", "Represents the US grade level needed to understand the text."),
                ("gunning_fog", "Gunning Fog Index", "Estimates years of formal education needed to understand the text."),
                ("smog_index", "SMOG Index", "Estimates years of education needed to understand text, emphasizing polysyllables."),
                ("automated_readability_index", "Automated Readability Index", "Represents the US grade level needed to understand the text based on characters."),
                ("coleman_liau_index", "Coleman-Liau Index", "Represents the US grade level needed to understand the text based on characters vs words."),
                ("linsear_write_formula", "Linsear Write Formula", "Designed for the US Air Force to calculate the readability of technical manuals."),
                ("dale_chall_readability_score", "Dale-Chall Readability Score", "Based on the percentage of words not on a list of common words.")
            ]
            readable_metrics = []
            for key, name, desc in metrics:
                if key in readability:
                    value = readability[key]
                    readable_metrics.append((name, value, desc))
            table.setRowCount(len(readable_metrics))
            for i, (name, value, desc) in enumerate(readable_metrics):
                table.setItem(i, 0, QTableWidgetItem(name))
                table.setItem(i, 1, QTableWidgetItem("{:.2f}".format(value)))
                table.item(i, 0).setToolTip(desc)
            explain_label = QLabel()
            explain_label.setText("""
                <p><b>Readability Metrics Explanation:</b></p>
                <p>Hover over metric names for explanations. Lower grade levels generally indicate more readable text.</p>
                <p>Flesch Reading Ease is scored 0-100 where higher scores mean more readable text.</p>
                <p>Other metrics approximate the US grade level or years of education needed to understand the text.</p>
            """)
            explain_label.setWordWrap(True)
            layout.addWidget(table)
            layout.addWidget(explain_label)
        else:
            error_msg = readability.get('error', 'Unknown error calculating readability')
            layout.addWidget(QLabel(f"Readability metrics unavailable: {error_msg}"))
        tab.setLayout(layout)
        return tab
    def _create_pos_chart(self, pos_categories):
        """Create a chart for POS distribution"""
        try:
            from matplotlib.figure import Figure
            from matplotlib.backends.backend_qt5agg import FigureCanvasQTAgg
            figure = Figure(figsize=(8, 5), dpi=100)
            canvas = FigureCanvasQTAgg(figure)
            ax = figure.add_subplot(111)
            labels = []
            values = []
            for category, count in pos_categories.items():
                if count > 0:
                    labels.append(category)
                    values.append(count)
            wedges, texts, autotexts = ax.pie(
                values, 
                labels=labels, 
                autopct='%1.1f%%', 
                startangle=90,
                explode=[0.05] * len(values),
                shadow=True
            )
            for text in texts:
                text.set_fontsize(9)
            for autotext in autotexts:
                autotext.set_fontsize(8)
                autotext.set_fontweight('bold')
            ax.axis('equal')
            figure.suptitle('Part of Speech Distribution', fontsize=12)
            figure.tight_layout()
            if hasattr(self._parent, 'cleanup_figure'):
                self._parent.cleanup_figure(figure)
            return canvas
        except Exception as e:
            logger.error(f"Error creating POS chart: {str(e)}")
            label = QLabel("Failed to create POS chart")
            return label
    def _on_thread_complete(self, stats):
        """Handler saat thread menyelesaikan analisis"""
        self.stats_data = stats
        self.analysis_complete.emit(stats)
    def _on_thread_error(self, error_msg):
        """Handler saat thread mengalami error"""
        logger.error(f"Text analysis error: {error_msg}")
        self.error_occurred.emit(error_msg)
    def cleanup(self):
        """Membersihkan semua resource yang digunakan"""
        if self.current_thread and self.current_thread.isRunning():
            logger.info("Stopping text analysis thread")
            self.current_thread.quit()
            if not self.current_thread.wait(1000):
                logger.warning("Text analysis thread did not quit in time - forcing termination")
                self.current_thread.terminate()
                self.current_thread.wait()
        self.current_thread = None
        self.stats_data = None
class TextViewerService(QObject):
    """Service untuk menampilkan teks lengkap"""
    text_displayed = Signal(str)
    error_occurred = Signal(str)
    def __init__(self, parent=None):
        super().__init__(parent)
        self.current_dialog = None
        self._parent = parent
    def view_full_text(self, text_data, parent_widget):
        """Menampilkan teks lengkap dalam dialog dengan loading indicator"""
        if not text_data:
            self.error_occurred.emit("No text data available to display.")
            return
        try:
            if self.current_dialog is not None:
                self.current_dialog.close()
            self.current_dialog = QDialog(parent_widget)
            self.current_dialog.setWindowModality(Qt.NonModal)
            self.current_dialog.setWindowTitle("Full Text")
            self.current_dialog.setMinimumSize(500, 400)
            self.current_dialog.setSizeGripEnabled(True)
            layout = QVBoxLayout()
            loading_label = QLabel("Loading text content...")
            loading_label.setAlignment(Qt.AlignCenter)
            loading_label.setStyleSheet("font-size: 14px; color: #666;")
            layout.addWidget(loading_label)
            text_edit = QTextEdit()
            text_edit.setReadOnly(True)
            text_edit.hide()
            layout.addWidget(text_edit)
            button_layout = QHBoxLayout()
            button_layout.addStretch()
            close_button = QPushButton("Close")
            close_button.clicked.connect(self.current_dialog.accept)
            button_layout.addWidget(close_button)
            layout.addLayout(button_layout)
            self.current_dialog.setLayout(layout)
            self.current_dialog.show()
            self.current_dialog.loading_label = loading_label
            self.current_dialog.text_edit = text_edit
            QTimer.singleShot(100, lambda: self._set_text_content(text_data))
        except Exception as e:
            logger.error(f"Error showing full text dialog: {str(e)}")
            self.error_occurred.emit(f"Failed to display text: {str(e)}")
    def _set_text_content(self, text_data):
        """Set konten teks setelah dialog muncul untuk mencegah UI freeze"""
        try:
            if not self.current_dialog:
                return
            self.current_dialog.text_edit.setPlainText(text_data)
            self.current_dialog.loading_label.hide()
            self.current_dialog.text_edit.show()
            self.text_displayed.emit(text_data[:100] + "..." if len(text_data) > 100 else text_data)
        except Exception as e:
            logger.error(f"Error setting text content: {str(e)}")
            if self.current_dialog:
                self.current_dialog.loading_label.setText(f"Error loading text: {str(e)}")
    def close_dialog(self):
        """Close current dialog if open"""
        if self.current_dialog is not None:
            self.current_dialog.close()
            self.current_dialog = None
    def cleanup(self):
        """Cleanup resources"""
        self.close_dialog()
class TopicAnalysisResultDialog(QDialog):
    def __init__(self, method: str, results: dict | list, parent=None, evolution_data=None):
        super().__init__(parent)
        self.method = method.upper()
        self.results = results
        self.evolution_data = evolution_data
        self.setup_ui()
    def setup_ui(self):
        self.setWindowTitle(f"Topic Analysis Results - {self.method}")
        self.setMinimumSize(700, 850)
        layout = QVBoxLayout()
        info_group = QGroupBox("Analysis Information")
        info_layout = QVBoxLayout()
        method_info = {
            'LDA': """
                <b>Latent Dirichlet Allocation (LDA)</b> is a probabilistic method that identifies topics based on word distribution patterns, suitable for analyzing long documents, articles, and discovering hidden themes. Score interpretation: Represents the probability distribution of words in topics.
            """,
            'NMF': """
                <b>Non-negative Matrix Factorization (NMF)</b> is a linear algebra method that decomposes text into non-negative components, ideal for short to medium-length documents, specific topics, and sparse or rare data patterns. Score interpretation: Represents the weight/importance of words in topics.
            """
        }
        info_text = QLabel(method_info.get(self.method, ""))
        info_text.setWordWrap(True)
        info_layout.addWidget(info_text)
        info_group.setLayout(info_layout)
        layout.addWidget(info_group)
        tab_widget = QTabWidget()
        overview_tab = self._create_overview_tab()
        tab_widget.addTab(overview_tab, "Topic Overview")
        detail_tab = self._create_detail_tab()
        tab_widget.addTab(detail_tab, "Detailed View")
        if self.evolution_data:
            evolution_tab = self._create_evolution_tab()
            tab_widget.addTab(evolution_tab, "Topic Evolution")
        layout.addWidget(tab_widget)
        button_layout = QHBoxLayout()
        export_btn = QPushButton("Export Results")
        export_btn.clicked.connect(self.export_results)
        button_layout.addWidget(export_btn)
        close_btn = QPushButton("Close")
        close_btn.clicked.connect(self.close)
        button_layout.addWidget(close_btn)
        layout.addLayout(button_layout)
        self.setLayout(layout)
    def _create_overview_tab(self):
        """Create the topic overview tab"""
        tab = QWidget()
        layout = QVBoxLayout()
        overview_table = QTableWidget()
        overview_table.setColumnCount(3)
        overview_table.setHorizontalHeaderLabels(["Topic", "Top Terms", "Weight"])
        overview_table.horizontalHeader().setSectionResizeMode(1, QHeaderView.ResizeMode.Stretch)
        topics = self.results
        overview_table.setRowCount(len(topics))
        for i, topic in enumerate(topics):
            topic_name = QTableWidgetItem(topic['topic'])
            overview_table.setItem(i, 0, topic_name)
            terms = ", ".join(topic['terms'][:5])
            terms_item = QTableWidgetItem(terms)
            overview_table.setItem(i, 1, terms_item)
            weight = QTableWidgetItem(f"{topic['weight']:.2f}")
            overview_table.setItem(i, 2, weight)
        layout.addWidget(overview_table)
        tab.setLayout(layout)
        return tab
    def _create_detail_tab(self):
        """Create the detailed view tab"""
        tab = QWidget()
        layout = QVBoxLayout()
        detail_tree = QTreeWidget()
        detail_tree.setHeaderLabels(["Topic / Term", "Weight", "Related Terms"])
        detail_tree.setColumnWidth(0, 200)
        for topic in self.results:
            topic_item = QTreeWidgetItem([topic['topic'], f"{topic['weight']:.2f}", ""])
            topic_item.setExpanded(True)
            for term in topic['terms']:
                term_item = QTreeWidgetItem([term, "", ""])
                related_terms = self.find_related_terms(term, self.results)
                term_item.setText(2, ", ".join(related_terms))
                topic_item.addChild(term_item)
            detail_tree.addTopLevelItem(topic_item)
        layout.addWidget(detail_tree)
        tab.setLayout(layout)
        return tab
    def _create_evolution_tab(self):
        """Create the topic evolution visualization tab"""
        tab = QWidget()
        layout = QVBoxLayout()
        try:
            if not self.evolution_data:
                raise ValueError("No evolution data available")
            from matplotlib.figure import Figure
            from matplotlib.backends.backend_qt5agg import FigureCanvasQTAgg
            # import numpy as np
            
            # Set minimum size untuk dialog
            # self.setMinimumSize(700, 800)
            self.setSizePolicy(QSizePolicy.Minimum, QSizePolicy.Minimum)
            
            # Data preparation
            segments = self.evolution_data.get('segments', [])
            topics = self.evolution_data.get('topics', [])
            topic_strength = self.evolution_data.get('topic_strength', [])
            
            if not segments or not topics or not topic_strength:
                raise ValueError("Incomplete evolution data")
                
            topic_strength_array = np.array(topic_strength).T
            
            # Visualization container (fixed size)
            viz_container = QWidget()
            viz_container.setSizePolicy(QSizePolicy.Expanding, QSizePolicy.Fixed)
            viz_container.setFixedHeight(320)
            viz_layout = QVBoxLayout(viz_container)
            viz_layout.setContentsMargins(0, 0, 0, 0)
            
            # Create figure
            figure = Figure(dpi=100)
            canvas = FigureCanvasQTAgg(figure)
            canvas.setSizePolicy(QSizePolicy.Expanding, QSizePolicy.Fixed)
            canvas.setFixedHeight(320)
            
            def update_plot():
                figure.clear()
                figure.subplots_adjust(left=0.12, right=0.85, bottom=0.15, top=0.95, wspace=0.2)
                ax = figure.add_subplot(111)
                im = ax.imshow(topic_strength_array, aspect='auto', cmap='YlOrRd')
                
                ax.set_xticks(np.arange(len(segments)))
                ax.set_yticks(np.arange(len(topics)))
                ax.set_xticklabels(segments, rotation=0, ha='center', fontsize=8)
                ax.set_yticklabels(topics, fontsize=8)
                
                ax.set_xlabel('Text Segments', fontsize=9, labelpad=5)
                ax.set_ylabel('Topics', fontsize=9, labelpad=5)
                
                cbar_ax = figure.add_axes([0.87, 0.15, 0.03, 0.8])
                cbar = figure.colorbar(im, cax=cbar_ax)
                cbar.set_label('Topic Strength', fontsize=9, labelpad=5)
                cbar.ax.tick_params(labelsize=8)
                
                for i in range(len(segments)):
                    for j in range(len(topics)):
                        value = topic_strength_array[j, i]
                        text_color = 'white' if value > 0.5 else 'black'
                        ax.text(i, j, f"{value:.2f}", 
                            ha="center", va="center",
                            color=text_color, fontsize=7)
                
                ax.tick_params(axis='both', which='major', labelsize=8, pad=2)
                canvas.draw()
            
            viz_layout.addWidget(canvas)
            layout.addWidget(viz_container)
            
            canvas.mpl_connect('resize_event', lambda evt: update_plot())
            update_plot()
            
            # Explanation (fixed size)
            explanation = QLabel(
                "<p><b>Topic Evolution</b> shows how topics change throughout the text. "
                "Darker colors indicate stronger topic presence in that segment.</p>"
            )
            explanation.setWordWrap(True)
            explanation.setFixedHeight(40)
            layout.addWidget(explanation)
            
            # Create expandable content widget for bottom section
            content_widget = QWidget()
            content_widget.setSizePolicy(QSizePolicy.Expanding, QSizePolicy.Expanding)
            content_layout = QHBoxLayout(content_widget)
            
            # Left column: Topic Keywords
            keywords_group = QGroupBox("Topic Keywords")
            keywords_group.setSizePolicy(QSizePolicy.Expanding, QSizePolicy.Expanding)
            keywords_layout = QVBoxLayout()
            
            topic_keywords = self.evolution_data.get('topic_keywords', {})
            for topic, keywords in topic_keywords.items():
                keyword_text = ", ".join(keywords[:7])
                label = QLabel(f"<b>{topic}:</b> {keyword_text}")
                label.setWordWrap(True)
                keywords_layout.addWidget(label)
            
            keywords_layout.addStretch()
            keywords_group.setLayout(keywords_layout)
            content_layout.addWidget(keywords_group)
            
            # Right column: Segment Details
            details_group = QGroupBox("Segment Details")
            details_group.setSizePolicy(QSizePolicy.Expanding, QSizePolicy.Expanding)
            details_layout = QVBoxLayout()
            
            segment_table = QTableWidget()
            segment_table.setSizePolicy(QSizePolicy.Expanding, QSizePolicy.Expanding)
            segment_table.setColumnCount(3)
            segment_table.setHorizontalHeaderLabels(["Segment", "Dominant Topic", "Strength"])
            segment_table.horizontalHeader().setSectionResizeMode(0, QHeaderView.ResizeMode.Stretch)
            
            segment_details = self.evolution_data.get('segment_details', [])
            if segment_details:
                segment_table.setRowCount(len(segment_details))
                for i, segment in enumerate(segment_details):
                    segment_label = f"Segment {i+1}"
                    if 'segment_text' in segment:
                        segment_label += f": {segment['segment_text']}"
                    segment_table.setItem(i, 0, QTableWidgetItem(segment_label))
                    
                    try:
                        distribution = segment.get('topic_distribution', [])
                        if distribution:
                            dominant_idx = max(range(len(distribution)), 
                                            key=lambda i: distribution[i])
                            dominant_topic = f"Topic {dominant_idx + 1}"
                            dominant_strength = distribution[dominant_idx]
                            
                            segment_table.setItem(i, 1, QTableWidgetItem(dominant_topic))
                            segment_table.setItem(i, 2, QTableWidgetItem(f"{dominant_strength:.3f}"))
                    except:
                        segment_table.setItem(i, 1, QTableWidgetItem("N/A"))
                        segment_table.setItem(i, 2, QTableWidgetItem("N/A"))
            
            details_layout.addWidget(segment_table)
            details_group.setLayout(details_layout)
            content_layout.addWidget(details_group)
            
            # Add the content widget to main layout with stretch
            layout.addWidget(content_widget, stretch=1)
            
        except Exception as e:
            error_msg = f"Failed to create topic evolution visualization: {str(e)}"
            logger.error(error_msg)
            error_label = QLabel(error_msg)
            error_label.setStyleSheet("color: red")
            error_label.setWordWrap(True)
            layout.addWidget(error_label)
        
        tab.setLayout(layout)
        return tab
    def find_related_terms(self, target_term: str, topics: list) -> list:
        """Find terms that appear in other topics"""
        related = set()
        try:
            for topic in topics:
                terms = topic.get('terms', [])
                if target_term in terms:
                    related.update(terms[:5])
            related.discard(target_term)
            return list(related)[:5]
        except Exception as e:
            logger.error(f"Error finding related terms: {e}")
            return []
    def export_results(self):
        """Export results to CSV including evolution data if available"""
        import csv
        file_path, _ = QFileDialog.getSaveFileName(
            self,
            "Export Results",
            "",
            "CSV files (*.csv)"
        )
        if file_path:
            try:
                with open(file_path, 'w', newline='', encoding='utf-8') as f:
                    writer = csv.writer(f)
                    writer.writerow(['Topic Analysis Results'])
                    writer.writerow(['Method', self.method])
                    writer.writerow([])
                    writer.writerow(['Topic', 'Weight', 'Terms', 'Related Terms'])
                    for topic in self.results:
                        for term in topic['terms']:
                            related = self.find_related_terms(term, self.results)
                            writer.writerow([
                                topic['topic'],
                                f"{topic['weight']:.2f}",
                                term,
                                ', '.join(related)
                            ])
                    if self.evolution_data:
                        writer.writerow([])
                        writer.writerow(['Topic Evolution Analysis'])
                        writer.writerow([])
                        segments = self.evolution_data.get('segments', [])
                        topics = self.evolution_data.get('topics', [])
                        header = ['Topic'] + segments
                        writer.writerow(header)
                        topic_strength = self.evolution_data.get('topic_strength', [])
                        if topics and topic_strength:
                            try:
                                import numpy as np
                                strength_array = np.array(topic_strength).T
                                for i, topic in enumerate(topics):
                                    if i < len(strength_array):
                                        row = [topic] + [f"{x:.3f}" for x in strength_array[i]]
                                        writer.writerow(row)
                            except:
                                writer.writerow(['Topic evolution data could not be formatted'])
                        writer.writerow([])
                        writer.writerow(['Topic Keywords'])
                        topic_keywords = self.evolution_data.get('topic_keywords', {})
                        for topic, keywords in topic_keywords.items():
                            writer.writerow([topic, ', '.join(keywords)])
                QMessageBox.information(self, "Success", "Results exported successfully!")
            except Exception as e:
                logger.error(f"Export error: {str(e)}")
                QMessageBox.critical(self, "Error", f"Failed to export results: {str(e)}")
class KeywordExtractionResultDialog(QDialog):
    """Dialog untuk menampilkan hasil keyword extraction"""
    def __init__(self, method: str, results: list, parent=None):
        super().__init__(parent)
        self.method = method.upper()
        self.results = results
        self.setup_ui()
        
    def setup_ui(self):
        self.setWindowTitle(f"Keyword Extraction Results - {self.method}")
        self.setMinimumSize(800, 700)
        
        # Main layout
        layout = QVBoxLayout()
        layout.setSpacing(10)
        
        try:
            # Create visualization container with strict size control
            viz_container = QWidget()
            viz_container.setSizePolicy(QSizePolicy.Expanding, QSizePolicy.Fixed)
            viz_container.setFixedHeight(350)
            viz_layout = QVBoxLayout(viz_container)
            viz_layout.setContentsMargins(0, 0, 0, 0)
            
            # Import visualization libraries
            from matplotlib.figure import Figure
            from matplotlib.backends.backend_qt5agg import FigureCanvasQTAgg
            import numpy as np
            
            # Create figure with fixed DPI
            figure = Figure(dpi=100)
            canvas = FigureCanvasQTAgg(figure)
            canvas.setSizePolicy(QSizePolicy.Expanding, QSizePolicy.Fixed)
            canvas.setFixedHeight(350)
            
            def update_plot():
                figure.clear()
                
                # Get top 15 keywords
                top_keywords = self.results[:15]
                keywords = [kw['keyword'] for kw in top_keywords]
                scores = [kw['score'] for kw in top_keywords]
                
                # Create horizontal bar chart
                ax = figure.add_subplot(111)
                
                # Adjust subplot parameters for better layout
                figure.subplots_adjust(left=0.3, right=0.95, bottom=0.15, top=0.95)
                
                # Create bars
                y_pos = np.arange(len(keywords))
                bars = ax.barh(y_pos, scores, align='center', 
                             color='#4285F4', alpha=0.7)
                
                # Customize axes
                ax.set_yticks(y_pos)
                ax.set_yticklabels(keywords, fontsize=9)
                ax.invert_yaxis()  # Show highest score at top
                
                # Customize appearance
                ax.set_xlabel('Relevance Score', fontsize=9, labelpad=10)
                ax.tick_params(axis='x', labelsize=8, pad=5)
                
                # Remove unnecessary spines
                ax.spines['top'].set_visible(False)
                ax.spines['right'].set_visible(False)
                
                # Add grid
                ax.xaxis.grid(True, linestyle='--', alpha=0.3)
                
                canvas.draw()
            
            # Add canvas to container
            viz_layout.addWidget(canvas)
            layout.addWidget(viz_container)
            
            # Connect resize event
            canvas.mpl_connect('resize_event', lambda evt: update_plot())
            update_plot()

            # Add explanation
            method_info = {
                'TFIDF': "TF-IDF scores indicate how important a word is to a document in a collection. "
                        "Higher scores indicate more important keywords.",
                'YAKE': "YAKE scores are inversely related to importance - lower scores indicate "
                       "more important keywords.",
                'RAKE': "RAKE scores indicate keyword relevance based on word frequency and co-occurrence. "
                       "Higher scores indicate more important keywords."
            }
            
            explanation = QLabel(
                f"<p><b>Keyword Extraction Results</b></p>"
                f"<p>{method_info.get(self.method, '')}</p>"
            )
            explanation.setWordWrap(True)
            explanation.setFixedHeight(50)
            layout.addWidget(explanation)

            # Create main content widget that will expand
            content_widget = QWidget()
            content_widget.setSizePolicy(QSizePolicy.Expanding, QSizePolicy.Expanding)
            content_layout = QVBoxLayout(content_widget)
            content_layout.setContentsMargins(0, 0, 0, 0)
            
            # Results table
            table_group = QGroupBox("Detailed Results")
            table_group.setSizePolicy(QSizePolicy.Expanding, QSizePolicy.Expanding)
            table_layout = QVBoxLayout()
            
            table = QTableWidget()
            table.setSizePolicy(QSizePolicy.Expanding, QSizePolicy.Expanding)
            table.setColumnCount(3)
            table.setHorizontalHeaderLabels(["Rank", "Keyword", "Score"])
            table.horizontalHeader().setSectionResizeMode(1, QHeaderView.Stretch)
            table.setRowCount(len(self.results))
            
            for i, result in enumerate(self.results):
                table.setItem(i, 0, QTableWidgetItem(str(i + 1)))
                table.setItem(i, 1, QTableWidgetItem(result['keyword']))
                table.setItem(i, 2, QTableWidgetItem(f"{result['score']:.4f}"))
            
            table_layout.addWidget(table)
            table_group.setLayout(table_layout)
            content_layout.addWidget(table_group)
            
            # Add the content widget to main layout with stretch
            layout.addWidget(content_widget, stretch=1)

            # Add buttons
            button_layout = QHBoxLayout()
            
            export_btn = QPushButton("Export Results")
            export_btn.clicked.connect(self.export_results)
            button_layout.addWidget(export_btn)
            
            close_btn = QPushButton("Close")
            close_btn.clicked.connect(self.accept)
            button_layout.addWidget(close_btn)
            
            layout.addLayout(button_layout)

        except Exception as e:
            error_msg = f"Failed to create visualization: {str(e)}"
            logger.error(error_msg)
            error_label = QLabel(error_msg)
            error_label.setStyleSheet("color: red")
            error_label.setWordWrap(True)
            layout.addWidget(error_label)

        self.setLayout(layout)

    def export_results(self):
        """Export results to CSV"""
        file_path, _ = QFileDialog.getSaveFileName(
            self,
            "Export Results",
            "",
            "CSV files (*.csv)"
        )
        
        if file_path:
            try:
                import csv
                with open(file_path, 'w', newline='', encoding='utf-8') as f:
                    writer = csv.writer(f)
                    writer.writerow(['Rank', 'Keyword', 'Score'])
                    for i, result in enumerate(self.results, 1):
                        writer.writerow([
                            i,
                            result['keyword'],
                            f"{result['score']:.4f}"
                        ])
                QMessageBox.information(
                    self, 
                    "Success", 
                    f"Successfully exported {len(self.results)} keywords!"
                )
            except Exception as e:
                QMessageBox.critical(
                    self, 
                    "Error", 
                    f"Failed to export results: {str(e)}"
                )
class SentimentAnalysisResultDialog(QDialog):
    def __init__(self, sentiment_mode: str, result: dict, parent=None):
        super().__init__(parent)
        self.sentiment_mode = sentiment_mode
        self.result = result
        self.setup_ui()
    def setup_ui(self):
        self.setWindowTitle(f"Sentiment Analysis Results - {self.sentiment_mode}")
        self.setMinimumSize(600, 500)
        layout = QVBoxLayout()
        tab_widget = QTabWidget()
        overview_tab = self._create_overview_tab()
        tab_widget.addTab(overview_tab, "Overall Results")
        if 'sentiment_timeline' in self.result:
            timeline_tab = self._create_timeline_tab()
            tab_widget.addTab(timeline_tab, "Sentiment Timeline")
        layout.addWidget(tab_widget)
        button_layout = QHBoxLayout()
        export_btn = QPushButton("Export Results")
        export_btn.clicked.connect(self.export_results)
        button_layout.addWidget(export_btn)
        close_btn = QPushButton("Close")
        close_btn.clicked.connect(self.accept)
        button_layout.addWidget(close_btn)
        layout.addLayout(button_layout)
        self.setLayout(layout)
    def _create_overview_tab(self):
        """Create the overview tab with existing sentiment analysis results"""
        tab = QWidget()
        layout = QVBoxLayout()
        results_group = QGroupBox("Analysis Results")
        results_layout = QVBoxLayout()
        sentiment_label = self.result["sentiment_label"]
        label_widget = QLabel(f"Overall Sentiment: {sentiment_label}")
        label_widget.setAlignment(Qt.AlignCenter)
        label_widget.setStyleSheet(f"""
            QLabel {{
                font-size: 14pt;
                padding: 10px;
                border-radius: 5px;
                background-color: {
                    '#e8f5e9' if sentiment_label == 'POSITIVE'
                    else '#ffebee' if sentiment_label == 'NEGATIVE'
                    else '#f5f5f5'
                };
                color: {
                    '#2e7d32' if sentiment_label == 'POSITIVE'
                    else '#c62828' if sentiment_label == 'NEGATIVE'
                    else '#424242'
                };
            }}
        """)
        results_layout.addWidget(label_widget)
        scores_table = QTableWidget()
        scores_table.setColumnCount(2)
        scores_table.setHorizontalHeaderLabels(["Component", "Score"])
        scores_table.horizontalHeader().setSectionResizeMode(0, QHeaderView.Stretch)
        scores_table.horizontalHeader().setSectionResizeMode(1, QHeaderView.ResizeToContents)
        if self.sentiment_mode in ["VADER", "VADER (Custom Lexicon)"]:
            main_score_type = "Compound Score"
        elif self.sentiment_mode in ["TextBlob", "TextBlob (Custom Lexicon)"]:
            main_score_type = "Polarity Score"
        else:
            main_score_type = "Confidence Score"
        table_data = []
        table_data.append((main_score_type, f"{self.result['compound_score']:.4f}"))
        if self.sentiment_mode in ["VADER", "VADER (Custom Lexicon)"]:
            table_data.append(("Positive Score", f"{self.result['positive_score']:.4f}"))
            table_data.append(("Neutral Score", f"{self.result['neutral_score']:.4f}"))
            table_data.append(("Negative Score", f"{self.result['negative_score']:.4f}"))
        elif self.sentiment_mode in ["Flair", "Flair (Custom Model)"]:
            table_data.append(("Positive Score", f"{self.result['positive_score']:.4f}"))
            table_data.append(("Negative Score", f"{self.result['negative_score']:.4f}"))
        if self.sentiment_mode in ["TextBlob", "TextBlob (Custom Lexicon)"]:
            try:
                subj_value = float(self.result["subjectivity"])
                table_data.append(("Subjectivity", f"{subj_value:.4f}"))
            except (ValueError, TypeError):
                pass
        scores_table.setRowCount(len(table_data))
        for i, (component, score) in enumerate(table_data):
            scores_table.setItem(i, 0, QTableWidgetItem(component))
            scores_table.setItem(i, 1, QTableWidgetItem(score))
        results_layout.addWidget(scores_table)
        results_group.setLayout(results_layout)
        layout.addWidget(results_group)
        interp_group = QGroupBox("Interpretation")
        interp_layout = QVBoxLayout()
        interpretation = self.get_interpretation()
        interp_label = QLabel(interpretation)
        interp_label.setWordWrap(True)
        interp_layout.addWidget(interp_label)
        interp_group.setLayout(interp_layout)
        layout.addWidget(interp_group)
        info_group = QGroupBox("Method Information")
        info_layout = QVBoxLayout()
        method_info = self.get_method_info()
        info_label = QLabel(method_info)
        info_label.setWordWrap(True)
        info_layout.addWidget(info_label)
        info_group.setLayout(info_layout)
        layout.addWidget(info_group)
        tab.setLayout(layout)
        return tab
    def _create_timeline_tab(self):
        """Create the sentiment timeline visualization tab"""
        tab = QWidget()
        layout = QVBoxLayout()
        
        try:
            timeline_data = self.result['sentiment_timeline']
            
            # Import visualization libraries
            from matplotlib.figure import Figure
            from matplotlib.backends.backend_qt5agg import FigureCanvasQTAgg
            from matplotlib.colors import LinearSegmentedColormap
            
            # Visualization container (fixed size)
            viz_container = QWidget()
            viz_container.setSizePolicy(QSizePolicy.Expanding, QSizePolicy.Fixed)
            viz_container.setFixedHeight(350)
            viz_layout = QVBoxLayout(viz_container)
            viz_layout.setContentsMargins(0, 0, 0, 0)
            
            # Create figure with fixed size
            figure = Figure(dpi=100)
            canvas = FigureCanvasQTAgg(figure)
            canvas.setSizePolicy(QSizePolicy.Expanding, QSizePolicy.Fixed)
            canvas.setFixedHeight(350)
            
            def update_plot():
                figure.clear()
                figure.subplots_adjust(left=0.1, right=0.85, bottom=0.15, top=0.95)
                ax = figure.add_subplot(111)
                
                sentences = range(len(timeline_data))
                sentiments = [s['score'] for s in timeline_data]
                
                scatter = ax.scatter(sentences, 
                                sentiments,
                                c=sentiments,
                                cmap='RdYlGn',
                                vmin=-1,
                                vmax=1,
                                s=50)
                
                ax.plot(sentences, sentiments, 'k--', alpha=0.5)
                ax.set_xlabel('Sentence Number', fontsize=9)
                ax.set_ylabel('Sentiment Score', fontsize=9)
                ax.tick_params(axis='both', which='major', labelsize=8)
                ax.grid(True, linestyle='--', alpha=0.3)
                
                cbar_ax = figure.add_axes([0.87, 0.15, 0.03, 0.8])
                cbar = figure.colorbar(scatter, cax=cbar_ax)
                cbar.set_label('Sentiment Intensity', fontsize=9)
                cbar.ax.tick_params(labelsize=8)
                
                canvas.draw()
            
            viz_layout.addWidget(canvas)
            layout.addWidget(viz_container)
            
            canvas.mpl_connect('resize_event', lambda evt: update_plot())
            update_plot()
            
            # Explanation (fixed size)
            explanation = QLabel(
                "<p><b>Sentiment Timeline</b> shows how sentiment changes throughout the text. "
                "Green indicates positive sentiment, red indicates negative sentiment.</p>"
            )
            explanation.setWordWrap(True)
            explanation.setFixedHeight(40)
            layout.addWidget(explanation)
            
            # Create expandable content widget for details
            content_widget = QWidget()
            content_widget.setSizePolicy(QSizePolicy.Expanding, QSizePolicy.Expanding)
            content_layout = QVBoxLayout(content_widget)
            
            # Sentence-level Analysis group
            details_group = QGroupBox("Sentence-level Analysis")
            details_group.setSizePolicy(QSizePolicy.Expanding, QSizePolicy.Expanding)
            details_layout = QVBoxLayout()
            
            # Create table
            table = QTableWidget()
            table.setSizePolicy(QSizePolicy.Expanding, QSizePolicy.Expanding)
            table.setColumnCount(4)
            
            # Set up headers
            headers = ["Sentence", "Sentiment", "Score"]
            if self.sentiment_mode in ["TextBlob", "TextBlob (Custom Lexicon)"]:
                headers.append("Subjectivity")
            elif self.sentiment_mode in ["VADER", "VADER (Custom Lexicon)"]:
                headers.append("Compound")
            elif self.sentiment_mode in ["Flair", "Flair (Custom Model)"]:
                headers.append("Confidence")
            
            table.setHorizontalHeaderLabels(headers)
            table.setRowCount(len(timeline_data))
            table.horizontalHeader().setSectionResizeMode(0, QHeaderView.Stretch)
            
            # Populate table
            for i, data in enumerate(timeline_data):
                # Truncate long sentences
                sentence_text = data['sentence']
                if len(sentence_text) > 100:
                    sentence_text = sentence_text[:97] + "..."
                
                sentence_item = QTableWidgetItem(sentence_text)
                label_item = QTableWidgetItem(data['label'])
                score_item = QTableWidgetItem(f"{data['score']:.3f}")
                
                # Set background color based on sentiment
                label_item.setBackground(QColor(
                    '#e8f5e9' if data['label'] == 'POSITIVE' else
                    '#ffebee' if data['label'] == 'NEGATIVE' else
                    '#f5f5f5'
                ))
                
                table.setItem(i, 0, sentence_item)
                table.setItem(i, 1, label_item)
                table.setItem(i, 2, score_item)
                
                # Add mode-specific metric
                if self.sentiment_mode in ["TextBlob", "TextBlob (Custom Lexicon)"]:
                    subj_item = QTableWidgetItem(f"{data.get('subjectivity', 0):.3f}")
                    table.setItem(i, 3, subj_item)
                elif self.sentiment_mode in ["VADER", "VADER (Custom Lexicon)"]:
                    compound_item = QTableWidgetItem(f"{data.get('score', 0):.3f}")
                    table.setItem(i, 3, compound_item)
                elif self.sentiment_mode in ["Flair", "Flair (Custom Model)"]:
                    confidence_item = QTableWidgetItem(f"{data.get('confidence', 0):.3f}")
                    table.setItem(i, 3, confidence_item)
            
            details_layout.addWidget(table)
            details_group.setLayout(details_layout)
            content_layout.addWidget(details_group)
            
            # Add the content widget to main layout with stretch
            layout.addWidget(content_widget, stretch=1)
            
        except Exception as e:
            error_msg = f"Failed to create timeline visualization: {str(e)}"
            logger.error(error_msg)
            error_label = QLabel(error_msg)
            error_label.setStyleSheet("color: red")
            error_label.setWordWrap(True)
            layout.addWidget(error_label)
        
        tab.setLayout(layout)
        return tab
    def export_results(self):
        """Export results to CSV including timeline data"""
        import csv
        file_path, _ = QFileDialog.getSaveFileName(
            self,
            "Export Results",
            "",
            "CSV files (*.csv)"
        )
        if file_path:
            try:
                with open(file_path, 'w', newline='', encoding='utf-8') as f:
                    writer = csv.writer(f)
                    writer.writerow(['Overall Analysis Results'])
                    writer.writerow(['Analysis Method', self.sentiment_mode])
                    writer.writerow(['Overall Sentiment', self.result["sentiment_label"]])
                    if self.sentiment_mode in ["VADER", "VADER (Custom Lexicon)"]:
                        score_type = "Compound Score"
                    elif self.sentiment_mode in ["TextBlob", "TextBlob (Custom Lexicon)"]:
                        score_type = "Polarity Score"
                    else:
                        score_type = "Confidence Score"
                    writer.writerow([score_type, f"{self.result['compound_score']:.4f}"])
                    if self.sentiment_mode in ["VADER", "VADER (Custom Lexicon)"]:
                        writer.writerow(['Positive Score', f"{self.result['positive_score']:.4f}"])
                        writer.writerow(['Neutral Score', f"{self.result['neutral_score']:.4f}"])
                        writer.writerow(['Negative Score', f"{self.result['negative_score']:.4f}"])
                    elif self.sentiment_mode in ["Flair", "Flair (Custom Model)"]:
                        writer.writerow(['Positive Score', f"{self.result['positive_score']:.4f}"])
                        writer.writerow(['Negative Score', f"{self.result['negative_score']:.4f}"])
                    if self.sentiment_mode in ["TextBlob", "TextBlob (Custom Lexicon)"]:
                        writer.writerow(['Subjectivity', f"{self.result.get('subjectivity', 'N/A')}"])
                    if 'sentiment_timeline' in self.result:
                        writer.writerow([])
                        writer.writerow(['Sentiment Timeline Analysis'])
                        headers = ['Sentence', 'Sentiment', 'Score']
                        if self.sentiment_mode in ["TextBlob", "TextBlob (Custom Lexicon)"]:
                            headers.append('Subjectivity')
                        elif self.sentiment_mode in ["VADER", "VADER (Custom Lexicon)"]:
                            headers.append('Compound')
                        elif self.sentiment_mode in ["Flair", "Flair (Custom Model)"]:
                            headers.append('Confidence')
                        writer.writerow(headers)
                        for data in self.result['sentiment_timeline']:
                            row = [
                                data['sentence'],
                                data['label'],
                                f"{data['score']:.3f}"
                            ]
                            if self.sentiment_mode in ["TextBlob", "TextBlob (Custom Lexicon)"]:
                                row.append(f"{data.get('subjectivity', 0):.3f}")
                            elif self.sentiment_mode in ["VADER", "VADER (Custom Lexicon)"]:
                                row.append(f"{data.get('score', 0):.3f}")
                            elif self.sentiment_mode in ["Flair", "Flair (Custom Model)"]:
                                row.append(f"{data.get('confidence', 0):.3f}")
                            writer.writerow(row)
                QMessageBox.information(self, "Success", "Results exported successfully!")
            except Exception as e:
                QMessageBox.critical(self, "Error", f"Failed to export results: {str(e)}")
    def get_interpretation(self) -> str:
        """Get human-readable interpretation of results"""
        score = self.result["compound_score"]
        timeline_insight = ""
        if 'sentiment_timeline' in self.result:
            timeline = self.result['sentiment_timeline']
            pos_count = sum(1 for item in timeline if item['label'] == 'POSITIVE')
            neg_count = sum(1 for item in timeline if item['label'] == 'NEGATIVE')
            neu_count = sum(1 for item in timeline if item['label'] == 'NEUTRAL')
            total = len(timeline)
            timeline_insight = f"\n\nSentiment Distribution in Text: "
            timeline_insight += f"{pos_count/total:.1%} positive, "
            timeline_insight += f"{neg_count/total:.1%} negative, "
            timeline_insight += f"{neu_count/total:.1%} neutral sentences."
        if self.sentiment_mode in ["VADER", "VADER (Custom Lexicon)"]:
            if score >= 0.05:
                base_interp = "The text expresses positive sentiment. This indicates favorable or positive content."
            elif score <= -0.05:
                base_interp = "The text expresses negative sentiment. This indicates unfavorable or negative content."
            else:
                base_interp = "The text is neutral. This indicates balanced or unclear sentiment."
        elif self.sentiment_mode in ["TextBlob", "TextBlob (Custom Lexicon)"]:
            subj = self.result.get("subjectivity", 0)
            sentiment = "positive" if score > 0 else "negative" if score < 0 else "neutral"
            subj_text = ("highly subjective" if subj >= 0.7 
                        else "highly objective" if subj <= 0.3 
                        else "moderately subjective/objective")
            base_interp = f"The text is {sentiment} and {subj_text}."
        else:
            confidence = score
            base_interp = f"The model predicts {self.result['sentiment_label'].lower()} sentiment with {confidence:.2%} confidence."
        return base_interp + timeline_insight
    def get_method_info(self) -> str:
        """Get method-specific information"""
        if self.sentiment_mode in ["VADER", "VADER (Custom Lexicon)"]:
            info = """
                <p><b>VADER (Valence Aware Dictionary and sEntiment Reasoner)</b></p>
                <p>VADER is optimized for social media and short texts. It considers:</p>
                <ul>
                    <li>Word order and punctuation</li>
                    <li>Slang and emoticons</li>
                    <li>Intensity modifiers</li>
                </ul>
                <p><b>Score Interpretation:</b></p>
                <ul>
                    <li>Compound Score: [-1 to 1] where:
                        <ul>
                            <li>≥ 0.05: Positive</li>
                            <li>≤ -0.05: Negative</li>
                            <li>Otherwise: Neutral</li>
                        </ul>
                    </li>
                    <li>Individual scores (pos/neg/neu) represent proportions of text</li>
                </ul>
                """
            if "Custom Lexicon" in self.sentiment_mode:
                info += """
                    <p><b>Custom Lexicon Features:</b></p>
                    <ul>
                        <li>Uses domain-specific vocabulary</li>
                        <li>Adapted for specific language patterns</li>
                        <li>May better handle specialized content</li>
                    </ul>
                    """
        elif self.sentiment_mode in ["TextBlob", "TextBlob (Custom Lexicon)"]:
            info = """
                <p><b>TextBlob Sentiment Analysis</b></p>
                <p>TextBlob provides both polarity and subjectivity analysis:</p>
                <ul>
                    <li>Polarity: [-1 to 1] indicating negative to positive</li>
                    <li>Subjectivity: [0 to 1] indicating objective to subjective</li>
                </ul>
                <p><b>Features:</b></p>
                <ul>
                    <li>Pattern-based analysis</li>
                    <li>Natural language processing capabilities</li>
                    <li>Good for general-purpose text analysis</li>
                </ul>
                """
            if "Custom Lexicon" in self.sentiment_mode:
                info += """
                    <p><b>Custom Lexicon Features:</b></p>
                    <ul>
                        <li>Custom word polarity definitions</li>
                        <li>Domain-specific sentiment patterns</li>
                        <li>Enhanced accuracy for specific use cases</li>
                    </ul>
                    """
        else:
            info = """
                <p><b>Flair Sentiment Analysis</b></p>
                <p>Flair uses state-of-the-art deep learning models:</p>
                <ul>
                    <li>Contextual string embeddings</li>
                    <li>Neural network architecture</li>
                    <li>Pre-trained on large datasets</li>
                </ul>
                <p><b>Score Interpretation:</b></p>
                <ul>
                    <li>Confidence score [0 to 1] indicates prediction strength</li>
                    <li>Higher scores mean more confident predictions</li>
                </ul>
                """
            if "Custom Model" in self.sentiment_mode:
                info += """
                    <p><b>Custom Model Features:</b></p>
                    <ul>
                        <li>Fine-tuned for specific domains</li>
                        <li>Adapted to custom language patterns</li>
                        <li>May handle non-English text better</li>
                    </ul>
                    """
        return info    
class TopicAnalysisService(QObject):
    """Service for handling topic analysis and keyword extraction operations"""
    analysis_started = Signal()
    analysis_progress = Signal(int)
    analysis_complete = Signal(dict)
    error_occurred = Signal(str)
    topic_evolution_ready = Signal(object)
    def __init__(self, parent=None):
        super().__init__(parent)
        self.parent = parent
        self.topic_thread = None
        self.keyword_thread = None
        self._cached_results = {}
        self._cached_evolution = {}
        self._mutex = QMutex()
        self.stop_words = set()
        self.vectorizer = CountVectorizer(
            token_pattern=r'(?u)\b[a-zA-Z][a-zA-Z0-9\'-]*[a-zA-Z0-9]\b',
            lowercase=True,
            strip_accents='unicode',
            max_features=1000,
            min_df=1,
            max_df=0.95
        )
        self.tfidf = TfidfVectorizer(**self.vectorizer.get_params())
    def analyze_topics(self, text, method='lda', num_topics=5, custom_stopwords=None):
        """Analyze topics in text"""
        try:
            if not text.strip():
                raise ValueError("No text provided for analysis")
            self.analysis_started.emit()
            if self.topic_thread and self.topic_thread.isRunning():
                self.topic_thread.requestInterruption()
                self.topic_thread.wait(500)
            self.topic_thread = TopicAnalysisThread(
                text, 
                num_topics, 
                method.lower(),
                custom_stopwords=custom_stopwords
            )
            self.topic_thread.finished.connect(
                lambda results: self._handle_topic_results(results, method)
            )
            self.topic_thread.error.connect(self.error_occurred.emit)
            self.topic_thread.topic_evolution_ready.connect(
                self._handle_topic_evolution
            )
            self.topic_thread.start()
        except Exception as e:
            self.error_occurred.emit(f"Failed to start topic analysis: {str(e)}")
    def extract_keywords(self, text, method='tfidf', num_keywords=10, custom_stopwords=None):
        """Extract keywords from text"""
        try:
            if not text.strip():
                raise ValueError("No text provided for keyword extraction")
            self.analysis_started.emit()
            if self.keyword_thread and self.keyword_thread.isRunning():
                self.keyword_thread.requestInterruption()
                self.keyword_thread.wait(500)
            self.keyword_thread = KeywordExtractionThread(
                text, 
                method, 
                num_keywords,
                custom_stopwords=custom_stopwords
            )
            self.keyword_thread.finished.connect(
                lambda results: self._handle_keyword_results(results, method)
            )
            self.keyword_thread.error.connect(self.error_occurred.emit)
            self.keyword_thread.start()
        except Exception as e:
            self.error_occurred.emit(f"Failed to start keyword extraction: {str(e)}")
    def _handle_topic_results(self, results, method):
        """Handle topic analysis results"""
        if results:
            try:
                cache_key = f"topics_{method}_{hash(str(results))}"
                with QMutexLocker(self._mutex):
                    self._cached_results[cache_key] = results
                self.analysis_complete.emit({
                    'type': 'topics',
                    'method': method,
                    'results': results
                })
            except Exception as e:
                logger.error(f"Error handling topic results: {str(e)}")
                self.error_occurred.emit(f"Error processing topic results: {str(e)}")
    def _handle_topic_evolution(self, evolution_data):
        """Handle topic evolution data"""
        if evolution_data and isinstance(evolution_data, dict):
            try:
                method = evolution_data.get('method', 'unknown')
                cache_key = f"evolution_{method}_{hash(str(evolution_data))}"
                with QMutexLocker(self._mutex):
                    self._cached_evolution[cache_key] = evolution_data
                    if len(self._cached_evolution) > 10:
                        oldest_key = next(iter(self._cached_evolution))
                        del self._cached_evolution[oldest_key]
                self.topic_evolution_ready.emit(evolution_data)
            except Exception as e:
                logger.error(f"Error handling topic evolution: {str(e)}")
    def _handle_keyword_results(self, results, method):
        """Handle keyword extraction results"""
        if results:
            try:
                cache_key = f"keywords_{method}_{hash(str(results))}"
                with QMutexLocker(self._mutex):
                    self._cached_results[cache_key] = results
                self.analysis_complete.emit({
                    'type': 'keywords',
                    'method': method,
                    'results': results
                })
            except Exception as e:
                logger.error(f"Error handling keyword results: {str(e)}")
                self.error_occurred.emit(f"Error processing keyword results: {str(e)}")
    def get_cached_evolution(self, method=None):
        """Get cached evolution data for specified method"""
        with QMutexLocker(self._mutex):
            if method:
                for key, data in self._cached_evolution.items():
                    if key.startswith(f"evolution_{method}_"):
                        return data
                return None
            else:
                if self._cached_evolution:
                    return next(reversed(self._cached_evolution.values()))
                return None
    def cancel_analysis(self):
        """Cancel ongoing analysis operations"""
        if self.topic_thread and self.topic_thread.isRunning():
            self.topic_thread.requestInterruption()
            self.topic_thread.wait(500)
        if self.keyword_thread and self.keyword_thread.isRunning():
            self.keyword_thread.requestInterruption()
            self.keyword_thread.wait(500)
    def cleanup(self):
        """Clean up resources"""
        self.cancel_analysis()
        with QMutexLocker(self._mutex):
            self._cached_results.clear()
            self._cached_evolution.clear()
"""CLASS ANALISIS"""
class BaseAnalyzer:
    def __init__(self):
        self._initialized = False
    @property
    def is_initialized(self):
        return self._initialized
    def cleanup(self):
        pass
class SentimentAnalyzer(BaseAnalyzer):
    def __init__(self, mode, text_data, analyzers):
        super().__init__()
        self.mode = mode
        self.text_data = text_data
        self.analyzers = analyzers
    def get_score_type(self):
        return AppConstants.SCORE_TYPES.get(self.mode, "Score")
class CustomVaderSentimentIntensityAnalyzer:
    def __init__(self, lexicon_file="vader_lexicon.txt", custom_lexicon_file=None):
        from vaderSentiment.vaderSentiment import SentimentIntensityAnalyzer
        self.analyzer = SentimentIntensityAnalyzer()
        if custom_lexicon_file:
            self.load_custom_lexicon(custom_lexicon_file)
    def load_custom_lexicon(self, custom_lexicon_file):
        try:
            with open(custom_lexicon_file, "r", encoding="utf-8") as file:
                for line in file:
                    parts = line.strip().split("\t")
                    if len(parts) == 2:
                        word, measure = parts
                        try:
                            self.analyzer.lexicon[word] = float(measure)
                        except ValueError:
                            pass
        except Exception as e:
            QMessageBox.warning(None, "Warning", f"Failed to load custom lexicon: {e}")
    def polarity_scores(self, text):
        return self.analyzer.polarity_scores(text)
class CustomTextBlobSentimentAnalyzer:
    def __init__(self, custom_lexicon_file=None):
        self.lexicon = {}
        self.intensifiers = {}
        self.negations = set()
        if custom_lexicon_file:
            self.load_custom_lexicon(custom_lexicon_file)
    def load_custom_lexicon(self, custom_lexicon_file):
        try:
            with open(custom_lexicon_file, "r", encoding="utf-8") as file:
                for line in file:
                    parts = line.strip().split("\t")
                    if len(parts) == 2:
                        word, measure = parts
                        if measure == "negation":
                            self.negations.add(word)
                        elif measure.startswith("intensifier:"):
                            try:
                                self.intensifiers[word] = float(measure.split(":")[1])
                            except ValueError:
                                continue
                        else:
                            try:
                                self.lexicon[word] = float(measure)
                            except ValueError:
                                continue
        except Exception as e:
            QMessageBox.warning(None, "Warning", f"Failed to load custom lexicon: {e}")
    def analyze(self, text):
        from textblob import TextBlob
        blob = TextBlob(text)
        total_polarity = 0.0
        words_with_context = []
        for sentence in blob.sentences:
            words = sentence.words
            for i, word in enumerate(words):
                word_lower = word.lower()
                current_score = self.lexicon.get(word_lower, 0.0)
                if i > 0 and words[i-1].lower() in self.negations:
                    current_score *= -1
                if i > 0 and words[i-1].lower() in self.intensifiers:
                    multiplier = self.intensifiers[words[i-1].lower()]
                    current_score *= multiplier
                words_with_context.append(current_score)
        valid_scores = [s for s in words_with_context if s != 0]
        avg_polarity = sum(valid_scores) / len(valid_scores) if valid_scores else 0.0
        return {
            'polarity': avg_polarity,
            'subjectivity': blob.sentiment.subjectivity,
            'processed_words': len(valid_scores)
        }
"""CLASS PROSES SPESIFIK"""
class TextLoadThread(QThread):
    """Thread untuk memuat dan memproses teks"""
    text_ready = Signal(str)
    error_occurred = Signal(str)
    def __init__(self, text_data):
        super().__init__()
        self._text_data = text_data
        self._is_running = False
    def run(self):
        """Process text in background thread"""
        self._is_running = True
        try:
            processed_text = self._process_text(self._text_data)
            if self._is_running:
                self.text_ready.emit(processed_text)
        except Exception as e:
            logger.error(f"Error processing text: {e}")
            self.error_occurred.emit(str(e))
        finally:
            self._is_running = False
    def _process_text(self, text):
        """Process text with proper formatting"""
        try:
            lines = text.split('\n')
            formatted_lines = []
            for line in lines:
                if not self._is_running:
                    break
                line = line.replace('&', '&amp;')
                line = line.replace('<', '&lt;')
                line = line.replace('>', '&gt;')
                formatted_lines.append(line)
            return '<br>'.join(formatted_lines)
        except Exception as e:
            logger.error(f"Text processing error: {e}")
            raise
    def cancel(self):
        """Cancel the running operation"""
        self._is_running = False
class TextStatsThread(QThread):
    """Thread untuk menghitung statistik teks tanpa memblokir UI"""
    stats_ready = Signal(dict)
    error_occurred = Signal(str)
    progress_updated = Signal(int)
    def __init__(self, text_data):
        super().__init__()
        self.text_data = text_data
        self.include_stopwords = False
        self.custom_stopwords = None
        self.setObjectName("TextStatsThread")
    def cleanup_resources(self):
        """Clean up resources used by this thread"""
        logger.debug(f"Cleaning up resources for {self.objectName()}")
        try:
            if hasattr(self, 'word_freq'):
                del self.word_freq
            if hasattr(self, 'filtered_tokens'):
                del self.filtered_tokens
            if hasattr(self, 'basic_tokens'):
                del self.basic_tokens
        except Exception as e:
            logger.error(f"Error cleaning up text stats thread: {e}")        
    def run(self):
        try:
            if self.isInterruptionRequested():
                return            
            self.progress_updated.emit(5)
            import nltk
            from nltk.tokenize import word_tokenize, sent_tokenize
            import re
            import string
            from collections import Counter
            import textstat
            if self.isInterruptionRequested():
                return            
            """
            nltk_resources = ['punkt', 'stopwords']
            for resource in nltk_resources:
                try:
                    nltk.data.find(f'tokenizers/{resource}')
                except LookupError:
                    nltk.download(resource, quiet=True)
            """
            self.progress_updated.emit(10)                    
            stats = {}
            stats['char_count'] = len(self.text_data)
            stats['word_count'] = len(re.findall(r'\b\w+\b', self.text_data))
            stats['sentence_count'] = len(sent_tokenize(self.text_data))
            stats['paragraph_count'] = len([p for p in self.text_data.split('\n\n') if p.strip()])
            if self.isInterruptionRequested():
                return            
            self.progress_updated.emit(20)
            stats['avg_word_length'] = sum(len(word) for word in re.findall(r'\b\w+\b', self.text_data)) / stats['word_count'] if stats['word_count'] > 0 else 0
            stats['avg_sentence_length'] = stats['word_count'] / stats['sentence_count'] if stats['sentence_count'] > 0 else 0
            if self.isInterruptionRequested():
                return            
            self.progress_updated.emit(30)
            tokens = word_tokenize(self.text_data.lower())
            basic_tokens = [token for token in tokens 
                            if token not in string.punctuation
                            and not token.isdigit()
                            and len(token) > 1]
            if self.isInterruptionRequested():
                return                            
            self.progress_updated.emit(40)
            if not self.include_stopwords:
                stopwords_list = set(nltk.corpus.stopwords.words('english'))
                if self.custom_stopwords:
                    stopwords_list.update(self.custom_stopwords)
                filtered_tokens = [token for token in basic_tokens if token not in stopwords_list]
                stats['filtered_word_count'] = len(filtered_tokens)
                stats['filtered_unique_words'] = len(set(filtered_tokens))
                word_freq = Counter(filtered_tokens)
                stats['stopwords_used'] = list(stopwords_list)
            else:
                word_freq = Counter(basic_tokens)
                stats['filtered_unique_words'] = len(set(basic_tokens))
            if self.isInterruptionRequested():
                return                
            self.progress_updated.emit(60)
            stats['word_freq'] = word_freq
            stats['stopwords_removed'] = not self.include_stopwords
            stats['most_common_words'] = word_freq.most_common(50)
            stats['unique_word_count'] = len(set(basic_tokens))
            if len(basic_tokens) > 0:
                stats['lexical_diversity'] = len(set(basic_tokens)) / len(basic_tokens)
            else:
                stats['lexical_diversity'] = 0
            if self.isInterruptionRequested():
                return                
            self.progress_updated.emit(70)
            stats['readability'] = self._calculate_readability(self.text_data)
            if self.isInterruptionRequested():
                return                
            self.progress_updated.emit(80)
            """
            try:
                nltk.data.find('taggers/averaged_perceptron_tagger')
            except LookupError:
                nltk.download('averaged_perceptron_tagger', quiet=True)
            """
            try:
                pos_tags = nltk.pos_tag(basic_tokens)
                pos_counts = Counter(tag for word, tag in pos_tags)
                stats['pos_counts'] = pos_counts
                if self.isInterruptionRequested():
                    return                
                self.progress_updated.emit(90)
                category_map = {
                    'Nouns': ['NN', 'NNS', 'NNP', 'NNPS'],
                    'Verbs': ['VB', 'VBD', 'VBG', 'VBN', 'VBP', 'VBZ'],
                    'Adjectives': ['JJ', 'JJR', 'JJS'],
                    'Adverbs': ['RB', 'RBR', 'RBS'],
                    'Pronouns': ['PRP', 'PRP$', 'WP', 'WP$'],
                    'Determiners': ['DT', 'PDT', 'WDT'],
                    'Conjunctions': ['CC', 'IN'],
                    'Other': []
                }
                pos_categories = {category: 0 for category in category_map}
                for tag, count in pos_counts.items():
                    for category, tag_list in category_map.items():
                        if tag in tag_list:
                            pos_categories[category] += count
                            break
                    else:
                        pos_categories['Other'] += count
                stats['pos_categories'] = pos_categories
            except Exception as e:
                stats['pos_error'] = str(e)
            if self.isInterruptionRequested():
                return            
            self.progress_updated.emit(100)
            self.stats_ready.emit(stats)
        except Exception as e:
            import traceback
            error_msg = f"Error calculating text statistics: {str(e)}\n{traceback.format_exc()}"
            self.error_occurred.emit(error_msg)
            self.progress_updated.emit(0)
    def _calculate_readability(self, text):
        """Calculate various readability metrics"""
        try:
            import textstat
            readability = {}
            readability['flesch_reading_ease'] = textstat.flesch_reading_ease(text)
            readability['flesch_kincaid_grade'] = textstat.flesch_kincaid_grade(text)
            readability['gunning_fog'] = textstat.gunning_fog(text)
            readability['smog_index'] = textstat.smog_index(text)
            readability['automated_readability_index'] = textstat.automated_readability_index(text)
            readability['coleman_liau_index'] = textstat.coleman_liau_index(text)
            readability['linsear_write_formula'] = textstat.linsear_write_formula(text)
            readability['dale_chall_readability_score'] = textstat.dale_chall_readability_score(text)
            readability['readability_consensus'] = textstat.text_standard(text, float_output=False)
            return readability
        except ImportError:
            return {"error": "textstat library not available"}
        except Exception as e:
            return {"error": str(e)}
class ColormapLoaderThread(QThread):
    """Thread for loading matplotlib colormaps asynchronously"""
    loading_complete = Signal(dict)
    loading_progress = Signal(str)
    error_occurred = Signal(str)
    def __init__(self, parent=None):
        super().__init__(parent)
        self.setObjectName("ColormapLoaderThread")
        self._is_running = True
    def run(self):
        try:
            import matplotlib.pyplot as plt
            colormaps = {}
            for name in plt.colormaps():
                if not self._is_running:
                    return
                try:
                    colormaps[name] = plt.get_cmap(name)
                    self.loading_progress.emit(name)
                except Exception as e:
                    logger.debug(f"Failed to load colormap {name}: {e}")
            self.loading_complete.emit(colormaps)
        except Exception as e:
            self.error_occurred.emit(f"Error loading colormaps: {e}")
            logger.error(f"Error in colormap loader thread: {e}")
    def stop(self):
        """Safely stop the thread"""
        self._is_running = False
        self.wait(500)
class FontLoaderThread(QThread):
    """Thread untuk loading matplotlib fonts secara asinkron dengan optimasi performa"""
    loading_complete = Signal(dict)
    loading_progress = Signal(str, int)
    loading_finished = Signal()
    error_occurred = Signal(str)
    def __init__(self, parent=None, max_fonts=None):
        super().__init__(parent)
        self.setObjectName("FontLoaderThread")
        self._stop_requested = False
        self._max_fonts = max_fonts
    def run(self):
        try:
            fonts_dict = {}
            start_time = time.time()
            self.loading_progress.emit("Initializing font system...", 0)
            import matplotlib.font_manager as fm
            weight_conversion = {
                100: "Thin", 200: "Extra Light", 300: "Light", 
                400: "Regular", 500: "Medium", 600: "Semi Bold", 
                700: "Bold", 800: "Extra Bold", 900: "Black"
            }
            default_fonts = ["Arial", "Times New Roman", "Verdana", "Courier New", "Tahoma", "Georgia"]
            for font_name in default_fonts:
                try:
                    path = fm.findfont(fm.FontProperties(family=font_name), fallback=False)
                    if path and os.path.exists(path):
                        fonts_dict[font_name] = path
                except Exception:
                    pass
            self.loading_progress.emit("Searching for fonts...", 5)
            self.loading_progress.emit("Loading font manager...", 10)
            all_fonts = fm.fontManager.ttflist
            total_fonts = len(all_fonts)
            self.loading_progress.emit(f"Found {total_fonts} font files", 15)
            processed = 0
            for i, font in enumerate(all_fonts):
                if self._stop_requested:
                    logger.debug("Font loading interrupted")
                    break
                try:
                    family = font.family_name if hasattr(font, "family_name") else font.name
                    style = font.style_name.lower() if hasattr(font, "style_name") else font.style.lower()
                    weight = weight_conversion.get(font.weight, "Regular")
                    if not family or family.strip() == "":
                        continue
                    display_parts = [family]
                    if "italic" in style or "oblique" in style:
                        display_parts.append("Italic")
                    elif style != "normal" and style != "regular":
                        display_parts.append(style.title())
                    if weight != "Regular":
                        display_parts.append(weight)
                    display_name = " ".join(display_parts)
                    if display_name and display_name.strip() and display_name not in fonts_dict:
                        fonts_dict[display_name] = font.fname
                        processed += 1
                    if i % 20 == 0 or i == total_fonts - 1:
                        progress = min(15 + int(85 * (i / total_fonts)), 100)
                        remaining = total_fonts - i - 1
                        self.loading_progress.emit(
                            f"Processing fonts: {i+1}/{total_fonts} ({remaining} remaining)", 
                            progress
                        )
                        self.msleep(1)
                except Exception as e:
                    continue
            elapsed = time.time() - start_time
            logger.info(f"Font loading completed in {elapsed:.2f}s - loaded {len(fonts_dict)} fonts")
            if not self._stop_requested:
                self.loading_complete.emit(fonts_dict)
                self.loading_finished.emit()
        except Exception as e:
            error_msg = f"Error loading fonts: {str(e)}"
            logger.error(error_msg)
            self.error_occurred.emit(error_msg)
            self.loading_finished.emit()
    def stop(self):
        """Request graceful stop of the thread"""
        self._stop_requested = True            
class FlairModelLoaderThread(QThread):
    model_loaded = Signal(object)
    error_occurred = Signal(str)
    finished = Signal()
    _cached_model = None
    def __init__(self, model_path="sentiment"):
        super().__init__()
        self.model_path = model_path
        self._is_loading = False
        self._interrupt_requested = False
    def run(self):
        try:
            from flair.models import TextClassifier
            if self.isInterruptionRequested():
                return
            self._is_loading = True
            self._interrupt_requested = False
            if FlairModelLoaderThread._cached_model is None:
                try:
                    FlairModelLoaderThread._cached_model = TextClassifier.load(self.model_path)
                except Exception as e:
                    if not self.isInterruptionRequested():
                        self.error_occurred.emit(str(e))
                    return
            if self.isInterruptionRequested():
                return
            self.model_loaded.emit(FlairModelLoaderThread._cached_model)
        except Exception as e:
            if not self.isInterruptionRequested():
                self.error_occurred.emit(str(e))
        finally:
            self._is_loading = False
            self.finished.emit()
    def requestInterruption(self):
        logger.info("FlairModelLoader: Interruption requested")
        self._interrupt_requested = True
        self._is_loading = False
        super().requestInterruption()
class TopicAnalysisThread(QThread):
    """Thread for topic analysis to prevent UI freezing"""
    finished = Signal(list)
    error = Signal(str)
    topic_evolution_ready = Signal(object)
    def __init__(self, text, num_topics, method='lda', custom_stopwords=None):
        super().__init__()
        self.text = text
        self.num_topics = num_topics
        self.method = method
        self._is_canceled = False        
        self.setObjectName(f"TopicAnalysisThread_{method}")
        self._resources = []
        from nltk.corpus import stopwords
        try:
            default_stopwords = set(stopwords.words('english'))
        except:
            default_stopwords = STOPWORDS
        default_stopwords = {word.lower().strip() for word in default_stopwords}
        if custom_stopwords:
            custom_stopwords = {word.lower().strip() for word in custom_stopwords if word.strip()}
            self.stopwords = default_stopwords.union(custom_stopwords)
        else:
            self.stopwords = default_stopwords
        self._is_running = True
    def cleanup_resources(self):
        """Clean up resources used by this thread"""
        logger.debug(f"Cleaning up resources for {self.objectName()}")
        try:
            if hasattr(self, 'lda_model'):
                del self.lda_model
            if hasattr(self, 'nmf_model'):
                del self.nmf_model
            if hasattr(self, 'vectorizer'):
                del self.vectorizer
            if hasattr(self, 'corpus'):
                del self.corpus
            if hasattr(self, 'dictionary'):
                del self.dictionary
            if hasattr(self, 'evolution_data'):
                del self.evolution_data
            for resource in self._resources:
                if hasattr(resource, 'close'):
                    try:
                        resource.close()
                    except:
                        pass
                elif hasattr(resource, 'clean_up'):
                    try:
                        resource.clean_up()
                    except:
                        pass
            self._resources.clear()
        except Exception as e:
            logger.error(f"Error cleaning up topic analysis thread: {e}")
    def run(self):
        try:
            if self._is_canceled or self.isInterruptionRequested():
                return
            if self.method == 'lda':
                result = self._process_lda()
                if self._is_running and not self._is_canceled and not self.isInterruptionRequested():
                    self._analyze_topic_evolution_lda()
            else:
                result = self._process_nmf()
                if self._is_running and not self._is_canceled and not self.isInterruptionRequested():
                    self._analyze_topic_evolution_nmf()
            if self._is_running:
                self.finished.emit(result)
            if self._is_canceled or self.isInterruptionRequested():
                return
        except Exception as e:
            self.error.emit(str(e))
    def _process_lda(self):
        """Process LDA analysis"""
        try:
            if self._is_canceled or self.isInterruptionRequested():
                return []
            from sklearn.feature_extraction.text import CountVectorizer
            from sklearn.decomposition import LatentDirichletAllocation
            vectorizer = CountVectorizer(
                max_features=1000,
                stop_words=list(self.stopwords),
                token_pattern=r'(?u)\b[a-zA-Z][a-zA-Z0-9\'-]*[a-zA-Z0-9]\b',
                lowercase=True,
                strip_accents='unicode'
            )
            if self._is_canceled or self.isInterruptionRequested():
                return []
            dtm = vectorizer.fit_transform([self.text])
            if self._is_canceled or self.isInterruptionRequested():
                return []
            lda = LatentDirichletAllocation(
                n_components=self.num_topics,
                random_state=42
            )
            if self._is_canceled or self.isInterruptionRequested():
                return []
            lda.fit(dtm)
            self.vectorizer = vectorizer
            self.lda_model = lda
            if self._is_canceled or self.isInterruptionRequested():
                return []
            terms = vectorizer.get_feature_names_out()
            topics = []
            for idx, topic in enumerate(lda.components_):
                if self._is_canceled or self.isInterruptionRequested():
                    return []
                top_terms = [terms[i] for i in topic.argsort()[:-10-1:-1]]
                topics.append({
                    'topic': f'Topic {idx + 1}',
                    'terms': top_terms,
                    'weight': topic.sum()
                })
            return topics
        except Exception as e:
            raise Exception(f"LDA analysis failed: {str(e)}")
    def _process_nmf(self):
        """Process NMF analysis"""
        try:
            if self._is_canceled or self.isInterruptionRequested():
                return []
            from sklearn.feature_extraction.text import TfidfVectorizer
            from sklearn.decomposition import NMF
            vectorizer = TfidfVectorizer(
                max_features=1000,
                stop_words=list(self.stopwords),
                token_pattern=r'(?u)\b[a-zA-Z][a-zA-Z0-9\'-]*[a-zA-Z0-9]\b',
                lowercase=True,
                strip_accents='unicode'
            )
            if self._is_canceled or self.isInterruptionRequested():
                return []
            dtm = vectorizer.fit_transform([self.text])
            if self._is_canceled or self.isInterruptionRequested():
                return []
            nmf = NMF(
                n_components=self.num_topics,
                random_state=42
            )
            if self._is_canceled or self.isInterruptionRequested():
                return []
            nmf.fit(dtm)
            self.vectorizer = vectorizer
            self.nmf_model = nmf
            if self._is_canceled or self.isInterruptionRequested():
                return []
            terms = vectorizer.get_feature_names_out()
            topics = []
            for idx, topic in enumerate(nmf.components_):
                if self._is_canceled or self.isInterruptionRequested():
                    return []
                top_terms = [terms[i] for i in topic.argsort()[:-10-1:-1]]
                topics.append({
                    'topic': f'Topic {idx + 1}',
                    'terms': top_terms,
                    'weight': topic.sum()
                })
            return topics
        except Exception as e:
            raise Exception(f"NMF analysis failed: {str(e)}")
    def _analyze_topic_evolution_lda(self):
        """Analyze topic evolution using LDA model"""
        try:
            if self._is_canceled or self.isInterruptionRequested() or not hasattr(self, 'lda_model'):
                return
            segments = self._split_into_segments(self.text, 5)
            if self._is_canceled or self.isInterruptionRequested():
                return
            vectorizer = self.vectorizer
            lda_model = self.lda_model
            feature_names = vectorizer.get_feature_names_out()
            segment_topics = []
            segment_matrices = []
            for i, segment in enumerate(segments):
                if self._is_canceled or self.isInterruptionRequested():
                    return
                segment_dtm = vectorizer.transform([segment])
                segment_topics_matrix = lda_model.transform(segment_dtm)
                segment_matrices.append(segment_topics_matrix[0])
                segment_topic_terms = []
                for topic_idx, topic in enumerate(lda_model.components_):
                    top_terms_idx = topic.argsort()[:-10-1:-1]
                    top_terms = [feature_names[i] for i in top_terms_idx]
                    segment_topic_terms.append(top_terms)
                segment_topics.append({
                    'segment_id': i,
                    'segment_text': segment[:100] + '...' if len(segment) > 100 else segment,
                    'topic_distribution': segment_topics_matrix[0].tolist(),
                    'topic_terms': segment_topic_terms
                })
            topics_keywords = {}
            for topic_idx, topic in enumerate(lda_model.components_):
                top_terms_idx = topic.argsort()[:-10-1:-1]
                top_terms = [feature_names[i] for i in top_terms_idx]
                topics_keywords[f'Topic {topic_idx+1}'] = top_terms
            evolution_data = {
                'segments': [f'Segment {i+1}' for i in range(len(segments))],
                'segment_texts': [s[:100] + '...' if len(s) > 100 else s for s in segments],
                'topics': [f'Topic {i+1}' for i in range(self.num_topics)],
                'topic_keywords': topics_keywords,
                'topic_strength': [m.tolist() for m in segment_matrices],
                'segment_details': segment_topics,
                'method': 'lda'
            }
            self.evolution_data = evolution_data
            self.topic_evolution_ready.emit(evolution_data)
        except Exception as e:
            logger.error(f"Error analyzing LDA topic evolution: {str(e)}")
    def _analyze_topic_evolution_nmf(self):
        """Analyze topic evolution using NMF model"""
        try:
            if self._is_canceled or self.isInterruptionRequested() or not hasattr(self, 'nmf_model'):
                return
            segments = self._split_into_segments(self.text, 5)
            if self._is_canceled or self.isInterruptionRequested():
                return
            vectorizer = self.vectorizer
            nmf_model = self.nmf_model
            feature_names = vectorizer.get_feature_names_out()
            segment_topics = []
            segment_matrices = []
            for i, segment in enumerate(segments):
                if self._is_canceled or self.isInterruptionRequested():
                    return
                segment_dtm = vectorizer.transform([segment])
                segment_topics_matrix = nmf_model.transform(segment_dtm)
                segment_matrices.append(segment_topics_matrix[0])
                segment_topic_terms = []
                for topic_idx, topic in enumerate(nmf_model.components_):
                    top_terms_idx = topic.argsort()[:-10-1:-1]
                    top_terms = [feature_names[i] for i in top_terms_idx]
                    segment_topic_terms.append(top_terms)
                segment_topics.append({
                    'segment_id': i,
                    'segment_text': segment[:100] + '...' if len(segment) > 100 else segment,
                    'topic_distribution': segment_topics_matrix[0].tolist(),
                    'topic_terms': segment_topic_terms
                })
            topics_keywords = {}
            for topic_idx, topic in enumerate(nmf_model.components_):
                top_terms_idx = topic.argsort()[:-10-1:-1]
                top_terms = [feature_names[i] for i in top_terms_idx]
                topics_keywords[f'Topic {topic_idx+1}'] = top_terms
            evolution_data = {
                'segments': [f'Segment {i+1}' for i in range(len(segments))],
                'segment_texts': [s[:100] + '...' if len(s) > 100 else s for s in segments],
                'topics': [f'Topic {i+1}' for i in range(self.num_topics)],
                'topic_keywords': topics_keywords,
                'topic_strength': [m.tolist() for m in segment_matrices],
                'segment_details': segment_topics,
                'method': 'nmf'
            }
            self.evolution_data = evolution_data
            self.topic_evolution_ready.emit(evolution_data)
        except Exception as e:
            logger.error(f"Error analyzing NMF topic evolution: {str(e)}")
    def _split_into_segments(self, text, num_segments):
        """Split text into roughly equal segments"""
        words = text.split()
        segment_size = max(1, len(words) // num_segments)
        segments = []
        for i in range(0, len(words), segment_size):
            segment = ' '.join(words[i:i + segment_size])
            segments.append(segment)
        while len(segments) > num_segments:
            segments.pop()
        while len(segments) < num_segments:
            segments.append("")
        return segments
    def requestInterruption(self):
        """Override method untuk menginisiasi interupsi thread"""
        super().requestInterruption()
        self._is_canceled = True
class KeywordExtractionThread(QThread):
    """Thread for keyword extraction to prevent UI freezing"""
    finished = Signal(list)
    error = Signal(str)
    def __init__(self, text, method, num_keywords, custom_stopwords=None):
        super().__init__()
        self.text = text
        self.method = method
        self.num_keywords = num_keywords
        self._is_canceled = False
        self.setObjectName(f"KeywordExtractionThread_{method}")
        self._resources = []        
        from nltk.corpus import stopwords
        try:
            default_stopwords = set(stopwords.words('english'))
        except:
            default_stopwords = STOPWORDS
        default_stopwords = {word.lower().strip() for word in default_stopwords}
        if custom_stopwords:
            custom_stopwords = {word.lower().strip() for word in custom_stopwords if word.strip()}
            self.stopwords = default_stopwords.union(custom_stopwords)
        else:
            self.stopwords = default_stopwords
        self._is_running = True
    def cleanup_resources(self):
        """Clean up resources used by this thread"""
        logger.debug(f"Cleaning up resources for {self.objectName()}")
        try:
            if hasattr(self, 'vectorizer'):
                del self.vectorizer
            if hasattr(self, 'feature_matrix'):
                del self.feature_matrix
            if hasattr(self, 'extractor'):
                del self.extractor
            for resource in self._resources:
                if hasattr(resource, 'close'):
                    try:
                        resource.close()
                    except:
                        pass
                elif hasattr(resource, 'clean_up'):
                    try:
                        resource.clean_up()
                    except:
                        pass
            self._resources.clear()
        except Exception as e:
            logger.error(f"Error cleaning up keyword extraction thread: {e}")        
    def run(self):
        try:
            if self._is_canceled or self.isInterruptionRequested():
                return
            import nltk
            nltk.data.path = ['lib/nltk_data'] + nltk.data.path
            if self.method == 'tfidf':
                keywords = self._extract_tfidf()
            elif self.method == 'yake':
                keywords = self._extract_yake()
            elif self.method == 'rake':
                keywords = self._extract_rake()
            else:
                self.error.emit(f"Unknown keyword extraction method: {self.method}")
                return
            if self._is_canceled or self.isInterruptionRequested():
                return
            self.finished.emit(keywords)
        except Exception as e:
            logger.error(f"Keyword extraction error: {e}")
            if not self._is_canceled and not self.isInterruptionRequested():
                self.error.emit(str(e))
        finally:
            self.cleanup_resources()
    def _extract_tfidf(self):
        """Extract keywords using TF-IDF"""
        try:
            if self._is_canceled or self.isInterruptionRequested():
                return []
            from sklearn.feature_extraction.text import TfidfVectorizer
            vectorizer = TfidfVectorizer(
                max_features=1000,
                stop_words=list(self.stopwords),
                token_pattern=r'(?u)\b[a-zA-Z][a-zA-Z0-9\'-]*[a-zA-Z0-9]\b',
                lowercase=True,
                strip_accents='unicode'
            )
            if self._is_canceled or self.isInterruptionRequested():
                return []
            response = vectorizer.fit_transform([self.text])
            if self._is_canceled or self.isInterruptionRequested():
                return []
            feature_names = vectorizer.get_feature_names_out()
            scores = response.toarray()[0]
            if self._is_canceled or self.isInterruptionRequested():
                return []
            pairs = list(zip(scores, feature_names))
            pairs.sort(reverse=True)
            return [{'keyword': word, 'score': score} 
                    for score, word in pairs[:self.num_keywords]]
        except Exception as e:
            raise Exception(f"TF-IDF extraction failed: {str(e)}")
    def _extract_yake(self):
        """Extract keywords using YAKE"""
        try:
            if self._is_canceled or self.isInterruptionRequested():
                return []
            import yake
            if self._is_canceled or self.isInterruptionRequested():
                return []
            kw_extractor = yake.KeywordExtractor(
                lan="en",
                n=2,
                windowsSize=2,
                top=self.num_keywords,
                stopwords=list(self.stopwords)
            )
            if self._is_canceled or self.isInterruptionRequested():
                return []
            keywords = kw_extractor.extract_keywords(self.text)
            if self._is_canceled or self.isInterruptionRequested():
                return []
            return [{'keyword': kw[0], 'score': 1-kw[1]} for kw in keywords]
        except Exception as e:
            raise Exception(f"YAKE extraction failed: {str(e)}")
    def _extract_rake(self):
        """Extract keywords using RAKE"""
        try:
            if self._is_canceled or self.isInterruptionRequested():
                return []
            from rake_nltk import Rake
            if self._is_canceled or self.isInterruptionRequested():
                return []
            rake = Rake(
                stopwords=list(self.stopwords),
                min_length=1,
                max_length=3
            )
            if self._is_canceled or self.isInterruptionRequested():
                return []
            rake.extract_keywords_from_text(self.text)
            if self._is_canceled or self.isInterruptionRequested():
                return []
            keywords = rake.get_ranked_phrases_with_scores()
            if self._is_canceled or self.isInterruptionRequested():
                return []
            keywords.sort(reverse=True)
            return [{'keyword': kw[1], 'score': kw[0]} 
                    for kw in keywords[:self.num_keywords]]
        except Exception as e:
            raise Exception(f"RAKE extraction failed: {str(e)}")
    def requestInterruption(self):
        """Override method untuk menginisiasi interupsi thread"""
        super().requestInterruption()
        self._is_canceled = True            
class SentimentAnalysisThread(QThread):
    sentiment_analyzed = Signal(dict)
    offline_warning = Signal(str)
    translation_failed = Signal(str)
    translation_progress = Signal(str)
    sentiment_timeline_ready = Signal(object)
    def __init__(self, text_data, sentiment_mode, vader_analyzer, flair_classifier, 
                 flair_classifier_cuslang, textblob_analyzer=None):
        super().__init__()
        self.text_data = text_data
        self.sentiment_mode = sentiment_mode
        self.vader_analyzer = vader_analyzer
        self.flair_classifier = flair_classifier
        self.flair_classifier_cuslang = flair_classifier_cuslang
        self.textblob_analyzer = textblob_analyzer
        self.cached_translation = None
        self.temp_dir = Path(tempfile.gettempdir()) / "textplora_cache"
        self.text_hash = hashlib.md5(text_data.encode()).hexdigest()
        self.temp_file = self.temp_dir / f"trans_{self.text_hash}.txt"
        self.temp_dir.mkdir(parents=True, exist_ok=True)
    def cleanup_resources(self):
        """Clean up resources used by this thread"""
        logger.debug(f"Cleaning up resources for {self.objectName()}")
        try:
            if hasattr(self, '_translator') and self._translator:
                del self._translator
            if hasattr(self, '_cache_file') and self._cache_file:
                try:
                    self._cache_file.close()
                except:
                    pass
            self.flair_classifier = None
            self.flair_classifier_cuslang = None
            self.vader_analyzer = None
            self.textblob_analyzer = None
        except Exception as e:
            logger.error(f"Error cleaning up sentiment thread: {e}")        
    async def _async_translate(self, text):
        try:
            translated = GoogleTranslator(source="auto", target="en").translate(text)
            return translated
        except Exception as e:
            logger.error(f"Translation error: {e}")
            return None
    def translate_text(self, text):
        try:
            loop = asyncio.get_event_loop()
        except RuntimeError:
            loop = asyncio.new_event_loop()
            asyncio.set_event_loop(loop)
        sentences = [s.strip() for s in text.split(".") if s.strip()]
        total_sentences = len(sentences)
        translated = []
        self.translation_progress.emit(f"Preparing to translate {total_sentences} sentences...")
        async def run_translations():
            for idx, sentence in enumerate(sentences, 1):
                if self.isInterruptionRequested():
                    logger.info("Translation interrupted")
                    return None                
                try:
                    self.translation_progress.emit(f"Translating sentence {idx}/{total_sentences}...")
                    result = await self._async_translate(sentence)
                    if result:
                        translated.append(result)
                        progress_pct = (idx / total_sentences) * 100
                        self.translation_progress.emit(f"Translation progress: {progress_pct:.1f}%")
                except Exception as e:
                    logger.error(f"Translation error: {e}")
                    continue
            return ". ".join(translated)
        try:
            result = loop.run_until_complete(run_translations())
            if result:
                self.translation_progress.emit("Translation completed!")
            return result
        except Exception as e:
            logger.error(f"Translation error: {e}")
            self.translation_progress.emit("Translation failed!")
            return None
        finally:
            if loop.is_running():
                loop.close()
    def analyze_textblob_timeline(self, text, use_custom=False):
        """Analyze sentiment timeline using TextBlob"""
        try:
            sentences = []
            if use_custom and self.textblob_analyzer:
                from nltk.tokenize import sent_tokenize
                raw_sentences = sent_tokenize(text)
                for i, sentence in enumerate(raw_sentences):
                    if self.isInterruptionRequested():
                        break
                    analysis = self.textblob_analyzer.analyze(sentence)
                    sentences.append({
                        'index': i,
                        'sentence': sentence,
                        'score': analysis['polarity'],
                        'subjectivity': analysis['subjectivity'],
                        'label': 'POSITIVE' if analysis['polarity'] > 0 
                                else 'NEGATIVE' if analysis['polarity'] < 0 
                                else 'NEUTRAL'
                    })
            else:
                from textblob import TextBlob
                blob = TextBlob(text)
                for i, sentence in enumerate(blob.sentences):
                    if self.isInterruptionRequested():
                        break
                    sentiment = sentence.sentiment
                    sentences.append({
                        'index': i,
                        'sentence': str(sentence),
                        'score': sentiment.polarity,
                        'subjectivity': sentiment.subjectivity,
                        'label': 'POSITIVE' if sentiment.polarity > 0 
                                else 'NEGATIVE' if sentiment.polarity < 0 
                                else 'NEUTRAL'
                    })
            return sentences
        except Exception as e:
            logger.error(f"Error in TextBlob timeline analysis: {str(e)}")
            return None
    def analyze_vader_timeline(self, text, use_custom=True):
        """Analyze sentiment timeline using VADER"""
        try:
            from nltk.tokenize import sent_tokenize
            sentences = sent_tokenize(text)
            timeline = []
            for i, sentence in enumerate(sentences):
                if self.isInterruptionRequested():
                    break
                scores = self.vader_analyzer.polarity_scores(sentence)
                timeline.append({
                    'index': i,
                    'sentence': sentence,
                    'score': scores['compound'],
                    'pos': scores['pos'],
                    'neg': scores['neg'],
                    'neu': scores['neu'],
                    'label': 'POSITIVE' if scores['compound'] >= 0.05 
                            else 'NEGATIVE' if scores['compound'] <= -0.05 
                            else 'NEUTRAL'
                })
            return timeline
        except Exception as e:
            logger.error(f"Error in VADER timeline analysis: {str(e)}")
            return None
    def analyze_flair_timeline(self, text, use_custom=False):
        """Analyze sentiment timeline using Flair"""
        try:
            from nltk.tokenize import sent_tokenize
            from flair.data import Sentence
            classifier = self.flair_classifier_cuslang if use_custom else self.flair_classifier
            if not classifier:
                raise ValueError(f"{'Custom' if use_custom else 'Default'} Flair model not initialized")
            sentences = sent_tokenize(text)
            timeline = []
            for i, sentence_text in enumerate(sentences):
                if self.isInterruptionRequested():
                    break
                sentence = Sentence(sentence_text)
                classifier.predict(sentence)
                label = sentence.labels[0]
                timeline.append({
                    'index': i,
                    'sentence': sentence_text,
                    'score': label.score,
                    'label': label.value,
                    'confidence': label.score
                })
            return timeline
        except Exception as e:
            logger.error(f"Error in Flair timeline analysis: {str(e)}")
            return None                
    def analyze_textblob(self, text, use_custom=False):
        if use_custom and self.textblob_analyzer:
            analysis = self.textblob_analyzer.analyze(text)
            polarity = analysis['polarity']
            subjectivity = analysis['subjectivity']
        else:
            from textblob import TextBlob
            blob = TextBlob(text)
            polarity = blob.sentiment.polarity
            subjectivity = blob.sentiment.subjectivity
        if abs(polarity) < 0.1:
            neutral_score = 1.0
            positive_score = negative_score = 0.0
        else:
            positive_score = (polarity + 1) / 2 if polarity > 0 else 0
            negative_score = (-polarity + 1) / 2 if polarity < 0 else 0
            neutral_score = 1 - (positive_score + negative_score)
        return {
            "positive_score": positive_score,
            "negative_score": negative_score,
            "neutral_score": neutral_score,
            "compound_score": polarity,
            "sentiment_label": "POSITIVE" if polarity > 0 else "NEGATIVE" if polarity < 0 else "NEUTRAL",
            "subjectivity": subjectivity
        }
    def get_cached_translation(self):
        """Get translation from cache if exists"""
        if self.temp_file.exists():
            try:
                with open(self.temp_file, "r", encoding="utf-8") as f:
                    return f.read()
            except Exception as e:
                logger.error(f"Error reading cached translation: {e}")
                return None
        return None
    def save_translation(self, translated_text):
        """Save translation to cache"""
        try:
            with open(self.temp_file, "w", encoding="utf-8") as f:
                f.write(translated_text)
        except Exception as e:
            logger.error(f"Error saving translation: {e}")
    def run(self):
        result = {
            "positive_score": 0,
            "neutral_score": 0,
            "negative_score": 0,
            "compound_score": 0,
            "sentiment_label": "N/A",
            "subjectivity": "N/A",
        }
        try:
            if self.isInterruptionRequested():
                return            
            text_to_analyze = self.text_data
            needs_translation = False
            if self.sentiment_mode in ["VADER", "Flair", "TextBlob"]:
                try:
                    from langdetect import detect
                    detected_lang = detect(self.text_data)
                    needs_translation = detected_lang != "en"
                except Exception as e:
                    logger.error(f"Language detection error: {e}")
                    needs_translation = False
            if self.isInterruptionRequested():
                return                    
            if needs_translation:
                if not is_connected():
                    self.offline_warning.emit("Your text is not in English.\nTranslation required but no internet connection.\nPlease use the mode with custom lexicon/model instead.")
                    return
                cached = self.get_cached_translation()
                if cached:
                    text_to_analyze = cached
                else:
                    if self.isInterruptionRequested():
                        return                    
                    translated_text = self.translate_text(self.text_data)
                    if not translated_text:
                        self.translation_failed.emit("Translation failed - analysis aborted")
                        return
                    if self.isInterruptionRequested():
                        return
                    self.save_translation(translated_text)
                    text_to_analyze = translated_text
            if self.isInterruptionRequested():
                return
            if self.sentiment_mode == "TextBlob":
                result.update(self.analyze_textblob(text_to_analyze))
            elif self.sentiment_mode == "TextBlob (Custom Lexicon)":
                result.update(self.analyze_textblob(self.text_data, True))
            elif self.sentiment_mode == "VADER":
                if self.isInterruptionRequested(): 
                    return                
                if self.vader_analyzer:
                    sentiment = self.vader_analyzer.polarity_scores(text_to_analyze)
                    total = sentiment["pos"] + sentiment["neg"] + sentiment["neu"]
                    positive_score = sentiment["pos"]
                    negative_score = sentiment["neg"] / total if total > 0 else 0
                    neutral_score = sentiment["neu"] / total if total > 0 else 0
                    result.update({
                        "positive_score": positive_score,
                        "negative_score": negative_score,
                        "neutral_score": neutral_score,
                        "compound_score": sentiment["compound"],
                        "sentiment_label": "POSITIVE" if sentiment["compound"] >= 0.05 else "NEGATIVE" if sentiment["compound"] <= -0.05 else "NEUTRAL",
                    })
                else:
                    result["sentiment_label"] = "Error: VADER not initialized"
            elif self.sentiment_mode == "VADER (Custom Lexicon)":
                if self.isInterruptionRequested(): 
                    return
                if self.vader_analyzer:
                    sentiment = self.vader_analyzer.polarity_scores(self.text_data)
                    result.update({
                        "positive_score": sentiment["pos"],
                        "negative_score": sentiment["neg"],
                        "neutral_score": sentiment["neu"],
                        "compound_score": sentiment["compound"],
                        "sentiment_label": "POSITIVE" if sentiment["compound"] >= 0.05 else "NEGATIVE" if sentiment["compound"] <= -0.05 else "NEUTRAL",
                    })
                else:
                    result["sentiment_label"] = "Error: Custom VADER not initialized"
            elif self.sentiment_mode == "Flair":
                if self.isInterruptionRequested(): 
                    return
                if self.flair_classifier:
                    from flair.data import Sentence
                    sent = Sentence(text_to_analyze)
                    self.flair_classifier.predict(sent)
                    if self.isInterruptionRequested(): 
                        return
                    result["sentiment_label"] = sent.labels[0].value
                    result["compound_score"] = sent.labels[0].score
                    for label in self.flair_classifier.label_dictionary.get_items():
                        if label == "POSITIVE":
                            result["positive_score"] = sent.labels[0].score if sent.labels[0].value == "POSITIVE" else 0
                        elif label == "NEGATIVE":
                            result["negative_score"] = sent.labels[0].score if sent.labels[0].value == "NEGATIVE" else 0
                        elif label == "NEUTRAL":
                            result["neutral_score"] = sent.labels[0].score if sent.labels[0].value == "NEUTRAL" else 0
            elif self.sentiment_mode == "Flair (Custom Model)":
                if self.isInterruptionRequested(): 
                    return
                if self.flair_classifier_cuslang:
                    from flair.data import Sentence
                    sent = Sentence(self.text_data)
                    self.flair_classifier_cuslang.predict(sent)
                    if self.isInterruptionRequested(): 
                        return
                    result["sentiment_label"] = sent.labels[0].value
                    result["compound_score"] = sent.labels[0].score
                    for label in self.flair_classifier_cuslang.label_dictionary.get_items():
                        if label == "POSITIVE":
                            result["positive_score"] = sent.labels[0].score if sent.labels[0].value == "POSITIVE" else 0
                        elif label == "NEGATIVE":
                            result["negative_score"] = sent.labels[0].score if sent.labels[0].value == "NEGATIVE" else 0
                        elif label == "NEUTRAL":
                            result["neutral_score"] = sent.labels[0].score if sent.labels[0].value == "NEUTRAL" else 0
            timeline_data = None
            try:
                if self.sentiment_mode == "TextBlob":
                    if needs_translation:
                        timeline_data = self.analyze_textblob_timeline(text_to_analyze, False)
                    else:
                        timeline_data = self.analyze_textblob_timeline(self.text_data, False)
                elif self.sentiment_mode == "TextBlob (Custom Lexicon)":
                    timeline_data = self.analyze_textblob_timeline(self.text_data, True)
                elif self.sentiment_mode == "VADER":
                    if needs_translation:
                        timeline_data = self.analyze_vader_timeline(text_to_analyze, False)
                    else:
                        timeline_data = self.analyze_vader_timeline(self.text_data, False)
                elif self.sentiment_mode == "VADER (Custom Lexicon)":
                    timeline_data = self.analyze_vader_timeline(self.text_data, True)
                elif self.sentiment_mode == "Flair":
                    if needs_translation:
                        timeline_data = self.analyze_flair_timeline(text_to_analyze, False)
                    else:
                        timeline_data = self.analyze_flair_timeline(self.text_data, False)
                elif self.sentiment_mode == "Flair (Custom Model)":
                    timeline_data = self.analyze_flair_timeline(self.text_data, True)
                if timeline_data:
                    result['sentiment_timeline'] = timeline_data
                    self.sentiment_timeline_ready.emit(timeline_data)
            except Exception as e:
                logger.error(f"Error analyzing sentiment timeline: {e}")
            if self.isInterruptionRequested():
                return
            self.sentiment_analyzed.emit(result)
        except Exception as e:
            logger.error(f"Sentiment analysis error: {e}")
            result["sentiment_label"] = f"Error: {str(e)}"
            if not self.isInterruptionRequested():
                self.sentiment_analyzed.emit(result)
        finally:
            self.cleanup_resources()
class WordCloudGeneratorThread(QThread):
    """Thread untuk generate word cloud"""
    progress_updated = Signal(int)
    generation_complete = Signal(object)
    error_occurred = Signal(str)
    def __init__(self, text_data, params, mask_path):
        super().__init__()
        self.text_data = text_data
        self.params = params
        self.mask_path = mask_path
        self._is_canceled = False
        self.setObjectName("WordCloudGeneratorThread")
    def get_nltk_stopwords(self):
        """Get stopwords from NLTK"""
        try:
            from nltk.corpus import stopwords
            return set(stopwords.words('english'))
        except Exception as e:
            logger.debug(f"Failed to load NLTK stopwords: {e}")
            return set()
    def run(self):
        try:
            if self._is_canceled or self.isInterruptionRequested():
                return
            if not self.text_data:
                self.error_occurred.emit("No text data provided")
                return
            self.progress_updated.emit(10)
            font_path = None
            if self.params.get('font_choice') and self.params['font_choice'] != "Default":
                font_map = self.params.get('font_map', {})
                if not font_map:
                    logger.warning("Font map is empty")
                    from matplotlib.font_manager import findfont, FontProperties
                    font_path = findfont(FontProperties())
                else:
                    font_name = self.params['font_choice']
                    font_path = font_map.get(font_name)
                    if not font_path:
                        logger.warning(f"Font path not found in font_map for {font_name}")
                        from matplotlib.font_manager import findfont, FontProperties
                        font_path = findfont(FontProperties())
                    elif not os.path.exists(font_path):
                        logger.warning(f"Font file not found at path: {font_path}")
                        from matplotlib.font_manager import findfont, FontProperties
                        font_path = findfont(FontProperties())
                    else:
                        logger.info(f"Using font: {font_name} from path: {font_path}")
            else:
                from matplotlib.font_manager import findfont, FontProperties
                font_path = findfont(FontProperties())
                logger.info("Using default font")
            self.progress_updated.emit(20)
            mask = None
            if self.mask_path and os.path.exists(self.mask_path):
                try:
                    from PIL import Image
                    import numpy as np
                    mask = np.array(Image.open(self.mask_path))
                except Exception as e:
                    logger.debug(f"Failed to load mask: {e}")
            self.progress_updated.emit(30)
            nltk_stopwords = self.get_nltk_stopwords()
            user_stopwords = self.params.get('stopwords', set())
            combined_stopwords = nltk_stopwords.union(user_stopwords) if nltk_stopwords else user_stopwords
            if self._is_canceled or self.isInterruptionRequested():
                return
            wc_params = {
                'width': 1000,
                'height': 600,
                'background_color': self.params.get('bg_color', 'white'),
                'max_words': int(self.params.get('max_words', 200)),
                'min_font_size': int(self.params.get('min_font_size', 4)),
                'stopwords': combined_stopwords
            }
            if font_path:
                try:
                    from PIL import ImageFont
                    test_font = ImageFont.truetype(font_path, 12)
                    wc_params['font_path'] = font_path
                    logger.info(f"Font verified and set: {font_path}")
                except Exception as e:
                    logger.warning(f"Font verification failed for {font_path}: {e}")
                    from matplotlib.font_manager import findfont, FontProperties
                    wc_params['font_path'] = findfont(FontProperties())
                    logger.info("Fallback to default font after verification failure")
            if mask is not None:
                wc_params['mask'] = mask
            colormap_name = self.params.get('colormap')
            original_colormap = colormap_name
            if not colormap_name or not colormap_name.startswith('custom_'):
                wc_params['colormap'] = colormap_name or 'viridis'
            else:
                custom_palettes = self.params.get('custom_color_palettes', {})
                if colormap_name in custom_palettes:
                    color_list = custom_palettes[colormap_name]
                    if color_list and len(color_list) > 0:
                        import matplotlib.pyplot as plt
                        from matplotlib.colors import ListedColormap
                        try:
                            wc_params['colormap'] = ListedColormap(list(color_list), name=colormap_name)
                        except Exception as e1:
                            logger.debug(f"ListedColormap failed: {e1}")
                            try:
                                from matplotlib.colors import LinearSegmentedColormap
                                wc_params['colormap'] = LinearSegmentedColormap.from_list(colormap_name, list(color_list))
                            except Exception as e2:
                                logger.debug(f"All colormap attempts failed: {e2}")
                                wc_params['colormap'] = 'viridis'
                    else:
                        logger.debug(f"Color list for {colormap_name} is empty or invalid")
                        wc_params['colormap'] = 'viridis'
                else:
                    logger.debug(f"Custom palette {colormap_name} not found in available palettes")
                    wc_params['colormap'] = 'viridis'
            self.progress_updated.emit(50)
            if self._is_canceled or self.isInterruptionRequested():
                return
            from wordcloud import WordCloud
            wordcloud = WordCloud(**wc_params).generate(self.text_data)
            if self._is_canceled or self.isInterruptionRequested():
                return
            wordcloud.original_colormap = original_colormap
            self.progress_updated.emit(90)
            if self._is_canceled or self.isInterruptionRequested():
                return
            self.generation_complete.emit(wordcloud)
        except Exception as e:
            if not self._is_canceled and not self.isInterruptionRequested():
                self.error_occurred.emit(f"Error generating word cloud: {e}")
        finally:
            self.cleanup_resources()
    def cleanup_resources(self):
        """Clean up resources used by this thread"""
        logger.debug(f"Cleaning up resources for {self.objectName()}")
        try:
            self._is_canceled = False
            if hasattr(self, 'fig'):
                import matplotlib.pyplot as plt
                plt.close(self.fig)
            if hasattr(self, 'wordcloud'):
                del self.wordcloud
            if hasattr(self, 'mask_array'):
                del self.mask_array
        except Exception as e:
            logger.error(f"Error cleaning up wordcloud thread: {e}")
    def cancel(self):
        """Cancel the thread operation"""
        self._is_canceled = True
        self.requestInterruption()
        self.progress_updated.emit(0)
class SummarizeThread(QThread):
    """Thread untuk summarize dan memproses teks"""
    summary_ready = Signal(str)
    error_occurred = Signal(str)
    def __init__(self, text_data):
        super().__init__()
        self._text_data = text_data
        self._is_running = False
    def run(self):
        """Summarize dan proses teks di background thread"""
        self._is_running = True
        try:
            summary = self._summarize_text(self._text_data)
            if not self._is_running:
                return
            processed_summary = self._process_text(summary)
            if self._is_running:
                self.summary_ready.emit(processed_summary)
        except Exception as e:
            logger.error(f"Error in summarize thread: {e}")
            self.error_occurred.emit(str(e))
        finally:
            self._is_running = False
    def _summarize_text(self, text):
        """Generate text summary"""
        try:
            from sumy.parsers.plaintext import PlaintextParser
            from sumy.nlp.tokenizers import Tokenizer
            from sumy.summarizers.lsa import LsaSummarizer
            from sumy.nlp.stemmers import Stemmer
            from sumy.utils import get_stop_words
            parser = PlaintextParser.from_string(text, Tokenizer("english"))
            stemmer = Stemmer("english")
            summarizer = LsaSummarizer(stemmer)
            summarizer.stop_words = get_stop_words("english")
            sentence_count = max(3, len(parser.document.sentences) // 3)
            summary = summarizer(parser.document, sentence_count)
            return " ".join([str(sentence) for sentence in summary])
        except Exception as e:
            logger.error(f"Summarization error: {e}")
            raise
    def _process_text(self, text):
        """Process and format summary text"""
        try:
            lines = text.split('\n')
            formatted_lines = []
            for line in lines:
                if not self._is_running:
                    break
                line = line.replace('&', '&amp;')
                line = line.replace('<', '&lt;')
                line = line.replace('>', '&gt;')
                formatted_lines.append(line)
            return '<br>'.join(formatted_lines)
        except Exception as e:
            logger.error(f"Text processing error: {e}")
            raise
    def cancel(self):
        """Cancel the running operation"""
        self._is_running = False
"""CLASS UI & DIALOG"""
class TextViewerDialog(QDialog):
    """Dialog untuk menampilkan teks lengkap"""
    def __init__(self, parent=None):
        super().__init__(parent)
        self.setWindowTitle("Full Text View")
        self.setModal(False)
        self.setMinimumSize(500, 300)
        self.parent = parent
        self.layout = QVBoxLayout(self)
        self.layout.setContentsMargins(10, 10, 10, 10)
        self.layout.setSpacing(10)
        self.text_browser = QTextBrowser(self)
        self.text_browser.setOpenExternalLinks(False)
        self.text_browser.setReadOnly(True)
        self.text_browser.setLineWrapMode(QTextEdit.WidgetWidth)
        self.text_browser.setVerticalScrollBarPolicy(Qt.ScrollBarAsNeeded)
        self.layout.addWidget(self.text_browser)
        button_layout = QHBoxLayout()
        button_layout.addStretch()
        self.close_button = QPushButton("Close", self)
        self.close_button.clicked.connect(self.close)
        button_layout.addWidget(self.close_button)
        self.layout.addLayout(button_layout)
        self._load_thread = None
    def set_text(self, text_data):
        """Load text asynchronously"""
        self.text_browser.clear()
        if self._load_thread and self._load_thread.isRunning():
            self._load_thread.cancel()
            self._load_thread.wait()
        if hasattr(self.parent, 'set_progress'):
            self.parent.set_progress('text_loading', visible=True)
        self._load_thread = TextLoadThread(text_data)
        self._load_thread.text_ready.connect(self._on_text_ready)
        self._load_thread.error_occurred.connect(self._on_error)
        self._load_thread.start()
    def _on_text_ready(self, processed_text):
        """Handle processed text"""
        self.text_browser.setHtml(processed_text)
        if hasattr(self.parent, 'set_progress'):
            self.parent.set_progress('text_loading', visible=False)
    def _on_error(self, error_msg):
        """Handle loading error"""
        self.text_browser.setPlainText(f"Error loading text: {error_msg}")
        if hasattr(self.parent, 'set_progress'):
            self.parent.set_progress('text_loading', visible=False)
    def closeEvent(self, event):
        """Handle dialog closure"""
        if self._load_thread and self._load_thread.isRunning():
            self._load_thread.cancel()
            self._load_thread.wait()
            if hasattr(self.parent, 'set_progress'):
                self.parent.set_progress('text_loading', visible=False)
        event.accept()
class SummaryDialog(QDialog):
    """Dialog untuk menampilkan text summary"""
    def __init__(self, parent=None):
        super().__init__(parent)
        self.setWindowTitle("Text Summary")
        self.setModal(False)
        self.setMinimumSize(500, 300)
        self.parent = parent
        self.layout = QVBoxLayout(self)
        self.layout.setContentsMargins(10, 10, 10, 10)
        self.layout.setSpacing(10)
        self.text_browser = QTextBrowser(self)
        self.text_browser.setOpenExternalLinks(False)
        self.text_browser.setReadOnly(True)
        self.text_browser.setLineWrapMode(QTextEdit.WidgetWidth)
        self.text_browser.setVerticalScrollBarPolicy(Qt.ScrollBarAsNeeded)
        self.layout.addWidget(self.text_browser)
        button_layout = QHBoxLayout()
        button_layout.addStretch()
        self.close_button = QPushButton("Close", self)
        self.close_button.clicked.connect(self.close)
        button_layout.addWidget(self.close_button)
        self.layout.addLayout(button_layout)
        self._load_thread = None
    def set_text(self, text_data):
        """Start summarization process"""
        self.text_browser.clear()
        if self._load_thread and self._load_thread.isRunning():
            self._load_thread.cancel()
            self._load_thread.wait()
        if hasattr(self.parent, 'set_progress'):
            self.parent.set_progress('summarizing', visible=True)
        self._load_thread = SummarizeThread(text_data)
        self._load_thread.summary_ready.connect(self._on_summary_ready)
        self._load_thread.error_occurred.connect(self._on_error)
        self._load_thread.start()
    def _on_summary_ready(self, processed_text):
        """Handle processed summary"""
        self.text_browser.setHtml(processed_text)
        if hasattr(self.parent, 'set_progress'):
            self.parent.set_progress('summarizing', visible=False)
    def _on_error(self, error_msg):
        """Handle loading error"""
        self.text_browser.setPlainText(f"Error generating summary: {error_msg}")
        if hasattr(self.parent, 'set_progress'):
            self.parent.set_progress('summarizing', visible=False)
    def closeEvent(self, event):
        """Handle dialog closure"""
        if self._load_thread and self._load_thread.isRunning():
            self._load_thread.cancel()
            self._load_thread.wait()
            if hasattr(self.parent, 'set_progress'):
                self.parent.set_progress('summarizing', visible=False)
        event.accept()
class TopicAnalysisTab(QWidget):
    def __init__(self, parent=None):
        super().__init__(parent)
        self.text = ""
        self.parent_widget = parent
        self.setFixedHeight(140)
        self.topic_service = TopicAnalysisService(self)
        self.topic_service.analysis_started.connect(self._on_analysis_started)
        self.topic_service.analysis_complete.connect(self._on_analysis_complete)
        self.topic_service.error_occurred.connect(self._on_analysis_error)
        self.topic_service.topic_evolution_ready.connect(self._on_topic_evolution_ready)        
        self.initUI()
    def initUI(self):
        """Initialize the Topic Analysis Tab UI"""
        main_layout = QVBoxLayout()
        main_layout.setContentsMargins(1, 1, 1, 1)
        main_layout.setSpacing(5)
        group_layout = QHBoxLayout()
        topic_group = QGroupBox("Topic Modeling")
        topic_layout = QGridLayout()
        topic_layout.setContentsMargins(10, 5, 10, 5)
        topic_layout.setVerticalSpacing(2)
        topic_layout.addWidget(QLabel("Method:"), 0, 0)
        self.topic_method = QComboBox()
        self.topic_method.addItems(['LDA', 'NMF'])
        self.topic_method.setToolTip(
            "LDA (For discovering topics in large text corpora)\n"
            "NMF (For extracting hidden topics using non-negative features)"
        )
        topic_layout.addWidget(self.topic_method, 0, 1)
        topic_layout.addWidget(QLabel("Number of Topics:"), 1, 0)
        self.num_topics = QSpinBox()
        self.num_topics.setRange(1, 20)
        self.num_topics.setValue(5)
        topic_layout.addWidget(self.num_topics, 1, 1)
        self.analyze_topics_btn = QPushButton("Analyze Topics")
        topic_layout.addWidget(self.analyze_topics_btn, 2, 0, 1, 2)
        topic_group.setLayout(topic_layout)
        group_layout.addWidget(topic_group)
        keyword_group = QGroupBox("Keyword Extraction")
        keyword_layout = QGridLayout()
        keyword_layout.setContentsMargins(10, 5, 10, 5)
        keyword_layout.setVerticalSpacing(2)
        keyword_layout.addWidget(QLabel("Method:"), 0, 0)
        self.keyword_method = QComboBox()
        self.keyword_method.addItems(['TF-IDF', 'YAKE', 'RAKE'])
        self.keyword_method.setToolTip(
            "TF-IDF (For identifying important words based on frequency)\n"
            "RAKE (For extracting keywords from short texts)\n"
            "YAKE (For extracting context-based keywords from documents)"
        )
        keyword_layout.addWidget(self.keyword_method, 0, 1)
        keyword_layout.addWidget(QLabel("Number of Keywords:"), 1, 0)
        self.num_keywords = QSpinBox()
        self.num_keywords.setRange(5, 50)
        self.num_keywords.setValue(10)
        keyword_layout.addWidget(self.num_keywords, 1, 1)
        self.extract_keywords_btn = QPushButton("Extract Keywords")
        keyword_layout.addWidget(self.extract_keywords_btn, 2, 0, 1, 2)
        keyword_group.setLayout(keyword_layout)
        group_layout.addWidget(keyword_group)
        main_layout.addLayout(group_layout)
        self.setLayout(main_layout)
        self.analyze_topics_btn.setEnabled(False)
        self.extract_keywords_btn.setEnabled(False)
        self.analyze_topics_btn.clicked.connect(self.analyze_topics)
        self.extract_keywords_btn.clicked.connect(self.extract_keywords)
        self.analyze_topics_btn.setEnabled(False)
        self.extract_keywords_btn.setEnabled(False)        
    def analyze_topics(self):
        if not self.text.strip():
            QMessageBox.warning(self, "Error", "Please load text first!")
            return
        try:
            if self.parent_widget:
                self.parent_widget.button_manager.save_states()
                self.parent_widget.button_manager.disable_other_buttons('analyze_topics_btn')
                self.parent_widget.set_progress('topics')
            custom_stopwords = None
            if self.parent_widget:
                custom_stopwords = self.parent_widget.import_stopwords()
                self.parent_widget.update_stopwords_for_topic(custom_stopwords)
            method = self.topic_method.currentText().lower()
            num_topics = self.num_topics.value()
            self.topic_service.analyze_topics(
                text=self.text, 
                method=method, 
                num_topics=num_topics,
                custom_stopwords=custom_stopwords
            )
        except Exception as e:
            self._on_analysis_error(str(e))
    def extract_keywords(self):
        if not self.text.strip():
            QMessageBox.warning(self, "Error", "Please load text first!")
            return
        try:
            if self.parent_widget:
                self.parent_widget.button_manager.save_states()
                self.parent_widget.button_manager.disable_other_buttons('extract_keywords_btn')
                self.parent_widget.set_progress('keywords')
            custom_stopwords = None
            if self.parent_widget:
                custom_stopwords = self.parent_widget.import_stopwords()
                self.parent_widget.update_stopwords_for_topic(custom_stopwords)
            method = self.keyword_method.currentText().lower().replace('-', '')
            num_keywords = self.num_keywords.value()
            self.topic_service.extract_keywords(
                text=self.text, 
                method=method, 
                num_keywords=num_keywords,
                custom_stopwords=custom_stopwords
            )
        except Exception as e:
            self._on_analysis_error(str(e))
    def _on_analysis_started(self):
        """Handle when analysis starts"""
        if hasattr(self.parent_widget, 'set_progress'):
            self.parent_widget.set_progress('topics')
        self.analyze_topics_btn.setEnabled(False)
        self.extract_keywords_btn.setEnabled(False)
        if self.parent_widget:
            self.parent_widget.set_progress('topics')
    def _on_analysis_complete(self, result):
        """Handle analysis completion"""
        if self.parent_widget:
            self.parent_widget.set_progress('topics', False)
            self.parent_widget.set_progress('keywords', False)
            self.parent_widget.button_manager.restore_states()
        if result['type'] == 'topics':
            self.topic_evolution_data = getattr(self, 'topic_evolution_data', None)
            self.show_results_dialog(result['method'], result['results'])
        elif result['type'] == 'keywords':
            self.show_keyword_dialog(result['method'], result['results'])
    def _on_topic_evolution_ready(self, evolution_data):
        """Handle topic evolution data"""
        self.topic_evolution_data = evolution_data
    def _on_analysis_error(self, error_msg):
        """Handle analysis errors"""
        if self.parent_widget:
            self.parent_widget.set_progress('topics', False)
            self.parent_widget.set_progress('keywords', False)
            self.parent_widget.button_manager.restore_states()
        QMessageBox.critical(self, "Analysis Error", str(error_msg))
    def set_text(self, text):
        """Set the text to analyze"""
        self.text = text
        has_text = bool(text.strip())
        self.analyze_topics_btn.setEnabled(has_text)
        self.extract_keywords_btn.setEnabled(has_text)
    def show_results_dialog(self, method: str, results: dict):
        """Show topic analysis results dialog"""
        evolution_data = getattr(self, 'topic_evolution_data', None)
        dialog = TopicAnalysisResultDialog(method, results, self, evolution_data)
        dialog.show()
        self.topic_evolution_data = None   
    def show_keyword_dialog(self, method: str, results: list):
        """Show dialog with keyword extraction results"""
        dialog = KeywordExtractionResultDialog(method, results, parent=self)
        dialog.show()
    def setup_results_dialog(self, dialog, html_content):
        """Setup common dialog elements"""
        dialog.setMinimumSize(500, 300)
        text_browser = QTextBrowser()
        text_browser.setHtml(html_content)
        close_btn = QPushButton("Close", dialog)
        close_btn.clicked.connect(dialog.close)
        layout = QVBoxLayout()
        layout.addWidget(text_browser)
        layout.addWidget(close_btn)
        dialog.setLayout(layout)
        dialog.setWindowModality(Qt.NonModal)
        dialog.show()
class LayoutManager:
    """Kelas untuk mengelola pembuatan layout yang umum digunakan"""
    @staticmethod
    def create_form_layout(parent, fields):
        """
        Membuat form layout dengan label dan input fields
        Args:
            parent: Widget parent
            fields: List of tuples (label_text, input_widget, [optional_tooltip])
        Returns:
            QGridLayout: Layout yang sudah dikonfigurasi
        """
        layout = QGridLayout()
        for row, field_info in enumerate(fields):
            label_text = field_info[0]
            input_widget = field_info[1]
            label = QLabel(label_text, parent)
            layout.addWidget(label, row, 0)
            layout.addWidget(input_widget, row, 1)
            if len(field_info) > 2 and field_info[2]:
                input_widget.setToolTip(field_info[2])
                label.setToolTip(field_info[2])
        return layout
class DialogFactory:
    """Factory class untuk membuat dialog yang umum digunakan"""
    @staticmethod
    def create_info_dialog(parent, title, content, modal=True, min_size=(500, 400)):
        """
        Membuat dialog informasi dengan QTextBrowser
        Args:
            parent: Widget parent
            title: Judul dialog
            content: Konten HTML untuk ditampilkan
            modal: Apakah dialog bersifat modal
            min_size: Ukuran minimum dialog (width, height)
        Returns:
            QDialog: Dialog yang sudah dikonfigurasi
        """
        dialog = QDialog(parent, Qt.WindowSystemMenuHint | Qt.WindowCloseButtonHint)
        dialog.setWindowTitle(title)
        dialog.setMinimumSize(*min_size)
        if not modal:
            dialog.setWindowFlags(dialog.windowFlags() | Qt.Window)
            dialog.setWindowModality(Qt.NonModal)
        dialog.setSizeGripEnabled(True)
        layout = QVBoxLayout()
        text_browser = QTextBrowser()
        text_browser.setHtml(content)
        text_browser.setOpenExternalLinks(True)
        text_browser.setReadOnly(True)
        layout.addWidget(text_browser)
        close_button = QPushButton("Close")
        close_button.clicked.connect(dialog.close)
        layout.addWidget(close_button)
        dialog.setLayout(layout)
        return dialog
    @staticmethod
    def create_input_dialog(parent, title, fields, on_accept, modal=True):
        """
        Membuat dialog input dengan multiple fields
        Args:
            parent: Widget parent
            title: Judul dialog
            fields: List of tuples (label_text, input_widget_type, default_value)
            on_accept: Callback function when dialog is accepted
            modal: Apakah dialog bersifat modal
        Returns:
            QDialog: Dialog yang sudah dikonfigurasi
        """
        dialog = QDialog(parent)
        dialog.setWindowTitle(title)
        if modal:
            dialog.setWindowModality(Qt.ApplicationModal)
        else:
            dialog.setWindowModality(Qt.NonModal)
        layout = QVBoxLayout()
        form_layout = QGridLayout()
        input_widgets = {}
        for row, (label_text, widget_type, default_value) in enumerate(fields):
            label = QLabel(label_text)
            form_layout.addWidget(label, row, 0)
            if widget_type == "line_edit":
                widget = QLineEdit(default_value)
            elif widget_type == "combo_box":
                widget = QComboBox()
                widget.addItems(default_value)
            elif widget_type == "spin_box":
                widget = QSpinBox()
                widget.setValue(default_value)
            else:
                widget = QLineEdit(str(default_value))
            form_layout.addWidget(widget, row, 1)
            input_widgets[label_text] = widget
        layout.addLayout(form_layout)
        button_layout = QHBoxLayout()
        ok_button = QPushButton("OK")
        cancel_button = QPushButton("Cancel")
        ok_button.clicked.connect(lambda: on_accept(dialog, input_widgets))
        cancel_button.clicked.connect(dialog.reject)
        button_layout.addWidget(ok_button)
        button_layout.addWidget(cancel_button)
        layout.addLayout(button_layout)
        dialog.setLayout(layout)
        return dialog                
"""CLASS OPERASI FILE"""        
class FileOperations:
    """Kelas untuk mengelola operasi file dengan validasi path"""
    @staticmethod
    def load_text_file(parent, file_path):
        """
        Load text from file with proper validation and enhanced error handling
        Args:
            parent: Parent widget for error messages
            file_path: Path to the file
        Returns:
            str: File content or empty string if error
        """
        try:
            ext = Path(file_path).suffix.lower()
            supported_exts = ['.txt', '.csv', '.md', '.json', '.html', '.xml']
            if ext not in supported_exts:
                logger.warning(f"Unsupported file extension: {ext}")
                reply = QMessageBox.question(
                    parent,
                    "Unsupported File Type",
                    f"The file type '{ext}' may not be supported. Try to open anyway?",
                    QMessageBox.Yes | QMessageBox.No,
                    QMessageBox.No
                )
                if reply == QMessageBox.No:
                    return None
            try:
                content = safe_read_file(file_path, encoding='utf-8')
            except UnicodeDecodeError:
                logger.info(f"UTF-8 decoding failed for {file_path}, trying other encodings")
                encodings = ['latin-1', 'cp1252', 'iso-8859-1']
                for encoding in encodings:
                    try:
                        content = safe_read_file(file_path, encoding=encoding)
                        logger.info(f"Successfully decoded with {encoding}")
                        break
                    except UnicodeDecodeError:
                        continue
                else:
                    raise UnicodeDecodeError("All encodings failed", b"", 0, 1, "Cannot decode file")
            if not content.strip():
                logger.warning(f"Empty file: {file_path}")
                QMessageBox.warning(
                    parent, 
                    "Empty File", 
                    "The file appears to be empty. Please select a file with content."
                )
                return None
            if len(content) > 10 * 1024 * 1024:
                logger.warning(f"Large file detected: {file_path} ({len(content)} bytes)")
                reply = QMessageBox.question(
                    parent,
                    "Large File",
                    "This file is very large and may cause performance issues. Continue loading?",
                    QMessageBox.Yes | QMessageBox.No,
                    QMessageBox.No
                )
                if reply == QMessageBox.No:
                    return None
            parent.statusBar().showMessage(f"Loaded: {os.path.basename(file_path)}")
            return content
        except FileNotFoundError:
            QMessageBox.critical(
                parent, 
                "File Not Found", 
                f"The file '{os.path.basename(file_path)}' could not be found. It may have been moved or deleted."
            )
            logger.error(f"File not found: {file_path}")
            return None
        except PermissionError:
            QMessageBox.critical(
                parent, 
                "Permission Denied", 
                f"You don't have permission to access '{os.path.basename(file_path)}'. Check file permissions."
            )
            logger.error(f"Permission denied: {file_path}")
            return None
        except UnicodeDecodeError as e:
            QMessageBox.critical(
                parent, 
                "Encoding Error", 
                f"Could not decode the file. The file may be binary or use an unsupported encoding."
            )
            logger.error(f"Unicode decode error for {file_path}: {e}")
            return None
        except Exception as e:
            QMessageBox.critical(
                parent, 
                "Error", 
                f"Failed to load file: {str(e)}\n\nTry checking if the file is accessible and not corrupted."
            )
            logger.error(f"Error loading file {file_path}: {e}")
            return None
    @staticmethod
    def save_text_file(parent, file_path, content):
        """
        Save text to file with proper validation
        Args:
            parent: Parent widget for error messages
            file_path: Path to save the file
            content: Content to save
        Returns:
            bool: True if successful, False otherwise
        """
        try:
            safe_write_file(file_path, content)
            parent.statusBar().showMessage(f"Saved: {os.path.basename(file_path)}")
            return True
        except Exception as e:
            QMessageBox.critical(parent, "Error", f"Failed to save file: {str(e)}")
            logger.error(f"Error saving file {file_path}: {e}")
            return False
    @staticmethod
    def open_file_dialog(parent, title="Open File", file_filter=None):
        """
        Show open file dialog with path validation
        Args:
            parent: Parent widget
            title: Dialog title
            file_filter: File filter string
        Returns:
            str: Selected file path or None
        """
        options = QFileDialog.Options()
        if file_filter is None:
            file_filter = (
                "Text Files (*.txt);;CSV Files (*.csv);;Markdown (*.md);;"
                "JSON Files (*.json);;HTML Files (*.html);;XML Files (*.xml);;"
                "All Files (*)"
            )
        file_path, _ = QFileDialog.getOpenFileName(
            parent, title, "", file_filter, options=options
        )
        if file_path:
            try:
                return PathValidator.sanitize_path(file_path)
            except ValueError as e:
                QMessageBox.critical(parent, "Security Error", str(e))
                logger.warning(f"Blocked access to file: {file_path}")
                return None
        return None
    @staticmethod
    def save_file_dialog(parent, title="Save File", file_filter="All Files (*)"):
        """
        Show save file dialog with path validation
        Args:
            parent: Parent widget
            title: Dialog title
            file_filter: File filter string
        Returns:
            str: Selected file path or None
        """
        options = QFileDialog.Options()
        file_path, _ = QFileDialog.getSaveFileName(
            parent, title, "", file_filter, options=options
        )
        if file_path:
            try:
                return PathValidator.sanitize_path(file_path)
            except ValueError as e:
                QMessageBox.critical(parent, "Security Error", str(e))
                logger.warning(f"Blocked access to file: {file_path}")
                return None
        return None
"""CLASS UTAMA APLIKASI"""       
class AppUI(QMainWindow):
    def __init__(self):
        super().__init__()
        self.user_name = os.getlogin()
        self._cached_fonts = {}
        self._cached_models = {}
        self._cached_colormaps = {}
        self.text_data = ""
        self.additional_stopwords = set()
        self.custom_color_palettes = {}
        self._font_loading = False  
        self.flair_classifier = None
        self.flair_classifier_cuslang = None
        self.flair_first_load = True
        self.mask_path = ""
        self.current_figure = None
        self.custom_lexicon_path = None
        self.custom_textblob_lexicon_path = None
        self.textblob_analyzer = None
        self.flair_loader_thread = None
        self.import_thread = None
        self.progress_states = {
            'keywords': {'color': 'red', 'text': 'Extracting keywords'},
            'file': {'color': 'blue', 'text': 'Loading file'},
            'topics': {'color': 'yellow', 'text': 'Analyzing topics'},
            'model': {'color': 'green', 'text': 'Loading model'},
            'wordcloud': {'color': 'purple', 'text': 'Generating word cloud'}
        }
        self.vectorizer = CountVectorizer(max_features=1000, stop_words='english')
        self.tfidf = TfidfVectorizer(max_features=1000, stop_words='english')
        self._about_dialog = None
        self.lda_model = None
        self.nmf_model = None
        self.bert_model = None
        self.font_choice = QComboBox()
        self.font_choice.addItem("Default")
        self.font_choice.addItems(["Arial", "Times New Roman", "Verdana"])
        self.thread_manager = ThreadManager(self)
        self.thread_manager.thread_added.connect(self._handle_thread_added)
        self.thread_manager.thread_removed.connect(self._handle_thread_removed)
        self.thread_manager.all_threads_stopped.connect(self._handle_all_threads_stopped)
        self.thread_manager.error_occurred.connect(self._handle_thread_error)
        self.resource_controller = ResourceController(self)
        self.resource_controller.font_loaded.connect(self._update_font_selector)
        self.resource_controller.palette_updated.connect(self._refresh_color_menu)
        self.resource_controller.resource_error.connect(self._handle_resource_error)
        self.resource_controller.register_resource(
            'model_cache', 
            self._cached_models,
            cleanup_hook=lambda x: x.clear()
        )
        model_cache = self.resource_controller.get_resource(
            'model_cache',
            creator_func=lambda: {}
        )  
        self.wordcloud_service = WordCloudService(self)
        self.wordcloud_service.generation_progress.connect(self._update_wordcloud_progress)
        self.wordcloud_service.generation_complete.connect(self._on_wordcloud_ready)
        self.wordcloud_service.error_occurred.connect(self.handle_wordcloud_error)
        self.sentiment_service = SentimentAnalysisService(self)
        self.sentiment_service.analysis_started.connect(self._on_sentiment_analysis_started)
        self.sentiment_service.analysis_complete.connect(self.on_sentiment_analyzed)
        self.sentiment_service.error_occurred.connect(self._handle_sentiment_error)
        self.text_analysis_service = TextAnalysisService(self)   
        self.text_analysis_service.analysis_started.connect(self._on_text_analysis_started)
        self.text_analysis_service.analysis_complete.connect(self._on_text_analysis_complete)
        self.text_analysis_service.error_occurred.connect(self._on_text_analysis_error)
        self.text_viewer_service = TextViewerService(self)
        self.text_viewer_service.error_occurred.connect(self._handle_text_viewer_error)
        self.button_manager = ButtonStateManager(self)
        self.resource_controller.initialize()
        self._exit_requested = False
        self._cleanup_in_progress = False
        self._cleanup_complete = False
        self._cleanup_state = {
            'in_progress': False,
            'cache_cleared': False,
            'threads_stopped': False,
            'resources_cleaned': False,
            'temp_files_cleaned': False,
            'monitoring_stopped': False,
            'ui_cleaned': False
        }
        self.setWindowIcon(QIcon(str(ICON_PATH)))
        QThreadPool.globalInstance().setMaxThreadCount(3)
        self.register_cached_functions()
        self.setup_statusbar()
        self._init_basic()
        self._setup_lazy_loading()
        self.init_analyzers()
        self.initUI()
        self.setup_timers()
        self._init_imports()
        from PySide6.QtCore import QCoreApplication
        QCoreApplication.instance().aboutToQuit.connect(self.cleanup)
    def register_cached_functions(self):
        """Register all application caches with the cache manager"""
        if not hasattr(self, 'cache_manager'):
            return
        cached_functions = [
            (self, 'get_wordcloud'),
            (self, 'get_most_frequent_words'),
            (None, load_stopwords)
        ]
        for obj, func_name in cached_functions:
            func = getattr(obj, func_name) if obj else func_name
            if hasattr(func, 'cache_clear'):
                self.cache_manager.register_cached_function(func)
        if hasattr(self, 'token_opts'):
            self.cache_manager.register_cached_property(self, 'token_opts')
        cache_mappings = {
            'cached_models': '_cached_models',
            'cached_fonts': '_cached_fonts',
            'cached_colormaps': '_cached_colormaps',
            'translation': '_translation_cache',
            'lazy_loader': LazyLoader._cache
        }
        for cache_id, attr_name in cache_mappings.items():
            if hasattr(self, attr_name):
                self.cache_manager.register(cache_id, getattr(self, attr_name))
        model_attrs = ['_flair_model', '_vader_analyzer']
        for attr in model_attrs:
            if hasattr(self, attr):
                self.cache_manager.register_model_cache(getattr(self, attr))
    def _handle_translation_progress(self, message):
        """Handle translation progress updates"""
        if message:
            self.statusBar().showMessage(message, 3000)
    def _on_sentiment_analysis_complete(self, result):
        """Handler when sentiment analysis is complete"""
        self.set_progress('sentiment', False)
        self.button_manager.restore_states()
        self.statusBar().clearMessage()
        if result and "sentiment_label" in result:
            pass
    def _handle_sentiment_error(self, error_msg):
        """Handler when sentiment analysis fails"""
        self.set_progress('sentiment', False)
        self.button_manager.restore_states()
        self.statusBar().clearMessage()
        QMessageBox.critical(self, "Error", error_msg)               
    def load_matplotlib_fonts_optimized(self):
        """Load matplotlib fonts selalu dari sistem, tanpa caching"""
        if self._font_loading:
            logger.debug("Font loading already in progress")
            return
        self._font_loading = True
        self.statusBar().showMessage("Loading fonts...", 0)
        QTimer.singleShot(100, self._load_fonts_from_system)
    def _load_fonts_from_system(self):
        """Load fonts from system using FontLoaderThread"""
        try:
            if hasattr(self, '_font_loader_thread') and self._font_loader_thread and self._font_loader_thread.isRunning():
                try:
                    self._font_loader_thread.stop()
                    self._font_loader_thread.wait(500)
                except Exception as e:
                    logger.error(f"Error stopping previous font loader thread: {e}")
            self._font_loader_thread = FontLoaderThread(self, None)
            self._font_loader_thread.loading_complete.connect(self._on_system_fonts_loaded)
            self._font_loader_thread.loading_progress.connect(self._on_font_loading_progress)
            self._font_loader_thread.error_occurred.connect(self._on_font_loading_error)
            self._font_loader_thread.finished.connect(self._on_font_loader_finished)
            self._font_loader_thread.start()
            logger.debug("Font loader thread started")
        except Exception as e:
            logger.error(f"Error starting font loader: {e}")
    def _on_font_loading_progress(self, message, percentage):
        """Update UI dengan progress font loading - hanya status tanpa progress bar"""
        self.statusBar().showMessage(message, 3000)
    def _on_font_loading_error(self, error_msg):
        """Handler untuk error font loading"""
        logger.error(f"Font loading error: {error_msg}")
        self.statusBar().showMessage(f"Error: {error_msg}", 3000)
    def _on_font_loader_finished(self):
        """Handler untuk thread font loader selesai"""
        logger.debug("Font loader thread finished")
        self._font_loading = False
    def _on_system_fonts_loaded(self, fonts_dict):
        """Handler untuk fonts yang berhasil dimuat dari sistem"""
        try:
            self._cached_fonts = fonts_dict
            self._update_font_ui(fonts_dict)
            if hasattr(self, 'resource_controller'):
                self.resource_controller._cached_fonts = fonts_dict.copy()
            if hasattr(self, 'wordcloud_service'):
                self.wordcloud_service.set_font_map(fonts_dict.copy())
            msg = f'Loaded {len(fonts_dict)} fonts from system'    
            logger.info(msg)
            self.statusBar().showMessage(msg, 3000)            
        except Exception as e:
            logger.error(f"Error processing loaded fonts: {e}")
    def _update_font_ui(self, fonts_dict):
        """Update UI komponen yang memerlukan font"""
        try:
            if not hasattr(self, 'font_choice'):
                logger.warning("Font choice dropdown not initialized")
                return
            current_text = self.font_choice.currentText()
            self.font_choice.blockSignals(True)
            self.font_choice.clear()
            self.font_choice.addItem("Default")
            families = {}
            for display_name in fonts_dict.keys():
                if display_name.startswith('__'):
                    continue
                parts = display_name.split()
                family = parts[0] if len(parts) == 1 else ' '.join(parts[:-1] if any(weight in parts[-1] for weight in 
                                                                                ["Bold", "Light", "Medium", "Black", "Thin", "Regular", "Italic"]) else parts)
                if family not in families:
                    families[family] = []
                families[family].append(display_name)
            sorted_families = sorted(families.keys())
            for family in sorted_families:
                base_font_added = False
                variants = sorted(families[family])
                for name in variants:
                    self.font_choice.addItem(name)
            if current_text == "Default":
                self.font_choice.setCurrentIndex(0)
            else:
                index = self.font_choice.findText(current_text)
                if index >= 0:
                    self.font_choice.setCurrentIndex(index)
                else:
                    self.font_choice.setCurrentIndex(0)
            self.font_choice.blockSignals(False)
            logger.debug(f"Updated UI with {len(fonts_dict)} fonts")
        except Exception as e:
            logger.error(f"Error updating font UI: {str(e)}")
            try:
                self.font_choice.blockSignals(True)
                self.font_choice.clear()
                self.font_choice.addItem("Default")
                self.font_choice.addItems(["Arial", "Times New Roman", "Verdana"])
                self.font_choice.blockSignals(False)
            except:
                pass
    def setup_statusbar(self):
        """Setup statusbar dengan indikator koneksi"""
        statusbar = self.statusBar()
        statusbar.setStyleSheet("""
            QStatusBar {
                border-top: 1px solid #c0c0c0;
                background: transparent;
                padding: 2px;
                font-size: 9pt;
            }
        """)
        self.connection_indicator = QLabel()
        self.connection_indicator.setFixedSize(16, 16)
        self.update_connection_status()
        self.connection_timer = QTimer(self)
        self.connection_timer.timeout.connect(self.update_connection_status)
        self.connection_timer.start(10000)
        statusbar.addPermanentWidget(self.connection_indicator)
        self.statusBar().setSizeGripEnabled(False)
    def update_connection_status(self):
        """Update indikator status koneksi"""
        if is_connected():
            self.connection_indicator.setPixmap(self.get_connection_icon("connected"))
            self.connection_indicator.setToolTip("Internet Connected")
        else:
            self.connection_indicator.setPixmap(self.get_connection_icon("disconnected"))
            self.connection_indicator.setToolTip("No Internet Connection")
    def get_connection_icon(self, status):
        """Generate ikon koneksi menggunakan QPainter"""
        pixmap = QPixmap(16, 16)
        pixmap.fill(Qt.transparent)
        try:
            painter = QPainter(pixmap)
            painter.setRenderHint(QPainter.Antialiasing)
            if status == "connected":
                painter.setPen(Qt.NoPen)
                painter.setBrush(QColor("#2ecc71"))
                painter.drawEllipse(2, 2, 12, 12)
            else:
                painter.setPen(Qt.NoPen)
                painter.setBrush(QColor("#e74c3c"))
                painter.drawEllipse(2, 2, 12, 12)
        finally:
            painter.end()
        return pixmap
    @cached_property 
    def token_opts(self):
        """Cached token configuration for vectorizers"""
        return {
            'token_pattern': r'(?u)\b[a-zA-Z][a-zA-Z0-9\'-]*[a-zA-Z0-9]\b',
            'lowercase': True,
            'strip_accents': 'unicode',
            'max_features': 1000,
            'min_df': 1,
            'max_df': 0.95
        }
    def _init_basic(self):
        """Initialize basic components"""
        self.text_data = ""
        self.additional_stopwords = set()
        self.custom_color_palettes = {}
    def _setup_lazy_loading(self):
        """Configure lazy loading for heavy dependencies"""
        self._vader = None
        self._flair = None
        self._textblob = None
    @cached_property
    def vader_analyzer(self):
        """Lazy load VADER analyzer"""
        if not self._vader:
            VaderAnalyzer = self.lazy_loader.load('vaderSentiment.vaderSentiment', 'SentimentIntensityAnalyzer')
            if VaderAnalyzer:
                self._vader = VaderAnalyzer()
        return self._vader
    @lru_cache(maxsize=32)
    def get_wordcloud(self, **kwargs):
        """Get cached wordcloud instance"""
        return OptimizedWordCloud(**kwargs)
    def _init_imports(self):
        """Pre-initialize heavy imports in background with proper cleanup"""
        if self.import_thread is not None:
            return
        self.import_thread = ImportThread()
        self.import_thread.finished.connect(self._cleanup_import_thread)
        self.add_managed_thread(self.import_thread)
        self.import_thread.start()
    def _cleanup_import_thread(self):
        """Cleanup import thread when finished"""
        if self.import_thread:
            try:
                self.import_thread.stop()
                self.remove_managed_thread(self.import_thread)
                self.import_thread.deleteLater()
            finally:
                self.import_thread = None
    def add_managed_thread(self, thread):
        """Add thread to managed threads list"""
        if hasattr(self, 'thread_manager'):
            self.thread_manager.add_thread(thread)
        else:
            self.threads_mutex.lock()
            try:
                self.active_threads.append(thread)
            finally:
                self.threads_mutex.unlock()
    def remove_managed_thread(self, thread):
        """Remove thread from managed threads list"""
        if hasattr(self, 'thread_manager'):
            self.thread_manager.remove_thread(thread)
        else:
            self.threads_mutex.lock()
            try:
                if thread in self.active_threads:
                    self.active_threads.remove(thread)
            finally:
                self.threads_mutex.unlock()
    def init_analyzers(self):
        """Initialize analyzers with caching"""
        from functools import lru_cache
        @lru_cache(maxsize=32)
        def get_vader_analyzer():
            from vaderSentiment.vaderSentiment import SentimentIntensityAnalyzer
            return SentimentIntensityAnalyzer()
        self.vader_analyzer = get_vader_analyzer()
        self.sentiment_mode = "TextBlob"
    def set_progress(self, state_key, visible=True, progress=None):
        """Set progress bar visibility and value"""
        if hasattr(self, 'unified_progress_bar'):
            self.unified_progress_bar.setVisible(visible)
            if progress is not None:
                self.unified_progress_bar.setValue(progress)
            if visible:
                self.unified_progress_bar.setValue(0)
        state = self.progress_states.get(state_key)
        if not state:
            return
        if visible:
            if state_key == 'file':
                pass
            elif state_key == 'model':
                self.sentiment_button.setEnabled(False)
            elif state_key == 'wordcloud':
                self.generate_wordcloud_button.setEnabled(False)
            elif state_key == 'topics':
                self.topic_tab.analyze_topics_btn.setEnabled(False)
            elif state_key == 'keywords':
                self.topic_tab.extract_keywords_btn.setEnabled(False)
            self.unified_progress_bar.setVisible(True)
            self.set_progress_style(state['color'])
            if progress is not None:
                self.progress_timer.stop()
                self.unified_progress_bar.setValue(int(progress))
            else:
                if not self.progress_timer.isActive():
                    self.unified_progress_bar.setValue(0)
                    self.progress_timer.start()
        else:
            has_text = bool(self.text_data)
            if state_key == 'file':
                self.load_file_button.setEnabled(True)
            elif state_key == 'model':
                self.sentiment_button.setEnabled(has_text)
            elif state_key == 'wordcloud':
                self.generate_wordcloud_button.setEnabled(has_text)
            elif state_key == 'topics':
                self.topic_tab.analyze_topics_btn.setEnabled(has_text)
            elif state_key == 'keywords':
                self.topic_tab.extract_keywords_btn.setEnabled(has_text)
            self.progress_timer.stop()
            self.unified_progress_bar.setValue(0)
            self.unified_progress_bar.setVisible(False)
        QApplication.processEvents()
    def set_progress_style(self, color):
        """Set progress bar color and style for indeterminate animation"""
        chunk_color = {
            'red': 'rgba(215, 0, 0, 0.2)',
            'blue': 'rgba(0, 120, 215, 0.2)', 
            'yellow': 'rgba(255, 255, 0, 0.2)',
            'green': 'rgba(0, 215, 0, 0.2)',
            'purple': 'rgba(160, 32, 240, 0.2)'
        }.get(color, 'rgba(0, 120, 215, 0.2)')
        style = f"""
            QProgressBar {{
                border: none;
                background: transparent;
                padding: 0px;
                margin: 0px;
            }}
            QProgressBar::chunk {{
                width: 1px;
                background: {chunk_color};
            }}
        """
        self.unified_progress_bar.setStyleSheet(style)
    def setup_timers(self):
        """Setup and optimize timers"""
        self.cleanup_timer = QTimer(self)
        self.cleanup_timer.timeout.connect(self.thread_manager.cleanup_finished_threads)
        self.cleanup_timer.start(30000)
    def initUI(self):
        """Use this if not include user name :
            WIN_TITLE = f"V0NHZW4gKyBUZXh0IEFuYWx5dGljcyAodjEuNikK"
            win_title = base64.b64decode(WIN_TITLE.encode()).decode()
        """
        WIN_TITLE = base64.b64decode("VGV4dHBsb3JhICh2MS42KSBbe31d".encode()).decode().format(self.user_name)
        self.setWindowTitle(WIN_TITLE)
        self.setFixedSize(550, 870)
        layout = QVBoxLayout()
        file_group = QGroupBox("File Input")
        file_layout = QGridLayout()
        filename_container = QFrame()
        filename_container.setStyleSheet("""
            QFrame { 
                border: opx solid #c0c0c0; 
                background: #ffffff;
                padding: 0px;
                margin: 0px;
            }
        """)
        filename_container.setFixedHeight(25)
        filename_layout = QHBoxLayout(filename_container)
        filename_layout.setContentsMargins(0, 0, 0, 0)
        filename_layout.setSpacing(0)
        text_container = QFrame()
        text_container.setLayout(QHBoxLayout())
        text_container.layout().setContentsMargins(5, 0, 5, 0)
        text_container.layout().setSpacing(0)
        self.file_name = QLineEdit()
        self.file_name.setReadOnly(True)
        self.file_name.setFrame(False)
        self.file_name.setPlaceholderText("No file selected")
        self.file_name.setStyleSheet("""
            QLineEdit {
                border: none;
                background: transparent;
                padding: 0px;
            }
        """)
        text_container.layout().addWidget(self.file_name)
        filename_layout.addWidget(text_container)
        self.unified_progress_bar = QProgressBar(filename_container)
        self.unified_progress_bar.setRange(0, 100)
        self.unified_progress_bar.setValue(0)
        self.unified_progress_bar.setTextVisible(False)
        self.unified_progress_bar.setMinimumWidth(filename_container.width())
        self.unified_progress_bar.setStyleSheet("""
            QProgressBar {
                border: none;
                background: transparent;
                padding: 0px;
                margin: 0px;
            }
            QProgressBar::chunk {
                width: 1px;
                background: rgba(0, 120, 215, 0.2);
            }
        """)
        self.unified_progress_bar.setFixedWidth(filename_container.width())
        self.unified_progress_bar.setFixedHeight(filename_container.height())
        self.unified_progress_bar.setGeometry(0, 0, filename_container.width(), filename_container.height())
        self.file_name.raise_()
        file_layout.addWidget(filename_container, 0, 0, 1, 6)
        self.load_file_button = QPushButton("Load Text File", self)
        self.load_file_button.clicked.connect(self.open_file)
        self.load_file_button.setToolTip(
            "Upload a text file for word cloud generation and sentiment analysis.\n"
            "Supports TXT, CSV, XLS/XLSX, PDF, DOC/DOCX.\n"
            "Ensure your text is well-formatted for better results."
        )
        file_layout.addWidget(self.load_file_button, 1, 0, 1, 2)
        self.view_fulltext_button = QPushButton("View Full Text", self)
        self.view_fulltext_button.clicked.connect(self.view_full_text)
        self.view_fulltext_button.setEnabled(False)
        self.view_fulltext_button.setToolTip(
            "Click to view the full text content in a separate window.\n"
            "Allows you to inspect the complete text before generating the word cloud.\n"
            "Useful for verifying text input and checking formatting."
        )
        file_layout.addWidget(self.view_fulltext_button, 1, 2, 1, 2)
        self.summarize_button = QPushButton("Summarize", self)
        self.summarize_button.clicked.connect(self.summarize_text)
        self.summarize_button.setEnabled(False)
        self.summarize_button.setToolTip(
            "Summarize the text content using the LSA (Latent Semantic Analysis) method.\n"
            "This extractive summarization technique selects the most important sentences,\n"
            "helping to generate concise word clouds and improve sentiment analysis."
        )
        file_layout.addWidget(self.summarize_button, 1, 4, 1, 2)
        file_group.setLayout(file_layout)
        layout.addWidget(file_group)
        wordcloud_group = QGroupBox("Word Cloud Generation")
        wordcloud_layout = QGridLayout()
        stopwords_container = QHBoxLayout()
        self.stopword_entry = QTextEdit(self)
        self.stopword_entry.setFixedHeight(50)
        self.stopword_entry.setPlaceholderText("Enter stopwords, separated by spaces or new lines (optional)")
        self.stopword_entry.setToolTip(
            "Enter stopwords (words to be excluded) separated by spaces or new lines.\n"
            "The words you enter will be ignored in the word cloud.\n"
            "Use custom stopwords to refine the visualization and focus on meaningful words."
        )        
        stopwords_container.addWidget(self.stopword_entry)     
        self.load_stopwords_btn = QPushButton("Load", self)
        self.load_stopwords_btn.setFixedSize(50, 50)
        self.load_stopwords_btn.clicked.connect(self.load_stopwords_file)
        stopwords_container.addWidget(self.load_stopwords_btn)
        wordcloud_layout.addLayout(stopwords_container, 0, 0, 1, 6)
        self.color_theme_label = QLabel("Color Theme:", self)
        wordcloud_layout.addWidget(self.color_theme_label, 1, 0, 1, 2)
        self.color_theme = QComboBox(self)
        self.color_theme.addItem("Loading colormaps...")
        self.color_theme.setEnabled(False)
        self.resource_controller.colormap_loading_started.connect(self._on_colormap_loading_started)
        self.resource_controller.colormap_loading_progress.connect(self._on_colormap_loading_progress)
        self.resource_controller.colormap_loading_finished.connect(self._on_colormap_loading_finished)
        self.color_theme.setToolTip(
            "Choose a color palette for the word cloud.\n"
            "Darker themes work well with light backgrounds, and vice-versa."
        )        
        wordcloud_layout.addWidget(self.color_theme, 1, 2, 1, 3)
        self.custom_palette_button = QPushButton("Custom", self)
        self.custom_palette_button.clicked.connect(self.create_custom_palette)
        wordcloud_layout.addWidget(self.custom_palette_button, 1, 5, 1, 1)
        self.bg_color_label = QLabel("Background Color:", self)
        wordcloud_layout.addWidget(self.bg_color_label, 2, 0, 1, 2)
        self.bg_color = QComboBox(self)
        self.bg_color.addItems(["white", "black", "gray", "blue", "red", "yellow"])
        self.bg_color.setToolTip(
            "Select the background color for the word cloud.\n"
            "Use contrast for better visibility.\n"
            "White or black backgrounds usually work best."
        )        
        wordcloud_layout.addWidget(self.bg_color, 2, 2, 1, 3)
        self.custom_bg_color_button = QPushButton("Custom", self)
        self.custom_bg_color_button.clicked.connect(self.select_custom_bg_color)
        wordcloud_layout.addWidget(self.custom_bg_color_button, 2, 5, 1, 1)
        self.title_label = QLabel("WordCloud Title:", self)
        wordcloud_layout.addWidget(self.title_label, 3, 0, 1, 2)
        self.title_entry = QLineEdit(self)
        self.title_entry.setPlaceholderText("Enter title (optional)")
        self.title_entry.setToolTip(
            "Enter a title for your word cloud (optional).\n"
            "This title will be displayed above the word cloud.\n"
            "Leave blank if no title is needed."
        )        
        wordcloud_layout.addWidget(self.title_entry, 3, 2, 1, 2)
        self.title_font_size = QSpinBox(self)
        self.title_font_size.setRange(8, 72)
        self.title_font_size.setValue(14)
        self.title_font_size.setToolTip(
            "Set the font size for the word cloud title.\n"
            "Larger values make the title more prominent.\n"
            "Recommended: 14-24 px for a balanced look.\n"
            "Too large titles may overlap with the word cloud."
        )        
        wordcloud_layout.addWidget(self.title_font_size, 3, 4, 1, 1)
        self.title_position = QComboBox(self)
        self.title_position.addItems(["Left", "Center", "Right"])
        self.title_position.setCurrentText("Center")
        self.title_position.setToolTip(
            "Choose where to display the title relative to the word cloud.\n"
            "Positioning affects the overall layout and readability.\n"
            "Recomended: Center"
        )        
        wordcloud_layout.addWidget(self.title_position, 3, 5, 1, 1)
        self.font_choice_label = QLabel("Font Choice:", self)
        wordcloud_layout.addWidget(self.font_choice_label, 4, 0, 1, 2)
        self.font_choice = QComboBox(self)
        self.font_choice.addItem("Default")
        self.font_choice.setToolTip(
            "Choose a font style for the word cloud.\n"
            "Different fonts affect readability and aesthetics.\n"
            "Sans-serif fonts are recommended for clarity."
        )        
        wordcloud_layout.addWidget(self.font_choice, 4, 2, 1, 4)
        QTimer.singleShot(100, self.load_matplotlib_fonts)
        self.min_font_size_label = QLabel("Minimum Font Size:", self)
        wordcloud_layout.addWidget(self.min_font_size_label, 5, 0, 1, 2)
        self.min_font_size_input = QSpinBox(self)
        self.min_font_size_input.setValue(11)
        self.min_font_size_input.setToolTip(
            "Set the smallest font size for words in the word cloud.\n"
            "Prevents low-frequency words from becoming too small to read.\n"
            "Recommended value: 10-12 px for readability."
        )        
        wordcloud_layout.addWidget(self.min_font_size_input, 5, 2, 1, 1)
        self.max_words_label = QLabel("Maximum Words:", self)
        wordcloud_layout.addWidget(self.max_words_label, 5, 3, 1, 2, Qt.AlignRight)
        self.max_words_input = QSpinBox(self)
        self.max_words_input.setMaximum(10000)
        self.max_words_input.setValue(200)
        self.max_words_input.setToolTip(
            "Set the maximum number of words displayed in the word cloud.\n"
            "Higher values provide more detail but may reduce clarity.\n"
            "Recommended: 100-200 words for balanced visualization"
        )        
        wordcloud_layout.addWidget(self.max_words_input, 5, 5, 1, 1)
        self.mask_label = QLabel("Mask Image:", self)
        wordcloud_layout.addWidget(self.mask_label, 6, 0, 1, 2)
        self.mask_path_label = QLineEdit("default (rectangle)", self)
        self.mask_path_label.setReadOnly(True)
        wordcloud_layout.addWidget(self.mask_path_label, 6, 2, 1, 4)
        self.mask_button = QPushButton("Load Mask Image", self)
        self.mask_button.clicked.connect(self.choose_mask)
        self.mask_button.setToolTip(
            "Upload an image to shape the word cloud (PNG/JPG/BMP).\n"
            "White areas will be ignored, and words will fill the dark areas.\n"
            "Use simple shapes for best results."
        )        
        wordcloud_layout.addWidget(self.mask_button, 7, 2, 1, 2)
        self.reset_mask_button = QPushButton("Remove Mask Image", self)
        self.reset_mask_button.clicked.connect(self.reset_mask)
        self.reset_mask_button.setToolTip(
            "Remove the selected mask image and revert to the default shape.\n"
            "The word cloud will be displayed in a rectangular format.\n"
            "Use this if you no longer want a custom shape for the word cloud."
        )        
        wordcloud_layout.addWidget(self.reset_mask_button, 7, 4, 1, 2)
        self.generate_wordcloud_button = QPushButton("Generate Word Cloud", self)
        self.generate_wordcloud_button.clicked.connect(self.generate_wordcloud)
        self.generate_wordcloud_button.setEnabled(False)
        wordcloud_layout.addWidget(self.generate_wordcloud_button, 8, 0, 1, 6)
        self.text_stats_button = QPushButton("View Text Statistics", self)
        self.text_stats_button.clicked.connect(self.view_text_stats)
        self.text_stats_button.setEnabled(False)
        wordcloud_layout.addWidget(self.text_stats_button, 9, 0, 1, 6)
        wordcloud_group.setLayout(wordcloud_layout)
        layout.addWidget(wordcloud_group)
        sentiment_group = QGroupBox("Sentiment Analysis")
        sentiment_layout = QGridLayout()
        self.sentiment_mode_label = QLabel("Analysis Mode:", self)
        sentiment_layout.addWidget(self.sentiment_mode_label, 0, 0, 1, 2)
        self.sentiment_mode_combo = QComboBox(self)
        self.sentiment_mode_combo.addItems(AppConstants.SENTIMENT_MODES)
        self.sentiment_mode_combo.setToolTip(
            "TextBlob (Best for formal texts like news articles and reports.)\n"
            "VADER (Best for social media, tweets, and informal texts with slang/emojis.)\n"
            "Flair (Best for long-form content and complex sentiment analysis using deep learning.)"
        )
        self.sentiment_mode_combo.currentTextChanged.connect(self.change_sentiment_mode)
        sentiment_layout.addWidget(self.sentiment_mode_combo, 0, 2, 1, 3)
        self.sentiment_mode_info_button = QPushButton("Info", self)
        self.sentiment_mode_info_button.clicked.connect(self.show_sentiment_mode_info)
        sentiment_layout.addWidget(self.sentiment_mode_info_button, 0, 5, 1, 1)
        self.custom_lexicon_button = QPushButton("Load Lexicon", self)
        self.custom_lexicon_button.clicked.connect(self.load_custom_lexicon)
        self.custom_lexicon_button.setEnabled(False)
        sentiment_layout.addWidget(self.custom_lexicon_button, 1, 0, 1, 3)
        self.custom_model_button = QPushButton("Load Model", self)
        self.custom_model_button.clicked.connect(self.load_custom_model)
        self.custom_model_button.setEnabled(False)
        sentiment_layout.addWidget(self.custom_model_button, 1, 3, 1, 3)
        self.sentiment_button = QPushButton("Analyze Sentiment", self)
        self.sentiment_button.clicked.connect(self.analyze_sentiment)
        self.sentiment_button.setEnabled(False)
        sentiment_layout.addWidget(self.sentiment_button, 2, 0, 1, 6)
        sentiment_group.setLayout(sentiment_layout)
        layout.addWidget(sentiment_group)
        self.topic_tab = TopicAnalysisTab(parent=self)
        layout.addWidget(self.topic_tab)
        bottom_group = QGroupBox()
        bottom_layout = QVBoxLayout()
        button_layout = QHBoxLayout()
        self.about_button = QPushButton("About", self)
        self.about_button.clicked.connect(self.show_about)
        self.about_button.setFixedWidth(100)
        button_layout.addWidget(self.about_button)
        button_layout.addStretch()
        self.panic_button = QPushButton("STOP", self)
        self.panic_button.setStyleSheet("background-color: #ff6666;")
        self.panic_button.clicked.connect(self.stop_all_processes)
        self.panic_button.setFixedWidth(100)
        button_layout.addWidget(self.panic_button)
        self.quit_button = QPushButton("Exit", self)
        self.quit_button.clicked.connect(self.close)
        self.quit_button.setFixedWidth(100)
        button_layout.addWidget(self.quit_button)
        bottom_layout.addLayout(button_layout)
        bottom_group.setLayout(bottom_layout)
        layout.addWidget(bottom_group)
        self.progress_timer = QTimer()
        self.progress_timer.timeout.connect(self._update_progress_animation)
        self.progress_timer.setInterval(20)
        container = QWidget()
        container.setLayout(layout)
        self.setCentralWidget(container)
    def _update_font_selector(self, font_name):
        """Update font selector saat font baru di-load"""
        if hasattr(self, 'font_choice'):
            current_text = self.font_choice.currentText()
            self.font_choice.addItem(font_name)
            self.font_choice.setCurrentText(current_text)
    def _refresh_color_menu(self, palette_name):
        """Refresh color menu saat palette baru ditambahkan"""
        if hasattr(self, 'color_theme'):
            current_text = self.color_theme.currentText()
            self.color_theme.clear()
            self.color_theme.addItems(self.resource_controller.res_get_palette_names())
            if current_text in self.resource_controller.res_get_palette_names():
                self.color_theme.setCurrentText(current_text)
            else:
                self.color_theme.setCurrentText(palette_name)
    def _handle_resource_error(self, error_type, message):
        """Handle error dari resource controller"""
        logger.error(f"Resource error ({error_type}): {message}")
        if error_type in ["model"]:
            QMessageBox.warning(self, "Resource Error", message)
    def get_available_fonts(self):
        """Wrapper untuk res_get_available_fonts"""
        return self.resource_controller.res_get_available_fonts()
    def create_custom_palette(self, color_list, palette_name=None):
        """Wrapper untuk res_create_custom_palette"""
        return self.resource_controller.res_create_custom_palette(color_list, palette_name)
    def get_palette_names(self):
        """Wrapper untuk res_get_palette_names"""
        return self.resource_controller.res_get_palette_names()
    def _unload_ml_models(self):
        """Wrapper untuk res_unload_ml_models"""
        return self.resource_controller.res_unload_ml_models()
    def _clear_model_cache(self):
        """Wrapper untuk res_clear_model_cache"""
        return self.resource_controller.res_clear_model_cache()
    def cleanup_figure(self, figure=None):
        """Wrapper untuk res_cleanup_figure"""
        return self.resource_controller.res_cleanup_figure(figure)
    def load_matplotlib_fonts(self):
        """Update font selector with available fonts"""
        if hasattr(self, 'font_map') and self.font_map and len(self.font_map) > 0:
            logger.debug("Fonts already loaded in UI, skipping")
            return
        self.font_choice.clear()
        self.font_choice.addItem("Default")
        try:
            if hasattr(self.resource_controller, 'font_loaded'):
                try:
                    self.resource_controller.font_loaded.disconnect(self._on_font_loaded)
                except (TypeError, RuntimeError):
                    pass
        except Exception as e:
            logger.debug(f"Error disconnecting font_loaded signal: {e}")
        if hasattr(self.resource_controller, 'font_loader_thread') and self.resource_controller.font_loader_thread:
            try:
                self.resource_controller.font_loader_thread.loading_complete.disconnect(self._on_all_fonts_loaded)
            except (TypeError, RuntimeError):
                pass
            self.resource_controller.font_loader_thread.loading_complete.connect(self._on_all_fonts_loaded)
        fonts = self.resource_controller.res_get_available_fonts()
        if fonts and len(fonts) > 0:
            self._on_all_fonts_loaded(fonts)
    def _on_all_fonts_loaded(self, fonts_dict):
        """Add all loaded fonts at once to the dropdown"""
        try:
            if not hasattr(self, 'font_choice'):
                logger.warning("Font choice dropdown not initialized")
                return
            if not fonts_dict or len(fonts_dict) == 0:
                logger.warning("Empty font dictionary received")
                return
            if hasattr(self, 'font_map') and self.font_map == fonts_dict:
                logger.debug("Fonts already loaded, skipping UI update")
                return
            current_text = self.font_choice.currentText()
            self.font_choice.clear()
            self.font_choice.addItem("Default")
            sorted_font_names = sorted(fonts_dict.keys())
            for font_name in sorted_font_names:
                self.font_choice.addItem(font_name)
            if current_text in ["Default"] or current_text in fonts_dict:
                self.font_choice.setCurrentText(current_text)
            self.font_map = fonts_dict
            logger.debug(f"Added {len(fonts_dict)} alphabetically sorted fonts to UI dropdown")
        except Exception as e:
            logger.error(f"Error adding fonts to UI: {str(e)}")
            self.font_choice.clear()
            self.font_choice.addItem("Default")
            self.font_choice.addItems(["Arial", "Times New Roman", "Verdana"])
    def analyze_sentiment(self):
        """Analyze sentiment with non-blocking UI"""
        if not self.text_data:
            QMessageBox.critical(self, "Error", "No text data available for sentiment analysis.\nPlease load a text file first.")
            return
        if (self.sentiment_mode == "Flair" and not self.flair_classifier and 
                self.sentiment_mode_combo.currentText() == "Flair"):
            QMessageBox.warning(self, "Model Loading", "Default Flair model is still loading. Please wait...")
            return
        if self.sentiment_mode == "Flair (Custom Model)" and not self.flair_classifier_cuslang:
            QMessageBox.warning(self, "Model Required", "Please load a custom Flair model first!")
            return
        if self.sentiment_mode == "VADER (Custom Lexicon)" and not self.custom_lexicon_path:
            QMessageBox.warning(self, "Lexicon Required", "Please load a custom lexicon first!")
            return
        if self.sentiment_mode == "TextBlob (Custom Lexicon)" and not self.custom_textblob_lexicon_path:
            QMessageBox.warning(self, "Lexicon Required", "Please load a custom lexicon first!")
            return
        self.button_manager.disable_other_buttons('sentiment_button')
        self.set_progress('sentiment')
        analyzers = {
            'vader_analyzer': self.vader_analyzer,
            'flair_classifier': self.flair_classifier,
            'flair_classifier_cuslang': self.flair_classifier_cuslang,
            'textblob_analyzer': self.textblob_analyzer
        }
        self.sentiment_service.analyze_text(
            self.text_data, 
            self.sentiment_mode, 
            analyzers,
            translation_handler=self._handle_translation_progress
        )
    def _on_sentiment_analysis_started(self):
        """Handle sentiment analysis started"""
        self.set_progress('model')
        self.statusBar().showMessage(f"Analyzing sentiment using {self.sentiment_mode}...", 2000)
    def _handle_sentiment_error(self, error_msg):
        """Handle sentiment analysis errors"""
        self.set_progress('model', False)
        self.button_manager.restore_states()
        QMessageBox.warning(self, "Sentiment Analysis Error", error_msg)
    def stop_all_processes(self):
        """Enhanced thread termination with proper resource management"""
        logger.info("Initiating controlled process termination")
        try:
            if hasattr(self, 'unified_progress_bar'):
                self.unified_progress_bar.setVisible(False)
                self.unified_progress_bar.setValue(0)
                if hasattr(self, 'progress_timer'):
                    self.progress_timer.stop()
            self.disable_buttons()
            self.statusBar().showMessage("Stopping all processes...", 2000)
            if hasattr(self, 'wordcloud_service'):
                try:
                    self.wordcloud_service.stop_generation()
                    if hasattr(self.wordcloud_service, 'generator_thread'):
                        thread = self.wordcloud_service.generator_thread
                        if thread and thread.isRunning():
                            thread.blockSignals(True)
                            thread.cancel()
                            thread.quit()
                            thread.wait(500)
                            if hasattr(thread, 'cleanup_resources'):
                                thread.cleanup_resources()
                except Exception as e:
                    logger.error(f"Error stopping wordcloud thread: {e}")
            special_threads = []
            if hasattr(self, 'sentiment_thread') and self.sentiment_thread:
                special_threads.append(('sentiment', self.sentiment_thread))
            if hasattr(self, 'topic_thread') and self.topic_thread:
                special_threads.append(('topic', self.topic_thread))
            if hasattr(self, 'stats_thread') and self.stats_thread:
                special_threads.append(('stats', self.stats_thread))
            for thread_name, thread in special_threads:
                try:
                    if thread and thread.isRunning():
                        thread.blockSignals(True)
                        if hasattr(thread, 'cancel'):
                            thread.cancel()
                        if hasattr(thread, 'requestInterruption'):
                            thread.requestInterruption()
                        thread.quit()
                        thread.wait(500)
                        if hasattr(thread, 'cleanup_resources'):
                            thread.cleanup_resources()
                except Exception as e:
                    logger.error(f"Error stopping {thread_name} thread: {e}")
            if hasattr(self, 'thread_manager'):
                try:
                    self.thread_manager.stop_all_threads(wait_timeout=800)
                except Exception as e:
                    logger.error(f"Error in ThreadManager.stop_all_threads: {e}")
            if hasattr(self, 'resource_controller'):
                try:
                    self.resource_controller.cleanup_all_graphics()
                except Exception as e:
                    logger.error(f"Error cleaning up graphics: {e}")
            self.statusBar().showMessage("All processes stopped", 2000)
        except Exception as e:
            logger.error(f"Critical error during process termination: {e}")
            self.statusBar().showMessage("Error stopping processes", 2000)
        finally:
            self.enable_buttons()
            if hasattr(self, 'unified_progress_bar'):
                self.unified_progress_bar.setVisible(False)
                self.unified_progress_bar.setValue(0)
            if hasattr(self, 'progress_timer'):
                self.progress_timer.stop()
            if hasattr(self, 'wordcloud_service'):
                if hasattr(self.wordcloud_service, 'generator_thread'):
                    self.wordcloud_service.generator_thread = None
            logger.info("Process termination completed")
    def _unregister_termination_callbacks(self):
        """Unregister callbacks that might interfere with termination"""
        try:
            if hasattr(self, 'progress_timer') and self.progress_timer.isActive():
                self.progress_timer.stop()
            if hasattr(self, 'wordcloud_service'):
                self.wordcloud_service.blockSignals(True)
        except Exception as e:
            logger.error(f"Error unregistering callbacks: {e}")
    def _close_active_dialogs(self):
        """Close any active dialogs"""
        try:
            for widget in QApplication.topLevelWidgets():
                if isinstance(widget, QDialog) and widget.isVisible():
                    try:
                        widget.reject()
                    except Exception as dialog_err:
                        logger.error(f"Error closing dialog: {dialog_err}")
        except Exception as e:
            logger.error(f"Error closing dialogs: {e}")
    def _prepare_resources_for_termination(self):
        """Prepare resources for safe termination"""
        try:
            if hasattr(self, 'temp_files'):
                for file in self.temp_files:
                    try:
                        file.close()
                    except:
                        pass
            if hasattr(self, 'network_manager'):
                self.network_manager.networkAccessible = QNetworkAccessManager.NotAccessible
        except Exception as e:
            logger.error(f"Error preparing resources: {e}")
    def _perform_post_termination_cleanup(self):
        """Perform cleanup after termination"""
        try:
            self.is_processing = False
            self.unified_progress_bar.setValue(0)
            if hasattr(self, 'current_thread'):
                self.current_thread = None
        except Exception as e:
            logger.error(f"Error in post-termination cleanup: {e}")
    def _restore_ui_state(self):
        """Restore UI state after termination"""
        try:
            self.button_manager.restore_states()
            self.set_progress('wordcloud', False)
            self.set_progress('model', False)
            self.set_progress('analysis', False)
        except Exception as e:
            logger.error(f"Error restoring UI state: {e}")
    def _collect_special_threads(self):
        """Collects threads that are not managed by ThreadManager"""
        special_threads = []
        logger.debug("Collecting special threads for termination")
        thread_attributes = [
            ('sentiment_thread', 'Sentiment Analysis'),
            ('wordcloud_thread', 'WordCloud Generation'),
            ('file_loader_thread', 'File Loading'),
            ('stats_thread', 'Text Statistics'),
            ('import_thread', 'Initial Import'),
            ('flair_loader_thread', 'Flair Model Loading'),
            ('custom_flair_loader_thread', 'Custom Flair Model Loading'),
            ('text_viewer_thread', 'Text Viewer'),
            ('summarize_thread', 'Text Summarizer')
        ]
        for attr_name, thread_desc in thread_attributes:
            if hasattr(self, attr_name):
                thread = getattr(self, attr_name, None)
                if thread and isinstance(thread, QThread) and thread.isRunning():
                    try:
                        thread.objectName()
                        logger.debug(f"Found running special thread: self.{attr_name} ({thread_desc})")
                        has_cleanup = hasattr(thread, 'cleanup_resources')
                        logger.debug(f"Thread {attr_name} has cleanup_resources: {has_cleanup}")
                        special_threads.append((f"self.{attr_name}", thread))
                    except RuntimeError:
                        logger.warning(f"Thread {attr_name} has invalid C++ object reference, skipping")
                        setattr(self, attr_name, None)
        if hasattr(self, 'topic_tab'):
            topic_thread_attributes = [
                ('topic_thread', 'Topic Analysis'),
                ('keyword_thread', 'Keyword Extraction')
            ]
            for attr_name, thread_desc in topic_thread_attributes:
                if hasattr(self.topic_tab, attr_name):
                    thread = getattr(self.topic_tab, attr_name, None)
                    if thread and isinstance(thread, QThread) and thread.isRunning():
                        try:
                            thread.objectName()
                            logger.debug(f"Found running special thread: topic_tab.{attr_name} ({thread_desc})")
                            has_cleanup = hasattr(thread, 'cleanup_resources')
                            logger.debug(f"Thread topic_tab.{attr_name} has cleanup_resources: {has_cleanup}")
                            special_threads.append((f"topic_tab.{attr_name}", thread))
                        except RuntimeError:
                            logger.warning(f"{thread_desc} thread has invalid C++ object reference, skipping")
                            setattr(self.topic_tab, attr_name, None)
        service_thread_paths = [
            ('sentiment_service', 'analysis_thread', 'Sentiment Service Analysis'),
            ('text_analysis_service', 'stats_thread', 'Text Analysis Service'),
            ('wordcloud_service', 'generator_thread', 'WordCloud Service')
        ]
        for service_name, thread_attr, desc in service_thread_paths:
            if hasattr(self, service_name):
                service = getattr(self, service_name, None)
                if service and hasattr(service, thread_attr):
                    thread = getattr(service, thread_attr, None)
                    if thread and isinstance(thread, QThread) and thread.isRunning():
                        try:
                            thread.objectName()
                            logger.debug(f"Found running service thread: {service_name}.{thread_attr} ({desc})")
                            has_cleanup = hasattr(thread, 'cleanup_resources')
                            logger.debug(f"Thread {service_name}.{thread_attr} has cleanup_resources: {has_cleanup}")
                            special_threads.append((f"{service_name}.{thread_attr}", thread))
                        except RuntimeError:
                            logger.warning(f"{desc} thread has invalid C++ object reference, skipping")
                            setattr(service, thread_attr, None)
        logger.info(f"Collected {len(special_threads)} special threads to stop.")
        return special_threads
    def disable_buttons(self):
        """Disable all buttons during processing"""
        process_buttons = [
            self.generate_wordcloud_button,
            self.sentiment_button, 
            self.custom_lexicon_button,
            self.custom_model_button,
            self.text_stats_button
        ]
        if hasattr(self, 'topic_tab'):
            if hasattr(self.topic_tab, 'analyze_topics_btn'):
                process_buttons.append(self.topic_tab.analyze_topics_btn)
            if hasattr(self.topic_tab, 'extract_keywords_btn'):
                process_buttons.append(self.topic_tab.extract_keywords_btn)
        for button in process_buttons:
            button.setEnabled(False)
        self.load_file_button.setEnabled(True)
        self.panic_button.setEnabled(True)
        QApplication.processEvents()
    def disable_processing_buttons(self):
        """Alias for disable_buttons for compatibility"""
        self.disable_buttons()
    def enable_buttons(self):
        """Re-enable buttons after processing"""
        has_text = bool(self.text_data)
        self.load_file_button.setEnabled(True)
        self.view_fulltext_button.setEnabled(has_text)
        self.generate_wordcloud_button.setEnabled(has_text)
        self.text_stats_button.setEnabled(has_text)
        self.sentiment_button.setEnabled(has_text)
        self.summarize_button.setEnabled(has_text)
        if hasattr(self, 'topic_tab'):
            if hasattr(self.topic_tab, 'analyze_topics_btn'):
                self.topic_tab.analyze_topics_btn.setEnabled(has_text)
            if hasattr(self.topic_tab, 'extract_keywords_btn'):
                self.topic_tab.extract_keywords_btn.setEnabled(has_text)
        if self.sentiment_mode == "VADER (Custom Lexicon)":
            self.custom_lexicon_button.setEnabled(True)
            self.custom_model_button.setEnabled(False)
        elif self.sentiment_mode == "Flair (Custom Model)":
            self.custom_lexicon_button.setEnabled(False)  
            self.custom_model_button.setEnabled(True)
        else:
            self.custom_lexicon_button.setEnabled(False)
            self.custom_model_button.setEnabled(False)
        QApplication.processEvents()    
    def _register_resources(self):
        """Mendaftarkan sumber daya aplikasi ke resource manager"""
        if hasattr(self, 'lda_model'):
            self.resource_manager.register('lda_model', self.lda_model)
        if hasattr(self, 'nmf_model'):
            self.resource_manager.register('nmf_model', self.nmf_model)
        if hasattr(self, 'bert_model'):
            self.resource_manager.register('bert_model', self.bert_model)
        if hasattr(self, 'current_figure'):
            self.resource_manager.register('current_figure', self.current_figure, 
                                          lambda fig: plt.close(fig) if fig else None)
    def _register_caches(self):
        """Register all application caches with the cache manager"""
        self.cache_manager.register_model_cache(self._cached_models)
        if hasattr(self, '_cached_colormaps'):
            self.cache_manager.register_graphics_cache(self._cached_colormaps)
        if hasattr(self, '_cached_fonts'):
            self.cache_manager.register_graphics_cache(self._cached_fonts)
        if hasattr(self, 'get_wordcloud'):
            self.cache_manager.register_cached_function(self.get_wordcloud)
        if hasattr(self, 'token_opts'):
            self.cache_manager.register_cached_property(self, 'token_opts')
    def show_about(self):
        """Show about dialog using DialogFactory with caching"""
        if not self._about_dialog:
            about_text = base64.b64decode(AppConstants.ABOUT_TEXT.encode()).decode()
            about_text = about_text.replace("{logo_path}", str(logo_path.as_uri()))
            self._about_dialog = DialogFactory.create_info_dialog(
                self, "About Textplora", about_text, 
                modal=True,
                min_size=(400, 300)
            )
            self._about_dialog.finished.connect(self._cleanup_about_dialog)
        self._about_dialog.show()
        self._about_dialog.raise_()
        self._about_dialog.activateWindow()
    def _cleanup_about_dialog(self):
        """Cleanup about dialog when closed"""
        if self._about_dialog:
            self._about_dialog.deleteLater()
            self._about_dialog = None
    def open_file(self):
        dialog = QFileDialog(self)
        dialog.setWindowTitle("Load Text File")
        dialog.setNameFilter("Supported Files (*.txt *.pdf *.doc *.docx *.csv *.xlsx *.xls);;All Files (*)")
        dialog.setOptions(QFileDialog.Options())
        if dialog.exec():
            file_path = dialog.selectedFiles()[0]
            if not file_path:
                return
            self.set_progress('file')
            try:
                self.file_loader_thread = FileLoaderThread(file_path)
                self.file_loader_thread.file_loaded.connect(self.on_file_loaded)
                self.file_loader_thread.file_error.connect(self.handle_file_error)
                self.thread_manager.add_thread(self.file_loader_thread)
                self.file_loader_thread.start()
            except Exception as e:
                self.unified_progress_bar.setVisible(False)
                QMessageBox.critical(self, "Error", f"Failed to load file: {e}")
    def handle_file_error(self, error_message):
        self.unified_progress_bar.setVisible(False)
        QMessageBox.critical(self, "File Load Error", error_message)
        self._reset_file_state()
    def on_file_loaded(self, file_path, text_data):
        """Handle file loading completion"""
        self.set_progress('file', False)
        try:
            if not text_data.strip():
                raise ValueError("File appears empty after loading")
            self.text_data = text_data
            self.file_name.setText(os.path.basename(file_path))
            self.topic_tab.set_text(self.text_data)
            self.generate_wordcloud_button.setEnabled(True)
            self.text_stats_button.setEnabled(True)
            self.view_fulltext_button.setEnabled(True)
            self.summarize_button.setEnabled(True)
            self.load_file_button.setEnabled(True)
            self.change_sentiment_mode(self.sentiment_mode)
        except ValueError as e:
            self.handle_file_error(str(e))
            self._reset_file_state()
        self.load_file_button.setEnabled(True)
    def _reset_file_state(self):
        self.text_data = ""
        self.file_name.setText("")
        self.generate_wordcloud_button.setEnabled(False)
        self.text_stats_button.setEnabled(False)
        self.view_fulltext_button.setEnabled(False)
        self.sentiment_button.setEnabled(False)
        self.summarize_button.setEnabled(False)
    def choose_mask(self):
        options = QFileDialog.Options()
        mask_path, _ = QFileDialog.getOpenFileName(
            self, "Select Mask Image", "", "Image Files (*.png *.jpg *.bmp)", options=options
        )
        if mask_path:
            self.wordcloud_service.choose_mask(mask_path)
            self.mask_path_label.setText(f"{mask_path}")
    def reset_mask(self):
        self.wordcloud_service.reset_mask()
        self.mask_path_label.setText("default (rectangle)")
    def import_stopwords(self):
        custom_words = self.stopword_entry.toPlainText().strip().lower()
        if custom_words:
            self.additional_stopwords = set(custom_words.split())
        return STOPWORDS.union(self.additional_stopwords)
    def view_text_stats(self):
        """View text statistics with option to include/exclude stopwords"""
        if not self.text_data:
            QMessageBox.critical(self, "Error", "No text data available.\nPlease load a text file first.")
            return
        stopwords_dialog = QMessageBox(self)
        stopwords_dialog.setWindowTitle("Stopwords Preference")
        stopwords_dialog.setText("Do you want to include stopwords in the text statistics?")
        stopwords_dialog.setInformativeText("Excluding stopwords focuses on meaningful content words.\nIncluding stopwords shows complete language usage patterns.")
        include_btn = stopwords_dialog.addButton("Include Stopwords", QMessageBox.AcceptRole)
        exclude_btn = stopwords_dialog.addButton("Exclude Stopwords", QMessageBox.RejectRole)
        cancel_btn = stopwords_dialog.addButton(QMessageBox.Cancel)
        stopwords_dialog.setDefaultButton(exclude_btn)
        stopwords_dialog.exec()
        clicked_button = stopwords_dialog.clickedButton()
        if clicked_button == cancel_btn:
            return
        custom_stopwords = self.import_stopwords()
        include_stopwords = (clicked_button == include_btn)
        self.button_manager.disable_other_buttons('text_stats_button')
        self.set_progress('stats')
        if not hasattr(self.text_analysis_service, '_progress_connected'):
            self.text_analysis_service.analysis_progress.connect(self._update_stats_progress)
            self.text_analysis_service._progress_connected = True
        self.text_analysis_service.analyze_text(
            self.text_data, 
            include_stopwords=include_stopwords,
            custom_stopwords=custom_stopwords
        )
    def _on_text_analysis_started(self):
        """Handler saat analisis teks dimulai"""
        logger.info("Text analysis started")
        self.set_progress('stats', True)        
    def _update_stats_progress(self, value):
        """Update progress bar untuk text stats"""
        if hasattr(self, 'unified_progress_bar'):
            self.unified_progress_bar.setValue(value)
            self.statusBar().showMessage(f"Analyzing text: {value}%")        
    def _on_text_analysis_complete(self, stats):
        """Handler saat analisis teks selesai"""
        self.set_progress('stats', False)
        self.button_manager.restore_states()
        if hasattr(self, 'unified_progress_bar'):
            self.unified_progress_bar.setValue(100)
        self.statusBar().clearMessage()
        self.text_analysis_service.show_stats_dialog(self, stats)
    def _on_text_analysis_error(self, error_msg):
        """Handler saat terjadi error pada analisis teks"""
        self.set_progress('stats', False)
        self.button_manager.restore_states()
        if hasattr(self, 'unified_progress_bar'):
            self.unified_progress_bar.setValue(0)
        self.statusBar().clearMessage()
        QMessageBox.critical(self, "Error", f"Failed to analyze text: {error_msg}")    
    def generate_wordcloud(self):
        """Generate word cloud menggunakan service"""
        if not self.text_data:
            QMessageBox.critical(self, "Error", "Load text file first!")
            return
        self.button_manager.disable_other_buttons('generate_wordcloud_button')
        self.set_progress('wordcloud')
        font_map = getattr(self, '_cached_fonts', {})
        if not font_map:
            logger.warning("No cached fonts available, loading fonts first...")
            self.load_matplotlib_fonts_optimized()
            font_map = getattr(self, '_cached_fonts', {})
        custom_palettes = getattr(self.resource_controller, 'custom_color_palettes', {})
        selected_colormap = self.color_theme.currentText()
        if selected_colormap.startswith("custom_"):
            logger.debug(f"Selected custom palette: {selected_colormap}")
            if selected_colormap in custom_palettes:
                logger.debug(f"Colors in palette: {custom_palettes[selected_colormap]}")
            else:
                logger.warning(f"Selected palette {selected_colormap} not found in available palettes")
        params = {
            'font_choice': self.font_choice.currentText(),
            'font_map': font_map,
            'colormap': selected_colormap,
            'bg_color': self.bg_color.currentText(),
            'max_words': self.max_words_input.value(),
            'min_font_size': self.min_font_size_input.value(),
            'stopwords': self.import_stopwords(),
            'title_text': self.title_entry.text(),
            'title_font_size': self.title_font_size.value(),
            'title_position': self.title_position.currentText(),
            'custom_color_palettes': custom_palettes
        }
        self.wordcloud_service.set_font_map(font_map)
        self.wordcloud_service.generate_wordcloud(self.text_data, params)
    def _update_wordcloud_progress(self, progress):
        """Update progress bar dari signal service"""
        if hasattr(self, 'unified_progress_bar'):
            if progress <= 0:
                self.unified_progress_bar.setVisible(False)
                self.unified_progress_bar.setValue(0)
                if hasattr(self, 'progress_timer'):
                    self.progress_timer.stop()
            else:
                self.unified_progress_bar.setValue(progress)
                if progress == 100:
                    QTimer.singleShot(200, lambda: self.set_progress('wordcloud', False))
    def _on_wordcloud_ready(self, figure):
        self.current_figure = figure
        self.button_manager.restore_states()
        self.set_progress('wordcloud', False)
    def handle_wordcloud_error(self, error_msg):
        """Handle word cloud generation errors"""
        import matplotlib.pyplot as plt
        if self.current_figure:
            try:
                plt.close(self.current_figure)
            except:
                pass
            finally:
                self.current_figure = None
        QMessageBox.critical(self, "Error", f"Failed to generate word cloud: {error_msg}")
        self.button_manager.restore_states()
        self.set_progress('wordcloud', False)
    def create_custom_palette(self):
        dialog = QDialog(self)
        dialog.setWindowTitle("Create Custom Color Palette")
        dialog.setMinimumSize(500, 400)
        main_layout = QVBoxLayout()
        title_label = QLabel("<h3>Custom Color Palette Creator</h3>")
        title_label.setAlignment(Qt.AlignCenter)
        main_layout.addWidget(title_label)
        instructions = QLabel("Create a custom color palette for your word cloud by selecting colors in order.")
        instructions.setWordWrap(True)
        main_layout.addWidget(instructions)
        preview_group = QGroupBox("Selected Colors")
        preview_layout = QVBoxLayout()
        color_list = []
        preview_frame = QFrame()
        preview_frame.setMinimumHeight(40)
        preview_frame.setFrameShape(QFrame.Box)
        preview_frame.setStyleSheet("background-color: #f0f0f0; border: 1px solid #cccccc;")
        preview_layout.addWidget(preview_frame)
        color_blocks_layout = QHBoxLayout()
        preview_frame.setLayout(color_blocks_layout)
        color_labels = QLabel("")
        preview_layout.addWidget(color_labels)
        preview_group.setLayout(preview_layout)
        main_layout.addWidget(preview_group)
        button_layout = QHBoxLayout()
        add_color_button = QPushButton("+ Add Color")
        button_layout.addWidget(add_color_button)
        remove_color_button = QPushButton("- Remove Last")
        remove_color_button.setEnabled(False)
        button_layout.addWidget(remove_color_button)
        clear_button = QPushButton("Clear All")
        clear_button.setEnabled(False)
        button_layout.addWidget(clear_button)
        main_layout.addLayout(button_layout)
        def update_preview():
            while color_blocks_layout.count():
                item = color_blocks_layout.takeAt(0)
                if item.widget():
                    item.widget().deleteLater()
            if not color_list:
                preview_frame.setStyleSheet("background-color: #f0f0f0; border: 1px solid #cccccc;")
                color_labels.setText("")
                remove_color_button.setEnabled(False)
                clear_button.setEnabled(False)
                return
            for color in color_list:
                block = QFrame()
                block.setMinimumWidth(20)
                block.setFrameShape(QFrame.Box)
                block.setStyleSheet(f"background-color: {color}; border: 1px solid #888888;")
                color_blocks_layout.addWidget(block)
            color_labels.setText(", ".join(color_list))
            remove_color_button.setEnabled(True)
            clear_button.setEnabled(True)
        def add_color():
            color = QColorDialog.getColor()
            if color.isValid():
                color_list.append(color.name())
                update_preview()
        def remove_last_color():
            if color_list:
                color_list.pop()
                update_preview()
        def clear_colors():
            color_list.clear()
            update_preview()
        add_color_button.clicked.connect(add_color)
        remove_color_button.clicked.connect(remove_last_color)
        clear_button.clicked.connect(clear_colors)
        save_group = QGroupBox("Save Palette")
        save_layout = QVBoxLayout()
        name_layout = QHBoxLayout()
        name_layout.addWidget(QLabel("Palette Name:"))
        palette_name_edit = QLineEdit()
        palette_name_edit.setPlaceholderText("Enter a name for your custom palette")
        name_layout.addWidget(palette_name_edit)
        save_layout.addLayout(name_layout)
        save_button = QPushButton("Save Palette")
        save_layout.addWidget(save_button)
        save_group.setLayout(save_layout)
        main_layout.addWidget(save_group)
        button_box = QDialogButtonBox(QDialogButtonBox.Close)
        button_box.rejected.connect(dialog.reject)
        main_layout.addWidget(button_box)
        def save_palette():
            name = palette_name_edit.text().strip()
            if not name:
                QMessageBox.warning(dialog, "Input Required", "Please enter a name for your palette.")
                return
            if not color_list:
                QMessageBox.warning(dialog, "No Colors", "Please add at least one color to your palette.")
                return
            existing_names = self.resource_controller.res_get_palette_names()
            if name in existing_names and not name.startswith('custom_'):
                reply = QMessageBox.question(dialog, 
                                        "Name Already Exists", 
                                        f"The name '{name}' might conflict with an existing colormap. A prefix 'custom_' will be added. Continue?",
                                        QMessageBox.Yes | QMessageBox.No, 
                                        QMessageBox.No)
                if reply == QMessageBox.No:
                    return
            colors_to_save = color_list.copy()
            print(f"Saving palette {name} with colors: {colors_to_save}")
            palette_name = self.resource_controller.res_create_custom_palette(colors_to_save, name)
            if palette_name:
                index = self.color_theme.findText(palette_name)
                if index == -1:
                    self.color_theme.addItem(palette_name)
                self.color_theme.setCurrentText(palette_name)
                QMessageBox.information(dialog, "Success", 
                                    f"Custom palette '{palette_name}' has been created and selected.")
                palette_name_edit.clear()
                clear_colors()
            else:
                QMessageBox.critical(dialog, "Error", 
                                "Failed to create custom palette. Please try again.")
        save_button.clicked.connect(save_palette)
        dialog.setLayout(main_layout)
        dialog.exec()
    def save_custom_palette(self, color_list, dialog):
        palette_name, ok = QInputDialog.getText(self, "Save Palette", "Enter palette name:")
        if ok and palette_name:
            palette_name = self.wordcloud_service.create_custom_palette(color_list, palette_name)
            self.color_theme.addItem(palette_name)
            dialog.accept()
    def select_custom_bg_color(self):
        color = QColorDialog.getColor()
        if color.isValid():
            self.bg_color.addItem(color.name())
            self.bg_color.setCurrentText(color.name())
    def view_full_text(self):
        """Show full text in dialog"""
        try:
            if not hasattr(self, 'text_data') or not self.text_data:
                QMessageBox.warning(self, "Warning", "No text available to display")
                return
            if not hasattr(self, '_text_viewer_dialog'):
                self._text_viewer_dialog = TextViewerDialog(self)
            self._text_viewer_dialog.set_text(self.text_data)
            self._text_viewer_dialog.show()
        except Exception as e:
            logger.error(f"Error showing full text: {e}")
            QMessageBox.critical(self, "Error", f"Failed to display text: {str(e)}")
            self.set_progress('text_loading', visible=False)
    def _handle_text_viewer_error(self, error_msg):
        """Handler untuk error dari text viewer service"""
        QMessageBox.critical(self, "Error", error_msg)
    def handle_offline_warning(self, msg):
        self.unified_progress_bar.setVisible(False)
        QMessageBox.warning(self, "Offline Mode", msg)
    def handle_translation_failure(self, error_msg):
        """Handle translation failure gracefully"""
        self.unified_progress_bar.setVisible(False)
        QMessageBox.warning(self, "Translation Error", error_msg)
        self.enable_buttons()
    def on_sentiment_analyzed(self, result):
        self.set_progress('model', False)
        self.button_manager.restore_states()
        dialog = SentimentAnalysisResultDialog(
            sentiment_mode=self.sentiment_mode,
            result=result,
            parent=self
        )
        dialog.show()        
    @lru_cache(maxsize=128)
    def get_most_frequent_words(self, text, n):
        stopwords = self.import_stopwords().union(STOPWORDS)
        words = [word.lower() for word in text.split() if word.lower() not in stopwords]
        word_counts = Counter(words)
        most_common = word_counts.most_common(n)
        return [word for word, count in most_common]
    def change_sentiment_mode(self, mode):
        """Handle sentiment mode changes with proper button states"""
        self.sentiment_mode = mode
        status_text = ""
        has_text = bool(self.text_data.strip())
        self.custom_lexicon_button.setEnabled(False)
        self.custom_model_button.setEnabled(False)
        self.sentiment_button.setEnabled(False)
        if mode == "Flair":
            if self.flair_classifier:
                status_text = "Flair: Ready"
                self.sentiment_button.setEnabled(has_text)
            else:
                status_text = "Loading..."
                self.load_flair_model()
        elif mode == "Flair (Custom Model)":
            if not self.flair_classifier:
                self.load_flair_model()
                status_text = "Initializing Flair environment..."
            else:
                self.custom_model_button.setEnabled(True)
                if self.flair_classifier_cuslang:  
                    status_text = "Ready"
                    self.sentiment_button.setEnabled(has_text)
                else:
                    status_text = "Load custom model first!"
        elif mode == "VADER (Custom Lexicon)":
            self.custom_lexicon_button.setEnabled(True)
            status_text = "Load custom lexicon first!" if not self.custom_lexicon_path else "Ready"
            self.sentiment_button.setEnabled(has_text and bool(self.custom_lexicon_path))
            if not self.custom_lexicon_path:
                QMessageBox.warning(
                    self, "Lexicon Required", 
                    "Please load a custom lexicon before analyzing sentiment."
                )
        elif mode == "TextBlob (Custom Lexicon)":
            self.custom_lexicon_button.setEnabled(True)
            status_text = "Load custom lexicon first!" if not self.custom_textblob_lexicon_path else "Ready"
            self.sentiment_button.setEnabled(has_text and bool(self.custom_textblob_lexicon_path))
            if not self.custom_textblob_lexicon_path:
                QMessageBox.warning(
                    self, "Lexicon Required", 
                    "Please load a custom lexicon before analyzing sentiment."
                )
        else:
            status_text = "Ready" if has_text else "Load text first!"
            self.sentiment_button.setEnabled(has_text)
        self.sentiment_button.setToolTip(
            f"{mode} - {status_text}\nClick to analyze sentiment using {mode}"
        )
    def load_custom_lexicon(self):
        options = QFileDialog.Options()
        lexicon_path, _ = QFileDialog.getOpenFileName(
            self, "Select Custom Lexicon File", "", "Text Files (*.txt)", options=options
        )
        if lexicon_path:
            try:
                self.button_manager.disable_other_buttons('custom_lexicon_button')
                self.custom_lexicon_button.setText("Loading Lexicon...")
                self.custom_lexicon_button.setEnabled(False)
                self.sentiment_button.setEnabled(False)
                self.lexicon_loader_thread = CustomFileLoaderThread(
                    lexicon_path, 
                    "TextBlob (Custom Lexicon)" if self.sentiment_mode == "TextBlob (Custom Lexicon)" else "lexicon"
                )
                self.lexicon_loader_thread.file_loaded.connect(self.on_lexicon_loaded)
                self.thread_manager.add_thread(self.lexicon_loader_thread)
                self.lexicon_loader_thread.start()
            except Exception as e:
                self.button_manager.restore_states()
                self.custom_lexicon_button.setText("Load Lexicon")
                QMessageBox.critical(self, "Error", f"Failed to start lexicon loading: {str(e)}")
    def on_lexicon_loaded(self, result, success):
        self.unified_progress_bar.setVisible(False)
        self.button_manager.restore_states()
        self.custom_lexicon_button.setEnabled(True)
        self.custom_lexicon_button.setText("Load Lexicon")
        try:
            if success:
                if self.sentiment_mode == "TextBlob (Custom Lexicon)":
                    self.custom_textblob_lexicon_path = result
                    self.textblob_analyzer = CustomTextBlobSentimentAnalyzer(self.custom_textblob_lexicon_path)
                    QMessageBox.information(self, "Success", "Custom TextBlob lexicon loaded!")
                else:
                    self.custom_lexicon_path = result
                    self.vader_analyzer = CustomVaderSentimentIntensityAnalyzer(custom_lexicon_file=self.custom_lexicon_path)
                    QMessageBox.information(self, "Success", "Custom VADER lexicon loaded!")
                has_text = bool(self.text_data.strip())
                self.sentiment_button.setEnabled(has_text)
            else:
                raise ValueError(str(result))
        except Exception as e:
            if self.sentiment_mode == "TextBlob (Custom Lexicon)":
                self.custom_textblob_lexicon_path = None
                self.textblob_analyzer = None
            else:
                self.custom_lexicon_path = None
                self.vader_analyzer = None
            QMessageBox.critical(self, "Error", f"Failed to load lexicon: {str(e)}")
            self.sentiment_button.setEnabled(False)
            self.button_manager.restore_states()
            self.custom_lexicon_button.setEnabled(True)
    def load_custom_model(self):
        options = QFileDialog.Options()
        model_path, _ = QFileDialog.getOpenFileName(
            self, "Select Custom Model File", "", "Model Files (*.pt)", options=options
        )
        if model_path:
            try:
                self.button_manager.disable_other_buttons('custom_model_button')
                self.custom_model_button.setText("Loading Model...")
                self.custom_model_button.setEnabled(False)
                self.sentiment_button.setEnabled(False)
                self.model_loader_thread = CustomFileLoaderThread(model_path, "model")
                self.model_loader_thread.file_loaded.connect(self.on_model_loaded)
                self.thread_manager.add_thread(self.model_loader_thread)
                self.model_loader_thread.start()
            except Exception as e:
                self.button_manager.restore_states()
                self.custom_model_button.setText("Load Model")
                QMessageBox.critical(self, "Error", f"Failed to start model loading: {str(e)}")
    def on_model_loaded(self, result, success):
        """Handle when custom Flair model is loaded"""
        self.unified_progress_bar.setVisible(False)
        self.custom_model_button.setText("Load Model")
        self.button_manager.restore_states()
        self.custom_model_button.setEnabled(True)
        try:
            if success:
                from flair.models import TextClassifier
                from flair.data import Sentence
                if not isinstance(result, TextClassifier):
                    raise ValueError("Invalid model type. Please load a valid Flair TextClassifier model.")
                test_sentence = Sentence("test")
                result.predict(test_sentence)
                if not test_sentence.labels:
                    raise ValueError("Model didn't produce any labels")
                label = test_sentence.labels[0].value
                if label not in ['POSITIVE', 'NEGATIVE', 'NEUTRAL']:
                    raise ValueError(f"Model produces incompatible labels: {label}")
                self.flair_classifier_cuslang = result
                QMessageBox.information(self, "Success", "Custom model loaded successfully!")
                self.button_manager.restore_states()
                has_text = bool(self.text_data.strip())
                if has_text:
                    self.sentiment_button.setEnabled(True)
                    self.view_fulltext_button.setEnabled(True)
                    self.summarize_button.setEnabled(True)
                    self.text_stats_button.setEnabled(True)
                    if hasattr(self, 'topic_tab'):
                        if hasattr(self.topic_tab, 'analyze_topics_btn'):
                            self.topic_tab.analyze_topics_btn.setEnabled(True)
                        if hasattr(self.topic_tab, 'extract_keywords_btn'):
                            self.topic_tab.extract_keywords_btn.setEnabled(True)
                self.custom_model_button.setEnabled(True)
            else:
                raise ValueError(str(result))
        except Exception as e:
            self.flair_classifier_cuslang = None
            QMessageBox.critical(self, "Error", f"Failed to load custom model: {str(e)}")
            self.sentiment_button.setEnabled(False)
            self.button_manager.restore_states()
            self.custom_model_button.setEnabled(True)
    def load_flair_model(self):
        """Load Flair model asynchronously using SentimentAnalysisService"""
        try:
            self.button_manager.disable_other_buttons('sentiment_button')
            self.sentiment_button.setText("Loading Flair...")
            self.sentiment_button.setEnabled(False)
            self.set_progress('model')
            self.statusBar().showMessage("Loading Flair Model... Please wait.")
            if not hasattr(self, '_flair_loading_connected'):
                self.sentiment_service.model_loaded.connect(self._on_flair_model_loaded)
                self.sentiment_service.model_loading_failed.connect(self._on_flair_model_error)
                self._flair_loading_connected = True
            self.sentiment_service.load_model('flair')
        except Exception as e:
            self.button_manager.restore_states()
            self.sentiment_button.setText("Analyze Sentiment")
            self.set_progress('model', False)
            QMessageBox.critical(self, "Error", f"Failed to load Flair: {str(e)}")
    def _on_flair_model_loaded(self, model_type, model):
        """Handle when Flair model is loaded successfully"""
        if model_type.lower() != 'flair':
            return
        self.set_progress('model', False)
        self.button_manager.restore_states()
        self.sentiment_button.setText("Analyze Sentiment")
        self.statusBar().showMessage("Flair Model Loaded Successfully! Ready to analyze sentiment.", 3000)
        if model:
            self.flair_classifier = model
            has_text = bool(self.text_data.strip())
            if self.flair_first_load:
                self.flair_first_load = False
                QMessageBox.information(self, "Ready", "Flair library loaded successfully!")
            if self.sentiment_mode == "Flair":
                self.sentiment_button.setEnabled(has_text)
            elif self.sentiment_mode == "Flair (Custom Model)":
                self.custom_model_button.setEnabled(True)
                if self.flair_classifier_cuslang:
                    self.sentiment_button.setEnabled(has_text)
                else:
                    QMessageBox.information(self, "Next Step", "Please load your custom model")
                    self.sentiment_button.setEnabled(False)
        else:
            self.sentiment_button.setEnabled(False)
            QMessageBox.critical(self, "Error", "Failed to load Flair model")
    def _on_flair_model_error(self, model_type, error):
        """Handle Flair model loading error"""
        if model_type.lower() != 'flair':
            return
        self.set_progress('model', False)
        self.button_manager.restore_states()
        self.sentiment_button.setText("Analyze Sentiment")
        self.statusBar().showMessage("Failed to load Flair model!", 3000)
        QMessageBox.critical(self, "Loading Error", f"Failed to load Flair model: {error}")
    def load_custom_flair_model(self):
        """Load custom Flair model asynchronously"""
        try:
            model_path = FileHandler.open_file_dialog(
                self, "Select Custom Flair Model", "Model Files (*.pt);;All Files (*)"
            )
            if not model_path:
                return
            self.button_manager.disable_other_buttons('custom_model_button')
            self.custom_model_button.setText("Loading...")
            self.custom_model_button.setEnabled(False)
            self.set_progress('model')
            self.statusBar().showMessage("Loading Custom Flair Model... Please wait.")
            if not hasattr(self, '_custom_flair_loading_connected'):
                self.sentiment_service.model_loaded.connect(self._on_custom_flair_model_loaded)
                self.sentiment_service.model_loading_failed.connect(self._on_custom_flair_model_error)
                self._custom_flair_loading_connected = True
            self.sentiment_service.load_model('flair', model_path, custom=True)
        except Exception as e:
            self.button_manager.restore_states()
            self.custom_model_button.setText("Load Custom Model")
            self.set_progress('model', False)
            QMessageBox.critical(self, "Error", f"Failed to load custom model: {str(e)}")
    def _on_custom_flair_model_loaded(self, model_type, model):
        """Handle when custom Flair model is loaded successfully"""
        if model_type.lower() != 'flair':
            return
        self.set_progress('model', False)
        self.button_manager.restore_states()
        self.custom_model_button.setText("Load Custom Model")
        self.statusBar().showMessage("Custom Flair Model Loaded Successfully! Please load your custom model.", 3000)
        if model:
            self.flair_classifier_cuslang = model
            self.custom_model_path = model_path
            has_text = bool(self.text_data.strip())
            QMessageBox.information(self, "Success", "Custom model loaded successfully!")
            if self.sentiment_mode == "Flair (Custom Model)":
                self.sentiment_button.setEnabled(has_text)
        else:
            QMessageBox.critical(self, "Error", "Failed to load custom model")
    def _on_custom_flair_model_error(self, model_type, error):
        """Handle custom Flair model loading error"""
        if model_type.lower() != 'flair':
            return
        self.set_progress('model', False)
        self.button_manager.restore_states()
        self.custom_model_button.setText("Load Custom Model")
        self.statusBar().showMessage("Failed to load Custom Flair Model!", 3000)
        QMessageBox.critical(self, "Loading Error", f"Failed to load custom model: {error}")
    def cleanup_flair_thread(self):
        self.progress_timer.stop()
        self.unified_progress_bar.setValue(0)
        self.unified_progress_bar.setVisible(False)
        self.thread_manager.remove_thread(self.flair_loader_thread)
        self.flair_loader_thread = None        
    def show_sentiment_mode_info(self):
        """Show sentiment mode info dialog using DialogFactory"""
        mode_info = base64.b64decode(AppConstants.MODE_INFO.encode()).decode()
        dialog = DialogFactory.create_info_dialog(
            self, "Sentiment Analysis Modes", mode_info, modal=False
        )
        dialog.show()
    def update_topic_analysis(self, text):
        """Update topic analysis tab with new text"""
        self.topic_tab.set_text(text)
    def _update_progress_animation(self):
        """Smooth indeterminate progress animation"""
        if hasattr(self, 'unified_progress_bar') and self.unified_progress_bar.isVisible():
            current = self.unified_progress_bar.value()
            next_value = (current + 2) % 100
            self.unified_progress_bar.setValue(next_value)
    def resizeEvent(self, event):
        """Handle resize to keep progress bar full width"""
        super().resizeEvent(event)
        if hasattr(self, 'unified_progress_bar'):
            container = self.unified_progress_bar.parent()
            self.unified_progress_bar.setFixedWidth(container.width())
            self.unified_progress_bar.setGeometry(0, 0, container.width(), container.height())
    def update_stopwords_for_topic(self, custom_stopwords=None):
        """Initialize and update stopwords with consistent preprocessing"""
        from nltk.corpus import stopwords
        try:
            default_stopwords = set(stopwords.words('english'))
        except:
            default_stopwords = STOPWORDS
        default_stopwords = {word.lower().strip() for word in default_stopwords}
        if custom_stopwords:
            custom_stopwords = {word.lower().strip() for word in custom_stopwords if word.strip()}
            self.stop_words = default_stopwords.union(custom_stopwords)
        else:
            self.stop_words = default_stopwords
        vectorizer_config = {
            'token_pattern': r'(?u)\b[a-zA-Z][a-zA-Z0-9\'-]*[a-zA-Z0-9]\b',
            'stop_words': list(self.stop_words),
            'lowercase': True,
            'strip_accents': 'unicode',
            'max_features': 1000,
            'min_df': 1,
            'max_df': 1
        }
        self.vectorizer = CountVectorizer(**vectorizer_config)
        self.tfidf = TfidfVectorizer(**vectorizer_config)
    def _extract_tfidf(self, text, num_keywords):
        """Extract keywords using TF-IDF with consistent preprocessing"""
        if not hasattr(self, 'stop_words'):
            self.update_stopwords_for_topic()
        vectorizer = TfidfVectorizer(
            token_pattern=r'(?u)\b[a-zA-Z][a-zA-Z0-9\'-]*[a-zA-Z0-9]\b',
            stop_words=list(self.stop_words),
            lowercase=True,
            strip_accents='unicode',
            min_df=1,
            max_df=1
        )
        try:
            sentences = [s.strip() for s in text.split('.') if s.strip()]
            if not sentences:
                sentences = [text]
            response = vectorizer.fit_transform(sentences)
            feature_names = vectorizer.get_feature_names_out()
            scores = response.mean(axis=0).A[0]
            pairs = list(zip(scores, feature_names))
            pairs.sort(reverse=True)
            return [{'keyword': word, 'score': score} 
                    for score, word in pairs[:num_keywords]]
        except ValueError as e:
            logger.error(f"TF-IDF extraction error: {str(e)}")
            return [{'keyword': 'Error: No valid keywords found', 'score': 0.0}]
    def _process_topic_modeling(self, text, num_topics, model_type='lda'):
        """Generic topic modeling processor"""
        try:
            sentences = [sent.strip() for sent in text.split('.') if sent.strip()] or [text]
            vectorizer = self.vectorizer if model_type == 'lda' else self.tfidf
            dtm = vectorizer.fit_transform(sentences)
            if model_type == 'nmf':
                max_possible_topics = min(dtm.shape[0], dtm.shape[1])
                if num_topics > max_possible_topics:
                    raise ValueError(f"Number of topics ({num_topics}) cannot exceed {max_possible_topics} for NMF with this dataset")
            num_topics = min(num_topics, max(2, dtm.shape[1] - 1))
            model_class = LatentDirichletAllocation if model_type == 'lda' else NMF
            model_opts = {
                'n_components': num_topics,
                'random_state': 42,
                'max_iter': 20 if model_type == 'lda' else 500
            }
            if model_type == 'nmf':
                model_opts['init'] = 'nndsvd'
                model_opts['tol'] = 1e-4
            if model_type == 'lda':
                model_opts['learning_method'] = 'online'
            model = model_class(**model_opts)
            model.fit(dtm)
            terms = vectorizer.get_feature_names_out()
            return [{
                'topic': f'Topic {idx + 1}',
                'terms': [terms[i] for i in topic.argsort()[:-10-1:-1]],
                'weight': topic.sum()
            } for idx, topic in enumerate(model.components_)]
        except Exception as e:
            logger.error(f"{model_type.upper()} Analysis error: {str(e)}")
            raise ValueError(f"Topic analysis failed: {str(e)}")
    def analyze_topics_lda(self, text, num_topics=5):
        """LDA topic analysis wrapper"""
        return self._process_topic_modeling(text, num_topics, 'lda')
    def analyze_topics_nmf(self, text, num_topics=5):
        """NMF topic analysis wrapper"""
        return self._process_topic_modeling(text, num_topics, 'nmf')
    def extract_keywords(self, text, method='tfidf', num_keywords=10):
        """Keyword extraction using various methods"""
        if method == 'tfidf':
            return self._extract_tfidf(text, num_keywords)
        elif method == 'yake':
            return self._extract_yake(text, num_keywords)
        elif method == 'rake':
            return self._extract_rake(text, num_keywords)
        else:
            raise ValueError(f"Unknown keyword extraction method: {method}")
    def _extract_yake(self, text, num_keywords):
        """YAKE extraction"""
        kw_extractor = yake.KeywordExtractor(
            lan="en",
            n=2,
            windowsSize=2,
            top=num_keywords,
            stopwords=list(self.stop_words) if hasattr(self, 'stop_words') else None
        )
        try:
            keywords = kw_extractor.extract_keywords(text)
            return [{'keyword': kw[0], 'score': 1-kw[1]} for kw in keywords]
        except Exception:
            return [{'keyword': 'Error: YAKE extraction failed', 'score': 0.0}]
    def _extract_rake(self, text, num_keywords):
        """RAKE extraction"""
        try:
            rake = rake_nltk.Rake(stopwords=list(self.stop_words) if hasattr(self, 'stop_words') else None)
            rake.extract_keywords_from_text(text)
            keywords = rake.get_ranked_phrases_with_scores()
            keywords.sort(reverse=True)
            return [{'keyword': kw[1], 'score': kw[0]} for kw in keywords[:num_keywords]]
        except Exception:
            return [{'keyword': 'Error: RAKE extraction failed', 'score': 0.0}]
    def load_stopwords_file(self):
        """Load stopwords from a text file"""
        try:
            options = QFileDialog.Options()
            file_path, _ = QFileDialog.getOpenFileName(
                self, 
                "Load Stopwords File", 
                "", 
                "Text Files (*.txt);;All Files (*)", 
                options=options
            )
            if file_path:
                with open(file_path, 'r', encoding='utf-8') as file:
                    stopwords = file.read().strip()
                    current_text = self.stopword_entry.toPlainText().strip()
                    if (current_text):
                        stopwords = current_text + "\n" + stopwords
                    self.stopword_entry.setPlainText(stopwords)
                    QMessageBox.information(self, "Success", "Stopwords loaded successfully!")
        except Exception as e:
            QMessageBox.critical(self, "Error", f"Failed to load stopwords file: {str(e)}")
    def show_token_warning(self, tokens):
        """Show non-blocking warning message about detected tokens"""
        msg = QMessageBox(self)
        msg.setIcon(QMessageBox.Information)
        msg.setWindowTitle("Analysis Notification")
        msg.setWindowModality(Qt.NonModal)
        msg.setTextFormat(Qt.TextFormat.RichText)
        message = f"""
        <b>Inconsistent tokens detected:</b> "{tokens}"<br><br>
        These tokens were detected by the tokenizer but are not found in your stopwords list.  
        This is not necessarily an issue, but it may indicate:<br>
        • Some unimportant words slipping through.<br>
        • Potential noise in your dataset.<br><br>
        <b>No immediate action is required.</b> However, if you'd like to refine your stopwords list, you may consider the following:<br>
        1. Review these tokens.<br>
        2. Add them to the stopwords list if they are irrelevant.<br>
        3. Regenerate the analysis for improved results.<br><br>
        <i>You can add them directly via the stopwords entry field.</i>
        """
        msg.setText(message)
        msg.setStandardButtons(QMessageBox.Ok)
        msg.show()
    def summarize_text(self):
        """Show summary dialog and start summarization"""
        try:
            if not hasattr(self, 'text_data') or not self.text_data:
                QMessageBox.warning(self, "Warning", "No text available to summarize")
                return
            if not hasattr(self, '_summary_dialog'):
                self._summary_dialog = SummaryDialog(self)
            self._summary_dialog.set_text(self.text_data)
            self._summary_dialog.show()
        except Exception as e:
            logger.error(f"Error showing summary dialog: {e}")
            QMessageBox.critical(self, "Error", f"Failed to show summary: {str(e)}")
            self.set_progress('summarizing', visible=False)
    def get_score_type(self):
        return AppConstants.SCORE_TYPES.get(self.mode, "Score")
    def create_worker_thread(self, worker, on_finished=None, on_error=None, thread_name=None):
        """
        Create and configure a worker thread with proper error handling and cleanup
        Args:
            worker (QRunnable): The worker to run in the thread
            on_finished (callable, optional): Callback when thread finishes successfully
            on_error (callable, optional): Callback when thread encounters an error
            thread_name (str, optional): Name for the thread for debugging purposes
        Returns:
            QThread: The configured thread object
        """
        thread = QThread()
        if thread_name:
            thread.setObjectName(thread_name)
        worker.moveToThread(thread)
        thread.started.connect(worker.run)
        worker.finished.connect(thread.quit)
        worker.finished.connect(worker.deleteLater)
        thread.finished.connect(thread.deleteLater)
        if on_finished:
            worker.finished.connect(on_finished)
        if on_error:
            worker.error.connect(on_error)
        thread.finished.connect(lambda: self.thread_manager.remove_thread(thread))
        self.thread_manager.add_thread(thread)
        return thread
    def register_cached_functions(self):
        """Register caches with cache manager"""
        if not hasattr(self, 'cache_manager'):
            return
        try:
            caches = {
                'get_wordcloud': self.get_wordcloud,
                'get_most_frequent_words': self.get_most_frequent_words,
                'load_stopwords': load_stopwords,
                'token_opts': 'property'
            }
            for cache_id, cache_obj in caches.items():
                if cache_obj == 'property':
                    self.cache_manager.register_cached_property(self, cache_id)
                elif hasattr(cache_obj, 'cache_clear'):
                    self.cache_manager.register_cached_function(cache_obj)
            models = ['_flair_model', '_vader_analyzer']
            for model_name in models:
                if hasattr(self, model_name):
                    self.cache_manager.register_model_cache(getattr(self, model_name))
        except Exception as e:
            logger.error(f"Error registering caches: {e}")         
    def detect_file_format(self, file_path, content):
        """
        Deteksi format file dan berikan penanganan khusus jika diperlukan
        Args:
            file_path (str): Path file
            content (str): Isi file
        Returns:
            tuple: (processed_content, file_type)
        """
        ext = Path(file_path).suffix.lower()
        if ext == '.csv':
            delimiters = [',', ';', '\t', '|']
            delimiter = ','
            for d in delimiters:
                if d in content[:1000]:
                    delimiter = d
                    break
            reply = QMessageBox.question(
                self,
                "CSV File Detected",
                f"This appears to be a CSV file with delimiter '{delimiter}'.\n\n"
                "How would you like to process it?",
                "Extract Text Column", "Join All Columns", "Raw Text",
                defaultButtonNumber=0
            )
            if reply == 0:
                column_index, ok = QInputDialog.getInt(
                    self, "Select Column", "Enter the column number (0-based):", 0, 0, 100, 1
                )
                if ok:
                    try:
                        import csv
                        from io import StringIO
                        rows = []
                        reader = csv.reader(StringIO(content), delimiter=delimiter)
                        for row in reader:
                            if len(row) > column_index:
                                rows.append(row[column_index])
                        return ("\n".join(rows), "csv_column")
                    except Exception as e:
                        logger.error(f"Error processing CSV column: {e}")
            elif reply == 1:
                try:
                    import csv
                    from io import StringIO
                    rows = []
                    reader = csv.reader(StringIO(content), delimiter=delimiter)
                    for row in reader:
                        rows.append(" ".join(row))
                    return ("\n".join(rows), "csv_joined")
                except Exception as e:
                    logger.error(f"Error joining CSV columns: {e}")
        elif ext in ['.html', '.xml']:
            reply = QMessageBox.question(
                self,
                f"{ext.upper()[1:]} File Detected",
                f"This appears to be a {ext.upper()[1:]} file.\n\n"
                "Would you like to extract plain text from it?",
                QMessageBox.Yes | QMessageBox.No,
                QMessageBox.Yes
            )
            if reply == QMessageBox.Yes:
                try:
                    from bs4 import BeautifulSoup
                    soup = BeautifulSoup(content, 'html.parser')
                    for script in soup(["script", "style"]):
                        script.extract()
                    text = soup.get_text(separator="\n")
                    lines = [line.strip() for line in text.splitlines() if line.strip()]
                    text = "\n".join(lines)
                    return (text, f"{ext[1:]}_extracted")
                except ImportError:
                    QMessageBox.warning(
                        self,
                        "Module Missing",
                        "BeautifulSoup module is required to extract text from HTML/XML.\n"
                        "Using raw text instead."
                    )
                except Exception as e:
                    logger.error(f"Error extracting text from {ext}: {e}")
        return (content, "raw_text")
    def open_file_dialog(self):
        """Open file dialog with path validation"""
        options = QFileDialog.Options()
        file_filter = (
            "Text Files (*.txt);;PDF Files (*.pdf);;Word Documents (*.doc *.docx);;"
            "CSV Files (*.csv);;Excel Files (*.xlsx *.xls);;All Files (*)"
        )
        file_path, _ = QFileDialog.getOpenFileName(
            self, "Open File", "", file_filter, options=options
        )
        if file_path:
            try:
                validated_path = PathValidator.sanitize_path(file_path)
                ext = Path(validated_path).suffix.lower()
                supported_exts = ['.txt', '.pdf', '.doc', '.docx', '.csv', '.xlsx', '.xls']
                if ext not in supported_exts:
                    QMessageBox.warning(
                        self,
                        "Unsupported File Type",
                        f"The file type '{ext}' is not supported. Please select a supported file type."
                    )
                    return
                if ext == '.txt':
                    content = FileOperations.load_text_file(self, validated_path)
                elif ext in ['.csv']:
                    content = self.load_csv_file(validated_path)
                elif ext in ['.xlsx', '.xls']:
                    content = self.load_excel_file(validated_path)
                elif ext in ['.pdf']:
                    content = self.load_pdf_file(validated_path)
                elif ext in ['.doc', '.docx']:
                    content = self.load_word_file(validated_path)
                else:
                    content = None
                if content is not None:
                    self.text_data = content
                    self.file_path = validated_path
                    self.text_edit.setPlainText(content)
                    self.statusBar().showMessage(
                        f"Loaded: {os.path.basename(validated_path)}"
                    )
            except ValueError as e:
                QMessageBox.critical(self, "Security Error", str(e))
                logger.warning(f"Blocked access to file: {file_path}")
    def load_csv_file(self, file_path):
        """Load and process CSV file"""
        try:
            import pandas as pd
            with open(file_path, 'r', encoding='utf-8') as f:
                sample = f.read(1024)
            delimiters = [',', ';', '\t', '|']
            delimiter = ','
            for d in delimiters:
                if d in sample:
                    delimiter = d
                    break
            reply = QMessageBox.question(
                self,
                "CSV File Options",
                f"How would you like to process this CSV file?\nDetected delimiter: '{delimiter}'",
                "Extract Column", "Join All Columns", "Cancel",
                defaultButtonNumber=0
            )
            if reply == 0:
                df = pd.read_csv(file_path, delimiter=delimiter, nrows=0)
                columns = df.columns.tolist()
                column, ok = QInputDialog.getItem(
                    self, "Select Column", "Choose text column to extract:", 
                    columns, 0, False
                )
                if ok and column:
                    df = pd.read_csv(file_path, delimiter=delimiter, usecols=[column])
                    return "\n".join(df[column].astype(str).tolist())
                return None
            elif reply == 1:
                df = pd.read_csv(file_path, delimiter=delimiter)
                result = []
                for _, row in df.iterrows():
                    result.append(" ".join(row.astype(str).tolist()))
                return "\n".join(result)
            else:
                return None
        except Exception as e:
            QMessageBox.critical(self, "Error", f"Failed to load CSV file: {str(e)}")
            logger.error(f"Error loading CSV file {file_path}: {e}")
            return None
    def load_excel_file(self, file_path):
        """Load and process Excel file"""
        try:
            import pandas as pd
            xl = pd.ExcelFile(file_path)
            if len(xl.sheet_names) > 1:
                sheet, ok = QInputDialog.getItem(
                    self, "Select Sheet", "Choose Excel sheet to load:", 
                    xl.sheet_names, 0, False
                )
                if not ok:
                    return None
            else:
                sheet = xl.sheet_names[0]
            df = pd.read_excel(file_path, sheet_name=sheet)
            reply = QMessageBox.question(
                self,
                "Excel File Options",
                "How would you like to process this Excel file?",
                "Extract Column", "Join All Columns", "Cancel",
                defaultButtonNumber=0
            )
            if reply == 0:
                columns = df.columns.tolist()
                column, ok = QInputDialog.getItem(
                    self, "Select Column", "Choose text column to extract:", 
                    columns, 0, False
                )
                if ok and column:
                    return "\n".join(df[column].astype(str).tolist())
                return None
            elif reply == 1:
                result = []
                for _, row in df.iterrows():
                    result.append(" ".join(row.astype(str).tolist()))
                return "\n".join(result)
            else:
                return None
        except Exception as e:
            QMessageBox.critical(self, "Error", f"Failed to load Excel file: {str(e)}")
            logger.error(f"Error loading Excel file {file_path}: {e}")
            return None
    def load_pdf_file(self, file_path):
        """Load and process PDF file"""
        try:
            try:
                import PyPDF2
            except ImportError:
                QMessageBox.critical(
                    self, 
                    "Missing Dependency", 
                    "PyPDF2 is required to read PDF files. Please install it with 'pip install PyPDF2'."
                )
                return None
            with open(file_path, 'rb') as file:
                reader = PyPDF2.PdfReader(file)
                if len(reader.pages) > 1:
                    options = ["All Pages", "Select Page Range", "Cancel"]
                    reply = QMessageBox.question(
                        self,
                        "PDF Options",
                        f"This PDF has {len(reader.pages)} pages. How would you like to proceed?",
                        *options,
                        defaultButtonNumber=0
                    )
                    if reply == 0:
                        text = ""
                        for page in reader.pages:
                            text += page.extract_text() + "\n\n"
                        return text
                    elif reply == 1:
                        start, ok1 = QInputDialog.getInt(
                            self, "Start Page", "Enter start page (1-based):", 
                            1, 1, len(reader.pages), 1
                        )
                        if not ok1:
                            return None
                        end, ok2 = QInputDialog.getInt(
                            self, "End Page", "Enter end page:", 
                            min(start + 9, len(reader.pages)), start, len(reader.pages), 1
                        )
                        if not ok2:
                            return None
                        text = ""
                        for i in range(start - 1, end):
                            text += reader.pages[i].extract_text() + "\n\n"
                        return text
                    else:
                        return None
                else:
                    return reader.pages[0].extract_text()
        except Exception as e:
            QMessageBox.critical(self, "Error", f"Failed to load PDF file: {str(e)}")
            logger.error(f"Error loading PDF file {file_path}: {e}")
            return None
    def load_word_file(self, file_path):
        """Load and process Word document"""
        try:
            try:
                import docx
            except ImportError:
                QMessageBox.critical(
                    self, 
                    "Missing Dependency", 
                    "python-docx is required to read Word documents. Please install it with 'pip install python-docx'."
                )
                return None
            doc = docx.Document(file_path)
            text = "\n".join([paragraph.text for paragraph in doc.paragraphs if paragraph.text])
            if doc.tables:
                reply = QMessageBox.question(
                    self,
                    "Word Document Tables",
                    f"This document contains {len(doc.tables)} tables. Include table content?",
                    QMessageBox.Yes | QMessageBox.No,
                    QMessageBox.Yes
                )
                if reply == QMessageBox.Yes:
                    for table in doc.tables:
                        for row in table.rows:
                            row_text = " | ".join([cell.text for cell in row.cells if cell.text])
                            if row_text:
                                text += "\n" + row_text
            return text
        except Exception as e:
            QMessageBox.critical(self, "Error", f"Failed to load Word document: {str(e)}")
            logger.error(f"Error loading Word document {file_path}: {e}")
            return None
    def create_custom_colormap(self, colors):
        """Membuat colormap kustom"""
        try:
            if not isinstance(colors, list):
                colors = list(colors)
            colormap_name = f"custom_{'_'.join(c.replace('#','') for c in colors)}"
            n_bins = 256
            colormap = LinearSegmentedColormap.from_list(colormap_name, colors, N=n_bins)
            return colormap
        except Exception as e:
            logger.error(f"Error creating custom colormap: {e}")
            return None
    def apply_custom_palette(self):
        """Menerapkan palette warna kustom"""
        try:
            colors = self.custom_color_palettes.get(self.color_theme.currentText(), [])
            if colors:
                colormap = self.create_custom_colormap(colors)
                if colormap:
                    self._cached_colormaps[self.color_theme.currentText()] = colormap
                    return colormap
        except Exception as e:
            logger.error(f"Error applying custom palette: {e}")
        return None
    def _handle_thread_added(self, thread):
        """Handle thread added signal"""
        pass
    def _handle_thread_removed(self, thread):
        """Handle thread removed signal"""
        pass
    def _handle_all_threads_stopped(self):
        """Handle all threads stopped signal"""
        pass
    def _handle_thread_error(self, error_msg):
        """Handle thread error signal"""
        logger.error(f"Thread error: {error_msg}")
    def add_managed_thread(self, thread):
        """Legacy method for backward compatibility"""
        self.thread_manager.add_thread(thread)
    def remove_managed_thread(self, thread):
        """Legacy method for backward compatibility"""
        self.thread_manager.remove_thread(thread)
    def _on_colormap_loading_started(self):
        """Handle colormap loading started"""
        self.color_theme.clear()
        self.color_theme.addItem("Loading colormaps...")
        self.color_theme.setEnabled(False)
        self.statusBar().showMessage("Loading color palettes...", 2000)
    def _on_colormap_loading_progress(self, colormap_name):
        """Handle colormap loading progress"""
        self.statusBar().showMessage(f"Loading color palette: {colormap_name}", 500)
    def _on_colormap_loading_finished(self, colormap_names):
        """Handle colormap loading finished"""
        self.color_theme.clear()
        self.color_theme.addItems(colormap_names)
        self.color_theme.setEnabled(True)
        self.statusBar().showMessage("Color palettes loaded", 2000)
        default_colormap = "viridis"
        if default_colormap in colormap_names:
            self.color_theme.setCurrentText(default_colormap)
    def update_cleanup_status(self, message):
        """Update status bar dengan pesan cleanup"""
        if hasattr(self, 'statusBar'):
            self.statusBar().showMessage(message, 3000)
        logger.info(message)
    def show_cleanup_error(self, error_message):
        """Tampilkan error message saat cleanup"""
        logger.error(f"Cleanup error: {error_message}")
        if hasattr(self, 'statusBar'):
            self.statusBar().showMessage(f"Error: {error_message}", 5000)            
    def _stop_active_operations(self):
        """Stop all running operations"""
        logger.info("Stopping all active operations...")
        try:
            threads = ['sentiment_thread', 'wordcloud_thread', 'text_stats_thread', 
                      'topic_analysis_thread', 'flair_loader_thread']
            active_threads = 0
            for thread_name in threads:
                thread = getattr(self, thread_name, None)
                if thread and thread.isRunning():
                    active_threads += 1
                    if hasattr(thread, 'cancel'):
                        thread.cancel()
                    elif hasattr(thread, 'requestInterruption'):
                        thread.requestInterruption()
            logger.info(f"Stopping all active threads ({active_threads})")
            if hasattr(self, 'thread_manager'):
                self.thread_manager.stop_all_threads()
        except Exception as e:
            logger.error(f"Error stopping operations: {e}")
    def _cleanup_resources(self):
        """Cleanup all application resources"""
        try:
            if hasattr(self, 'cache_manager'):
                self.cache_manager.cleanup_on_exit()
            import matplotlib.pyplot as plt
            plt.close('all')
            if hasattr(self, '_api_session'):
                self._api_session.close()
            caches_to_clear = [
                ('_translation_cache', 'clear'),
                ('_download_cache', 'clear'),
                ('_font_cache', 'clear'),
                ('get_wordcloud', 'cache_clear'),
                ('get_most_frequent_words', 'cache_clear')
            ]
            for attr_name, method_name in caches_to_clear:
                if hasattr(self, attr_name):
                    attr = getattr(self, attr_name)
                    if hasattr(attr, method_name):
                        getattr(attr, method_name)()
        except Exception as e:
            logger.error(f"Error cleaning resources: {e}")            
    def _remove_translation_cache(self):
        """Remove translation cache files"""
        if hasattr(self, 'sentiment_thread') and self.sentiment_thread:
            if hasattr(self.sentiment_thread, 'temp_file') and self.sentiment_thread.temp_file.exists():
                try:
                    self.sentiment_thread.temp_file.unlink()
                    logger.debug("Translation cache file removed")
                except Exception as e:
                    logger.error(f"Error removing translation cache: {e}")
    def _close_network_connections(self):
        """Tutup semua koneksi jaringan yang mungkin masih terbuka"""
        try:
            if hasattr(self, 'sentiment_thread') and self.sentiment_thread:
                if hasattr(self.sentiment_thread, 'cached_translation'):
                    self.sentiment_thread.cached_translation = None
            try:
                import socket
                socket.setdefaulttimeout(1)
            except:
                pass
            try:
                import asyncio
                for task in asyncio.all_tasks(loop=asyncio.get_event_loop()):
                    if not task.done() and not task.cancelled():
                        task.cancel()
            except (ImportError, RuntimeError, AttributeError):
                pass
        except Exception as e:
            logger.error(f"Error closing network connections: {e}")
    def _close_matplotlib_figures(self):
        """Menutup semua jendela matplotlib"""
        try:
            import matplotlib.pyplot as plt
            plt.close('all')
            if hasattr(self, 'wordcloud_service'):
                self.wordcloud_service.cleanup()
            logger.info("Successfully closed all matplotlib figures")
        except Exception as e:
            logger.error(f"Error closing matplotlib figures: {e}")
    def _clean_temp_files(self):
        """Bersihkan file temporary aplikasi dan direktori cache"""
        try:
            temp_dir = Path(tempfile.gettempdir()) / "textplora_cache"
            if temp_dir.exists():
                try:
                    file_count = sum(1 for _ in temp_dir.glob('**/*') if _.is_file())
                    shutil.rmtree(temp_dir, ignore_errors=True)
                    logger.info(f"Cleaned temp directory: {temp_dir} ({file_count} files removed)")
                except Exception as e:
                    logger.error(f"Error cleaning temp directory: {e}")
            if hasattr(self, 'sentiment_thread') and self.sentiment_thread:
                if hasattr(self.sentiment_thread, 'temp_file') and self.sentiment_thread.temp_file.exists():
                    try:
                        self.sentiment_thread.temp_file.unlink()
                        logger.info(f"Removed translation cache: {self.sentiment_thread.temp_file}")
                    except Exception as e:
                        logger.error(f"Error removing translation cache: {e}")
            app_dir = Path(os.path.dirname(__file__))
            cache_dirs = [
                app_dir / 'temp',
                app_dir / 'cache' / 'translations',
                app_dir / 'cache' / 'downloads'
            ]
            for directory in cache_dirs:
                if directory.exists():
                    try:
                        file_count = sum(1 for _ in directory.glob('**/*') if _.is_file())
                        shutil.rmtree(directory, ignore_errors=True)
                        logger.info(f"Cleaned cache directory: {directory} ({file_count} files removed)")
                    except Exception as e:
                        logger.error(f"Error cleaning cache directory {directory}: {e}")
            try:
                additional_cache_files = []
                for file_path in additional_cache_files:
                    if os.path.exists(file_path):
                        os.remove(file_path)
                        logger.info(f"Removed additional cache file: {file_path}")
            except Exception as e:
                logger.error(f"Error cleaning additional cache files: {e}")
            logger.info("Temporary files cleanup completed")
        except Exception as e:
            logger.error(f"Error in _clean_temp_files: {e}")
    def cleanup_font_resources(self):
        """Membersihkan resource font saat aplikasi ditutup"""
        logger.debug("AppUI: Cleaning up font resources")
        if hasattr(self, '_font_loader_thread') and self._font_loader_thread:
            if self._font_loader_thread.isRunning():
                logger.debug("AppUI: Stopping font loader thread")
                self._font_loader_thread.requestInterruption()
                self._font_loader_thread.wait(2000)
                if self._font_loader_thread.isRunning():
                    logger.warning("AppUI: Font loader thread did not stop gracefully, terminating")
                    self._font_loader_thread.terminate()
        if hasattr(self, '_fonts_loaded') and self._fonts_loaded and self._cached_fonts:
            if not getattr(self, '_fonts_saved_to_cache', False):
                logger.debug("AppUI: Saving fonts to cache before exit")
                self._save_fonts_to_cache(self._cached_fonts)
        logger.debug("AppUI: Font cleanup completed")
    def cleanup(self):
        """Main cleanup method - mengkonsolidasikan semua proses pembersihan"""
        if self._cleanup_complete or self._cleanup_in_progress:
            logger.debug("Cleanup already performed or in progress, skipping...")
            return
        logger.info("Starting application cleanup...")
        self._cleanup_in_progress = True
        try:
            self._stop_active_operations()
            if hasattr(self, 'thread_manager'):
                self.thread_manager.stop_all_threads(wait_timeout=1000)
            if hasattr(self, 'cache_manager'):
                self.cache_manager.cleanup_on_exit()
            if hasattr(self, 'resource_controller'):
                self.resource_controller.cleanup_all_resources()
            self._clean_temp_files()
            self._remove_translation_cache()
            self._close_network_connections()
            self._close_matplotlib_figures()
            gc.collect()
            logger.info("Cleanup completed")
            self._cleanup_complete = True
        except Exception as e:
            logger.error(f"Cleanup error: {e}")
        finally:
            self._cleanup_in_progress = False                   
    def closeEvent(self, event):
        """Handle application shutdown"""
        try:
            if self._cleanup_complete:
                event.accept()
                return
            if self._cleanup_in_progress:
                logger.debug("Cleanup already in progress, waiting...")
                event.accept()
                return
            reply = QMessageBox.question(self, "Confirm Exit", 
                "Are you sure you want to exit?",
                QMessageBox.Yes | QMessageBox.No, QMessageBox.No)
            if reply == QMessageBox.No:
                event.ignore()
                return
            self._cleanup_in_progress = True
            logger.info("Starting cleanup during application shutdown...")
            self._close_matplotlib_figures()
            self._clean_temp_files()
            self._remove_translation_cache()
            self.cleanup()
            self._cleanup_complete = True
            logger.info("Cleanup completed. Application shutting down.")
            event.accept()
        except Exception as e:
            logger.error(f"Error during shutdown: {e}")
            self._cleanup_in_progress = False
            event.accept()     
"""FUNGSI MAIN/ENTRY POINT"""
if __name__ == "__main__":
    import matplotlib
    matplotlib.use("QtAgg")
    app = QApplication(sys.argv)
    loop = QEventLoop(app)
    asyncio.set_event_loop(loop)
    ex = AppUI()
    ex.show()
    app.aboutToQuit.connect(ex.cleanup)
    with loop:
        sys.exit(loop.run_forever())
