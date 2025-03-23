import sys
import os
import re
from pathlib import Path
from functools import lru_cache
import torch
os.environ["QT_API"] = "pyside6"
import asyncio
from collections import Counter
from qasync import QEventLoop
from deep_translator import GoogleTranslator
from PySide6.QtWidgets import (
    QApplication, QMainWindow, QFileDialog, QMessageBox, QLabel, QPushButton,
    QVBoxLayout, QGridLayout, QWidget, QLineEdit, QComboBox, QSpinBox, QDialog,
    QTextEdit, QProgressBar, QFrame, QColorDialog, QInputDialog, QTextBrowser,
)
from PySide6.QtCore import QThread, Signal, Qt, QTimer, QMutex, QRunnable, QThreadPool
from PySide6.QtGui import QIcon
import matplotlib
matplotlib.use("QtAgg")
from wordcloud import WordCloud, STOPWORDS
import numpy as np
from PIL import Image
import socket
import base64

# App icon
APP_DIR = Path(__file__).parent
ICON_PATH = APP_DIR / "icon.ico"

ABOUT_TEXT = "PGgyPvCfjJ8gV0NHZW4gLSBXb3JkIENsb3VkIEdlbmVyYXRvciArIFNlbnRpbWVudCBBbmFseXNpczwvaDI+CjxwPjxiPlZlcnNpb246PC9iPiAxLjU8L3A+CjxwPiZjb3B5OyAyMDI1IE0uIEFkaWIgWmF0YSBJbG1hbTwvcD4KPHA+PGEgaHJlZj0iaHR0cHM6Ly9naXRodWIuY29tL3phdGFpbG0vd2Nsb3VkZ3VpIj7wn5OMIEdpdEh1YiBSZXBvc2l0b3J5PC9hPjwvcD4KPGgzPvCflI0gV2hhdCBpcyBXQ0dlbj88L2gzPgo8cD5XQ0dlbiBpcyBhbiBhZHZhbmNlZCB5ZXQgZWFzeS10by11c2UgYXBwbGljYXRpb24gZGVzaWduZWQgZm9yIGNyZWF0aW5nIHZpc3VhbGx5IGFwcGVhbGluZyB3b3JkIGNsb3VkcyBmcm9tIHRleHQgZGF0YS48L3A+CjxoMz7wn5qAIEtleSBGZWF0dXJlczwvaDM+Cjx1bD4KICAgIDxsaT48Yj5TdXBwb3J0cyBtdWx0aXBsZSBmaWxlIGZvcm1hdHM6PC9iPiBUWFQsIFBERiwgRE9DL0RPQ1gsIENTViwgWExTWDwvbGk+CiAgICA8bGk+PGI+RnVsbHkgY3VzdG9taXphYmxlIHdvcmQgY2xvdWRzOjwvYj4gQ2hvb3NlIGNvbG9ycywgZm9udHMsIHNoYXBlcywgYW5kIHRoZW1lczwvbGk+CiAgICA8bGk+PGI+U3RvcHdvcmQgZmlsdGVyaW5nOjwvYj4gUmVtb3ZlIGNvbW1vbiB3b3JkcyBmb3IgY2xlYXJlciBpbnNpZ2h0czwvbGk+CiAgICA8bGk+PGI+U21hcnQgdGV4dCBwcm9jZXNzaW5nOjwvYj4gSGFuZGxlcyBsYXJnZSBkYXRhc2V0cyBlZmZpY2llbnRseTwvbGk+CiAgICA8bGk+PGI+RXhwb3J0ICYgc2F2ZSBvcHRpb25zOjwvYj4gSGlnaC1yZXNvbHV0aW9uIGltYWdlIHNhdmluZyBpbiBtdWx0aXBsZSBmb3JtYXRzPC9saT4KICAgIDxsaT48Yj5TZW50aW1lbnQgYW5hbHlzaXM6PC9iPiBPcHRpb25hbCBhbmFseXNpcyB1c2luZyBUZXh0QmxvYiwgVkFERVIsIGFuZCBGbGFpcjwvbGk+CjwvdWw+CjxoMz7wn5OWIEhvdyBXQ0dlbiBIZWxwcyBZb3U8L2gzPgo8cD5XaGV0aGVyIHlvdeKAmXJlIGFuYWx5emluZyBjdXN0b21lciBmZWVkYmFjaywgY29uZHVjdGluZyBhY2FkZW1pYyByZXNlYXJjaCwgb3IgdmlzdWFsaXppbmcgdGV4dC1iYXNlZCBpbnNpZ2h0czwgV0NHZW4gc2ltcGxpZmllcyB0aGUgcHJvY2VzcyBhbmQgZW5oYW5jZXMgeW91ciB3b3JrZmxvdyB3aXRoIGl0cyBpbnR1aXRpdmUgZGVzaWduIGFuZCBwb3dlcmZ1bCBmZWF0dXJlcy48L3A+CjxoMz7wn5OcIExpY2Vuc2U8L2gzPgo8cD5XQ0dlbiBpcyBmcmVlIGZvciBwZXJzb25hbCBhbmQgZWR1Y2F0aW9uYWwgdXNlLiBGb3IgY29tbWVyY2lhbCBhcHBsaWNhdGlvbnMsIHBsZWFzZSByZWZlciB0byB0aGUgbGljZW5zaW5nIHRlcm1zLjwvcD4KPGgzPvCfk5ogSG93IHRvIENpdGUgV0NHZW48L2gzPgo8cD5JZiB5b3UgdXNlIFdDR2VuIGluIHlvdXIgcmVzZWFyY2ggb3IgcHVibGljYXRpb24sIHBsZWFzZSBjaXRlIGl0IGFzIGZvbGxvd3MgKEFQQSA3KTo8L3A+CjxwPklsbWFtLCBNLiBBLiBaLiAoMjAyNSkuIDxpPldDR2VuIC0gV29yZCBDbG91ZCBHZW5lcmF0b3IgKyBTZW50aW1lbnQgQW5hbHlzaXM8L2k+IChWZXJzaW9uIDEuNSkgW1NvZnR3YXJlXS4gWmVub2RvLiA8YSBocmVmPSJodHRwczovL2RvaS5vcmcvMTAuNTI4MS96ZW5vZG8uMTUwMzQ4NDMiPmh0dHBzOi8vZG9pLm9yZy8xMC41MjgxL3plbm9kby4xNDkzMjY1MDwvYT48L3A+Cg=="

MODE_INFO = "PGgyPlNlbnRpbWVudCBBbmFseXNpcyBNb2RlczwvaDI+CjxwPlNlbGVjdCB0aGUgbW9zdCBzdWl0YWJsZSBzZW50aW1lbnQgYW5hbHlzaXMgbWV0aG9kIGJhc2VkIG9uIHlvdXIgdGV4dCB0eXBlIGFuZCBhbmFseXNpcyBuZWVkcy48L3A+CjxoMz7wn5OdIFRleHRCbG9iPC9oMz4KPHA+PGI+QmVzdCBmb3I6PC9iPiBGb3JtYWwgdGV4dHMsIHdlbGwtc3RydWN0dXJlZCBkb2N1bWVudHMsIG5ld3MgYXJ0aWNsZXMsIHJlc2VhcmNoIHBhcGVycywgYW5kIHJlcG9ydHMuPC9wPgo8cD5UZXh0QmxvYiBpcyBhIGxleGljb24tYmFzZWQgc2VudGltZW50IGFuYWx5c2lzIHRvb2wgdGhhdCBwcm92aWRlcyBhIHNpbXBsZSB5ZXQgZWZmZWN0aXZlIGFwcHJvYWNoIGZvciBldmFsdWF0aW5nIHRoZSBzZW50aW1lbnQgb2Ygc3RydWN0dXJlZCB0ZXh0LiBJdCBhc3NpZ25zIGEgcG9sYXJpdHkgc2NvcmUgKHBvc2l0aXZlLCBuZWdhdGl2ZSwgb3IgbmV1dHJhbCkgYW5kIGNhbiBhbHNvIGFuYWx5emUgc3ViamVjdGl2aXR5IGxldmVscy48L3A+CjxoMz7wn5KsIFZBREVSIChWYWxlbmNlIEF3YXJlIERpY3Rpb25hcnkgYW5kIHNFbnRpbWVudCBSZWFzb25lcik8L2gzPgo8cD48Yj5CZXN0IGZvcjo8L2I+IFNvY2lhbCBtZWRpYSBwb3N0cywgdHdlZXRzLCBzaG9ydCBjb21tZW50cywgY2hhdCBtZXNzYWdlcywgYW5kIGluZm9ybWFsIHJldmlld3MuPC9wPgo8cD5WQURFUiBpcyBzcGVjaWZpY2FsbHkgZGVzaWduZWQgZm9yIGFuYWx5emluZyBzaG9ydCwgaW5mb3JtYWwgdGV4dHMgdGhhdCBvZnRlbiBjb250YWluIHNsYW5nLCBlbW9qaXMsIGFuZCBwdW5jdHVhdGlvbi1iYXNlZCBlbW90aW9ucy4gSXQgaXMgYSBydWxlLWJhc2VkIHNlbnRpbWVudCBhbmFseXNpcyBtb2RlbCB0aGF0IGVmZmljaWVudGx5IGRldGVybWluZXMgc2VudGltZW50IGludGVuc2l0eSBhbmQgd29ya3MgZXhjZXB0aW9uYWxseSB3ZWxsIGZvciByZWFsLXRpbWUgYXBwbGljYXRpb25zLjwvcD4KPGgzPvCflKwgRmxhaXI8L2gzPgo8cD48Yj5CZXN0IGZvcjo8L2I+IExvbmctZm9ybSBjb250ZW50LCBwcm9kdWN0IHJldmlld3MsIHByb2Zlc3Npb25hbCBkb2N1bWVudHMsIGFuZCBBSS1iYXNlZCBkZWVwIHNlbnRpbWVudCBhbmFseXNpcy48L3A+CjxwPkZsYWlyIHV0aWxpemVzIGRlZXAgbGVhcm5pbmcgdGVjaG5pcXVlcyBmb3Igc2VudGltZW50IGFuYWx5c2lzLCBtYWtpbmcgaXQgaGlnaGx5IGFjY3VyYXRlIGZvciBjb21wbGV4IHRleHRzLiBJdCBpcyBpZGVhbCBmb3IgYW5hbHl6aW5nIGxhcmdlLXNjYWxlIHRleHR1YWwgZGF0YSwgY2FwdHVyaW5nIGNvbnRleHQgbW9yZSBlZmZlY3RpdmVseSB0aGFuIHRyYWRpdGlvbmFsIHJ1bGUtYmFzZWQgbW9kZWxzLiBIb3dldmVyLCBpdCByZXF1aXJlcyBtb3JlIGNvbXB1dGF0aW9uYWwgcmVzb3VyY2VzIGNvbXBhcmVkIHRvIFRleHRCbG9iIGFuZCBWQURFUi48L3A+CjxoMz7wn4yQIEltcG9ydGFudCBOb3RlIGZvciBMYW5ndWFnZSBTdXBwb3J0PC9oMz4KPHA+V2hpbGUgdGhpcyBhcHBsaWNhdGlvbiBzdXBwb3J0cyBub24tRW5nbGlzaCB0ZXh0IHRocm91Z2ggYXV0b21hdGljIHRyYW5zbGF0aW9uLCBpdCBpcyA8Yj5oaWdobHkgcmVjb21tZW5kZWQ8L2I+IHRvIHVzZSA8Yj5tYW51YWxseSB0cmFuc2xhdGVkIGFuZCByZWZpbmVkIEVuZ2xpc2ggdGV4dDwvYj4gZm9yIHRoZSBtb3N0IGFjY3VyYXRlIHNlbnRpbWVudCBhbmFseXNpcy4gVGhlIGJ1aWx0LWluIGF1dG9tYXRpYyB0cmFuc2xhdGlvbiBmZWF0dXJlIG1heSBub3QgYWx3YXlzIGZ1bmN0aW9uIGNvcnJlY3RseSwgbGVhZGluZyB0byBwb3RlbnRpYWwgbWlzaW50ZXJwcmV0YXRpb25zIG9yIGluYWNjdXJhdGUgc2VudGltZW50IHJlc3VsdHMuPC9wPgo8cD5Gb3IgdGhlIGJlc3QgcGVyZm9ybWFuY2UsIGVuc3VyZSB0aGF0IG5vbi1FbmdsaXNoIHRleHQgaXMgcHJvcGVybHkgcmV2aWV3ZWQgYW5kIGFkanVzdGVkIGJlZm9yZSBzZW50aW1lbnQgYW5hbHlzaXMuIPCfmoA8L3A+CjxoMz7wn5OMIEN1c3RvbSBMZXhpY29uIEZvcm1hdCBFeGFtcGxlPC9oMz4KPHA+QmVsb3cgaXMgYW4gZXhhbXBsZSBvZiBhIGN1c3RvbSBsZXhpY29uIGZvcm1hdCBmb3Igc2VudGltZW50IGFuYWx5c2lzOjwvcD4KPHByZSBzdHlsZT0nYmFja2dyb3VuZC1jb2xvcjojZjRmNGY0OyBwYWRkaW5nOjEwcHg7IGJvcmRlci1yYWRpdXM6NXB4Oyc+CmV4Y2VsbGVudCAgIDEuNQphd2Z1bCAgICAgIC0xLjUKbm90ICAgICAgICBuZWdhdGlvbiAgICAgICAgICMgTWFyayBhcyBuZWdhdGlvbiB3b3JkCmludGVuc2VseSAgaW50ZW5zaWZpZXI6MS43ICAjIEN1c3RvbSBpbnRlbnNpZmllciB3aXRoIG11bHRpcGxpZXIKPC9wcmU+CjxwPlRoaXMgY3VzdG9tIGxleGljb24gYWxsb3dzIGZpbmUtdHVuaW5nIG9mIHNlbnRpbWVudCBzY29yZXMgYnkgYWRkaW5nIGN1c3RvbSB3b3JkcywgbmVnYXRpb25zLCBhbmQgaW50ZW5zaWZpZXJzIHRvIGltcHJvdmUgc2VudGltZW50IGFuYWx5c2lzIGFjY3VyYWN5LjwvcD4K"

def sanitize_path(path):
    path = os.path.normpath(path)
    base_dir = os.path.abspath(os.path.dirname(__file__))
    abs_path = os.path.abspath(path)
    if not abs_path.startswith(base_dir):
        raise ValueError("Path traversal attempt detected")
    return path

@lru_cache(maxsize=128)
def load_stopwords():
    return set(STOPWORDS)

class ChunkProcessor:
    CHUNK_SIZE = 1024 * 1024  # 1MB chunks
    
    @staticmethod
    def process_file_chunks(file_path, callback):
        total_size = os.path.getsize(file_path)
        processed = 0
        
        with open(file_path, 'r', encoding='utf-8') as f:
            while True:
                chunk = f.read(ChunkProcessor.CHUNK_SIZE)
                if not chunk:
                    break
                    
                callback(chunk)
                processed += len(chunk)
                progress = (processed / total_size) * 100
                yield progress

class StartupThread(QThread):
    def run(self):
        try:
            import matplotlib.pyplot as plt
            from matplotlib.font_manager import FontProperties
        except Exception as e:
            QMessageBox.warning(None, "Error", f"StartupThread error: {e}")

class FileLoaderThread(QThread):
    file_loaded = Signal(str, str)
    file_error = Signal(str)

    def __init__(self, file_path):
        super().__init__()
        self.file_path = file_path

    def run(self):
        try:
            if self.isInterruptionRequested():
                return

            if not os.path.exists(self.file_path):
                raise FileNotFoundError(f"File not found: {self.file_path}")

            text_data = ""
            
            if self.file_path.endswith(".txt"):
                with open(self.file_path, "r", encoding="utf-8") as file:
                    while True:
                        if self.isInterruptionRequested():
                            return
                        line = file.readline()
                        if not line:
                            break
                        text_data += line

            elif self.file_path.endswith(".pdf"):
                text_data = self.extract_text_from_pdf(self.file_path)
                
            elif self.file_path.endswith((".doc", ".docx")):
                text_data = self.extract_text_from_word(self.file_path)
                
            elif self.file_path.endswith((".xlsx", ".xls")):
                text_data = self.extract_text_from_excel(self.file_path)
                
            elif self.file_path.endswith(".csv"):
                text_data = self.extract_text_from_csv(self.file_path)
                
            else:
                raise ValueError(f"Unsupported file format: {os.path.splitext(self.file_path)[1]}")

            if self.isInterruptionRequested():
                return

            if not text_data.strip():
                raise ValueError("File is empty or contains no extractable text")

            self.file_loaded.emit(self.file_path, text_data)

        except Exception as e:
            self.file_error.emit(f"Error loading {os.path.basename(self.file_path)}: {str(e)}")

    def extract_text_from_pdf(self, pdf_path):
        try:
            import pypdf
            text = ""
            with open(pdf_path, "rb") as file:
                reader = pypdf.PdfReader(file)
                if not reader.pages:
                    raise ValueError("PDF contains no readable pages")
                
                for page in reader.pages:
                    if self.isInterruptionRequested():
                        return ""
                    page_text = page.extract_text()
                    if page_text:
                        text += page_text + "\n"
            return text
        except Exception as e:
            raise RuntimeError(f"PDF processing failed: {str(e)}") from e

    def extract_text_from_word(self, word_path):
        try:
            import docx
            doc = docx.Document(word_path)
            if not doc.paragraphs:
                raise ValueError("Word document contains no readable text")
            
            text = ""
            for para in doc.paragraphs:
                if self.isInterruptionRequested():
                    return ""
                if para.text.strip():
                    text += para.text + "\n"
            return text
        except Exception as e:
            raise RuntimeError(f"Word processing failed: {str(e)}") from e

    def extract_text_from_excel(self, excel_path):
        try:
            import pandas as pd
            text_data = ""
            df = pd.read_excel(excel_path, sheet_name=None)
            if not df:
                raise ValueError("Excel file is empty or unreadable")
            
            for sheet_name, sheet_df in df.items():
                if self.isInterruptionRequested():
                    return ""
                text_data += f"\n--- {sheet_name} ---\n"
                text_data += sheet_df.to_string(index=False, header=True)
            return text_data
        except Exception as e:
            raise RuntimeError(f"Excel processing failed: {str(e)}") from e

    def extract_text_from_csv(self, csv_path):
        try:
            import pandas as pd
            df = pd.read_csv(csv_path)
            if df.empty:
                raise ValueError("CSV file is empty")
            
            if self.isInterruptionRequested():
                return ""
                
            return df.to_string(index=False, header=True)
        except Exception as e:
            raise RuntimeError(f"CSV processing failed: {str(e)}") from e

class CustomFileLoaderThread(QThread):
    file_loaded = Signal(object, bool)

    def __init__(self, file_path, file_type):
        super().__init__()
        self.file_path = file_path
        self.file_type = file_type

    def run(self):
        try:
            if self.file_type == "lexicon":
                from vaderSentiment.vaderSentiment import SentimentIntensityAnalyzer
                self.vader_analyzer = CustomVaderSentimentIntensityAnalyzer(custom_lexicon_file=self.file_path)
                self.file_loaded.emit(self.file_path, True)

            elif self.file_type == "TextBlob (Custom Lexicon)":
                from textblob import TextBlob
                self.textblob_analyzer = CustomTextBlobSentimentAnalyzer(self.file_path)
                self.file_loaded.emit(self.file_path, True)

            elif self.file_type == "model":
                from flair.models import TextClassifier
                from flair.data import Sentence  # Tambahkan import ini
                try:
                    model = TextClassifier.load(self.file_path)
                    # Verifikasi model adalah sentiment classifier
                    test = Sentence("test")
                    model.predict(test)
                    if not test.labels or test.labels[0].value not in ['POSITIVE', 'NEGATIVE', 'NEUTRAL']:
                        raise ValueError("Model is not a valid sentiment classifier")
                    self.file_loaded.emit(model, True)
                except Exception as e:
                    self.file_loaded.emit(f"Invalid sentiment model: {str(e)}", False)

        except Exception as e:
            self.file_loaded.emit(str(e), False)

class CustomVaderSentimentIntensityAnalyzer:
    def __init__(self, lexicon_file="vader_lexicon.txt", custom_lexicon_file=None):
        from vaderSentiment.vaderSentiment import SentimentIntensityAnalyzer
        self.analyzer = SentimentIntensityAnalyzer()
        if custom_lexicon_file:
            self.load_custom_lexicon(custom_lexicon_file)

    def load_custom_lexicon(self, custom_lexicon_file):
        try:
            with open(custom_lexicon_file, "r", encoding="utf-8") as file:
                for line in file:
                    parts = line.strip().split("\t")
                    if len(parts) == 2:
                        word, measure = parts
                        try:
                            self.analyzer.lexicon[word] = float(measure)
                        except ValueError:
                            pass
        except Exception as e:
            QMessageBox.warning(None, "Warning", f"Failed to load custom lexicon: {e}")

    def polarity_scores(self, text):
        return self.analyzer.polarity_scores(text)

class CustomTextBlobSentimentAnalyzer:
    def __init__(self, custom_lexicon_file=None):
        self.lexicon = {}
        self.intensifiers = {}
        self.negations = set()
        if custom_lexicon_file:
            self.load_custom_lexicon(custom_lexicon_file)

    def load_custom_lexicon(self, custom_lexicon_file):
        try:
            with open(custom_lexicon_file, "r", encoding="utf-8") as file:
                for line in file:
                    parts = line.strip().split("\t")
                    if len(parts) == 2:
                        word, measure = parts
                        if measure == "negation":
                            self.negations.add(word)
                        elif measure.startswith("intensifier:"):
                            try:
                                self.intensifiers[word] = float(measure.split(":")[1])
                            except ValueError:
                                continue
                        else:
                            try:
                                self.lexicon[word] = float(measure)
                            except ValueError:
                                continue
        except Exception as e:
            QMessageBox.warning(None, "Warning", f"Failed to load custom lexicon: {e}")

    def analyze(self, text):
        from textblob import TextBlob
        blob = TextBlob(text)
        total_polarity = 0.0
        words_with_context = []

        for sentence in blob.sentences:
            words = sentence.words
            for i, word in enumerate(words):
                word_lower = word.lower()
                current_score = self.lexicon.get(word_lower, 0.0)
                if i > 0 and words[i-1].lower() in self.negations:
                    current_score *= -1
                if i > 0 and words[i-1].lower() in self.intensifiers:
                    multiplier = self.intensifiers[words[i-1].lower()]
                    current_score *= multiplier
                words_with_context.append(current_score)

        valid_scores = [s for s in words_with_context if s != 0]
        avg_polarity = sum(valid_scores) / len(valid_scores) if valid_scores else 0.0

        return {
            'polarity': avg_polarity,
            'subjectivity': blob.sentiment.subjectivity,
            'processed_words': len(valid_scores)
        }

class SentimentAnalysisThread(QThread):
    sentiment_analyzed = Signal(dict)
    offline_warning = Signal(str)

    def __init__(self, text_data, sentiment_mode, vader_analyzer, flair_classifier, flair_classifier_cuslang, textblob_analyzer=None):
        super().__init__()
        self.text_data = text_data
        self.sentiment_mode = sentiment_mode
        self.vader_analyzer = vader_analyzer
        self.flair_classifier = flair_classifier
        self.flair_classifier_cuslang = flair_classifier_cuslang
        self.textblob_analyzer = textblob_analyzer

    async def _async_translate(self, text):
        try:
            translated = GoogleTranslator(source="auto", target="en").translate(text)
            return translated
        except Exception as e:
            return None

    def translate_text(self, text):
        try:
            loop = asyncio.get_event_loop()
        except RuntimeError:
            loop = asyncio.new_event_loop()
            asyncio.set_event_loop(loop)

        sentences = [s.strip() for s in text.split(".") if s.strip()]
        translated = []

        async def run_translations():
            for sentence in sentences:
                try:
                    result = await self._async_translate(sentence)
                    if result:
                        translated.append(result)
                except Exception as e:
                    continue
            return ". ".join(translated)

        try:
            return loop.run_until_complete(run_translations())
        except Exception as e:
            return None
        finally:
            if loop.is_running():
                loop.close()

    def analyze_textblob(self, text, use_custom=False):
        if use_custom and self.textblob_analyzer:
            analysis = self.textblob_analyzer.analyze(text)
            polarity = analysis['polarity']
            subjectivity = analysis['subjectivity']
        else:
            from textblob import TextBlob
            blob = TextBlob(text)
            polarity = blob.sentiment.polarity
            subjectivity = blob.sentiment.subjectivity

        if abs(polarity) < 0.05:
            neutral_score = 1.0
            positive_score = negative_score = 0.0
        else:
            positive_score = (polarity + 1) / 2 if polarity > 0 else 0
            negative_score = (-polarity + 1) / 2 if polarity < 0 else 0
            neutral_score = 1 - (positive_score + negative_score)

        return {
            "positive_score": positive_score,
            "negative_score": negative_score,
            "neutral_score": neutral_score,
            "compound_score": polarity,
            "sentiment_label": "POSITIVE" if polarity > 0 else "NEGATIVE" if polarity < 0 else "NEUTRAL",
            "subjectivity": subjectivity
        }

    def run(self):
        result = {
            "positive_score": 0,
            "neutral_score": 0,
            "negative_score": 0,
            "compound_score": 0,
            "sentiment_label": "N/A",
            "subjectivity": "N/A (only available in TextBlob mode)",
        }

        try:
            text_to_analyze = self.text_data
            needs_translation = False

            if self.sentiment_mode in ["VADER", "Flair", "TextBlob"]:
                try:
                    from langdetect import detect
                    detected_lang = detect(self.text_data)
                    needs_translation = detected_lang != "en"
                except:
                    needs_translation = False

            if needs_translation:
                if not is_connected():
                    self.offline_warning.emit("Translation required but no internet connection")
                    return

                translated_text = self.translate_text(self.text_data)
                if translated_text:
                    text_to_analyze = translated_text

            if self.sentiment_mode == "TextBlob":
                result.update(self.analyze_textblob(text_to_analyze))
            elif self.sentiment_mode == "TextBlob (Custom Lexicon)":
                result.update(self.analyze_textblob(self.text_data, True))

            elif self.sentiment_mode == "VADER":
                if self.vader_analyzer:
                    sentiment = self.vader_analyzer.polarity_scores(text_to_analyze)
                    total = sentiment["pos"] + sentiment["neg"] + sentiment["neu"]
                    positive_score = sentiment["pos"] / total if total > 0 else 0
                    negative_score = sentiment["neg"] / total if total > 0 else 0
                    neutral_score = sentiment["neu"] / total if total > 0 else 0

                    result.update({
                        "positive_score": positive_score,
                        "negative_score": negative_score,
                        "neutral_score": neutral_score,
                        "compound_score": sentiment["compound"],
                        "sentiment_label": "POSITIVE" if sentiment["compound"] >= 0.05 else "NEGATIVE" if sentiment["compound"] <= -0.05 else "NEUTRAL",
                    })
                else:
                    result["sentiment_label"] = "Error: VADER not initialized"

            elif self.sentiment_mode == "VADER (Custom Lexicon)":
                if self.vader_analyzer:
                    sentiment = self.vader_analyzer.polarity_scores(self.text_data)
                    result.update({
                        "positive_score": sentiment["pos"],
                        "negative_score": sentiment["neg"],
                        "neutral_score": sentiment["neu"],
                        "compound_score": sentiment["compound"],
                        "sentiment_label": "POSITIVE" if sentiment["compound"] >= 0.05 else "NEGATIVE" if sentiment["compound"] <= -0.05 else "NEUTRAL",
                    })
                else:
                    result["sentiment_label"] = "Error: Custom VADER not initialized"

            elif self.sentiment_mode == "Flair":
                from flair.data import Sentence
                sentence = Sentence(text_to_analyze)
                self.flair_classifier.predict(sentence)
                sentiment = sentence.labels[0]
                confidence = sentiment.score
                sentiment_label = sentiment.value

                if confidence < 0.55:
                    sentiment_label = "NEUTRAL"
                    positive_score = negative_score = 0.0
                    neutral_score = 1.0
                else:
                    positive_score = confidence if sentiment_label == "POSITIVE" else 0
                    negative_score = confidence if sentiment_label == "NEGATIVE" else 0
                    neutral_score = 0

                result.update({
                    "compound_score": confidence,
                    "sentiment_label": sentiment_label,
                    "positive_score": positive_score,
                    "negative_score": negative_score,
                    "neutral_score": neutral_score,
                })

            elif self.sentiment_mode == "Flair (Custom Model)":
                try:
                    from flair.data import Sentence
                    sentence = Sentence(text_to_analyze)
                    
                    if self.flair_classifier_cuslang is None:
                        raise ValueError("Custom model not loaded")
                        
                    self.flair_classifier_cuslang.predict(sentence)
                    if not sentence.labels:
                        raise ValueError("Model produced no labels")
                        
                    sentiment = sentence.labels[0]
                    confidence = sentiment.score
                    sentiment_label = sentiment.value

                    if sentiment_label not in ['POSITIVE', 'NEGATIVE', 'NEUTRAL']:
                        raise ValueError(f"Invalid label: {sentiment_label}")

                    if confidence < 0.55:
                        sentiment_label = "NEUTRAL"
                        positive_score = negative_score = 0.0
                        neutral_score = 1.0
                    else:
                        positive_score = confidence if sentiment_label == "POSITIVE" else 0
                        negative_score = confidence if sentiment_label == "NEGATIVE" else 0
                        neutral_score = 1.0 if sentiment_label == "NEUTRAL" else 0

                    result.update({
                        "compound_score": confidence,
                        "sentiment_label": sentiment_label,
                        "positive_score": positive_score,
                        "negative_score": negative_score, 
                        "neutral_score": neutral_score,
                    })
                except ImportError as e:
                    result["sentiment_label"] = "Error: Flair not installed properly"
                except Exception as e:
                    result["sentiment_label"] = f"Error: {str(e)}"

            self.sentiment_analyzed.emit(result)

        except Exception as e:
            result["sentiment_label"] = f"Error: {str(e)}"
            self.sentiment_analyzed.emit(result)

class FlairModelLoaderThread(QThread):
    model_loaded = Signal(object)
    error_occurred = Signal(str)

    _cached_model = None

    def __init__(self, model_path="sentiment"):
        super().__init__()
        self.model_path = model_path

    def run(self):
        try:
            import time
            from flair.models import TextClassifier

            if FlairModelLoaderThread._cached_model is None:
                time.sleep(0.1)
                model = TextClassifier.load(self.model_path)
                FlairModelLoaderThread._cached_model = model
            else:
                model = FlairModelLoaderThread._cached_model
            self.model_loaded.emit(model)
        except Exception as e:
            self.error_occurred.emit(str(e))

class WordCloudGenerator(QMainWindow):
    def __init__(self):
        super().__init__()
        
        # Initialize components
        self.thread_pool = QThreadPool()
        self.thread_pool.setMaxThreadCount(5)
        
        self.file_path = ""
        self.text_data = ""
        self.mask_path = ""
        self.additional_stopwords = set()
        self.custom_color_palettes = {}
        self.current_figure = None
        self.flair_first_load = True        

        from vaderSentiment.vaderSentiment import SentimentIntensityAnalyzer
        self.vader_analyzer = SentimentIntensityAnalyzer()
        self.sentiment_mode = "TextBlob"
        self.custom_lexicon_path = None
        self.custom_textblob_lexicon_path = None
        self.textblob_analyzer = None
        self.custom_model_path = None
        self.flair_classifier = None
        self.flair_classifier_cuslang = None

        self.active_threads = []
        self.threads_mutex = QMutex()

        self.startup_thread = StartupThread()
        
        self.threads_mutex.lock()
        self.active_threads.append(self.startup_thread)
        self.threads_mutex.unlock()
        
        self.startup_thread.start()

        self.initUI()

        self.cleanup_timer = QTimer(self)
        self.cleanup_timer.timeout.connect(self.cleanup_finished_threads)
        self.cleanup_timer.start(5000)

    def cleanup_finished_threads(self):
        self.threads_mutex.lock()
        self.active_threads = [t for t in self.active_threads if t.isRunning()]
        self.threads_mutex.unlock()

    def initUI(self):
        self.setWindowTitle("WCGen + Sentiment Analysis (v1.5)")
        self.setFixedSize(550, 750)
        self.setWindowIcon(QIcon(str(ICON_PATH)))

        layout = QGridLayout()

        self.file_name = QLineEdit(self)
        self.file_name.setReadOnly(True)
        self.file_name.setFixedHeight(30)
        self.file_name.setPlaceholderText("No file selected")
        layout.addWidget(self.file_name, 0, 0, 1, 6)

        self.load_file_button = QPushButton("Load Text File", self)
        self.load_file_button.clicked.connect(self.pilih_file)
        self.load_file_button.setFixedHeight(30)
        self.load_file_button.setToolTip(
            "Upload a text file for word cloud generation and sentiment analysis.\n"
            "Supports TXT, CSV, XLS/XLSX, PDF, DOC/DOCX.\n"
            "Ensure your text is well-formatted for better results."
        )        
        layout.addWidget(self.load_file_button, 1, 0, 1, 4)

        self.view_fulltext_button = QPushButton("View Full Text", self)
        self.view_fulltext_button.setFixedHeight(30)
        self.view_fulltext_button.clicked.connect(self.view_full_text)
        self.view_fulltext_button.setEnabled(False)
        self.view_fulltext_button.setToolTip(
            "Click to view the full text content in a separate window.\n"
            "Allows you to inspect the complete text before generating the word cloud.\n"
            "Useful for verifying text input and checking formatting."
        )        
        layout.addWidget(self.view_fulltext_button, 1, 4, 1, 2)

        self.progress_bar = QProgressBar(self)
        self.progress_bar.setRange(0, 0)
        self.progress_bar.setFixedHeight(4)
        self.progress_bar.setStyleSheet("""
            QProgressBar { border-radius: 1px; background-color: white; text-align: center; margin-left: 5px; margin-right: 5px; }
            QProgressBar::chunk { background-color: red; width: 1px; margin: 0px; }
        """)
        self.progress_bar.setContentsMargins(10, 0, 10, 0)
        self.progress_bar.setVisible(False)
        layout.addWidget(self.progress_bar, 2, 0, 1, 6)

        self.stopword_entry = QTextEdit(self)
        self.stopword_entry.setFixedHeight(75)
        self.stopword_entry.setPlaceholderText("Enter stopwords, separated by spaces or new lines (optional)")
        self.stopword_entry.setToolTip(
            "Enter stopwords (words to be excluded) separated by spaces or new lines.\n"
            "The words you enter will be ignored in the word cloud.\n"
            "Use custom stopwords to refine the visualization and focus on meaningful words."
        )
        layout.addWidget(self.stopword_entry, 3, 0, 1, 6)

        hline = QFrame()
        hline.setFrameShape(QFrame.HLine)
        hline.setFrameShadow(QFrame.Sunken)
        layout.addWidget(hline, 4, 0, 1, 6)

        self.color_theme_label = QLabel("Color Theme:", self)
        layout.addWidget(self.color_theme_label, 5, 0, 1, 2)

        self.color_theme = QComboBox(self)
        self.color_theme.setFixedHeight(30)
        QTimer.singleShot(100, self.load_colormaps)
        self.color_theme.setToolTip(
            "Choose a color palette for the word cloud.\n"
            "Darker themes work well with light backgrounds, and vice-versa."
        )
        layout.addWidget(self.color_theme, 5, 2, 1, 3)

        self.custom_palette_button = QPushButton("Custom", self)
        self.custom_palette_button.setFixedHeight(30)
        self.custom_palette_button.setToolTip("Create a custom color palette")
        self.custom_palette_button.clicked.connect(self.create_custom_palette)
        layout.addWidget(self.custom_palette_button, 5, 5, 1, 1)

        self.bg_color_label = QLabel("Background Color:", self)
        layout.addWidget(self.bg_color_label, 7, 0, 1, 2)

        self.bg_color = QComboBox(self)
        self.bg_color.setFixedHeight(30)
        self.bg_color.addItems(["white", "black", "gray", "blue", "red", "yellow"])
        self.bg_color.setToolTip(
            "Select the background color for the word cloud.\n"
            "Use contrast for better visibility.\n"
            "White or black backgrounds usually work best."
        )
        layout.addWidget(self.bg_color, 7, 2, 1, 3)

        self.custom_bg_color_button = QPushButton("Custom", self)
        self.custom_bg_color_button.setFixedHeight(30)
        self.custom_bg_color_button.setToolTip("Select a custom background color")
        self.custom_bg_color_button.clicked.connect(self.select_custom_bg_color)
        layout.addWidget(self.custom_bg_color_button, 7, 5, 1, 1)

        hline = QFrame()
        hline.setFrameShape(QFrame.HLine)
        hline.setFrameShadow(QFrame.Sunken)
        layout.addWidget(hline, 8, 0, 1, 6)

        self.title_label = QLabel("WordCloud Title:", self)
        layout.addWidget(self.title_label, 9, 0, 1, 2)

        self.title_entry = QLineEdit(self)
        self.title_entry.setFixedHeight(30)
        self.title_entry.setPlaceholderText("Enter title (optional)")
        self.title_entry.setToolTip(
            "Enter a title for your word cloud (optional).\n"
            "This title will be displayed above the word cloud.\n"
            "Leave blank if no title is needed."
        )
        layout.addWidget(self.title_entry, 9, 2, 1, 2)

        self.title_font_size = QSpinBox(self)
        self.title_font_size.setRange(8, 72)
        self.title_font_size.setValue(14)
        self.title_font_size.setFixedHeight(30)
        self.title_font_size.setToolTip(
            "Set the font size for the word cloud title.\n"
            "Larger values make the title more prominent.\n"
            "Recommended: 14-24 px for a balanced look.\n"
            "Too large titles may overlap with the word cloud."
        )
        layout.addWidget(self.title_font_size, 9, 4, 1, 1)

        self.title_position = QComboBox(self)
        self.title_position.addItems(["Left", "Center", "Right"])
        self.title_position.setCurrentText("Center")
        self.title_position.setFixedHeight(30)
        self.title_position.setToolTip(
            "Choose where to display the title relative to the word cloud.\n"
            "Positioning affects the overall layout and readability.\n"
            "Recomended: Center"
        )
        layout.addWidget(self.title_position, 9, 5, 1, 1)

        self.font_choice_label = QLabel("Font Choice:", self)
        layout.addWidget(self.font_choice_label, 10, 0, 1, 2)

        self.font_choice = QComboBox(self)
        self.font_choice.setFixedHeight(30)
        self.font_choice.addItem("Default")
        self.font_choice.setToolTip(
            "Choose a font style for the word cloud.\n"
            "Different fonts affect readability and aesthetics.\n"
            "Sans-serif fonts are recommended for clarity."
        )
        layout.addWidget(self.font_choice, 10, 2, 1, 4)
        QTimer.singleShot(100, self.load_matplotlib_fonts)

        hline = QFrame()
        hline.setFrameShape(QFrame.HLine)
        hline.setFrameShadow(QFrame.Sunken)
        layout.addWidget(hline, 11, 0, 1, 6)

        self.min_font_size_label = QLabel("Minimum Font Size:", self)
        layout.addWidget(self.min_font_size_label, 12, 0, 1, 2)

        self.min_font_size_input = QSpinBox(self)
        self.min_font_size_input.setFixedHeight(30)
        self.min_font_size_input.setValue(11)
        self.min_font_size_input.setToolTip(
            "Set the smallest font size for words in the word cloud.\n"
            "Prevents low-frequency words from becoming too small to read.\n"
            "Recommended value: 10-12 px for readability."
        )
        layout.addWidget(self.min_font_size_input, 12, 2, 1, 1)

        self.max_words_label = QLabel("Maximum Words:", self)
        layout.addWidget(self.max_words_label, 12, 3, 1, 2, Qt.AlignRight)

        self.max_words_input = QSpinBox(self)
        self.max_words_input.setFixedHeight(30)
        self.max_words_input.setMaximum(10000)
        self.max_words_input.setValue(200)
        self.max_words_input.setToolTip(
            "Set the maximum number of words displayed in the word cloud.\n"
            "Higher values provide more detail but may reduce clarity.\n"
            "Recommended: 100-200 words for balanced visualization"
        )
        layout.addWidget(self.max_words_input, 12, 5, 1, 1)

        hline = QFrame()
        hline.setFrameShape(QFrame.HLine)
        hline.setFrameShadow(QFrame.Sunken)
        layout.addWidget(hline, 13, 0, 1, 6)

        self.mask_label = QLabel("Mask Image:", self)
        layout.addWidget(self.mask_label, 14, 0, 1, 2)

        self.mask_path_label = QLineEdit("default (rectangle)", self)
        self.mask_path_label.setReadOnly(True)
        self.mask_path_label.setFixedHeight(30)
        layout.addWidget(self.mask_path_label, 14, 2, 1, 4)

        self.mask_button = QPushButton("Load Mask Image", self)
        self.mask_button.setFixedHeight(30)
        self.mask_button.setToolTip(
            "Upload an image to shape the word cloud (PNG/JPG/BMP).\n"
            "White areas will be ignored, and words will fill the dark areas.\n"
            "Use simple shapes for best results."
        )
        self.mask_button.clicked.connect(self.choose_mask)
        layout.addWidget(self.mask_button, 15, 2, 1, 2)

        self.reset_mask_button = QPushButton("Remove Mask Image", self)
        self.reset_mask_button.setFixedHeight(30)
        self.reset_mask_button.setToolTip(
            "Remove the selected mask image and revert to the default shape.\n"
            "The word cloud will be displayed in a rectangular format.\n"
            "Use this if you no longer want a custom shape for the word cloud."
        )
        self.reset_mask_button.clicked.connect(self.reset_mask)
        layout.addWidget(self.reset_mask_button, 15, 4, 1, 2)

        hline = QFrame()
        hline.setFrameShape(QFrame.HLine)
        hline.setFrameShadow(QFrame.Sunken)
        layout.addWidget(hline, 16, 0, 1, 6)

        self.generate_wordcloud_button = QPushButton("Generate Word Cloud", self)
        self.generate_wordcloud_button.setFixedHeight(50)
        self.generate_wordcloud_button.setToolTip("Generate the word cloud")
        self.generate_wordcloud_button.clicked.connect(self.generate_wordcloud)
        self.generate_wordcloud_button.setEnabled(False)
        layout.addWidget(self.generate_wordcloud_button, 17, 0, 1, 6)

        self.wordcloud_progress_bar = QProgressBar(self)
        self.wordcloud_progress_bar.setRange(0, 0)
        self.wordcloud_progress_bar.setFixedHeight(4)
        self.wordcloud_progress_bar.setStyleSheet("""
            QProgressBar { border-radius: 1px; background-color: white; text-align: center; margin-left: 5px; margin-right: 5px; }
            QProgressBar::chunk { background-color: blue; width: 1px; margin: 0px; }
        """)
        self.wordcloud_progress_bar.setVisible(False)
        layout.addWidget(self.wordcloud_progress_bar, 18, 0, 1, 6)

        self.text_stats_button = QPushButton("View Text Statistics", self)
        self.text_stats_button.setFixedHeight(30)
        self.text_stats_button.setToolTip("View text statistics")
        self.text_stats_button.clicked.connect(self.text_analysis_report)
        self.text_stats_button.setEnabled(False)
        layout.addWidget(self.text_stats_button, 19, 0, 1, 3)

        self.save_wc_button = QPushButton("Save Word Cloud", self)
        self.save_wc_button.setFixedHeight(30)
        self.save_wc_button.setToolTip("Save the generated word cloud")
        self.save_wc_button.clicked.connect(self.simpan_wordcloud)
        self.save_wc_button.setEnabled(False)
        layout.addWidget(self.save_wc_button, 19, 3, 1, 3)

        hline = QFrame()
        hline.setFrameShape(QFrame.HLine)
        hline.setFrameShadow(QFrame.Sunken)
        layout.addWidget(hline, 20, 0, 1, 6)

        self.sentiment_mode_label = QLabel("Sentiment Analysis Mode:", self)
        layout.addWidget(self.sentiment_mode_label, 21, 0, 1, 2)

        self.sentiment_mode_combo = QComboBox(self)
        self.sentiment_mode_combo.setFixedHeight(30)
        self.sentiment_mode_combo.addItems([
            "TextBlob", "TextBlob (Custom Lexicon)", "VADER", "VADER (Custom Lexicon)", "Flair", "Flair (Custom Model)",
        ])
        self.sentiment_mode_combo.setToolTip("Select sentiment analysis mode")
        self.sentiment_mode_combo.setCurrentText("TextBlob")
        self.sentiment_mode_combo.currentTextChanged.connect(self.change_sentiment_mode)
        layout.addWidget(self.sentiment_mode_combo, 21, 2, 1, 3)

        self.sentiment_mode_info_button = QPushButton("Info", self)
        self.sentiment_mode_info_button.setFixedHeight(30)
        self.sentiment_mode_info_button.setToolTip("Show description for each sentiment analysis mode")
        self.sentiment_mode_info_button.clicked.connect(self.show_sentiment_mode_info)
        layout.addWidget(self.sentiment_mode_info_button, 21, 5, 1, 1)

        self.custom_lexicon_button = QPushButton("Load Lexicon", self)
        self.custom_lexicon_button.setFixedHeight(30)
        self.custom_lexicon_button.setToolTip("Load a custom lexicon file for TextBlob or VADER")
        self.custom_lexicon_button.clicked.connect(self.load_custom_lexicon)
        self.custom_lexicon_button.setEnabled(False)
        layout.addWidget(self.custom_lexicon_button, 22, 2, 1, 2)

        self.custom_model_button = QPushButton("Load Model", self)
        self.custom_model_button.setFixedHeight(30)
        self.custom_model_button.setToolTip("Load a custom model file for Flair")
        self.custom_model_button.clicked.connect(self.load_custom_model)
        self.custom_model_button.setEnabled(False)
        layout.addWidget(self.custom_model_button, 22, 4, 1, 2)

        self.model_progress_bar = QProgressBar(self)
        self.model_progress_bar.setRange(0, 0)
        self.model_progress_bar.setFixedHeight(4)
        self.model_progress_bar.setStyleSheet("""
            QProgressBar { border-radius: 1px; background-color: white; text-align: center; margin-left: 5px; margin-right: 5px; }
            QProgressBar::chunk { background-color: orange; width: 1px; margin: 0px; }
        """)
        self.model_progress_bar.setVisible(False)
        layout.addWidget(self.model_progress_bar, 22, 0, 1, 2)

        self.sentiment_button = QPushButton("Analyze Sentiment", self)
        self.sentiment_button.setFixedHeight(50)
        self.sentiment_button.setToolTip("Analyze sentiment using current mode")
        self.sentiment_button.clicked.connect(self.analyze_sentiment)
        self.sentiment_button.setEnabled(False)
        layout.addWidget(self.sentiment_button, 24, 0, 1, 6)

        hline = QFrame()
        hline.setFrameShape(QFrame.HLine)
        hline.setFrameShadow(QFrame.Sunken)
        layout.addWidget(hline, 25, 0, 1, 6)

        self.about_button = QPushButton("About", self)
        self.about_button.setFixedHeight(30)
        self.about_button.setToolTip("Show information about the application")
        self.about_button.clicked.connect(self.show_about)
        layout.addWidget(self.about_button, 26, 0, 1, 2)

        self.panic_button = QPushButton("STOP", self)
        self.panic_button.setFixedHeight(30)
        self.panic_button.setToolTip("Stop all ongoing processes")
        self.panic_button.clicked.connect(self.stop_all_processes)
        layout.addWidget(self.panic_button, 26, 2, 1, 2)

        self.quit_button = QPushButton("Quit", self)
        self.quit_button.setFixedHeight(30)
        self.quit_button.setToolTip("Quit WCGen :(")
        self.quit_button.clicked.connect(self.close)
        layout.addWidget(self.quit_button, 26, 4, 1, 2)

        container = QWidget()
        container.setLayout(layout)
        self.setCentralWidget(container)

    def load_matplotlib_fonts(self):
        try:
            from matplotlib import font_manager
            self.font_map = {}

            weight_conversion = {
                100: "Thin", 200: "Extra Light", 300: "Light", 400: "Regular",
                500: "Medium", 600: "Semi Bold", 700: "Bold", 800: "Extra Bold", 900: "Black",
            }

            for font in font_manager.fontManager.ttflist:
                try:
                    family = font.family_name if hasattr(font, "family_name") else font.name
                    style = font.style_name.lower() if hasattr(font, "style_name") else font.style.lower()
                    weight = weight_conversion.get(font.weight, str(font.weight))

                    display_parts = [family]
                    if "italic" in style or "oblique" in style:
                        display_parts.append("Italic")
                    elif "normal" not in style:
                        display_parts.append(style.title())

                    if weight != "Regular":
                        display_parts.append(weight)

                    display_name = " ".join(display_parts)
                    if display_name not in self.font_map:
                        self.font_map[display_name] = font.fname

                except Exception as e:
                    continue

            self.font_choice.clear()
            self.font_choice.addItem("Default")
            self.font_choice.addItems(sorted(self.font_map.keys()))

        except Exception as e:
            QMessageBox.warning(self, "Font Error", f"Failed to load fonts: {str(e)}")
            self.font_choice.addItems(["Arial", "Times New Roman", "Verdana"])

    def load_colormaps(self):
        try:
            import matplotlib.pyplot as plt
            self.color_theme.addItems(plt.colormaps())
        except ImportError as e:
            self.color_theme.clear()
            self.color_theme.addItem("Default")
            QMessageBox.warning(self, "Dependency Error", f"Failed to load color maps: {str(e)}")

    def analyze_sentiment(self):
        if not self.text_data:
            QMessageBox.critical(self, "Error", "No text data available for sentiment analysis.\nPlease load a text file first.")
            return

        if self.sentiment_mode in ["Flair", "TextBlob"]:
            try:
                from langdetect import detect
            except:
                pass

        if (self.sentiment_mode == "Flair" and not self.flair_classifier and self.sentiment_mode_combo.currentText() == "Flair"):
            QMessageBox.warning(self, "Model Loading", "Default Flair model is still loading. Please wait...")
            return

        if self.sentiment_mode == "Flair (Custom Model)":
            if not self.flair_classifier_cuslang:
                QMessageBox.warning(self, "Model Required", "Please load a custom Flair model first!")
                return

        classifier = self.flair_classifier_cuslang if self.sentiment_mode == "Flair (Custom Model)" else self.flair_classifier

        self.progress_bar.show()
        self.sentiment_thread = SentimentAnalysisThread(
            self.text_data, self.sentiment_mode, self.vader_analyzer, classifier, self.flair_classifier_cuslang, self.textblob_analyzer
        )
        self.sentiment_thread.offline_warning.connect(self.handle_offline_warning)
        self.sentiment_thread.sentiment_analyzed.connect(self.on_sentiment_analyzed)
        self.active_threads.append(self.sentiment_thread)
        self.sentiment_thread.start()

    def stop_all_processes(self):
        self.threads_mutex.lock()
        threads_to_stop = self.active_threads.copy()
        self.threads_mutex.unlock()

        SOFT_STOP_TIMEOUT = 1000
        FORCE_STOP_TIMEOUT = 500

        stopped_threads = []
        failed_to_stop = []

        for thread in threads_to_stop:
            try:
                if thread.isRunning():
                    thread.requestInterruption()
                    thread.quit()
                    
                    if not thread.wait(SOFT_STOP_TIMEOUT):
                        thread.terminate()
                        if not thread.wait(FORCE_STOP_TIMEOUT):
                            failed_to_stop.append(thread)
                            continue
                    
                    stopped_threads.append(thread)
                    
            except RuntimeError as e:
                failed_to_stop.append(thread)
            except Exception as e:
                failed_to_stop.append(thread)

        self.threads_mutex.lock()
        self.active_threads = [
            t for t in self.active_threads 
            if t not in stopped_threads and t not in failed_to_stop
        ]
        self.threads_mutex.unlock()

        self.progress_bar.setVisible(False)
        self.model_progress_bar.setVisible(False)
        self.wordcloud_progress_bar.setVisible(False)
        
        status_message = []
        if stopped_threads:
            status_message.append(f"Stopped {len(stopped_threads)} processes")
        if failed_to_stop:
            status_message.append(f"Failed to stop {len(failed_to_stop)} processes")
        
        final_message = "All processes stopped" if not status_message else "\n".join(status_message)
        
        QMessageBox.information(
            self,
            "Process Termination Report",
            f"""<b>Process termination completed:</b>
            <li>✅ Successfully stopped: {len(stopped_threads)}</li>
            <li>❌ Failed to stop: {len(failed_to_stop)}</li>
            <i>Note: Some background processes may still complete if termination failed</i>"""
        )

        self.cleanup_finished_threads()

    def closeEvent(self, event):
        import matplotlib.pyplot as plt
        
        reply = QMessageBox.question(
            self,
            "Quit Confirmation",
            "<b>Are you sure you want to quit?</b><br>"
            "All ongoing processes will be terminated.",
            QMessageBox.StandardButton.Yes | QMessageBox.StandardButton.No,
            QMessageBox.StandardButton.No
        )

        if reply == QMessageBox.StandardButton.No:
            event.ignore()
            return

        self.cleanup()
        event.accept()

    def cleanup(self):
        """Clean up resources before exit"""
        self.thread_pool.clear()
        if torch.cuda.is_available():
            torch.cuda.empty_cache()

        if self.current_figure:
            try:
                plt.close(self.current_figure)
                self.current_figure = None
            except Exception:
                pass

        self.threads_mutex.lock()
        remaining_threads = self.active_threads.copy()
        self.threads_mutex.unlock()

        TERMINATION_TIMEOUT = 2000

        for thread in remaining_threads:
            try:
                if thread.isRunning():
                    thread.requestInterruption()
                    thread.quit()
                    
                    if not thread.wait(TERMINATION_TIMEOUT):
                        thread.terminate()
                        thread.wait(TERMINATION_TIMEOUT)
            except Exception as e:
                pass

        for widget in QApplication.topLevelWidgets():
            if widget is not self:
                try:
                    widget.close()
                    widget.deleteLater()
                except Exception as e:
                    pass

        try:
            if self.flair_classifier:
                del self.flair_classifier
            if self.flair_classifier_cuslang:
                del self.flair_classifier_cuslang
            if self.textblob_analyzer:
                del self.textblob_analyzer
        except Exception as e:
            pass

        QApplication.processEvents()

    def show_about(self):
        about_text = base64.b64decode(ABOUT_TEXT.encode()).decode()

        dialog = QDialog(self)
        dialog.setWindowModality(Qt.NonModal)
        dialog.setWindowTitle("About WCGen")
        dialog.setMinimumSize(500, 400)
        dialog.setSizeGripEnabled(True)
        layout = QVBoxLayout()

        text_browser = QTextBrowser()
        text_browser.setHtml(about_text)
        text_browser.setOpenExternalLinks(True)
        text_browser.setReadOnly(True)
        layout.addWidget(text_browser)

        close_button = QPushButton("Close")
        close_button.clicked.connect(dialog.accept)
        layout.addWidget(close_button)

        dialog.setLayout(layout)
        dialog.show()

    def pilih_file(self):
        options = QFileDialog.Options()
        file_path, _ = QFileDialog.getOpenFileName(
            self, "Load Text File", "", "Supported Files (*.txt *.pdf *.doc *.docx *.csv *.xlsx *.xls);;All Files (*)", options=options
        )
        if not file_path:
            return

        self.progress_bar.setVisible(True)

        try:
            self.file_loader_thread = FileLoaderThread(file_path)
            self.file_loader_thread.file_loaded.connect(self.on_file_loaded)
            self.file_loader_thread.file_error.connect(self.handle_file_error)

            self.threads_mutex.lock()
            self.active_threads.append(self.file_loader_thread)
            self.threads_mutex.unlock()

            self.file_loader_thread.start()

        except Exception as e:
            self.progress_bar.setVisible(False)
            QMessageBox.critical(self, "Error", f"Failed to load file: {e}")

    def handle_file_error(self, error_message):
        self.progress_bar.setVisible(False)
        QMessageBox.critical(self, "File Load Error", error_message)
        self._reset_file_state()

    def on_file_loaded(self, file_path, text_data):
        self.progress_bar.setVisible(False)

        try:
            if not text_data.strip():
                raise ValueError("File appears empty after loading")

            self.text_data = text_data
            self.file_name.setText(os.path.basename(file_path))

            self.generate_wordcloud_button.setEnabled(True)
            self.save_wc_button.setEnabled(True)
            self.text_stats_button.setEnabled(True)
            self.view_fulltext_button.setEnabled(True)

            self.change_sentiment_mode(self.sentiment_mode)

        except ValueError as e:
            self.handle_file_error(str(e))
            self._reset_file_state()

    def _reset_file_state(self):
        self.text_data = ""
        self.file_name.setText("")
        self.generate_wordcloud_button.setEnabled(False)
        self.save_wc_button.setEnabled(False)
        self.text_stats_button.setEnabled(False)
        self.view_fulltext_button.setEnabled(False)
        self.sentiment_button.setEnabled(False)

    def choose_mask(self):
        options = QFileDialog.Options()
        mask_path, _ = QFileDialog.getOpenFileName(
            self, "Select Mask Image", "", "Image Files (*.png *.jpg *.bmp)", options=options
        )
        if mask_path:
            self.mask_path = mask_path
            self.mask_path_label.setText(f"{mask_path}")

    def reset_mask(self):
        self.mask_path = ""
        self.mask_path_label.setText("default (rectangle)")

    def import_stopwords(self):
        custom_words = self.stopword_entry.toPlainText().strip().lower()
        if custom_words:
            self.additional_stopwords = set(custom_words.split())
        return STOPWORDS.union(self.additional_stopwords)

    def text_analysis_report(self):
        if not self.text_data:
            QMessageBox.critical(self, "Error", "Load text file first!")
            return

        stopwords = self.import_stopwords()
        words = [word.lower() for word in self.text_data.split() if word.lower() not in stopwords]
        word_counts = Counter(words)
        sorted_word_counts = sorted(word_counts.items(), key=lambda item: item[1], reverse=True)

        text_length = len(self.text_data)
        word_count = len(words)
        char_count_excl_spaces = len(self.text_data.replace(" ", ""))
        avg_word_length = char_count_excl_spaces / word_count if word_count > 0 else 0
        most_frequent_words = [word for word, count in sorted_word_counts[:5]]

        if hasattr(self, "stats_dialog") and self.stats_dialog is not None:
            self.stats_dialog.close()

        self.stats_dialog = QDialog(self)
        self.stats_dialog.setWindowModality(Qt.NonModal)
        self.stats_dialog.setWindowTitle("Text Analysis Report")
        self.stats_dialog.setMinimumSize(500, 400)

        layout = QVBoxLayout()

        text_browser = QTextBrowser()

        html_content = f"""
        <h3>Text Analysis Overview</h3>
        <table border="1" cellspacing="0" cellpadding="2" width="100%">
            <tr><th align="left">Metric</th><th align="left">Value</th></tr>
            <tr><td>Text Length</td><td>{text_length} characters</td></tr>
            <tr><td>Word Count</td><td>{word_count}</td></tr>
            <tr><td>Character Count (excluding spaces)</td><td>{char_count_excl_spaces}</td></tr>
            <tr><td>Average Word Length</td><td>{avg_word_length:.2f}</td></tr>
            <tr><td>Most Frequent Words</td><td>{", ".join(most_frequent_words)}</td></tr>
        </table>
        <br>
        <h3>Word Count</h3>
        <table border="1" cellspacing="0" cellpadding="2" width="100%">
            <tr><th align="left">Word</th><th align="left">Count</th></tr>
        """
        for word, count in sorted_word_counts:
            html_content += f"<tr><td>{word}</td><td>{count}</td></tr>"
        html_content += "</table>"

        text_browser.setHtml(html_content)
        text_browser.setOpenExternalLinks(True)
        text_browser.setReadOnly(True)
        layout.addWidget(text_browser)

        close_button = QPushButton("Close")
        close_button.clicked.connect(self.stats_dialog.accept)
        layout.addWidget(close_button)

        self.stats_dialog.setLayout(layout)
        self.stats_dialog.show()

    def generate_wordcloud(self):
        import matplotlib.pyplot as plt

        if not self.text_data:
            QMessageBox.critical(self, "Error", "Load text file first!")
            return

        font_path = None
        selected_font = self.font_choice.currentText()
        if selected_font != "Default":
            font_path = self.font_map.get(selected_font)
            if not font_path or not os.path.exists(font_path):
                QMessageBox.warning(self, "Font Error", f"Font file not found: {font_path}")
                font_path = None
                return

        stopwords = self.import_stopwords()
        mask = None
        if self.mask_path:
            try:
                mask = np.array(Image.open(self.mask_path))
            except Exception as e:
                QMessageBox.critical(self, "Error", f"Failed to load mask image: {e}")
                return

        try:
            if self.current_figure:
                plt.close(self.current_figure)
                self.current_figure = None

            self.wordcloud_progress_bar.setVisible(True)
            QApplication.processEvents()

            wc = WordCloud(
                width=800, height=400, background_color=self.bg_color.currentText(), stopwords=stopwords,
                colormap=self.color_theme.currentText(), max_words=self.max_words_input.value(),
                min_font_size=self.min_font_size_input.value(), mask=mask, font_path=font_path
            ).generate(self.text_data)

            plt.ion()
            self.current_figure = plt.figure(figsize=(10, 5))
            plt.imshow(wc, interpolation="bilinear")

            title_text = self.title_entry.text().strip()
            if title_text:
                title_font = None
                if selected_font != "Default" and font_path:
                    title_font = FontProperties(fname=font_path, size=self.title_font_size.value())
                plt.title(title_text, loc=self.title_position.currentText().lower(), fontproperties=title_font)

            plt.axis("off")
            plt.show(block=False)
        except Exception as e:
            if self.current_figure:
                plt.close(self.current_figure)
                self.current_figure = None
            QMessageBox.critical(self, "Error", f"Failed to generate word cloud: {e}")
        finally:
            self.wordcloud_progress_bar.setVisible(False)

    def simpan_wordcloud(self):
        options = QFileDialog.Options()
        save_path, _ = QFileDialog.getSaveFileName(
            self, "Save WordCloud", "", "PNG file (*.png);;JPG file (*.jpg)", options=options
        )
        if not save_path:
            return

        stopwords = self.import_stopwords()
        mask = None

        font_path = None
        if self.font_choice.currentText() != "Default":
            from matplotlib import font_manager
            selected_font = self.font_choice.currentText()
            font_path = self.font_map.get(selected_font)

            try:
                font_manager.findfont(self.font_choice.currentText())
                font_path = self.font_choice.currentText()
            except:
                QMessageBox.warning(self, "Font Error", "Selected font not found, using default")

        if self.mask_path:
            try:
                mask = np.array(Image.open(self.mask_path))
            except Exception as e:
                QMessageBox.critical(self, "Error", f"Failed to load mask image: {e}")
                return

        try:
            wc = WordCloud(
                width=800, height=400, background_color=self.bg_color.currentText(), stopwords=stopwords,
                colormap=self.color_theme.currentText(), max_words=self.max_words_input.value(),
                min_font_size=self.min_font_size_input.value(), mask=mask,
                font_path=None if self.font_choice.currentText() == "Default" else self.font_choice.currentText()
            ).generate(self.text_data)
            wc.to_file(save_path)
            QMessageBox.information(self, "Succeed", "Word cloud saved!")
        except Exception as e:
            QMessageBox.critical(self, "Error", f"Failed to save word cloud: {e}")

    def create_custom_palette(self):
        dialog = QDialog(self)
        dialog.setWindowTitle("Create Custom Palette")
        dialog.setFixedSize(400, 150)

        layout = QVBoxLayout()
        color_list = []

        def add_color():
            color = QColorDialog.getColor()
            if color.isValid():
                color_list.append(color.name())
                color_label.setText(", ".join(color_list))

        color_label = QLabel("", self)
        layout.addWidget(color_label)

        add_color_button = QPushButton("Add Color", self)
        add_color_button.clicked.connect(add_color)
        layout.addWidget(add_color_button)

        save_palette_button = QPushButton("Save Palette", self)
        save_palette_button.clicked.connect(lambda: self.save_custom_palette(color_list, dialog))
        layout.addWidget(save_palette_button)

        dialog.setLayout(layout)
        dialog.exec()

    def save_custom_palette(self, color_list, dialog):
        palette_name, ok = QInputDialog.getText(self, "Save Palette", "Enter palette name:")
        if ok and palette_name:
            self.custom_color_palettes[palette_name] = color_list
            self.color_theme.addItem(palette_name)
            dialog.accept()

    def select_custom_bg_color(self):
        color = QColorDialog.getColor()
        if color.isValid():
            self.bg_color.addItem(color.name())
            self.bg_color.setCurrentText(color.name())

    def view_full_text(self):
        if not self.text_data:
            QMessageBox.critical(self, "Error", "No text data available to display.")
            return

        if hasattr(self, "text_dialog") and self.text_dialog is not None:
            self.text_dialog.close()

        self.text_dialog = QDialog(self)
        self.text_dialog.setWindowModality(Qt.NonModal)
        self.text_dialog.setWindowTitle("Full Text")
        self.text_dialog.setMinimumSize(500, 400)
        self.text_dialog.setSizeGripEnabled(True)

        layout = QVBoxLayout()

        text_browser = QTextBrowser()
        text_browser.setPlainText(self.text_data)
        text_browser.setOpenExternalLinks(True)
        text_browser.setReadOnly(True)
        layout.addWidget(text_browser)

        close_button = QPushButton("Close")
        close_button.clicked.connect(self.text_dialog.accept)
        layout.addWidget(close_button)

        self.text_dialog.setLayout(layout)
        self.text_dialog.show()

    def handle_offline_warning(self, msg):
        self.progress_bar.setVisible(False)
        QMessageBox.warning(self, "Offline Mode", msg)

    def on_sentiment_analyzed(self, result):
        self.progress_bar.setVisible(False)
        text_length = len(self.text_data)
        word_count = len(self.text_data.split())
        char_count_excl_spaces = len(self.text_data.replace(" ", ""))
        avg_word_length = char_count_excl_spaces / word_count if word_count > 0 else 0
        most_frequent_words = self.get_most_frequent_words(self.text_data, 5)

        self.show_sentiment_analysis(
            self.sentiment_mode, result["positive_score"], result["neutral_score"], result["negative_score"],
            result["compound_score"], result["sentiment_label"], text_length, result["subjectivity"],
            word_count, char_count_excl_spaces, avg_word_length, most_frequent_words
        )

    def show_sentiment_analysis(
        self, analysis_mode, positive_score, neutral_score, negative_score, compound_score, sentiment_label,
        text_length, subjectivity, word_count, char_count_excl_spaces, avg_word_length, most_frequent_words
    ):
        dialog = QDialog(self)
        dialog.setWindowModality(Qt.NonModal)
        dialog.setWindowTitle("Sentiment Analysis Result")
        dialog.setMinimumSize(500, 270)

        layout = QVBoxLayout()

        text_browser = QTextBrowser()

        sentiment_result = f"""
        <table border="1" cellspacing="0" cellpadding="2" width="100%">
            <tr><th align="left">Metric</th><th align="left">Value</th></tr>
            <tr><td>Analysis Mode</td><td>{analysis_mode}</td></tr>
            <tr><td>Sentiment Label</td><td><b>{sentiment_label}</b></td></tr>
            <tr><td>Positive Sentiment</td><td>{positive_score:.2f}</td></tr>
            <tr><td>Neutral Sentiment</td><td>{neutral_score:.2f}</td></tr>
            <tr><td>Negative Sentiment</td><td>{negative_score:.2f}</td></tr>
            <tr><td>Compound Score</td><td>{compound_score:.2f}</td></tr>
            <tr><td>Subjectivity</td><td>{subjectivity if isinstance(subjectivity, str) else f"{subjectivity:.2f}"}</td></tr>
        </table>
        """

        text_browser.setHtml(sentiment_result)
        text_browser.setOpenExternalLinks(True)
        text_browser.setReadOnly(True)
        layout.addWidget(text_browser)

        close_button = QPushButton("Close")
        close_button.clicked.connect(dialog.accept)
        layout.addWidget(close_button)

        dialog.setLayout(layout)
        dialog.show()

    def get_most_frequent_words(self, text, n):
        stopwords = self.import_stopwords().union(STOPWORDS)
        words = [word.lower() for word in text.split() if word.lower() not in stopwords]
        word_counts = Counter(words)
        most_common = word_counts.most_common(n)
        return [word for word, count in most_common]

    def change_sentiment_mode(self, mode):
        self.sentiment_mode = mode
        status_text = ""
        has_text = bool(self.text_data.strip())

        self.custom_lexicon_button.setEnabled(mode == "VADER (Custom Lexicon)")
        self.custom_model_button.setEnabled(mode == "Flair (Custom Model)")

        if mode == "Flair":
            status_text = "Flair: Ready" if self.flair_classifier else "Loading..."
            self.sentiment_button.setEnabled(bool(self.flair_classifier) and has_text)
            self.custom_model_button.setEnabled(False)
            if not self.flair_classifier:
                self.load_flair_model()

        elif mode == "Flair (Custom Model)":
            if not self.flair_classifier:
                self.custom_model_button.setEnabled(False)
                self.load_flair_model()
                status_text = "Initializing Flair environment..."
                self.sentiment_button.setEnabled(False)
            else:
                self.custom_model_button.setEnabled(True)
                if self.flair_classifier_cuslang:
                    status_text = "Ready"
                    self.sentiment_button.setEnabled(has_text)
                else:
                    status_text = "Load custom model first!"
                    self.sentiment_button.setEnabled(False)

        elif mode == "VADER (Custom Lexicon)":
            status_text = "Load custom lexicon first!" if not self.custom_lexicon_path else "Ready"
            self.custom_lexicon_button.setEnabled(True)
            self.sentiment_button.setEnabled(has_text and bool(self.custom_lexicon_path))
            if not self.custom_lexicon_path:
                QMessageBox.warning(
                    self, "Lexicon Required", "You have selected 'VADER (Custom Lexicon)'. Please load a custom lexicon before analyzing sentiment."
                )

        elif mode == "TextBlob (Custom Lexicon)":
            status_text = "Load custom lexicon first!" if not self.custom_textblob_lexicon_path else "Ready"
            self.custom_lexicon_button.setEnabled(True)
            self.sentiment_button.setEnabled(has_text and bool(self.custom_textblob_lexicon_path))
            if not self.custom_textblob_lexicon_path:
                QMessageBox.warning(
                    self, "Lexicon Required", "You have selected 'TextBlob (Custom Lexicon)'. Please load a custom lexicon before analyzing sentiment."
                )

        else:
            status_text = "Ready" if has_text else "Load text first!"
            self.sentiment_button.setEnabled(has_text)

        self.sentiment_button.setToolTip(
            f"{mode} - {status_text}\nClick to analyze sentiment\nSubjectivity scores only available in TextBlob modes"
        )

    def load_custom_lexicon(self):
        options = QFileDialog.Options()
        lexicon_path, _ = QFileDialog.getOpenFileName(
            self, "Select Custom Lexicon File", "", "Text Files (*.txt)", options=options
        )
        if lexicon_path:
            self.model_progress_bar.setVisible(True)
            self.custom_lexicon_button.setEnabled(False)

            self.lexicon_loader_thread = CustomFileLoaderThread(
                lexicon_path, "TextBlob (Custom Lexicon)" if self.sentiment_mode == "TextBlob (Custom Lexicon)" else "lexicon"
            )
            self.lexicon_loader_thread.file_loaded.connect(self.on_lexicon_loaded)

            self.threads_mutex.lock()
            self.active_threads.append(self.lexicon_loader_thread)
            self.threads_mutex.unlock()

            self.lexicon_loader_thread.start()

    def on_lexicon_loaded(self, result, success):
        self.model_progress_bar.setVisible(False)
        self.custom_lexicon_button.setEnabled(True)
        self.sentiment_button.setEnabled(False)

        if success:
            if self.sentiment_mode == "TextBlob (Custom Lexicon)":
                self.custom_textblob_lexicon_path = result
                self.textblob_analyzer = CustomTextBlobSentimentAnalyzer(self.custom_textblob_lexicon_path)
                QMessageBox.information(self, "Success", "Custom TextBlob lexicon loaded successfully!")
                self.change_sentiment_mode(self.sentiment_mode)
            else:
                self.custom_lexicon_path = result
                self.vader_analyzer = CustomVaderSentimentIntensityAnalyzer(custom_lexicon_file=self.custom_lexicon_path)
                QMessageBox.information(self, "Success", "Custom lexicon loaded successfully! VADER will now use this lexicon.")
                self.change_sentiment_mode(self.sentiment_mode)
        else:
            QMessageBox.critical(self, "Error", f"Failed to load custom lexicon: {result}")

    def load_custom_model(self):
        options = QFileDialog.Options()
        model_path, _ = QFileDialog.getOpenFileName(
            self, "Select Custom Model File", "", "Model Files (*.pt)", options=options
        )
        if model_path:
            self.custom_model_button.setEnabled(False)
            self.model_progress_bar.setVisible(True)

            self.model_loader_thread = CustomFileLoaderThread(model_path, "model")
            self.model_loader_thread.file_loaded.connect(self.on_model_loaded)
            self.active_threads.append(self.model_loader_thread)

            self.threads_mutex.lock()
            self.active_threads.append(self.model_loader_thread)
            self.threads_mutex.unlock()

            self.model_loader_thread.start()

    def on_model_loaded(self, result, success):
        self.custom_model_button.setEnabled(True)
        self.model_progress_bar.setVisible(False)

        if success:
            try:
                from flair.models import TextClassifier
                from flair.data import Sentence  # Tambahkan import ini

                if not isinstance(result, TextClassifier):
                    QMessageBox.critical(self, "Error", "Invalid model type. Please load a valid Flair TextClassifier model.")
                    return

                # Test model dengan data dummy
                test_sentence = Sentence("This is a test sentence")
                result.predict(test_sentence)
                
                # Verifikasi label output
                if not test_sentence.labels:
                    raise ValueError("Model didn't produce any labels")
                    
                # Verifikasi label sesuai format sentiment (POSITIVE/NEGATIVE)
                label = test_sentence.labels[0].value
                if label not in ['POSITIVE', 'NEGATIVE', 'NEUTRAL']:
                    raise ValueError(f"Model produces incompatible labels: {label}. Expected: POSITIVE/NEGATIVE/NEUTRAL")

                self.flair_classifier_cuslang = result
                QMessageBox.information(self, "Success", "Custom model loaded successfully! Flair will now use this model.")

            except Exception as e:
                QMessageBox.critical(self, "Model Test Failed", 
                    f"Model validation failed: {str(e)}\n\n"
                    "Please ensure this is a valid sentiment analysis model that produces POSITIVE/NEGATIVE/NEUTRAL labels.\n"
                    "BERT models may need to be fine-tuned specifically for sentiment analysis tasks."
                )
                self.flair_classifier_cuslang = None
        else:
            QMessageBox.critical(self, "Error", f"Failed to load custom model: {result}")
            self.flair_classifier_cuslang = None
        
        self.change_sentiment_mode(self.sentiment_mode)

    def load_flair_model(self):
        if hasattr(self, "flair_loader_thread") and self.flair_loader_thread and self.flair_loader_thread.isRunning():
            self.flair_loader_thread.wait()
            self.flair_loader_thread.quit()
            self.flair_loader_thread.wait()

        self.model_progress_bar.setVisible(True)
        self.model_progress_bar.setRange(0, 0)
        self.custom_model_button.setEnabled(False)
        self.sentiment_button.setEnabled(False)
        QApplication.processEvents()

        self.flair_loader_thread = FlairModelLoaderThread()
        self.flair_loader_thread.model_loaded.connect(self.on_flair_model_loaded)
        self.flair_loader_thread.error_occurred.connect(self.on_flair_model_error)
        self.flair_loader_thread.finished.connect(self.cleanup_flair_thread)

        self.threads_mutex.lock()
        self.active_threads.append(self.flair_loader_thread)
        self.threads_mutex.unlock()
        self.flair_loader_thread.start()

    def on_flair_model_loaded(self, model):
        if model:
            self.model_progress_bar.setVisible(False)
            self.flair_classifier = model

            if self.flair_first_load:
                self.flair_first_load = False
                QMessageBox.information(self, "Ready", "Flair library loaded successfully!")

            if self.sentiment_mode == "Flair (Custom Model)":
                self.custom_model_button.setEnabled(True)
                if not self.flair_classifier_cuslang:
                    QMessageBox.warning(
                        self, "Custom Model Required", "Please load your custom model using the 'Load Model' button"
                    )

            if self.flair_loader_thread is not None:
                self.flair_loader_thread.quit()
                self.flair_loader_thread.wait()
                self.flair_loader_thread = None

            self.change_sentiment_mode(self.sentiment_mode)
        else:
            QMessageBox.critical(self, "Error", "Flair model failed to load. Please try again.")

    def on_flair_model_error(self, error):
        self.model_progress_bar.setVisible(False)
        QMessageBox.critical(self, "Loading Error", f"Failed to load Flair model: {error}")

    def cleanup_flair_thread(self):
        self.threads_mutex.lock()
        if self.flair_loader_thread in self.active_threads:
            self.active_threads.remove(self.flair_loader_thread)
        self.threads_mutex.unlock()
        QTimer.singleShot(100, lambda: setattr(self, "flair_loader_thread", None))

    def show_sentiment_mode_info(self):
        mode_info = base64.b64decode(MODE_INFO.encode()).decode()

        dialog = QDialog(self)
        dialog.setWindowModality(Qt.NonModal)
        dialog.setWindowTitle("Sentiment Analysis Modes")
        dialog.setMinimumSize(500, 400)
        dialog.setSizeGripEnabled(True)

        layout = QVBoxLayout()
        text_browser = QTextBrowser()
        text_browser.setHtml(mode_info)
        text_browser.setOpenExternalLinks(True)
        text_browser.setReadOnly(True)
        layout.addWidget(text_browser)

        close_button = QPushButton("Close")
        close_button.clicked.connect(dialog.accept)
        layout.addWidget(close_button)

        dialog.setLayout(layout)
        dialog.show()

def is_connected():
    try:
        socket.create_connection(("8.8.8.8", 53), timeout=3)
        return True
    except OSError:
        return False

if __name__ == "__main__":
    app = QApplication(sys.argv)
    loop = QEventLoop(app)
    asyncio.set_event_loop(loop)
    ex = WordCloudGenerator()
    ex.show()

    with loop:
        sys.exit(loop.run_forever())