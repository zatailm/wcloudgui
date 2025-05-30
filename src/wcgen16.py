# Textplora 1.6

import sys
import os
import re
from pathlib import Path
from functools import lru_cache, cached_property
from contextlib import contextmanager
import asyncio
from collections import Counter
import socket
import base64
import importlib
import hashlib
import tempfile
import shutil
import logging
import logging.handlers
from typing import Optional, Union, Dict, Any

import torch
import psutil
os.environ["QT_API"] = "pyside6"

from wordcloud import WordCloud, STOPWORDS
from sklearn.feature_extraction.text import CountVectorizer, TfidfVectorizer
from sklearn.decomposition import LatentDirichletAllocation, NMF
import yake
import rake_nltk
from qasync import QEventLoop
from deep_translator import GoogleTranslator
from PySide6.QtWidgets import (
    QApplication, QMainWindow, QFileDialog, QMessageBox, QLabel, QPushButton,
    QVBoxLayout, QGridLayout, QWidget, QLineEdit, QComboBox, QSpinBox, QDialog,
    QTextEdit, QProgressBar, QFrame, QColorDialog, QInputDialog, QTextBrowser,
    QHBoxLayout, QGroupBox
)
from PySide6.QtCore import QThread, Signal, Qt, QTimer, QMutex, QThreadPool, QObject
from PySide6.QtGui import QIcon, QGuiApplication, QPixmap, QPainter, QColor
import matplotlib
matplotlib.use("QtAgg")
import numpy as np
from PIL import Image
from matplotlib.font_manager import FontProperties

from sumy.parsers.plaintext import PlaintextParser
from sumy.nlp.tokenizers import Tokenizer
from sumy.summarizers.lsa import LsaSummarizer

os.environ["QT_SCALE_FACTOR"] = "0.75"

# Setup logging configuration
def setup_logging():
    """Setup application logging"""
    log_dir = Path.home() / ".textplora" / "logs"
    log_dir.mkdir(parents=True, exist_ok=True)
    
    log_file = log_dir / "textplora.log"
    
    # Configure logging format
    fmt = logging.Formatter(
        '%(asctime)s - %(name)s - %(levelname)s - %(message)s'
    )
    
    # File handler with rotation
    file_handler = logging.handlers.RotatingFileHandler(
        log_file, maxBytes=1024*1024, backupCount=5
    )
    file_handler.setFormatter(fmt)
    
    # Console handler
    console_handler = logging.StreamHandler()
    console_handler.setFormatter(fmt)
    
    # Setup root logger
    root_logger = logging.getLogger()
    root_logger.setLevel(logging.INFO)
    root_logger.addHandler(file_handler)
    root_logger.addHandler(console_handler)
    
    return root_logger

logger = setup_logging()

# Custom exceptions
class TextploraError(Exception):
    """Base exception for Textplora application"""
    pass

class FileLoadError(TextploraError):
    """Error when loading files"""
    pass 

class ModelLoadError(TextploraError):
    """Error when loading ML models"""
    pass

class ProcessingError(TextploraError):
    """Error during text processing"""
    pass

def setup_dll_path():
    if getattr(sys, 'frozen', False):
        application_path = os.path.dirname(sys.executable)
        dll_path = os.path.join(application_path, 'lib')
        if dll_path not in os.environ['PATH']:
            os.environ['PATH'] = dll_path + os.pathsep + os.environ['PATH']
        
setup_dll_path()

if getattr(sys, 'frozen', False):
    APP_DIR = Path(sys.executable).parent
else:
    APP_DIR = Path(__file__).parent

ICON_PATH = APP_DIR / "icon.ico"

logo_path = APP_DIR / "logo.png"

# Kelas untuk menyimpan konstanta aplikasi
class AppConstants:
    """Class to store application constants and configuration values"""
    
    # Konstanta aplikasi
    ABOUT_TEXT = "PGh0bWw+Cjxib2R5IHN0eWxlPSJmb250LWZhbWlseTogQXJpYWwsIHNhbnMtc2VyaWY7IHRleHQtYWxpZ246IGNlbnRlcjsgcGFkZGluZzogMTBweDsiPgogICAgPGltZyBzcmM9Intsb2dvX3BhdGh9IiB3aWR0aD0iMTAwIiBoZWlnaHQ9IjEwMCIgYWx0PSJUZXh0cGxvcmEgTG9nbyI+CiAgICA8aDI+VGV4dHBsb3JhPC9oMj4KICAgIDxwPjxiPlZlcnNpb246PC9iPiAxLjY8L3A+CiAgICA8cD4mY29weTsgMjAyNSBNLiBBZGliIFphdGEgSWxtYW08L3A+CiAgICA8aHI+CiAgICA8cD5UZXh0cGxvcmEgKGZvcm1lcmx5IFdDR2VuKSBpcyBhIFB5dGhvbi1iYXNlZCBhcHBsaWNhdGlvbiB0aGF0IGhlbHBzIHlvdSBhbmFseXplIGZlZWRiYWNrLCBjb25kdWN0IHJlc2VhcmNoLCBhbmQgZXh0cmFjdCBtZWFuaW5nZnVsIGluc2lnaHRzIHdpdGggZWFzZS48L3A+CiAgICA8aHI+CiAgICA8cD48Yj5NYWluIGxpYnJhcmllczo8L2I+IFB5U2lkZTYsIG1hdHBsb3RsaWIsIHdvcmRjbG91ZCwgc2tsZWFybiwgdmFkZXJTZW50aW1lbnQsIHRleHRibG9iLCBmbGFpciwgc3VteSwgeWFrZSwgcmFrZV9ubHRrLCBQaWxsb3cuPC9wPgogICAgPGhyPgogICAgPHA+PGI+R2l0SHViIFJlcG9zaXRvcnk6PC9iPjwvcD4KICAgIDxwPjxhIGhyZWY9Imh0dHBzOi8vZ2l0aHViLmNvbS96YXRhaWxtL3djbG91ZGd1aSI+aHR0cHM6Ly9naXRodWIuY29tL3phdGFpbG0vd2Nsb3VkZ3VpPC9hPjwvcD4KICAgIDxocj4KICAgIDxwPjxiPkxpY2Vuc2U6PC9iPjwvcD4KICAgIDxwPkZyZWUgZm9yIHBlcnNvbmFsICYgZWR1Y2F0aW9uYWwgdXNlIChDQyBCWS1OQyA0LjApLjwvcD4KICAgIDxwPkxlYXJuIG1vcmU6IDxhIGhyZWY9Imh0dHBzOi8vY3JlYXRpdmVjb21tb25zLm9yZy9saWNlbnNlcy9ieS1uYy80LjAvIj5DQyBCWS1OQyA0LjA8L2E+PC9wPgo8L2JvZHk+CjwvaHRtbD4="
    
    MODE_INFO = "PGgyPlNlbnRpbWVudCBBbmFseXNpcyBNb2RlczwvaDI+CjxwPlNlbGVjdCB0aGUgbW9zdCBzdWl0YWJsZSBzZW50aW1lbnQgYW5hbHlzaXMgbWV0aG9kIGJhc2VkIG9uIHlvdXIgdGV4dCB0eXBlIGFuZCBhbmFseXNpcyBuZWVkcy48L3A+CjxoMz5UZXh0QmxvYjwvaDM+CjxwPjxiPkJlc3QgZm9yOjwvYj4gRm9ybWFsIHRleHRzLCB3ZWxsLXN0cnVjdHVyZWQgZG9jdW1lbnRzLCBuZXdzIGFydGljbGVzLCByZXNlYXJjaCBwYXBlcnMsIGFuZCByZXBvcnRzLjwvcD4KPHA+VGV4dEJsb2IgaXMgYSBsZXhpY29uLWJhc2VkIHNlbnRpbWVudCBhbmFseXNpcyB0b29sIHRoYXQgcHJvdmlkZXMgYSBzaW1wbGUgeWV0IGVmZmVjdGl2ZSBhcHByb2FjaCBmb3IgZXZhbHVhdGluZyB0aGUgc2VudGltZW50IG9mIHN0cnVjdHVyZWQgdGV4dC4gSXQgYXNzaWducyBhIHBvbGFyaXR5IHNjb3JlIChwb3NpdGl2ZSwgbmVnYXRpdmUsIG9yIG5ldXRyYWwpIGFuZCBjYW4gYWxzbyBhbmFseXplIHN1YmplY3Rpdml0eSBsZXZlbHMuPC9wPgo8aDM+VkFERVIgKFZhbGVuY2UgQXdhcmUgRGljdGlvbmFyeSBhbmQgc0VudGltZW50IFJlYXNvbmVyKTwvaDM+CjxwPjxiPkJlc3QgZm9yOjwvYj4gU29jaWFsIG1lZGlhIHBvc3RzLCB0d2VldHMsIHNob3J0IGNvbW1lbnRzLCBjaGF0IG1lc3NhZ2VzLCBhbmQgaW5mb3JtYWwgcmV2aWV3cy48L3A+CjxwPlZBREVSIGlzIHNwZWNpZmljYWxseSBkZXNpZ25lZCBmb3IgYW5hbHl6aW5nIHNob3J0LCBpbmZvcm1hbCB0ZXh0cyB0aGF0IG9mdGVuIGNvbnRhaW4gc2xhbmcsIGVtb2ppcywgYW5kIHB1bmN0dWF0aW9uLWJhc2VkIGVtb3Rpb25zLiBJdCBpcyBhIHJ1bGUtYmFzZWQgc2VudGltZW50IGFuYWx5c2lzIG1vZGVsIHRoYXQgZWZmaWNpZW50bHkgZGV0ZXJtaW5lcyBzZW50aW1lbnQgaW50ZW5zaXR5IGFuZCB3b3JrcyBleGNlcHRpb25hbGx5IHdlbGwgZm9yIHJlYWwtdGltZSBhcHBsaWNhdGlvbnMuPC9wPgo8aDM+RmxhaXI8L2gzPgo8cD48Yj5CZXN0IGZvcjo8L2I+IExvbmctZm9ybSBjb250ZW50LCBwcm9kdWN0IHJldmlld3MsIHByb2Zlc3Npb25hbCBkb2N1bWVudHMsIGFuZCBBSS1iYXNlZCBkZWVwIHNlbnRpbWVudCBhbmFseXNpcy48L3A+CjxwPkZsYWlyIHV0aWxpemVzIGRlZXAgbGVhcm5pbmcgdGVjaG5pcXVlcyBmb3Igc2VudGltZW50IGFuYWx5c2lzLCBtYWtpbmcgaXQgaGlnaGx5IGFjY3VyYXRlIGZvciBjb21wbGV4IHRleHRzLiBJdCBpcyBpZGVhbCBmb3IgYW5hbHl6aW5nIGxhcmdlLXNjYWxlIHRleHR1YWwgZGF0YSwgY2FwdHVyaW5nIGNvbnRleHQgbW9yZSBlZmZlY3RpdmVseSB0aGFuIHRyYWRpdGlvbmFsIHJ1bGUtYmFzZWQgbW9kZWxzLiBIb3dldmVyLCBpdCByZXF1aXJlcyBtb3JlIGNvbXB1dGF0aW9uYWwgcmVzb3VyY2VzIGNvbXBhcmVkIHRvIFRleHRCbG9iIGFuZCBWQURFUi48L3A+CjxoMj5JbXBvcnRhbnQgTm90ZSBmb3IgTGFuZ3VhZ2UgU3VwcG9ydDwvaDI+CjxwPldoaWxlIHRoaXMgYXBwbGljYXRpb24gc3VwcG9ydHMgbm9uLUVuZ2xpc2ggdGV4dCB0aHJvdWdoIGF1dG9tYXRpYyB0cmFuc2xhdGlvbiwgaXQgaXMgPGI+aGlnaGx5IHJlY29tbWVuZGVkPC9iPiB0byB1c2UgPGI+bWFudWFsbHkgdHJhbnNsYXRlZCBhbmQgcmVmaW5lZCBFbmdsaXNoIHRleHQ8L2I+IGZvciB0aGUgbW9zdCBhY2N1cmF0ZSBzZW50aW1lbnQgYW5hbHlzaXMuIFRoZSBidWlsdC1pbiBhdXRvbWF0aWMgdHJhbnNsYXRpb24gZmVhdHVyZSBtYXkgbm90IGFsd2F5cyBmdW5jdGlvbiBjb3JyZWN0bHksIGxlYWRpbmcgdG8gcG90ZW50aWFsIG1pc2ludGVycHJldGF0aW9ucyBvciBpbmFjY3VyYXRlIHNlbnRpbWVudCByZXN1bHRzLiBFbnN1cmUgeW91ciBjdXN0b20gbGV4aWNvbiBhbmQgbW9kZWwgaXMgdHJhaW5lZCBvbiBzaW1pbGFyIHRleHQgZG9tYWlucyBhcyB5b3VyIGFuYWx5c2lzIGRhdGEgZm9yIG9wdGltYWwgcGVyZm9ybWFuY2UuPC9wPgo8cD5Gb3IgdGhlIGJlc3QgcGVyZm9ybWFuY2UsIGVuc3VyZSB0aGF0IG5vbi1FbmdsaXNoIHRleHQgaXMgcHJvcGVybHkgcmV2aWV3ZWQgYW5kIGFkanVzdGVkIGJlZm9yZSBzZW50aW1lbnQgYW5hbHlzaXMuPC9wPgo8aDI+Q3VzdG9tIExleGljb24gRm9ybWF0IEV4YW1wbGU8L2gyPgo8cD5CZWxvdyBpcyBhbiBleGFtcGxlIG9mIGEgY3VzdG9tIGxleGljb24gZm9ybWF0IGZvciBzZW50aW1lbnQgYW5hbHlzaXM6PC9wPgo8cHJlIHN0eWxlPSdiYWNrZ3JvdW5kLWNvbG9yOiNmNGY0ZjQ7IGNvbG9yOiBibGFjazsgcGFkZGluZzoxMHB4OyBib3JkZXItcmFkaXVzOjVweDsnPgpleGNlbGxlbnQgICAxLjUKYXdmdWwgICAgICAtMS41Cm5vdCAgICAgICAgbmVnYXRpb24gICAgICAgICAjIE1hcmsgYXMgbmVnYXRpb24gd29yZAppbnRlbnNlbHkgIGludGVuc2lmaWVyOjEuNyAgIyBDdXN0b20gaW50ZW5zaWZpZXIgd2l0aCBtdWx0aXBsaWVyCjwvcHJlPgo8cD5UaGlzIGN1c3RvbSBsZXhpY29uIGFsbG93cyBmaW5lLXR1bmluZyBvZiBzZW50aW1lbnQgc2NvcmVzIGJ5IGFkZGluZyBjdXN0b20gd29yZHMsIG5lZ2F0aW9ucywgYW5kIGludGVuc2lmaWVycyB0byBpbXByb3ZlIHNlbnRpbWVudCBhbmFseXNpcyBhY2N1cmFjeS48L3A+CjxoMj5DdXN0b20gRmxhaXIgTW9kZWwgUmVxdWlyZW1lbnRzPC9oMj4KPHA+Rm9yIGN1c3RvbSBGbGFpciBtb2RlbHMgKDxiPi5wdDwvYj4gZmlsZXMpLCBlbnN1cmU6PC9wPgo8cHJlIHN0eWxlPSdiYWNrZ3JvdW5kLWNvbG9yOiNmNGY0ZjQ7IGNvbG9yOiBibGFjazsgcGFkZGluZzoxMHB4OyBib3JkZXItcmFkaXVzOjVweDsnPgoxLiBMYWJlbHMgbXVzdCBiZTogUE9TSVRJVkUsIE5FR0FUSVZFLCBvciBORVVUUkFMCjIuIE1vZGVsIG91dHB1dHMgY29uZmlkZW5jZSBzY29yZXMgKDAtMSkKMy4gVGVzdGVkIHdpdGggRmxhaXIgdjAuMTIrIGNvbXBhdGliaWxpdHkKPC9wcmU+CjxwPlRoZSBhcHBsaWNhdGlvbiB3aWxsIGF1dG9tYXRpY2FsbHkgdmFsaWRhdGUgbW9kZWwgbGFiZWxzIGR1cmluZyBsb2FkaW5nLjwvcD4="
    
    # Mapping untuk tipe skor berdasarkan mode analisis sentimen
    SCORE_TYPES = {
        "TextBlob": "Polarity Score", 
        "TextBlob (Custom Lexicon)": "Polarity Score",
        "VADER": "Compound Score",
        "VADER (Custom Lexicon)": "Compound Score",        
        "Flair": "Confidence Score",
        "Flair (Custom Model)": "Confidence Score"
    }
    
    # Daftar mode analisis sentimen yang didukung
    SENTIMENT_MODES = list(SCORE_TYPES.keys())
    
    # Timeout untuk manajemen thread
    THREAD_TIMEOUT = {
        "SOFT": 1000,
        "FORCE": 500,
        "TERMINATION": 2000
    }

def sanitize_path(path):
    """
    Memvalidasi path file untuk mencegah serangan path traversal
    
    Args:
        path (str): Path yang akan divalidasi
        
    Returns:
        str: Path yang telah dinormalisasi
        
    Raises:
        ValueError: Jika path mencoba melakukan traversal ke direktori di luar base_dir
    """
    if not path:
        return path
        
    path = os.path.normpath(path)
    
    # Jika path absolut, periksa apakah berada dalam direktori yang diizinkan
    if os.path.isabs(path):
        # Daftar direktori yang diizinkan (tambahkan sesuai kebutuhan)
        allowed_dirs = [
            os.path.abspath(os.path.expanduser("~")),  # Home directory
            os.path.abspath(APP_DIR),                  # Application directory
            os.path.abspath(tempfile.gettempdir())     # Temp directory
        ]
        
        abs_path = os.path.abspath(path)
        if not any(abs_path.startswith(allowed_dir) for allowed_dir in allowed_dirs):
            logger.warning(f"Path traversal attempt detected: {path}")
            raise ValueError(f"Path not allowed: {path}")
    
    return path

@lru_cache(maxsize=128)
def load_stopwords():
    if getattr(sys, 'frozen', False):
        stopwords_path = os.path.join(os.path.dirname(sys.executable), "lib", "wordcloud", "stopwords")
    else:
        from wordcloud import STOPWORDS
        return set(STOPWORDS)
        
    try:
        with open(stopwords_path, 'r', encoding='utf-8') as f:
            return set(word.strip() for word in f)
    except Exception as e:
        logger.warning(f"Could not load stopwords from {stopwords_path}: {e}")
        return set()

def is_connected():
    """Check if there is internet connection"""
    try:
        socket.create_connection(("8.8.8.8", 53), timeout=3)
        return True
    except OSError:
        return False

class MainClass(QMainWindow):
    def __init__(self):
        super().__init__()
        
        self.user_name = os.getlogin()
        
        # Inisialisasi statusbar dengan indikator koneksi
        self.setup_statusbar()
        
        # Inisialisasi dasar
        self.active_threads = []
        self.threads_mutex = QMutex()
        
        # Gunakan konstanta dari AppConstants
        self.sentiment_mode = AppConstants.SENTIMENT_MODES[0]  # Default ke mode pertama
        
        self.perf_monitor = PerformanceMonitor()
        
        self.resource_manager = ResourceManager()
        self.lazy_loader = LazyLoader()
        
        # Inisialisasi cache manager
        self.cache_manager = CacheManager(threshold_percent=75, check_interval=60000)
        
        # Daftarkan fungsi-fungsi yang menggunakan lru_cache
        self.register_cached_functions()
        
        # Mulai pemantauan memori
        self.cache_manager.start_monitoring()
        
        self.vectorizer = CountVectorizer(max_features=1000, stop_words='english')
        self.tfidf = TfidfVectorizer(max_features=1000, stop_words='english')
        self.lda_model = None
        self.nmf_model = None
        self.bert_model = None
        
        self._cached_models = {}
        self._cached_fonts = {}
        self._cached_colormaps = {}
        
        self._init_basic()
        self._setup_lazy_loading()
        
        QThreadPool.globalInstance().setMaxThreadCount(3)
        
        self.text_data = ""
        self.additional_stopwords = set()
        self.custom_color_palettes = {}
        self.flair_classifier = None
        self.flair_classifier_cuslang = None
        self.flair_first_load = True
        self.mask_path = ""
        self.current_figure = None
        self.custom_lexicon_path = None
        self.custom_textblob_lexicon_path = None
        self.textblob_analyzer = None
        self.flair_loader_thread = None
        
        self.import_thread = None
        self.active_threads = []
        self.threads_mutex = QMutex()
        
        self.setWindowIcon(QIcon(str(ICON_PATH)))
        self.setup_thread_management()
        self.init_analyzers()
        
        self.progress_states = {
            'keywords': {'color': 'red', 'text': 'Extracting keywords'},
            'file': {'color': 'blue', 'text': 'Loading file'},
            'topics': {'color': 'yellow', 'text': 'Analyzing topics'},
            'model': {'color': 'green', 'text': 'Loading model'},
            'wordcloud': {'color': 'purple', 'text': 'Generating word cloud'}
        }
        
        self.initUI()
        self.setup_timers()
        self._init_imports()

        self.warning_emitter = WarningEmitter()
        self.warning_handler = CustomWarningHandler(self.warning_emitter)
        self.warning_emitter.token_warning.connect(self.show_token_warning)

        import warnings
        warnings.showwarning = self.warning_handler.handle_warning        

        self.button_manager = ButtonStateManager(self)

    def setup_statusbar(self):
        """Setup statusbar dengan indikator koneksi"""
        statusbar = self.statusBar()
        statusbar.setStyleSheet("""
            QStatusBar {
                border-top: 1px solid #c0c0c0;
                background: transparent;
                padding: 2px;
                font-size: 9pt;
            }
        """)
        
        # Buat indikator koneksi
        self.connection_indicator = QLabel()
        self.connection_indicator.setFixedSize(16, 16)
        self.update_connection_status()
        
        # Timer untuk update status koneksi
        self.connection_timer = QTimer(self)
        self.connection_timer.timeout.connect(self.update_connection_status)
        self.connection_timer.start(10000)  # Update setiap 3 detik
        
        # Tambahkan widget ke statusbar
        statusbar.addPermanentWidget(self.connection_indicator)
        
        # Set pesan default
        statusbar.showMessage("Ready")
    
    def update_connection_status(self):
        """Update indikator status koneksi"""
        if is_connected():
            self.connection_indicator.setPixmap(self.get_connection_icon("connected"))
            self.connection_indicator.setToolTip("Internet Connected")
        else:
            self.connection_indicator.setPixmap(self.get_connection_icon("disconnected"))
            self.connection_indicator.setToolTip("No Internet Connection")
    
    def get_connection_icon(self, status):
        """Generate ikon koneksi menggunakan QPainter"""
        pixmap = QPixmap(16, 16)
        pixmap.fill(Qt.transparent)
        
        try:
            painter = QPainter(pixmap)
            painter.setRenderHint(QPainter.Antialiasing)
            
            if status == "connected":
                # Gambar ikon koneksi aktif (lingkaran hijau)
                painter.setPen(Qt.NoPen)
                painter.setBrush(QColor("#2ecc71"))  # Hijau
                painter.drawEllipse(2, 2, 12, 12)
            else:
                # Gambar ikon tidak terkoneksi (lingkaran merah)
                painter.setPen(Qt.NoPen)
                painter.setBrush(QColor("#e74c3c"))  # Merah
                painter.drawEllipse(2, 2, 12, 12)
        finally:
            painter.end()
            
        return pixmap

    @cached_property 
    def token_opts(self):
        """Cached token configuration for vectorizers"""
        return {
            'token_pattern': r'(?u)\b[a-zA-Z][a-zA-Z0-9\'-]*[a-zA-Z0-9]\b',
            'lowercase': True,
            'strip_accents': 'unicode',
            'max_features': 1000,
            'min_df': 1,
            'max_df': 0.95
        }

    def _init_basic(self):
        """Initialize basic components"""
        self.text_data = ""
        self.additional_stopwords = set()
        self.custom_color_palettes = {}
        self.setup_thread_management()
        
    def _setup_lazy_loading(self):
        """Configure lazy loading for heavy dependencies"""
        self._vader = None
        self._flair = None
        self._textblob = None
        
    @cached_property
    def vader_analyzer(self):
        """Lazy load VADER analyzer"""
        if not self._vader:
            VaderAnalyzer = self.lazy_loader.load('vaderSentiment.vaderSentiment', 'SentimentIntensityAnalyzer')
            if VaderAnalyzer:
                self._vader = VaderAnalyzer()
        return self._vader

    @lru_cache(maxsize=32)
    def get_wordcloud(self, **kwargs):
        """Get cached wordcloud instance"""
        return OptimizedWordCloud(**kwargs)

    def _init_imports(self):
        """Pre-initialize heavy imports in background with proper cleanup"""
        if self.import_thread is not None:
            return
            
        self.import_thread = ImportThread()
        self.import_thread.finished.connect(self._cleanup_import_thread)
        self.add_managed_thread(self.import_thread)
        self.import_thread.start()
        
    def _cleanup_import_thread(self):
        """Cleanup import thread when finished"""
        if self.import_thread:
            try:
                self.import_thread.stop()
                self.remove_managed_thread(self.import_thread)
                self.import_thread.deleteLater()
            finally:
                self.import_thread = None

    def add_managed_thread(self, thread):
        """Add thread to managed threads list"""
        self.threads_mutex.lock()
        try:
            self.active_threads.append(thread)
        finally:
            self.threads_mutex.unlock()

    def remove_managed_thread(self, thread):
        """Remove thread from managed threads list"""
        self.threads_mutex.lock()
        try:
            if thread in self.active_threads:
                self.active_threads.remove(thread)
        finally:
            self.threads_mutex.unlock()

    def setup_thread_management(self):
        """Setup thread management and cleanup mechanisms"""
        self.thread_manager = ThreadManager()
        
        # Connect application aboutToQuit signal to thread cleanup
        QApplication.instance().aboutToQuit.connect(self.cleanup_threads)

    def init_analyzers(self):
        """Initialize analyzers with caching"""
        from functools import lru_cache
        
        @lru_cache(maxsize=32)
        def get_vader_analyzer():
            from vaderSentiment.vaderSentiment import SentimentIntensityAnalyzer
            return SentimentIntensityAnalyzer()
            
        self.vader_analyzer = get_vader_analyzer()
        self.sentiment_mode = "TextBlob"
        
    def set_progress(self, state_key, visible=True, progress=None):
        """Enhanced progress bar control"""
        if not hasattr(self, 'unified_progress_bar'):
            return
            
        state = self.progress_states.get(state_key)
        if not state:
            return

        if visible:
            if state_key == 'file':
                pass
            elif state_key == 'model':
                self.sentiment_button.setEnabled(False)
            elif state_key == 'wordcloud':
                self.generate_wordcloud_button.setEnabled(False)
            elif state_key == 'topics':
                self.topic_tab.analyze_topics_btn.setEnabled(False)
            elif state_key == 'keywords':
                self.topic_tab.extract_keywords_btn.setEnabled(False)
                
            self.unified_progress_bar.setVisible(True)
            self.set_progress_style(state['color'])
            
            if progress is not None:
                self.progress_timer.stop()
                self.unified_progress_bar.setValue(int(progress))
            else:
                if not self.progress_timer.isActive():
                    self.unified_progress_bar.setValue(0)
                    self.progress_timer.start()
        else:
            has_text = bool(self.text_data)
            if state_key == 'file':
                self.load_file_button.setEnabled(True)
            elif state_key == 'model':
                self.sentiment_button.setEnabled(has_text)
            elif state_key == 'wordcloud':
                self.generate_wordcloud_button.setEnabled(has_text)
            elif state_key == 'topics':
                self.topic_tab.analyze_topics_btn.setEnabled(has_text)
            elif state_key == 'keywords':
                self.topic_tab.extract_keywords_btn.setEnabled(has_text)
            
            self.progress_timer.stop()
            self.unified_progress_bar.setValue(0)
            self.unified_progress_bar.setVisible(False)
            
        QApplication.processEvents()

    def set_progress_style(self, color):
        """Set progress bar color and style for indeterminate animation"""
        chunk_color = {
            'red': 'rgba(215, 0, 0, 0.2)',
            'blue': 'rgba(0, 120, 215, 0.2)', 
            'yellow': 'rgba(255, 255, 0, 0.2)',
            'green': 'rgba(0, 215, 0, 0.2)',
            'purple': 'rgba(160, 32, 240, 0.2)'
        }.get(color, 'rgba(0, 120, 215, 0.2)')

        style = f"""
            QProgressBar {{
                border: none;
                background: transparent;
                padding: 0px;
                margin: 0px;
            }}
            
            QProgressBar::chunk {{
                width: 1px;
                background: {chunk_color};
            }}
        """
        self.unified_progress_bar.setStyleSheet(style)

    def setup_timers(self):
        """Setup and optimize timers"""
        self.cleanup_timer = QTimer(self)
        # Jalankan cleanup setiap 30 detik
        self.cleanup_timer.timeout.connect(self.thread_manager.cleanup_finished)
        self.cleanup_timer.start(30000)  # 30 detik
        
    def initUI(self):
        """Use this if not include user name :
            # WIN_TITLE = f"V0NHZW4gKyBUZXh0IEFuYWx5dGljcyAodjEuNikK"
            # win_title = base64.b64decode(WIN_TITLE.encode()).decode()
        """
        WIN_TITLE = base64.b64decode("VGV4dHBsb3JhICh2MS42KSBbe31d".encode()).decode().format(self.user_name)
        self.setWindowTitle(WIN_TITLE)
        self.setFixedSize(550, 870)  # Tambah sedikit tinggi untuk statusbar

        layout = QVBoxLayout()

        file_group = QGroupBox("File Input")
        file_layout = QGridLayout()

        filename_container = QFrame()
        filename_container.setStyleSheet("""
            QFrame { 
                border: 1px solid #c0c0c0; 
                background: transparent;
                padding: 0px;
                margin: 0px;
            }
        """)
        filename_container.setFixedHeight(25)
        filename_layout = QHBoxLayout(filename_container)
        filename_layout.setContentsMargins(0, 0, 0, 0)
        filename_layout.setSpacing(0)

        text_container = QFrame()
        text_container.setLayout(QHBoxLayout())
        text_container.layout().setContentsMargins(5, 0, 5, 0)
        text_container.layout().setSpacing(0)
        
        self.file_name = QLineEdit()
        self.file_name.setReadOnly(True)
        self.file_name.setFrame(False)
        self.file_name.setPlaceholderText("No file selected")
        self.file_name.setStyleSheet("""
            QLineEdit {
                border: none;
                background: transparent;
                padding: 0px;
            }
        """)
        text_container.layout().addWidget(self.file_name)
        filename_layout.addWidget(text_container)

        self.unified_progress_bar = QProgressBar(filename_container)
        self.unified_progress_bar.setRange(0, 100)
        self.unified_progress_bar.setValue(0)
        self.unified_progress_bar.setTextVisible(False)
        self.unified_progress_bar.setMinimumWidth(filename_container.width())
        self.unified_progress_bar.setStyleSheet("""
            QProgressBar {
                border: none;
                background: transparent;
                padding: 0px;
                margin: 0px;
            }
            
            QProgressBar::chunk {
                width: 1px;
                background: rgba(0, 120, 215, 0.2);
            }
        """)
        
        self.unified_progress_bar.setFixedWidth(filename_container.width())
        self.unified_progress_bar.setFixedHeight(filename_container.height())
        
        self.unified_progress_bar.setGeometry(0, 0, filename_container.width(), filename_container.height())
        self.file_name.raise_()

        file_layout.addWidget(filename_container, 0, 0, 1, 6)

        self.load_file_button = QPushButton("Load Text File", self)
        self.load_file_button.clicked.connect(self.open_file)
        self.load_file_button.setToolTip(
            "Upload a text file for word cloud generation and sentiment analysis.\n"
            "Supports TXT, CSV, XLS/XLSX, PDF, DOC/DOCX.\n"
            "Ensure your text is well-formatted for better results."
        )
        file_layout.addWidget(self.load_file_button, 1, 0, 1, 2)

        self.view_fulltext_button = QPushButton("View Full Text", self)
        self.view_fulltext_button.clicked.connect(self.view_full_text)
        self.view_fulltext_button.setEnabled(False)
        self.view_fulltext_button.setToolTip(
            "Click to view the full text content in a separate window.\n"
            "Allows you to inspect the complete text before generating the word cloud.\n"
            "Useful for verifying text input and checking formatting."
        )
        file_layout.addWidget(self.view_fulltext_button, 1, 2, 1, 2)

        self.summarize_button = QPushButton("Summarize", self)
        self.summarize_button.clicked.connect(self.summarize_text)
        self.summarize_button.setEnabled(False)
        self.summarize_button.setToolTip(
            "Summarize the text content using the LSA (Latent Semantic Analysis) method.\n"
            "This extractive summarization technique selects the most important sentences,\n"
            "helping to generate concise word clouds and improve sentiment analysis."
        )
        file_layout.addWidget(self.summarize_button, 1, 4, 1, 2)

        file_group.setLayout(file_layout)
        layout.addWidget(file_group)

        wordcloud_group = QGroupBox("Word Cloud Generation")
        wordcloud_layout = QGridLayout()

        stopwords_container = QHBoxLayout()
        self.stopword_entry = QTextEdit(self)
        self.stopword_entry.setFixedHeight(50)
        self.stopword_entry.setPlaceholderText("Enter stopwords, separated by spaces or new lines (optional)")
        self.stopword_entry.setToolTip(
            "Enter stopwords (words to be excluded) separated by spaces or new lines.\n"
            "The words you enter will be ignored in the word cloud.\n"
            "Use custom stopwords to refine the visualization and focus on meaningful words."
        )        
        stopwords_container.addWidget(self.stopword_entry)     
        
        self.load_stopwords_btn = QPushButton("Load", self)
        self.load_stopwords_btn.setFixedSize(50, 50)
        self.load_stopwords_btn.clicked.connect(self.load_stopwords_file)
        stopwords_container.addWidget(self.load_stopwords_btn)
        
        wordcloud_layout.addLayout(stopwords_container, 0, 0, 1, 6)

        self.color_theme_label = QLabel("Color Theme:", self)
        wordcloud_layout.addWidget(self.color_theme_label, 1, 0, 1, 2)

        self.color_theme = QComboBox(self)
        QTimer.singleShot(100, self.load_colormaps)
        self.color_theme.setToolTip(
            "Choose a color palette for the word cloud.\n"
            "Darker themes work well with light backgrounds, and vice-versa."
        )        
        wordcloud_layout.addWidget(self.color_theme, 1, 2, 1, 3)

        self.custom_palette_button = QPushButton("Custom", self)
        self.custom_palette_button.clicked.connect(self.create_custom_palette)
        wordcloud_layout.addWidget(self.custom_palette_button, 1, 5, 1, 1)

        self.bg_color_label = QLabel("Background Color:", self)
        wordcloud_layout.addWidget(self.bg_color_label, 2, 0, 1, 2)

        self.bg_color = QComboBox(self)
        self.bg_color.addItems(["white", "black", "gray", "blue", "red", "yellow"])
        self.bg_color.setToolTip(
            "Select the background color for the word cloud.\n"
            "Use contrast for better visibility.\n"
            "White or black backgrounds usually work best."
        )        
        wordcloud_layout.addWidget(self.bg_color, 2, 2, 1, 3)

        self.custom_bg_color_button = QPushButton("Custom", self)
        self.custom_bg_color_button.clicked.connect(self.select_custom_bg_color)
        wordcloud_layout.addWidget(self.custom_bg_color_button, 2, 5, 1, 1)

        self.title_label = QLabel("WordCloud Title:", self)
        wordcloud_layout.addWidget(self.title_label, 3, 0, 1, 2)

        self.title_entry = QLineEdit(self)
        self.title_entry.setPlaceholderText("Enter title (optional)")
        self.title_entry.setToolTip(
            "Enter a title for your word cloud (optional).\n"
            "This title will be displayed above the word cloud.\n"
            "Leave blank if no title is needed."
        )        
        wordcloud_layout.addWidget(self.title_entry, 3, 2, 1, 2)

        self.title_font_size = QSpinBox(self)
        self.title_font_size.setRange(8, 72)
        self.title_font_size.setValue(14)
        self.title_font_size.setToolTip(
            "Set the font size for the word cloud title.\n"
            "Larger values make the title more prominent.\n"
            "Recommended: 14-24 px for a balanced look.\n"
            "Too large titles may overlap with the word cloud."
        )        
        wordcloud_layout.addWidget(self.title_font_size, 3, 4, 1, 1)

        self.title_position = QComboBox(self)
        self.title_position.addItems(["Left", "Center", "Right"])
        self.title_position.setCurrentText("Center")
        self.title_position.setToolTip(
            "Choose where to display the title relative to the word cloud.\n"
            "Positioning affects the overall layout and readability.\n"
            "Recomended: Center"
        )        
        wordcloud_layout.addWidget(self.title_position, 3, 5, 1, 1)

        self.font_choice_label = QLabel("Font Choice:", self)
        wordcloud_layout.addWidget(self.font_choice_label, 4, 0, 1, 2)

        self.font_choice = QComboBox(self)
        self.font_choice.addItem("Default")
        self.font_choice.setToolTip(
            "Choose a font style for the word cloud.\n"
            "Different fonts affect readability and aesthetics.\n"
            "Sans-serif fonts are recommended for clarity."
        )        
        wordcloud_layout.addWidget(self.font_choice, 4, 2, 1, 4)
        QTimer.singleShot(100, self.load_matplotlib_fonts)

        self.min_font_size_label = QLabel("Minimum Font Size:", self)
        wordcloud_layout.addWidget(self.min_font_size_label, 5, 0, 1, 2)

        self.min_font_size_input = QSpinBox(self)
        self.min_font_size_input.setValue(11)
        self.min_font_size_input.setToolTip(
            "Set the smallest font size for words in the word cloud.\n"
            "Prevents low-frequency words from becoming too small to read.\n"
            "Recommended value: 10-12 px for readability."
        )        
        wordcloud_layout.addWidget(self.min_font_size_input, 5, 2, 1, 1)

        self.max_words_label = QLabel("Maximum Words:", self)
        wordcloud_layout.addWidget(self.max_words_label, 5, 3, 1, 2, Qt.AlignRight)

        self.max_words_input = QSpinBox(self)
        self.max_words_input.setMaximum(10000)
        self.max_words_input.setValue(200)
        self.max_words_input.setToolTip(
            "Set the maximum number of words displayed in the word cloud.\n"
            "Higher values provide more detail but may reduce clarity.\n"
            "Recommended: 100-200 words for balanced visualization"
        )        
        wordcloud_layout.addWidget(self.max_words_input, 5, 5, 1, 1)

        self.mask_label = QLabel("Mask Image:", self)
        wordcloud_layout.addWidget(self.mask_label, 6, 0, 1, 2)

        self.mask_path_label = QLineEdit("default (rectangle)", self)
        self.mask_path_label.setReadOnly(True)
        wordcloud_layout.addWidget(self.mask_path_label, 6, 2, 1, 4)

        self.mask_button = QPushButton("Load Mask Image", self)
        self.mask_button.clicked.connect(self.choose_mask)
        self.mask_button.setToolTip(
            "Upload an image to shape the word cloud (PNG/JPG/BMP).\n"
            "White areas will be ignored, and words will fill the dark areas.\n"
            "Use simple shapes for best results."
        )        
        wordcloud_layout.addWidget(self.mask_button, 7, 2, 1, 2)

        self.reset_mask_button = QPushButton("Remove Mask Image", self)
        self.reset_mask_button.clicked.connect(self.reset_mask)
        self.reset_mask_button.setToolTip(
            "Remove the selected mask image and revert to the default shape.\n"
            "The word cloud will be displayed in a rectangular format.\n"
            "Use this if you no longer want a custom shape for the word cloud."
        )        
        wordcloud_layout.addWidget(self.reset_mask_button, 7, 4, 1, 2)

        self.generate_wordcloud_button = QPushButton("Generate Word Cloud", self)
        self.generate_wordcloud_button.clicked.connect(self.generate_wordcloud)
        self.generate_wordcloud_button.setEnabled(False)
        wordcloud_layout.addWidget(self.generate_wordcloud_button, 8, 0, 1, 6)

        self.text_stats_button = QPushButton("View Text Statistics", self)
        self.text_stats_button.clicked.connect(self.text_analysis_report)
        self.text_stats_button.setEnabled(False)
        wordcloud_layout.addWidget(self.text_stats_button, 9, 0, 1, 3)

        self.save_wc_button = QPushButton("Save Word Cloud", self)
        self.save_wc_button.clicked.connect(self.save_wordcloud)
        self.save_wc_button.setEnabled(False)
        wordcloud_layout.addWidget(self.save_wc_button, 9, 3, 1, 3)

        wordcloud_group.setLayout(wordcloud_layout)
        layout.addWidget(wordcloud_group)

        sentiment_group = QGroupBox("Sentiment Analysis")
        sentiment_layout = QGridLayout()

        self.sentiment_mode_label = QLabel("Analysis Mode:", self)
        sentiment_layout.addWidget(self.sentiment_mode_label, 0, 0, 1, 2)

        self.sentiment_mode_combo = QComboBox(self)
        self.sentiment_mode_combo.addItems(AppConstants.SENTIMENT_MODES)
        self.sentiment_mode_combo.setToolTip(
            "TextBlob (Best for formal texts like news articles and reports.)\n"
            "VADER (Best for social media, tweets, and informal texts with slang/emojis.)\n"
            "Flair (Best for long-form content and complex sentiment analysis using deep learning.)"
        )
        self.sentiment_mode_combo.currentTextChanged.connect(self.change_sentiment_mode)
        sentiment_layout.addWidget(self.sentiment_mode_combo, 0, 2, 1, 3)

        self.sentiment_mode_info_button = QPushButton("Info", self)
        self.sentiment_mode_info_button.clicked.connect(self.show_sentiment_mode_info)
        sentiment_layout.addWidget(self.sentiment_mode_info_button, 0, 5, 1, 1)

        self.custom_lexicon_button = QPushButton("Load Lexicon", self)
        self.custom_lexicon_button.clicked.connect(self.load_custom_lexicon)
        self.custom_lexicon_button.setEnabled(False)
        sentiment_layout.addWidget(self.custom_lexicon_button, 1, 0, 1, 3)

        self.custom_model_button = QPushButton("Load Model", self)
        self.custom_model_button.clicked.connect(self.load_custom_model)
        self.custom_model_button.setEnabled(False)
        sentiment_layout.addWidget(self.custom_model_button, 1, 3, 1, 3)

        self.sentiment_button = QPushButton("Analyze Sentiment", self)
        self.sentiment_button.clicked.connect(self.analyze_sentiment)
        self.sentiment_button.setEnabled(False)
        sentiment_layout.addWidget(self.sentiment_button, 2, 0, 1, 6)

        sentiment_group.setLayout(sentiment_layout)
        layout.addWidget(sentiment_group)

        self.topic_tab = TopicAnalysisTab(parent=self)
        layout.addWidget(self.topic_tab)

        bottom_group = QGroupBox()
        bottom_layout = QVBoxLayout()
        
        button_layout = QHBoxLayout()
        
        # Tombol About
        self.about_button = QPushButton("About", self)
        self.about_button.clicked.connect(self.show_about)
        self.about_button.setFixedWidth(100)  # Atur ukuran tetap
        button_layout.addWidget(self.about_button)
        
        # Tambahkan tombol Clear Cache di sini
        self.clear_cache_button = QPushButton("Clear Cache", self)
        self.clear_cache_button.setToolTip("Manually clear all cached data to free memory")
        self.clear_cache_button.clicked.connect(self.clear_caches)
        self.clear_cache_button.setFixedWidth(100)  # Atur ukuran tetap
        button_layout.addWidget(self.clear_cache_button)
        
        # Spacer untuk mendorong tombol ke kanan
        button_layout.addStretch()
        
        # Tombol STOP
        self.panic_button = QPushButton("STOP", self)
        self.panic_button.setStyleSheet("background-color: #ff6666;")
        self.panic_button.clicked.connect(self.stop_all_processes)
        self.panic_button.setFixedWidth(100)  # Atur ukuran tetap
        button_layout.addWidget(self.panic_button)
        
        # Tombol Exit
        self.quit_button = QPushButton("Exit", self)
        self.quit_button.clicked.connect(self.close)
        self.quit_button.setFixedWidth(100)  # Atur ukuran tetap
        button_layout.addWidget(self.quit_button)
        
        bottom_layout.addLayout(button_layout)

        bottom_group.setLayout(bottom_layout)
        layout.addWidget(bottom_group)

        self.progress_timer = QTimer()
        self.progress_timer.timeout.connect(self._update_progress_animation)
        self.progress_timer.setInterval(20)

        container = QWidget()
        container.setLayout(layout)
        self.setCentralWidget(container)

    def load_matplotlib_fonts(self):
        try:
            from matplotlib import font_manager
            self.font_map = {}

            weight_conversion = {
                100: "Thin", 150: "Extra Thin",  200: "Extra Light",  250: "Light Thin",  300: "Light", 350: "Book", 
                400: "Regular",  450: "Text", 500: "Medium", 550: "Demi Medium", 600: "Semi Bold", 650: "Demi Bold", 
                700: "Bold", 750: "Extra Bold Light", 800: "Extra Bold", 850: "Heavy", 900: "Black", 950: "Extra Black", 
                1000: "Ultra Black" 
            }

            for font in font_manager.fontManager.ttflist:
                try:     
                    family = font.family_name if hasattr(font, "family_name") else font.name
                    style = font.style_name.lower() if hasattr(font, "style_name") else font.style.lower()
                    weight = weight_conversion.get(font.weight, str(font.weight))

                    display_parts = [family]
                    if "italic" in style or "oblique" in style:
                        display_parts.append("Italic")
                    elif "normal" not in style:
                        display_parts.append(style.title())

                    if weight != "Regular":
                        display_parts.append(weight)

                    display_name = " ".join(display_parts)
                    if display_name not in self.font_map:
                        self.font_map[display_name] = font.fname

                except Exception as e:
                    continue

            self.font_choice.clear()
            self.font_choice.addItem("Default")
            self.font_choice.addItems(sorted(self.font_map.keys()))

        except Exception as e:
            QMessageBox.warning(self, "Font Error", f"Failed to load fonts: {str(e)}")
            self.font_choice.addItems(["Arial", "Times New Roman", "Verdana"])

    def load_colormaps(self):
        try:
            import matplotlib.pyplot as plt
            self.color_theme.addItems(plt.colormaps())
        except ImportError as e:
            self.color_theme.clear()
            self.color_theme.addItem("Default")
            QMessageBox.warning(self, "Dependency Error", f"Failed to load color maps: {str(e)}")

    def analyze_sentiment(self):
        """Analyze sentiment with button state management"""
        if not self.text_data:
            QMessageBox.critical(self, "Error", "No text data available for sentiment analysis.\nPlease load a text file first.")
            return

        if self.sentiment_mode in ["Flair", "TextBlob"]:
            try:
                from langdetect import detect
            except:
                pass

        if (self.sentiment_mode == "Flair" and not self.flair_classifier and self.sentiment_mode_combo.currentText() == "Flair"):
            QMessageBox.warning(self, "Model Loading", "Default Flair model is still loading. Please wait...")
            return

        if self.sentiment_mode == "Flair (Custom Model)":
            if not self.flair_classifier_cuslang:
                QMessageBox.warning(self, "Model Required", "Please load a custom Flair model first!")
                return

        classifier = self.flair_classifier_cuslang if self.sentiment_mode == "Flair (Custom Model)" else self.flair_classifier

        self.button_manager.disable_other_buttons('sentiment_button')

        self.sentiment_thread = SentimentAnalysisThread(
            self.text_data, 
            self.sentiment_mode,
            self.vader_analyzer,
            self.flair_classifier,
            self.flair_classifier_cuslang,
            self.textblob_analyzer
        )  
        self.sentiment_thread.translation_failed.connect(self.handle_translation_failure)
        self.sentiment_thread.offline_warning.connect(self.handle_offline_warning)
        self.sentiment_thread.sentiment_analyzed.connect(self.on_sentiment_analyzed)

        self.thread_manager.add_thread(self.sentiment_thread)

        self.set_progress('model')
        self.sentiment_thread.start()

    def stop_all_processes(self):
        """Enhanced process termination including word cloud generation"""
        logger.info("Starting enhanced process termination...")
        self.disable_buttons()

        self.progress_timer.stop()
        self.unified_progress_bar.setValue(0)
        self.unified_progress_bar.setVisible(False)

        QApplication.processEvents()        

        try:
            self.threads_mutex.lock()
            all_threads = self.active_threads.copy()
            self.threads_mutex.unlock()

            stopped_threads = []
            failed_to_stop = []

            if hasattr(self, 'file_loader_thread') and self.file_loader_thread:
                if self.file_loader_thread.isRunning():
                    logger.info("Terminating File Loader thread...")
                    all_threads.append(self.file_loader_thread)

            if hasattr(self, 'sentiment_thread') and self.sentiment_thread:
                if self.sentiment_thread.isRunning():
                    logger.info("Terminating Sentiment Analysis thread...")
                    all_threads.append(self.sentiment_thread)

            if hasattr(self, 'flair_loader_thread') and self.flair_loader_thread:
                if self.flair_loader_thread.isRunning():
                    logger.info("Terminating Flair Model Loader thread...")
                    all_threads.append(self.flair_loader_thread)

            if hasattr(self, 'lexicon_loader_thread') and self.lexicon_loader_thread:
                if self.lexicon_loader_thread.isRunning():
                    logger.info("Terminating Lexicon Loader thread...")
                    all_threads.append(self.lexicon_loader_thread)

            if hasattr(self, 'model_loader_thread') and self.model_loader_thread:
                if self.model_loader_thread.isRunning():
                    logger.info("Terminating Model Loader thread...")
                    all_threads.append(self.model_loader_thread)

            if hasattr(self, 'import_thread') and self.import_thread:
                if self.import_thread.isRunning():
                    logger.info("Terminating Import thread...")
                    self.import_thread.stop()
                    all_threads.append(self.import_thread)

            if hasattr(self, 'topic_tab'):
                if hasattr(self.topic_tab, 'topic_thread') and self.topic_tab.topic_thread:
                    if self.topic_tab.topic_thread.isRunning():
                        logger.info("Terminating Topic Analysis thread...")
                        all_threads.append(self.topic_tab.topic_thread)
                        
                if hasattr(self.topic_tab, 'keyword_thread') and self.topic_tab.keyword_thread:
                    if self.topic_tab.keyword_thread.isRunning():
                        logger.info("Terminating Keyword Extraction thread...")
                        all_threads.append(self.topic_tab.keyword_thread)

            for thread in all_threads:
                try:
                    if not thread.isRunning():
                        continue
                        
                    logger.info(f"Terminating {type(thread).__name__}...")
                    thread.requestInterruption()
                    
                    if thread.wait(800):
                        logger.info(f"Graceful stop: {type(thread).__name__}")
                        stopped_threads.append(thread)
                    else:
                        logger.info(f"Forcing termination: {type(thread).__name__}")
                        thread.terminate()
                        if thread.wait(500):
                            stopped_threads.append(thread)
                        else:
                            failed_to_stop.append(thread)
                            
                except Exception as e:
                    logger.error(f"Termination error: {str(e)}")
                    failed_to_stop.append(thread)

            self.threads_mutex.lock()
            self.active_threads = [
                t for t in self.active_threads 
                if t not in stopped_threads and t.isRunning()
            ]
            self.threads_mutex.unlock()

            # report = f"""
            # <b>Termination Report:</b>
            # • Total threads: {len(all_threads)}
            # • Successfully stopped: {len(stopped_threads)}
            # • Failed to stop: {len(failed_to_stop)}
            # • Remaining active: {len(self.active_threads)}
            # """
            # QMessageBox.information(self, "Process Report", report)
            report = f"Threads: {len(all_threads)} • Stopped: {len(stopped_threads)} • Failed: {len(failed_to_stop)} • Active: {len(self.active_threads)}"
            self.statusBar().showMessage(report, 3000)

        except Exception as e:
            logger.error(f"Critical termination error: {str(e)}")
        finally:
            self.progress_timer.stop()
            self.unified_progress_bar.setValue(0)
            self.unified_progress_bar.setVisible(False)      
            self.enable_buttons()
            logger.info("Enhanced termination completed")

    def disable_buttons(self):
        """Disable all buttons during processing"""
        process_buttons = [
            self.generate_wordcloud_button,
            self.sentiment_button, 
            self.custom_lexicon_button,
            self.custom_model_button,
            self.save_wc_button,
            self.text_stats_button
        ]
        
        if hasattr(self, 'topic_tab'):
            if hasattr(self.topic_tab, 'analyze_topics_btn'):
                process_buttons.append(self.topic_tab.analyze_topics_btn)
            if hasattr(self.topic_tab, 'extract_keywords_btn'):
                process_buttons.append(self.topic_tab.extract_keywords_btn)
                
        for button in process_buttons:
            button.setEnabled(False)
            
        self.load_file_button.setEnabled(True)
        self.panic_button.setEnabled(True)
        
        QApplication.processEvents()

    def disable_processing_buttons(self):
        """Alias for disable_buttons for compatibility"""
        self.disable_buttons()

    def enable_buttons(self):
        """Re-enable buttons after processing"""
        has_text = bool(self.text_data)
        
        self.load_file_button.setEnabled(True)
        self.view_fulltext_button.setEnabled(has_text)
        self.generate_wordcloud_button.setEnabled(has_text)
        self.text_stats_button.setEnabled(has_text)
        self.save_wc_button.setEnabled(has_text)
        self.sentiment_button.setEnabled(has_text)
        self.summarize_button.setEnabled(has_text)
        
        if hasattr(self, 'topic_tab'):
            if hasattr(self.topic_tab, 'analyze_topics_btn'):
                self.topic_tab.analyze_topics_btn.setEnabled(has_text)
            if hasattr(self.topic_tab, 'extract_keywords_btn'):
                self.topic_tab.extract_keywords_btn.setEnabled(has_text)
        
        if self.sentiment_mode == "VADER (Custom Lexicon)":
            self.custom_lexicon_button.setEnabled(True)
            self.custom_model_button.setEnabled(False)
        elif self.sentiment_mode == "Flair (Custom Model)":
            self.custom_lexicon_button.setEnabled(False)  
            self.custom_model_button.setEnabled(True)
        else:
            self.custom_lexicon_button.setEnabled(False)
            self.custom_model_button.setEnabled(False)
            
        QApplication.processEvents()    

    def closeEvent(self, event):
        """Handle application close event with proper cleanup"""
        # Ask for confirmation if there are active threads
        if self.thread_manager.active_threads:
            reply = QMessageBox.question(
                self, 
                'Confirm Exit',
                'There are active processes running.\n Are you sure you want to exit?',
                QMessageBox.Yes | QMessageBox.No, 
                QMessageBox.No
            )
        else:
            reply = QMessageBox.question(
                self,
                "Confirm Exit",
                "<b>Are you sure you want to exit?</b><br>",
                QMessageBox.StandardButton.Yes | QMessageBox.StandardButton.No,
                QMessageBox.StandardButton.No
        )
            
        if reply == QMessageBox.No:
            event.ignore()
            return
        
        # Proceed with cleanup
        self.cleanup_threads()
        
        # Save application settings
        # self.save_settings()
        
        # Accept the close event
        event.accept()

    def cleanup(self):
        """Enhanced application cleanup"""
        logger.info("Performing application cleanup")
        
        # Hentikan pemantauan memori
        self.cache_manager.stop_monitoring()
        
        # Hentikan semua thread dengan ThreadManager
        self.thread_manager.stop_all_threads()
        
        # Bersihkan cache
        self.cache_manager.clear_all_caches()
        self.get_wordcloud.cache_clear()
        LazyLoader._cache.clear()
        self._cached_models.clear()
        self._cached_fonts.clear()
        self._cached_colormaps.clear()
        
        # Bersihkan resource
        self.resource_manager.cleanup()
        
        # Bersihkan CUDA cache jika tersedia
        if torch.cuda.is_available():
            torch.cuda.empty_cache()
        
        # Tutup semua figure matplotlib
        import matplotlib.pyplot as plt
        plt.close('all')
        
        # Bersihkan temporary files
        temp_dir = Path(tempfile.gettempdir()) / "textplora_cache"
        if temp_dir.exists():
            shutil.rmtree(temp_dir, ignore_errors=True)
        
        # Proses events yang tersisa
        QApplication.processEvents()

    def show_about(self):
        """Show about dialog using DialogFactory"""
        about_text = base64.b64decode(AppConstants.ABOUT_TEXT.encode()).decode()
        about_text = about_text.replace("{logo_path}", str(logo_path.as_uri()))
        
        dialog = DialogFactory.create_info_dialog(
            self, "About Textplora", about_text, modal=False
        )
        dialog.show()

    def open_file(self):
        dialog = QFileDialog(self)
        dialog.setWindowTitle("Load Text File")
        dialog.setNameFilter("Supported Files (*.txt *.pdf *.doc *.docx *.csv *.xlsx *.xls);;All Files (*)")
        dialog.setOptions(QFileDialog.Options())

        if dialog.exec_():
            file_path = dialog.selectedFiles()[0]
            if not file_path:
                return

            self.set_progress('file')

            try:
                self.file_loader_thread = FileLoaderThread(file_path)
                self.file_loader_thread.file_loaded.connect(self.on_file_loaded)
                self.file_loader_thread.file_error.connect(self.handle_file_error)

                self.thread_manager.add_thread(self.file_loader_thread)

                self.file_loader_thread.start()

            except Exception as e:
                self.unified_progress_bar.setVisible(False)
                QMessageBox.critical(self, "Error", f"Failed to load file: {e}")

    def handle_file_error(self, error_message):
        self.unified_progress_bar.setVisible(False)
        QMessageBox.critical(self, "File Load Error", error_message)
        self._reset_file_state()

    def on_file_loaded(self, file_path, text_data):
        """Handle file loading completion"""
        self.set_progress('file', False)

        try:
            if not text_data.strip():
                raise ValueError("File appears empty after loading")

            self.text_data = text_data
            self.file_name.setText(os.path.basename(file_path))
            self.topic_tab.set_text(self.text_data)

            self.generate_wordcloud_button.setEnabled(True)
            self.save_wc_button.setEnabled(True)
            self.text_stats_button.setEnabled(True)
            self.view_fulltext_button.setEnabled(True)
            self.summarize_button.setEnabled(True)
            
            self.load_file_button.setEnabled(True)

            self.change_sentiment_mode(self.sentiment_mode)

        except ValueError as e:
            self.handle_file_error(str(e))
            self._reset_file_state()
            
        self.load_file_button.setEnabled(True)

    def _reset_file_state(self):
        self.text_data = ""
        self.file_name.setText("")
        self.generate_wordcloud_button.setEnabled(False)
        self.save_wc_button.setEnabled(False)
        self.text_stats_button.setEnabled(False)
        self.view_fulltext_button.setEnabled(False)
        self.sentiment_button.setEnabled(False)
        self.summarize_button.setEnabled(False)

    def choose_mask(self):
        options = QFileDialog.Options()
        mask_path, _ = QFileDialog.getOpenFileName(
            self, "Select Mask Image", "", "Image Files (*.png *.jpg *.bmp)", options=options
        )
        if mask_path:
            self.mask_path = mask_path
            self.mask_path_label.setText(f"{mask_path}")

    def reset_mask(self):
        self.mask_path = ""
        self.mask_path_label.setText("default (rectangle)")

    def import_stopwords(self):
        custom_words = self.stopword_entry.toPlainText().strip().lower()
        if custom_words:
            self.additional_stopwords = set(custom_words.split())
        return STOPWORDS.union(self.additional_stopwords)

    def text_analysis_report(self):
        if not self.text_data:
            QMessageBox.critical(self, "Error", "Load text file first!")
            return

        self.button_manager.disable_other_buttons('text_stats_button')
        
        try:
            stopwords = self.import_stopwords()
            words = [word.lower() for word in self.text_data.split() if word.lower() not in stopwords]
            word_counts = Counter(words)
            sorted_word_counts = sorted(word_counts.items(), key=lambda item: item[1], reverse=True)

            text_length = len(self.text_data)
            word_count = len(words)
            char_count_excl_spaces = len(self.text_data.replace(" ", ""))
            avg_word_length = char_count_excl_spaces / word_count if word_count > 0 else 0
            most_frequent_words = [word for word, count in sorted_word_counts[:5]]

            if hasattr(self, "stats_dialog") and self.stats_dialog is not None:
                self.stats_dialog.close()

            self.stats_dialog = QDialog(self)
            self.stats_dialog.setWindowModality(Qt.NonModal)
            self.stats_dialog.setWindowTitle("Text Analysis Report")
            self.stats_dialog.setMinimumSize(500, 400)
            self.stats_dialog.setSizeGripEnabled(True)

            layout = QVBoxLayout()

            text_browser = QTextBrowser()

            txtstat_content = f"""
            <h3 style="text-align: center;">Text Analysis Overview</h3>
            <table border="1" cellspacing="0" cellpadding="2" width="100%" style="margin-top: 20px;">
                <tr style="background-color: #d3d3d3; color: black">
                    <th>Metric</th>
                    <th>Value</th>
                </tr>
                <tr><td>Text Length</td><td>{text_length} characters</td></tr>
                <tr><td>Word Count</td><td>{word_count}</td></tr>
                <tr><td>Character Count (excluding spaces)</td><td>{char_count_excl_spaces}</td></tr>
                <tr><td>Average Word Length</td><td>{avg_word_length:.2f}</td></tr>
                <tr><td>Most Frequent Words</td><td>{", ".join(most_frequent_words)}</td></tr>
            </table>

            <h3 style="margin-top: 20px; text-align: center;">Word Count</h3>  
            <table border="1" cellspacing="0" cellpadding="2" width="100%"style="margin-top: 20px;">
                <tr style="background-color: #d3d3d3; color: black">
                    <th>Word</th>
                    <th>Count</th>
                </tr>
            """
            for word, count in sorted_word_counts:
                txtstat_content += f"<tr><td>{word}</td><td>{count}</td></tr>"
            txtstat_content += "</table>"

            text_browser.setHtml(txtstat_content)
            text_browser.setOpenExternalLinks(True)
            text_browser.setReadOnly(True)
            layout.addWidget(text_browser)

            close_button = QPushButton("Close")
            close_button.clicked.connect(self.stats_dialog.accept)
            layout.addWidget(close_button)

            self.stats_dialog.setLayout(layout)
            self.stats_dialog.finished.connect(lambda: self.button_manager.restore_states())
            self.stats_dialog.show()
            
            self.button_manager.restore_states()
            
        except Exception as e:
            QMessageBox.critical(self, "Error", f"Failed to generate text analysis: {e}")
            self.button_manager.restore_states()

    def generate_wordcloud(self):
        """Generate word cloud with enhanced thread handling"""
        import matplotlib.pyplot as plt
        
        if not self.text_data:
            QMessageBox.critical(self, "Error", "Load text file first!")
            return

        try:
            if self.current_figure:
                plt.close(self.current_figure)
                self.current_figure = None

            font_path = None
            selected_font = self.font_choice.currentText()
            if selected_font != "Default":
                font_path = self.font_map.get(selected_font)
                if not font_path or not os.path.exists(font_path):
                    QMessageBox.warning(self, "Font Error", f"Font file not found: {font_path}")
                    return

            mask = None
            if self.mask_path:
                try:
                    mask = np.array(Image.open(self.mask_path))
                except Exception as e:
                    QMessageBox.critical(self, "Error", f"Failed to load mask image: {e}")
                    return

            colormap = self.color_theme.currentText()
            if (colormap in self.custom_color_palettes):
                colors = self.custom_color_palettes[colormap]
                colormap = LinearSegmentedColormap.from_list(colormap, colors)

            mem_available, cpu_available = self.perf_monitor.check_resources()
            if mem_available < 500:
                self.cleanup_cache()

            self.button_manager.disable_other_buttons('generate_wordcloud_button')
            self.set_progress('wordcloud')

            wc = MPWordCloud(
                width=800, height=400,
                background_color=self.bg_color.currentText(),
                stopwords=self.import_stopwords(),
                colormap=colormap,
                max_words=self.max_words_input.value(),
                min_font_size=self.min_font_size_input.value(),
                mask=mask,
                font_path=font_path
            )
            
            wc.generate_threaded(
                self.text_data,
                progress_callback=lambda p: self.unified_progress_bar.setValue(p),
                finished_callback=lambda wc_obj: self.display_wordcloud(wc_obj),
                error_callback=self.handle_wordcloud_error
            )

            self.current_wordcloud = wc

        except Exception as e:
            self.handle_wordcloud_error(str(e))

    def display_wordcloud(self, wordcloud):
        """Display generated wordcloud"""
        try:
            import matplotlib.pyplot as plt
            
            plt.ion()
            self.current_figure = plt.figure(figsize=(10, 5))
            plt.imshow(wordcloud, interpolation="bilinear")

            title_text = self.title_entry.text().strip()
            if title_text:
                title_font = None
                if self.font_choice.currentText() != "Default":
                    title_font = FontProperties(
                        fname=self.font_map.get(self.font_choice.currentText()),
                        size=self.title_font_size.value()
                    )
                plt.title(
                    title_text, 
                    loc=self.title_position.currentText().lower(),
                    fontproperties=title_font
                )

            plt.axis("off")
            plt.show(block=False)

            self.save_wc_button.setEnabled(True)
            
        except Exception as e:
            self.handle_wordcloud_error(str(e))
        finally:
            self.button_manager.restore_states()
            self.set_progress('wordcloud', False)

    def handle_wordcloud_error(self, error_msg):
        """Handle word cloud generation errors"""
        import matplotlib.pyplot as plt
        
        if self.current_figure:
            try:
                plt.close(self.current_figure)
            except:
                pass
            finally:
                self.current_figure = None
                
        QMessageBox.critical(self, "Error", f"Failed to generate word cloud: {error_msg}")
        self.button_manager.restore_states()
        self.set_progress('wordcloud', False)

    def save_wordcloud(self):
        self.button_manager.disable_other_buttons('save_wc_button')
        
        try:
            options = QFileDialog.Options()
            save_path, _ = QFileDialog.getSaveFileName(
                self, "Save WordCloud", "", "PNG file (*.png);;JPG file (*.jpg)", options=options
            )
            if not save_path:
                self.button_manager.restore_states()
                return

            stopwords = self.import_stopwords()
            mask = None

            font_path = None
            if self.font_choice.currentText() != "Default":
                from matplotlib import font_manager
                selected_font = self.font_choice.currentText()
                font_path = self.font_map.get(selected_font)

                try:
                    font_manager.findfont(self.font_choice.currentText())
                    font_path = self.font_choice.currentText()
                except:
                    QMessageBox.warning(self, "Font Error", "Selected font not found, using default")

            if self.mask_path:
                try:
                    mask = np.array(Image.open(self.mask_path))
                except Exception as e:
                    QMessageBox.critical(self, "Error", f"Failed to load mask image: {e}")
                    return

            try:
                wc = WordCloud(
                    width=800, height=400, background_color=self.bg_color.currentText(), stopwords=stopwords,
                    colormap=self.color_theme.currentText(), max_words=self.max_words_input.value(),
                    min_font_size=self.min_font_size_input.value(), mask=mask,
                    font_path=None if self.font_choice.currentText() == "Default" else self.font_choice.currentText()
                ).generate(self.text_data)
                wc.to_file(save_path)
                QMessageBox.information(self, "Succeed", "Word cloud saved!")
            except Exception as e:
                QMessageBox.critical(self, "Error", f"Failed to save word cloud: {e}")
        finally:
            self.button_manager.restore_states()

    def create_custom_palette(self):
        dialog = QDialog(self)
        dialog.setWindowTitle("Create Custom Palette")
        dialog.setFixedSize(400, 150)

        layout = QVBoxLayout()
        color_list = []

        def add_color():
            color = QColorDialog.getColor()
            if color.isValid():
                color_list.append(color.name())
                color_label.setText(", ".join(color_list))

        color_label = QLabel("", self)
        layout.addWidget(color_label)

        add_color_button = QPushButton("Add Color", self)
        add_color_button.clicked.connect(add_color)
        layout.addWidget(add_color_button)

        save_palette_button = QPushButton("Save Palette", self)
        save_palette_button.clicked.connect(lambda: self.save_custom_palette(color_list, dialog))
        layout.addWidget(save_palette_button)

        dialog.setLayout(layout)
        dialog.exec()

    def save_custom_palette(self, color_list, dialog):
        palette_name, ok = QInputDialog.getText(self, "Save Palette", "Enter palette name:")
        if ok and palette_name:
            self.custom_color_palettes[palette_name] = color_list
            self.color_theme.addItem(palette_name)
            dialog.accept()

    def select_custom_bg_color(self):
        color = QColorDialog.getColor()
        if color.isValid():
            self.bg_color.addItem(color.name())
            self.bg_color.setCurrentText(color.name())

    def view_full_text(self):
        if not self.text_data:
            QMessageBox.critical(self, "Error", "No text data available to display.")
            return
        
        try:
            if hasattr(self, "text_dialog") and self.text_dialog is not None:
                self.text_dialog.close()

            self.text_dialog = QDialog(self)
            self.text_dialog.setWindowModality(Qt.NonModal)
            self.text_dialog.setWindowTitle("Full Text")
            self.text_dialog.setMinimumSize(500, 400)
            self.text_dialog.setSizeGripEnabled(True)

            layout = QVBoxLayout()

            text_browser = QTextBrowser()
            text_browser.setPlainText(self.text_data)
            text_browser.setOpenExternalLinks(True)
            text_browser.setReadOnly(True)
            layout.addWidget(text_browser)

            close_button = QPushButton("Close")
            close_button.clicked.connect(self.text_dialog.accept)
            layout.addWidget(close_button)

            self.text_dialog.setLayout(layout)
            self.text_dialog.show()
            
        except Exception as e:
            QMessageBox.critical(self, "Error", f"Failed to display text: {e}")

    def handle_offline_warning(self, msg):
        self.unified_progress_bar.setVisible(False)
        QMessageBox.warning(self, "Offline Mode", msg)

    def handle_translation_failure(self, error_msg):
        """Handle translation failure gracefully"""
        self.unified_progress_bar.setVisible(False)
        QMessageBox.warning(self, "Translation Error", error_msg)
        self.enable_buttons()

    def on_sentiment_analyzed(self, result):
        self.set_progress('model', False)
        self.button_manager.restore_states()
        
        dialog = QDialog(self)
        dialog.setWindowTitle(f"Sentiment Analysis Results - {self.sentiment_mode}")
        dialog.setMinimumSize(400, 350)
        dialog.setSizeGripEnabled(True)

        layout = QVBoxLayout()
        text_browser = QTextBrowser()

        score_type = ""
        if self.sentiment_mode in ["VADER", "VADER (Custom Lexicon)"]:
            score_type = "Compound Score"
        elif self.sentiment_mode in ["TextBlob", "TextBlob (Custom Lexicon)"]:
            score_type = "Polarity Score"
        elif self.sentiment_mode in ["Flair", "Flair (Custom Model)"]:
            score_type = "Confidence Score"

        sentiment_result = f"""
        <h3 style="text-align: center;">Sentiment Analysis Results</h3>
        <table border="1" cellspacing="0" cellpadding="2" width="100%" style="margin-top: 20px;">
            <tr style="background-color: #d3d3d3; color: black">
                <th>Metric</th>
                <th>Value</th>
            </tr>
            <tr><td>Analysis Mode</td><td>{self.sentiment_mode}</td></tr>
            <tr><td>Sentiment Label</td><td><b>{result["sentiment_label"]}</b></td></tr>
            <tr><td>Positive Sentiment</td><td>{result["positive_score"]:.2f}</td></tr>
            <tr><td>Neutral Sentiment</td><td>{result["neutral_score"]:.2f}</td></tr>
            <tr><td>Negative Sentiment</td><td>{result["negative_score"]:.2f}</td></tr>
            <tr><td>{score_type}</td><td>{result["compound_score"]:.2f}</td></tr>
        """

        if self.sentiment_mode in ["TextBlob", "TextBlob (Custom Lexicon)"]:
            try:
                subj_value = float(result["subjectivity"])
                sentiment_result += f'<tr><td>Subjectivity</td><td>{subj_value:.2f}</td></tr>'
            except (ValueError, TypeError):
                sentiment_result += f'<tr><td>Subjectivity</td><td>{result["subjectivity"]}</td></tr>'

        sentiment_result += "</table>"

        sentiment_notes = ""

        if self.sentiment_mode in ["TextBlob", "TextBlob (Custom Lexicon)"]:
            sentiment_notes = """
            <p><b>Note:</b></p>
            <ul>
                <li>Sentiment analysis is performed using <b>TextBlob</b>, which assigns a sentiment polarity score between -1 and 1.</li>
                <li><b>Polarity:</b> A value close to <b>-1</b> indicates a strongly negative sentiment, while a value close to <b>1</b> indicates a strongly positive sentiment.</li>
                <li><b>Neutral Range:</b> Scores near <b>0</b> suggest a neutral sentiment.</li>
                <li><b>Subjectivity Score:</b> Ranges from <b>0</b> (very objective) to <b>1</b> (very subjective).</li>
                <li><b>Objective vs Subjective:</b> 
                    <ul>
                        <li>A low subjectivity score (<b>≤ 0.3</b>) suggests that the text is more factual and objective.</li>
                        <li>A high subjectivity score (<b>≥ 0.7</b>) indicates that the text contains opinions, emotions, or subjective statements.</li>
                    </ul>
                </li>
                <li>Results may vary depending on the text's context, sarcasm, and linguistic nuances.</li>
            </ul>
            """

        elif self.sentiment_mode in ["VADER", "VADER (Custom Lexicon)"]:
            sentiment_notes = """
            <p><b>Note:</b></p>
            <ul>
                <li>Sentiment analysis is performed using <b>VADER</b>, which is optimized for social media, short texts, and informal language.</li>
                <li><b>Compound Score:</b> The overall sentiment score, ranging from <b>-1</b> (very negative) to <b>1</b> (very positive).</li>
                <li><b>Thresholds:</b>
                    <ul>
                        <li>Compound score <b>≥ 0.05</b>: Positive sentiment.</li>
                        <li>Compound score <b>≤ -0.05</b>: Negative sentiment.</li>
                        <li>Compound score between <b>-0.05</b> and <b>0.05</b>: Neutral sentiment.</li>
                    </ul>
                </li>
                <li>VADER considers punctuation, capitalization, and emoticons to enhance sentiment detection.</li>
            </ul>
            """

        elif self.sentiment_mode in ["Flair", "Flair (Custom Model)"]:
            sentiment_notes = """
            <p><b>Note:</b></p>
            <ul>
                <li>Sentiment analysis is performed using <b>Flair</b>, a deep learning-based model for text classification.</li>
                <li><b>Model:</b> Flair uses a pre-trained <b>Bidirectional LSTM</b> trained on large sentiment datasets.</li>
                <li><b>Sentiment Labels:</b> The output consists of two possible classifications:
                    <ul>
                        <li><b>POSITIVE</b>: Indicates an overall positive sentiment.</li>
                        <li><b>NEGATIVE</b>: Indicates an overall negative sentiment.</li>
                    </ul>
                </li>
                <li><b>Confidence Score:</b> Flair provides a probability score (0 to 1) indicating how confident the model is in its classification.</li>
                <li>Unlike lexicon-based approaches (e.g., TextBlob, VADER), Flair captures <b>context and word relationships</b>, making it more robust for longer and complex texts.</li>
            </ul>
            """

        text_browser.setHtml(sentiment_result + sentiment_notes)
        text_browser.setOpenExternalLinks(True)
        text_browser.setReadOnly(True)
        layout.addWidget(text_browser)

        close_button = QPushButton("Close")
        close_button.clicked.connect(dialog.accept)
        layout.addWidget(close_button)

        dialog.setLayout(layout)
        dialog.show()

    @lru_cache(maxsize=128)
    def get_most_frequent_words(self, text, n):
        stopwords = self.import_stopwords().union(STOPWORDS)
        words = [word.lower() for word in text.split() if word.lower() not in stopwords]
        word_counts = Counter(words)
        most_common = word_counts.most_common(n)
        return [word for word, count in most_common]

    def change_sentiment_mode(self, mode):
        """Handle sentiment mode changes with proper button states"""
        self.sentiment_mode = mode
        status_text = ""
        has_text = bool(self.text_data.strip())

        self.custom_lexicon_button.setEnabled(False)
        self.custom_model_button.setEnabled(False)
        self.sentiment_button.setEnabled(False)

        if mode == "Flair":
            if self.flair_classifier:
                status_text = "Flair: Ready"
                self.sentiment_button.setEnabled(has_text)
            else:
                status_text = "Loading..."
                self.load_flair_model()

        elif mode == "Flair (Custom Model)":
            if not self.flair_classifier:
                self.load_flair_model()
                status_text = "Initializing Flair environment..."
            else:
                self.custom_model_button.setEnabled(True)
                if self.flair_classifier_cuslang:  
                    status_text = "Ready"
                    self.sentiment_button.setEnabled(has_text)
                else:
                    status_text = "Load custom model first!"

        elif mode == "VADER (Custom Lexicon)":
            self.custom_lexicon_button.setEnabled(True)
            status_text = "Load custom lexicon first!" if not self.custom_lexicon_path else "Ready"
            self.sentiment_button.setEnabled(has_text and bool(self.custom_lexicon_path))
            if not self.custom_lexicon_path:
                QMessageBox.warning(
                    self, "Lexicon Required", 
                    "Please load a custom lexicon before analyzing sentiment."
                )

        elif mode == "TextBlob (Custom Lexicon)":
            self.custom_lexicon_button.setEnabled(True)
            status_text = "Load custom lexicon first!" if not self.custom_textblob_lexicon_path else "Ready"
            self.sentiment_button.setEnabled(has_text and bool(self.custom_textblob_lexicon_path))
            if not self.custom_textblob_lexicon_path:
                QMessageBox.warning(
                    self, "Lexicon Required", 
                    "Please load a custom lexicon before analyzing sentiment."
                )

        else:
            status_text = "Ready" if has_text else "Load text first!"
            self.sentiment_button.setEnabled(has_text)

        self.sentiment_button.setToolTip(
            f"{mode} - {status_text}\nClick to analyze sentiment using {mode}"
        )

    def load_custom_lexicon(self):
        options = QFileDialog.Options()
        lexicon_path, _ = QFileDialog.getOpenFileName(
            self, "Select Custom Lexicon File", "", "Text Files (*.txt)", options=options
        )
        if lexicon_path:
            try:
                self.button_manager.disable_other_buttons('custom_lexicon_button')
                self.custom_lexicon_button.setText("Loading Lexicon...")
                self.custom_lexicon_button.setEnabled(False)
                self.sentiment_button.setEnabled(False)

                self.lexicon_loader_thread = CustomFileLoaderThread(
                    lexicon_path, 
                    "TextBlob (Custom Lexicon)" if self.sentiment_mode == "TextBlob (Custom Lexicon)" else "lexicon"
                )
                self.lexicon_loader_thread.file_loaded.connect(self.on_lexicon_loaded)
                self.thread_manager.add_thread(self.lexicon_loader_thread)
                self.lexicon_loader_thread.start()

            except Exception as e:
                self.button_manager.restore_states()
                self.custom_lexicon_button.setText("Load Lexicon")
                QMessageBox.critical(self, "Error", f"Failed to start lexicon loading: {str(e)}")

    def on_lexicon_loaded(self, result, success):
        self.unified_progress_bar.setVisible(False)
        self.button_manager.restore_states()
        self.custom_lexicon_button.setText("Load Lexicon")
        
        try:
            if success:
                if self.sentiment_mode == "TextBlob (Custom Lexicon)":
                    self.custom_textblob_lexicon_path = result
                    self.textblob_analyzer = CustomTextBlobSentimentAnalyzer(self.custom_textblob_lexicon_path)
                    QMessageBox.information(self, "Success", "Custom TextBlob lexicon loaded!")
                else:
                    self.custom_lexicon_path = result
                    self.vader_analyzer = CustomVaderSentimentIntensityAnalyzer(custom_lexicon_file=self.custom_lexicon_path)
                    QMessageBox.information(self, "Success", "Custom VADER lexicon loaded!")
                
                has_text = bool(self.text_data.strip())
                self.sentiment_button.setEnabled(has_text)
                
            else:
                raise ValueError(str(result))
                
        except Exception as e:
            if self.sentiment_mode == "TextBlob (Custom Lexicon)":
                self.custom_textblob_lexicon_path = None
                self.textblob_analyzer = None
            else:
                self.custom_lexicon_path = None
                self.vader_analyzer = None
                
            QMessageBox.critical(self, "Error", f"Failed to load lexicon: {str(e)}")
            self.sentiment_button.setEnabled(False)

    def load_custom_model(self):
        options = QFileDialog.Options()
        model_path, _ = QFileDialog.getOpenFileName(
            self, "Select Custom Model File", "", "Model Files (*.pt)", options=options
        )
        if model_path:
            try:
                self.button_manager.disable_other_buttons('custom_model_button')
                self.custom_model_button.setText("Loading Model...")
                self.custom_model_button.setEnabled(False)
                self.sentiment_button.setEnabled(False)

                self.model_loader_thread = CustomFileLoaderThread(model_path, "model")
                self.model_loader_thread.file_loaded.connect(self.on_model_loaded)
                self.thread_manager.add_thread(self.model_loader_thread)
                self.model_loader_thread.start()

            except Exception as e:
                self.button_manager.restore_states()
                self.custom_model_button.setText("Load Model")
                QMessageBox.critical(self, "Error", f"Failed to start model loading: {str(e)}")

    def on_model_loaded(self, result, success):
        """Handle when custom Flair model is loaded"""
        self.unified_progress_bar.setVisible(False)
        self.custom_model_button.setText("Load Model")
        
        try:
            if success:
                from flair.models import TextClassifier
                from flair.data import Sentence

                if not isinstance(result, TextClassifier):
                    raise ValueError("Invalid model type. Please load a valid Flair TextClassifier model.")

                test_sentence = Sentence("test")
                result.predict(test_sentence)
                
                if not test_sentence.labels:
                    raise ValueError("Model didn't produce any labels")
                    
                label = test_sentence.labels[0].value
                if label not in ['POSITIVE', 'NEGATIVE', 'NEUTRAL']:
                    raise ValueError(f"Model produces incompatible labels: {label}")

                self.flair_classifier_cuslang = result
                QMessageBox.information(self, "Success", "Custom model loaded successfully!")
                
                self.button_manager.restore_states()
                has_text = bool(self.text_data.strip())
                
                if has_text:
                    self.sentiment_button.setEnabled(True)
                    self.view_fulltext_button.setEnabled(True)
                    self.summarize_button.setEnabled(True)
                    self.text_stats_button.setEnabled(True)
                    if hasattr(self, 'topic_tab'):
                        if hasattr(self.topic_tab, 'analyze_topics_btn'):
                            self.topic_tab.analyze_topics_btn.setEnabled(True)
                        if hasattr(self.topic_tab, 'extract_keywords_btn'):
                            self.topic_tab.extract_keywords_btn.setEnabled(True)
                
                self.custom_model_button.setEnabled(True)
                
            else:
                raise ValueError(str(result))

        except Exception as e:
            self.flair_classifier_cuslang = None
            QMessageBox.critical(self, "Error", f"Failed to load custom model: {str(e)}")
            self.sentiment_button.setEnabled(False)
            self.button_manager.restore_states()
            self.custom_model_button.setEnabled(True)

    def load_flair_model(self):
        
        try:
            self.button_manager.disable_other_buttons('sentiment_button')
            self.sentiment_button.setText("Loading Flair...")
            self.sentiment_button.setEnabled(False)
            
            self.flair_loader_thread = FlairModelLoaderThread()
            self.flair_loader_thread.model_loaded.connect(self.on_flair_model_loaded)
            self.flair_loader_thread.error_occurred.connect(self.on_flair_model_error)
            self.flair_loader_thread.finished.connect(self.cleanup_flair_thread)

            self.thread_manager.add_thread(self.flair_loader_thread)
            self.add_managed_thread(self.flair_loader_thread)

            self.flair_loader_thread.start()
            
        except Exception as e:
            self.button_manager.restore_states()
            self.sentiment_button.setText("Analyze Sentiment")
            QMessageBox.critical(self, "Error", f"Failed to load Flair: {str(e)}")

    def on_flair_model_loaded(self, model):
        """Handle when default Flair model is loaded"""
        self.unified_progress_bar.setVisible(False)
        self.button_manager.restore_states()
        self.sentiment_button.setText("Analyze Sentiment")
        
        if model:
            self.flair_classifier = model
            has_text = bool(self.text_data.strip())
            
            if self.flair_first_load:
                self.flair_first_load = False
                QMessageBox.information(self, "Ready", "Flair library loaded successfully!")

            if self.sentiment_mode == "Flair":
                self.sentiment_button.setEnabled(has_text)
                
            elif self.sentiment_mode == "Flair (Custom Model)":
                self.custom_model_button.setEnabled(True)
                if self.flair_classifier_cuslang:
                    self.sentiment_button.setEnabled(has_text)
                else:
                    QMessageBox.information(self, "Next Step", "Please load your custom model")
                    self.sentiment_button.setEnabled(False)
            
        else:
            self.sentiment_button.setEnabled(False)
            QMessageBox.critical(self, "Error", "Failed to load Flair model")

    def on_flair_model_error(self, error):
        self.unified_progress_bar.setVisible(False)
        QMessageBox.critical(self, "Loading Error", f"Failed to load Flair model: {error}")

    def cleanup_flair_thread(self):
        self.progress_timer.stop()
        self.unified_progress_bar.setValue(0)
        self.unified_progress_bar.setVisible(False)
        self.thread_manager.remove_thread(self.flair_loader_thread)
        self.flair_loader_thread = None        

    def show_sentiment_mode_info(self):
        """Show sentiment mode info dialog using DialogFactory"""
        mode_info = base64.b64decode(AppConstants.MODE_INFO.encode()).decode()
        
        dialog = DialogFactory.create_info_dialog(
            self, "Sentiment Analysis Modes", mode_info, modal=False
        )
        dialog.show()

    def update_topic_analysis(self, text):
        """Update topic analysis tab with new text"""
        self.topic_tab.set_text(text)

    def _update_progress_animation(self):
        """Smooth indeterminate progress animation"""
        if hasattr(self, 'unified_progress_bar') and self.unified_progress_bar.isVisible():
            current = self.unified_progress_bar.value()
            next_value = (current + 2) % 100
            self.unified_progress_bar.setValue(next_value)

    def resizeEvent(self, event):
        """Handle resize to keep progress bar full width"""
        super().resizeEvent(event)
        if hasattr(self, 'unified_progress_bar'):
            container = self.unified_progress_bar.parent()
            self.unified_progress_bar.setFixedWidth(container.width())
            self.unified_progress_bar.setGeometry(0, 0, container.width(), container.height())

    def update_stopwords_for_topic(self, custom_stopwords=None):
        """Initialize and update stopwords with consistent preprocessing"""
        from nltk.corpus import stopwords
        
        try:
            default_stopwords = set(stopwords.words('english'))
        except:
            default_stopwords = STOPWORDS
            
        default_stopwords = {word.lower().strip() for word in default_stopwords}
        
        if custom_stopwords:
            custom_stopwords = {word.lower().strip() for word in custom_stopwords if word.strip()}
            self.stop_words = default_stopwords.union(custom_stopwords)
        else:
            self.stop_words = default_stopwords

        vectorizer_config = {
            'token_pattern': r'(?u)\b[a-zA-Z][a-zA-Z0-9\'-]*[a-zA-Z0-9]\b',
            'stop_words': list(self.stop_words),
            'lowercase': True,
            'strip_accents': 'unicode',
            'max_features': 1000,
            'min_df': 1,
            'max_df': 1
        }
        
        self.vectorizer = CountVectorizer(**vectorizer_config)
        self.tfidf = TfidfVectorizer(**vectorizer_config)

    def _extract_tfidf(self, text, num_keywords):
        """Extract keywords using TF-IDF with consistent preprocessing"""
        if not hasattr(self, 'stop_words'):
            self.update_stopwords_for_topic()
            
        vectorizer = TfidfVectorizer(
            token_pattern=r'(?u)\b[a-zA-Z][a-zA-Z0-9\'-]*[a-zA-Z0-9]\b',
            stop_words=list(self.stop_words),
            lowercase=True,
            strip_accents='unicode',
            min_df=1,
            max_df=1
        )
        
        try:
            sentences = [s.strip() for s in text.split('.') if s.strip()]
            if not sentences:
                sentences = [text]
                
            response = vectorizer.fit_transform(sentences)
            feature_names = vectorizer.get_feature_names_out()
            
            scores = response.mean(axis=0).A[0]
            
            pairs = list(zip(scores, feature_names))
            pairs.sort(reverse=True)
            
            return [{'keyword': word, 'score': score} 
                    for score, word in pairs[:num_keywords]]
                    
        except ValueError as e:
            logger.error(f"TF-IDF extraction error: {str(e)}")
            return [{'keyword': 'Error: No valid keywords found', 'score': 0.0}]

    def _process_topic_modeling(self, text, num_topics, model_type='lda'):
        """Generic topic modeling processor"""
        try:
            sentences = [sent.strip() for sent in text.split('.') if sent.strip()] or [text]
            
            vectorizer = self.vectorizer if model_type == 'lda' else self.tfidf
            dtm = vectorizer.fit_transform(sentences)
            
            if model_type == 'nmf':
                max_possible_topics = min(dtm.shape[0], dtm.shape[1])
                if num_topics > max_possible_topics:
                    raise ValueError(f"Number of topics ({num_topics}) cannot exceed {max_possible_topics} for NMF with this dataset")
            
            num_topics = min(num_topics, max(2, dtm.shape[1] - 1))
            
            model_class = LatentDirichletAllocation if model_type == 'lda' else NMF
            model_opts = {
                'n_components': num_topics,
                'random_state': 42,
                'max_iter': 20 if model_type == 'lda' else 500
            }
            if model_type == 'nmf':
                model_opts['init'] = 'nndsvd'
                model_opts['tol'] = 1e-4

            if model_type == 'lda':
                model_opts['learning_method'] = 'online'
                
            model = model_class(**model_opts)
            model.fit(dtm)
            
            terms = vectorizer.get_feature_names_out()
            return [{
                'topic': f'Topic {idx + 1}',
                'terms': [terms[i] for i in topic.argsort()[:-10-1:-1]],
                'weight': topic.sum()
            } for idx, topic in enumerate(model.components_)]
            
        except Exception as e:
            logger.error(f"{model_type.upper()} Analysis error: {str(e)}")
            raise ValueError(f"Topic analysis failed: {str(e)}")

    def analyze_topics_lda(self, text, num_topics=5):
        """LDA topic analysis wrapper"""
        return self._process_topic_modeling(text, num_topics, 'lda')

    def analyze_topics_nmf(self, text, num_topics=5):
        """NMF topic analysis wrapper"""
        return self._process_topic_modeling(text, num_topics, 'nmf')

    def extract_keywords(self, text, method='tfidf', num_keywords=10):
        """Keyword extraction using various methods"""
        if method == 'tfidf':
            return self._extract_tfidf(text, num_keywords)
        elif method == 'yake':
            return self._extract_yake(text, num_keywords)
        elif method == 'rake':
            return self._extract_rake(text, num_keywords)
        else:
            raise ValueError(f"Unknown keyword extraction method: {method}")

    def _extract_yake(self, text, num_keywords):
        """YAKE extraction"""
        kw_extractor = yake.KeywordExtractor(
            lan="en",
            n=2,
            windowsSize=2,
            top=num_keywords,
            stopwords=list(self.stop_words) if hasattr(self, 'stop_words') else None
        )
        try:
            keywords = kw_extractor.extract_keywords(text)
            return [{'keyword': kw[0], 'score': 1-kw[1]} for kw in keywords]
        except Exception:
            return [{'keyword': 'Error: YAKE extraction failed', 'score': 0.0}]

    def _extract_rake(self, text, num_keywords):
        """RAKE extraction"""
        try:
            rake = rake_nltk.Rake(stopwords=list(self.stop_words) if hasattr(self, 'stop_words') else None)
            rake.extract_keywords_from_text(text)
            keywords = rake.get_ranked_phrases_with_scores()
            keywords.sort(reverse=True)
            return [{'keyword': kw[1], 'score': kw[0]} for kw in keywords[:num_keywords]]
        except Exception:
            return [{'keyword': 'Error: RAKE extraction failed', 'score': 0.0}]
    
    def load_stopwords_file(self):
        """Load stopwords from a text file"""
        try:
            options = QFileDialog.Options()
            file_path, _ = QFileDialog.getOpenFileName(
                self, 
                "Load Stopwords File", 
                "", 
                "Text Files (*.txt);;All Files (*)", 
                options=options
            )
            
            if file_path:
                with open(file_path, 'r', encoding='utf-8') as file:
                    stopwords = file.read().strip()
                    current_text = self.stopword_entry.toPlainText().strip()
                    if (current_text):
                        stopwords = current_text + "\n" + stopwords
                    self.stopword_entry.setPlainText(stopwords)
                    QMessageBox.information(self, "Success", "Stopwords loaded successfully!")
                    
        except Exception as e:
            QMessageBox.critical(self, "Error", f"Failed to load stopwords file: {str(e)}")

    def show_token_warning(self, tokens):
        """Show non-blocking warning message about detected tokens"""
        msg = QMessageBox(self)
        msg.setIcon(QMessageBox.Information)
        msg.setWindowTitle("Analysis Notification")
        msg.setWindowModality(Qt.NonModal)
        msg.setTextFormat(Qt.TextFormat.RichText)
        
        message = f"""
        <b>Inconsistent tokens detected:</b> "{tokens}"<br><br>

        These tokens were detected by the tokenizer but are not found in your stopwords list.  
        This is not necessarily an issue, but it may indicate:<br>
        • Some unimportant words slipping through.<br>
        • Potential noise in your dataset.<br><br>

        <b>No immediate action is required.</b> However, if you'd like to refine your stopwords list, you may consider the following:<br>
        1. Review these tokens.<br>
        2. Add them to the stopwords list if they are irrelevant.<br>
        3. Regenerate the analysis for improved results.<br><br>

        <i>You can add them directly via the stopwords entry field.</i>
        """

        msg.setText(message)
        msg.setStandardButtons(QMessageBox.Ok)
        
        msg.show()

    def cleanup_cache(self):
        """Clear cached resources when memory is low"""
        try:
            self.get_wordcloud.cache_clear()
            
            self._cached_models.clear()
            self._cached_fonts.clear()
            self._cached_colormaps.clear()
                
        except Exception as e:
            logger.error(f"Cache cleanup error: {e}")

    def summarize_text(self):
        if not self.text_data:
            QMessageBox.critical(self, "Error", "No text data available for summarization.\nPlease load a text file first.")
            return

        self.disable_buttons()
        self.summarize_thread = SummarizeThread(self.text_data)
        self.summarize_thread.summary_ready.connect(self.show_summary)
        self.summarize_thread.error_occurred.connect(self.handle_summarize_error)
        self.summarize_thread.finished.connect(self.enable_buttons)
        self.summarize_thread.start()

    def show_summary(self, summary):
        dialog = QDialog(self)
        dialog.setWindowTitle("Text Summary")
        dialog.setMinimumSize(500, 300)
        dialog.setSizeGripEnabled(True)
        layout = QVBoxLayout()

        text_browser = QTextBrowser()
        text_browser.setPlainText(summary)
        layout.addWidget(text_browser)

        close_button = QPushButton("Close")
        close_button.clicked.connect(dialog.accept)
        layout.addWidget(close_button)

        dialog.setLayout(layout)
        dialog.show()

    def handle_summarize_error(self, error_msg):
        QMessageBox.critical(self, "Error", f"Failed to summarize text: {error_msg}")
        self.enable_buttons()

    def get_score_type(self):
        return AppConstants.SCORE_TYPES.get(self.mode, "Score")

    def cleanup_threads(self):
        """Clean up all running threads before application exit"""
        logger.info("Cleaning up threads before exit")
        
        # Stop all threads with proper timeout
        self.thread_manager.stop_all_threads(wait_timeout=AppConstants.THREAD_TIMEOUT["SOFT"])
        
        # Log any remaining threads that couldn't be stopped
        if self.thread_manager.active_threads:
            thread_names = [t.objectName() or "unnamed" for t in self.thread_manager.active_threads]
            logger.warning(f"Some threads could not be stopped: {thread_names}")

    def create_worker_thread(self, worker, on_finished=None, on_error=None, thread_name=None):
        """
        Create and configure a worker thread with proper error handling and cleanup
        
        Args:
            worker (QRunnable): The worker to run in the thread
            on_finished (callable, optional): Callback when thread finishes successfully
            on_error (callable, optional): Callback when thread encounters an error
            thread_name (str, optional): Name for the thread for debugging purposes
            
        Returns:
            QThread: The configured thread object
        """
        thread = QThread()
        if thread_name:
            thread.setObjectName(thread_name)
        
        # Move worker to thread
        worker.moveToThread(thread)
        
        # Connect signals
        thread.started.connect(worker.run)
        worker.finished.connect(thread.quit)
        worker.finished.connect(worker.deleteLater)
        thread.finished.connect(thread.deleteLater)
        
        # Connect custom callbacks if provided
        if on_finished:
            worker.finished.connect(on_finished)
        if on_error:
            worker.error.connect(on_error)
        
        # Setup cleanup when thread finishes
        thread.finished.connect(lambda: self.thread_manager.remove_thread(thread))
        
        # Add to managed threads
        self.thread_manager.add_thread(thread)
        
        return thread

    def register_cached_functions(self):
        """Mendaftarkan semua fungsi yang menggunakan lru_cache ke cache manager"""
        # Daftarkan metode dengan lru_cache
        for name in dir(self):
            if name.startswith('__'):  # Lewati metode internal
                continue
                
            try:
                attr = getattr(self.__class__, name)
                # Hanya periksa apakah ini adalah cached_property
                if isinstance(attr, cached_property):
                    self.cache_manager.register_cached_property(self, name)
            except (AttributeError, TypeError):
                pass
            
            try:
                # Periksa metode instance yang memiliki cache_clear
                attr = getattr(self, name, None)
                if callable(attr) and hasattr(attr, 'cache_clear'):
                    self.cache_manager.register_cached_function(attr)
            except (AttributeError, TypeError):
                pass
        
        # Daftarkan fungsi global dengan lru_cache
        for name, func in globals().items():
            if callable(func) and hasattr(func, 'cache_clear'):
                self.cache_manager.register_cached_function(func)

    def clear_caches(self):
        """Membersihkan semua cache secara manual"""
        self.cache_manager.clear_all_caches()
        self.statusBar().showMessage("All caches cleared", 3000)

    def create_utility_buttons(self):
        """Create utility buttons"""
        utility_group = QGroupBox("Utilities", self)
        layout = QVBoxLayout()
        
        # Tombol untuk membersihkan cache
        clear_cache_btn = QPushButton("Clear Cache", self)
        clear_cache_btn.setToolTip("Manually clear all cached data to free memory")
        clear_cache_btn.clicked.connect(self.clear_caches)
        layout.addWidget(clear_cache_btn)
        
        # ... tombol utilitas lainnya ...
        
        utility_group.setLayout(layout)
        return utility_group

    def detect_file_format(self, file_path, content):
        """
        Deteksi format file dan berikan penanganan khusus jika diperlukan
        
        Args:
            file_path (str): Path file
            content (str): Isi file
            
        Returns:
            tuple: (processed_content, file_type)
        """
        ext = Path(file_path).suffix.lower()
        
        # Deteksi format berdasarkan ekstensi
        if ext == '.csv':
            # Coba deteksi delimiter
            delimiters = [',', ';', '\t', '|']
            delimiter = ','  # default
            
            for d in delimiters:
                if d in content[:1000]:
                    delimiter = d
                    break
            
            # Tanyakan kepada pengguna apakah ingin memproses sebagai CSV
            reply = QMessageBox.question(
                self,
                "CSV File Detected",
                f"This appears to be a CSV file with delimiter '{delimiter}'.\n\n"
                "How would you like to process it?",
                "Extract Text Column", "Join All Columns", "Raw Text",
                defaultButtonNumber=0
            )
            
            if reply == 0:  # Extract Text Column
                # Tampilkan dialog untuk memilih kolom
                column_index, ok = QInputDialog.getInt(
                    self, "Select Column", "Enter the column number (0-based):", 0, 0, 100, 1
                )
                if ok:
                    try:
                        import csv
                        from io import StringIO
                        
                        rows = []
                        reader = csv.reader(StringIO(content), delimiter=delimiter)
                        for row in reader:
                            if len(row) > column_index:
                                rows.append(row[column_index])
                        
                        return ("\n".join(rows), "csv_column")
                    except Exception as e:
                        logger.error(f"Error processing CSV column: {e}")
                        # Fallback ke raw text
            
            elif reply == 1:  # Join All Columns
                try:
                    import csv
                    from io import StringIO
                    
                    rows = []
                    reader = csv.reader(StringIO(content), delimiter=delimiter)
                    for row in reader:
                        rows.append(" ".join(row))
                    
                    return ("\n".join(rows), "csv_joined")
                except Exception as e:
                    logger.error(f"Error joining CSV columns: {e}")
                    # Fallback ke raw text
        
        elif ext in ['.html', '.xml']:
            # Tanyakan kepada pengguna apakah ingin mengekstrak teks dari HTML/XML
            reply = QMessageBox.question(
                self,
                f"{ext.upper()[1:]} File Detected",
                f"This appears to be a {ext.upper()[1:]} file.\n\n"
                "Would you like to extract plain text from it?",
                QMessageBox.Yes | QMessageBox.No,
                QMessageBox.Yes
            )
            
            if reply == QMessageBox.Yes:
                try:
                    from bs4 import BeautifulSoup
                    
                    soup = BeautifulSoup(content, 'html.parser')
                    # Hapus script dan style
                    for script in soup(["script", "style"]):
                        script.extract()
                    
                    # Ambil teks
                    text = soup.get_text(separator="\n")
                    
                    # Bersihkan teks
                    lines = [line.strip() for line in text.splitlines() if line.strip()]
                    text = "\n".join(lines)
                    
                    return (text, f"{ext[1:]}_extracted")
                except ImportError:
                    QMessageBox.warning(
                        self,
                        "Module Missing",
                        "BeautifulSoup module is required to extract text from HTML/XML.\n"
                        "Using raw text instead."
                    )
                except Exception as e:
                    logger.error(f"Error extracting text from {ext}: {e}")
                    # Fallback ke raw text
        
        # Default: gunakan teks mentah
        return (content, "raw_text")

    def open_file_dialog(self):
        """Open file dialog with path validation"""
        options = QFileDialog.Options()
        file_filter = (
            "Text Files (*.txt);;PDF Files (*.pdf);;Word Documents (*.doc *.docx);;"
            "CSV Files (*.csv);;Excel Files (*.xlsx *.xls);;All Files (*)"
        )
        
        file_path, _ = QFileDialog.getOpenFileName(
            self, "Open File", "", file_filter, options=options
        )
        
        if file_path:
            try:
                validated_path = sanitize_path(file_path)
                
                # Periksa ekstensi file
                ext = Path(validated_path).suffix.lower()
                supported_exts = ['.txt', '.pdf', '.doc', '.docx', '.csv', '.xlsx', '.xls']
                
                if ext not in supported_exts:
                    QMessageBox.warning(
                        self,
                        "Unsupported File Type",
                        f"The file type '{ext}' is not supported. Please select a supported file type."
                    )
                    return
                
                # Proses file berdasarkan ekstensi
                if ext == '.txt':
                    content = FileOperations.load_text_file(self, validated_path)
                elif ext in ['.csv']:
                    content = self.load_csv_file(validated_path)
                elif ext in ['.xlsx', '.xls']:
                    content = self.load_excel_file(validated_path)
                elif ext in ['.pdf']:
                    content = self.load_pdf_file(validated_path)
                elif ext in ['.doc', '.docx']:
                    content = self.load_word_file(validated_path)
                else:
                    content = None
                    
                if content is not None:
                    self.text_data = content
                    self.file_path = validated_path
                    self.text_edit.setPlainText(content)
                    
                    # Update status bar
                    self.statusBar().showMessage(
                        f"Loaded: {os.path.basename(validated_path)}"
                    )
            except ValueError as e:
                QMessageBox.critical(self, "Security Error", str(e))
                logger.warning(f"Blocked access to file: {file_path}")

    def load_csv_file(self, file_path):
        """Load and process CSV file"""
        try:
            import pandas as pd
            
            # Coba deteksi delimiter
            with open(file_path, 'r', encoding='utf-8') as f:
                sample = f.read(1024)
                
            delimiters = [',', ';', '\t', '|']
            delimiter = ','  # default
            
            for d in delimiters:
                if d in sample:
                    delimiter = d
                    break
            
            # Tanyakan kepada pengguna bagaimana memproses CSV
            reply = QMessageBox.question(
                self,
                "CSV File Options",
                f"How would you like to process this CSV file?\nDetected delimiter: '{delimiter}'",
                "Extract Column", "Join All Columns", "Cancel",
                defaultButtonNumber=0
            )
            
            if reply == 0:  # Extract Column
                # Baca header untuk menampilkan opsi kolom
                df = pd.read_csv(file_path, delimiter=delimiter, nrows=0)
                columns = df.columns.tolist()
                
                column, ok = QInputDialog.getItem(
                    self, "Select Column", "Choose text column to extract:", 
                    columns, 0, False
                )
                
                if ok and column:
                    df = pd.read_csv(file_path, delimiter=delimiter, usecols=[column])
                    return "\n".join(df[column].astype(str).tolist())
                return None
                
            elif reply == 1:  # Join All Columns
                df = pd.read_csv(file_path, delimiter=delimiter)
                # Gabungkan semua kolom dengan spasi
                result = []
                for _, row in df.iterrows():
                    result.append(" ".join(row.astype(str).tolist()))
                return "\n".join(result)
                
            else:  # Cancel
                return None
                
        except Exception as e:
            QMessageBox.critical(self, "Error", f"Failed to load CSV file: {str(e)}")
            logger.error(f"Error loading CSV file {file_path}: {e}")
            return None

    def load_excel_file(self, file_path):
        """Load and process Excel file"""
        try:
            import pandas as pd
            
            # Baca workbook
            xl = pd.ExcelFile(file_path)
            
            # Jika ada lebih dari satu sheet, tanyakan kepada pengguna
            if len(xl.sheet_names) > 1:
                sheet, ok = QInputDialog.getItem(
                    self, "Select Sheet", "Choose Excel sheet to load:", 
                    xl.sheet_names, 0, False
                )
                
                if not ok:
                    return None
            else:
                sheet = xl.sheet_names[0]
            
            # Baca sheet yang dipilih
            df = pd.read_excel(file_path, sheet_name=sheet)
            
            # Tanyakan kepada pengguna bagaimana memproses Excel
            reply = QMessageBox.question(
                self,
                "Excel File Options",
                "How would you like to process this Excel file?",
                "Extract Column", "Join All Columns", "Cancel",
                defaultButtonNumber=0
            )
            
            if reply == 0:  # Extract Column
                columns = df.columns.tolist()
                column, ok = QInputDialog.getItem(
                    self, "Select Column", "Choose text column to extract:", 
                    columns, 0, False
                )
                
                if ok and column:
                    return "\n".join(df[column].astype(str).tolist())
                return None
                
            elif reply == 1:  # Join All Columns
                # Gabungkan semua kolom dengan spasi
                result = []
                for _, row in df.iterrows():
                    result.append(" ".join(row.astype(str).tolist()))
                return "\n".join(result)
                
            else:  # Cancel
                return None
                
        except Exception as e:
            QMessageBox.critical(self, "Error", f"Failed to load Excel file: {str(e)}")
            logger.error(f"Error loading Excel file {file_path}: {e}")
            return None

    def load_pdf_file(self, file_path):
        """Load and process PDF file"""
        try:
            # Cek apakah PyPDF2 tersedia
            try:
                import PyPDF2
            except ImportError:
                QMessageBox.critical(
                    self, 
                    "Missing Dependency", 
                    "PyPDF2 is required to read PDF files. Please install it with 'pip install PyPDF2'."
                )
                return None
            
            with open(file_path, 'rb') as file:
                reader = PyPDF2.PdfReader(file)
                
                # Jika PDF memiliki banyak halaman, tanyakan kepada pengguna
                if len(reader.pages) > 1:
                    options = ["All Pages", "Select Page Range", "Cancel"]
                    reply = QMessageBox.question(
                        self,
                        "PDF Options",
                        f"This PDF has {len(reader.pages)} pages. How would you like to proceed?",
                        *options,
                        defaultButtonNumber=0
                    )
                    
                    if reply == 0:  # All Pages
                        text = ""
                        for page in reader.pages:
                            text += page.extract_text() + "\n\n"
                        return text
                        
                    elif reply == 1:  # Select Page Range
                        start, ok1 = QInputDialog.getInt(
                            self, "Start Page", "Enter start page (1-based):", 
                            1, 1, len(reader.pages), 1
                        )
                        
                        if not ok1:
                            return None
                            
                        end, ok2 = QInputDialog.getInt(
                            self, "End Page", "Enter end page:", 
                            min(start + 9, len(reader.pages)), start, len(reader.pages), 1
                        )
                        
                        if not ok2:
                            return None
                        
                        text = ""
                        for i in range(start - 1, end):
                            text += reader.pages[i].extract_text() + "\n\n"
                        return text
                        
                    else:  # Cancel
                        return None
                else:
                    # Hanya satu halaman
                    return reader.pages[0].extract_text()
                    
        except Exception as e:
            QMessageBox.critical(self, "Error", f"Failed to load PDF file: {str(e)}")
            logger.error(f"Error loading PDF file {file_path}: {e}")
            return None

    def load_word_file(self, file_path):
        """Load and process Word document"""
        try:
            # Cek apakah python-docx tersedia
            try:
                import docx
            except ImportError:
                QMessageBox.critical(
                    self, 
                    "Missing Dependency", 
                    "python-docx is required to read Word documents. Please install it with 'pip install python-docx'."
                )
                return None
            
            doc = docx.Document(file_path)
            
            # Ekstrak teks dari paragraf
            text = "\n".join([paragraph.text for paragraph in doc.paragraphs if paragraph.text])
            
            # Jika dokumen memiliki tabel, tanyakan kepada pengguna
            if doc.tables:
                reply = QMessageBox.question(
                    self,
                    "Word Document Tables",
                    f"This document contains {len(doc.tables)} tables. Include table content?",
                    QMessageBox.Yes | QMessageBox.No,
                    QMessageBox.Yes
                )
                
                if reply == QMessageBox.Yes:
                    # Ekstrak teks dari tabel
                    for table in doc.tables:
                        for row in table.rows:
                            row_text = " | ".join([cell.text for cell in row.cells if cell.text])
                            if row_text:
                                text += "\n" + row_text
            
            return text
                    
        except Exception as e:
            QMessageBox.critical(self, "Error", f"Failed to load Word document: {str(e)}")
            logger.error(f"Error loading Word document {file_path}: {e}")
            return None

class LazyLoader:
    """
    Optimized lazy loading system for managing heavy dependencies.
    
    This class implements a singleton pattern with caching to ensure
    efficient loading and reuse of module imports. It helps reduce
    memory usage and startup time by loading modules only when needed.
    
    Attributes:
        _instances (dict): Singleton instances storage
        _cache (dict): Module cache storage
        
    Methods:
        load(module_name, class_name=None): 
            Lazily loads a module or class and caches it
            
    Example:
        >>> loader = LazyLoader()
        >>> TextBlob = loader.load('textblob', 'TextBlob')
        >>> vader = loader.load('vaderSentiment.vaderSentiment')
    """
    _instances = {}
    _cache = {}
    
    @classmethod
    def load(cls, module_name: str, class_name: str = None) -> object:
        """
        Lazily load and cache a module or specific class.
        
        Args:
            module_name (str): Name of the module to import
            class_name (str, optional): Specific class to import from module
            
        Returns:
            object: Loaded module or class
            
        Raises:
            ImportError: If module/class cannot be loaded
        """
        key = f"{module_name}.{class_name}" if class_name else module_name
        if key not in cls._cache:
            try:
                module = __import__(module_name, fromlist=[class_name] if class_name else [])
                cls._cache[key] = getattr(module, class_name) if class_name else module
            except ImportError as e:
                logger.error(f"Failed to load {key}: {e}")
                return None
        return cls._cache[key]

class ThreadSafeSet:
    """Thread-safe set with context manager"""
    def __init__(self):
        self._items = set()
        self._mutex = QMutex()
    
    def locked(self):
        self._mutex.lock()
        return self._items

    def unlock(self):
        self._mutex.unlock()
            
    def add(self, item):
        items = self.locked()
        try:
            items.add(item)
        finally:
            self.unlock()
            
    def remove(self, item):
        items = self.locked()
        try:
            items.discard(item)
        finally:
            self.unlock()
            
    def clear(self):
        items = self.locked()
        try:
            items.clear()
        finally:
            self.unlock()

class ResourceManager:
    """Resource management with caching"""
    def __init__(self):
        self._resources = {}
        self._mutex = QMutex()
        
    def get(self, resource_id, creator_func):
        self._mutex.lock()
        try:
            if resource_id not in self._resources:
                self._resources[resource_id] = creator_func()
            return self._resources[resource_id]
        finally:
            self._mutex.unlock()
            
    def cleanup(self):
        self._mutex.lock()
        try:
            for resource in self._resources.values():
                if hasattr(resource, 'cleanup'):
                    resource.cleanup()
            self._resources.clear()
        finally:
            self._mutex.unlock()

class OptimizedWordCloud(WordCloud):
    """Optimized WordCloud with caching"""
    @lru_cache(maxsize=32)
    def _get_font_properties(self, font_path):
        return super()._get_font_properties(font_path)
        
    @lru_cache(maxsize=128)
    def _process_text(self, text):
        return super()._process_text(text)
        
    @lru_cache(maxsize=16)
    def _get_colormap(self, colormap_name):
        if isinstance(colormap_name, list):
            from matplotlib.colors import LinearSegmentedColormap
            return LinearSegmentedColormap.from_list('custom', colormap_name)
        return colormap_name

class ThreadManager:
    def __init__(self):
        self.active_threads = []
        self.mutex = QMutex()
        
    def add_thread(self, thread):
        """Menambahkan thread ke daftar thread aktif"""
        self.mutex.lock()
        try:
            if not isinstance(thread, QThread):
                raise TypeError("Thread harus merupakan instance dari QThread")
                
            thread.finished.connect(lambda: self.remove_thread(thread))
            self.active_threads.append(thread)
        finally:
            self.mutex.unlock()
            
    def remove_thread(self, thread):
        """Menghapus thread dari daftar thread aktif"""
        self.mutex.lock()
        try:
            if thread in self.active_threads:
                self.active_threads.remove(thread)
        finally:
            self.mutex.unlock()
            
    def cleanup_finished(self):
        """Membersihkan thread yang sudah selesai dari daftar"""
        self.mutex.lock()
        try:
            # Salin list untuk menghindari modifikasi saat iterasi
            for thread in self.active_threads[:]:
                if not thread.isRunning() or thread.isFinished():
                    self.remove_thread(thread)
                    try:
                        thread.deleteLater()
                    except:
                        pass
        finally:
            self.mutex.unlock()
            
    def stop_all_threads(self, wait_timeout=1000):
        """Menghentikan semua thread yang aktif"""
        self.mutex.lock()
        try:
            for thread in self.active_threads[:]:  # Copy list untuk iterasi yang aman
                if thread.isRunning():
                    thread.requestInterruption()
                    thread.quit()
                    if not thread.wait(wait_timeout):
                        logger.warning(f"Thread {thread.objectName() or 'unnamed'} tidak merespons, memaksa berhenti")
                        thread.terminate()
                self.remove_thread(thread)
                try:
                    thread.deleteLater()
                except:
                    pass
        finally:
            self.mutex.unlock()

class BaseAnalyzer:
    def __init__(self):
        self._initialized = False
        
    @property
    def is_initialized(self):
        return self._initialized
        
    def cleanup(self):
        pass

class SentimentAnalyzer(BaseAnalyzer):
    def __init__(self, mode, text_data, analyzers):
        super().__init__()
        self.mode = mode
        self.text_data = text_data
        self.analyzers = analyzers
        
    def get_score_type(self):
        return AppConstants.SCORE_TYPES.get(self.mode, "Score")

class CustomVaderSentimentIntensityAnalyzer:
    def __init__(self, lexicon_file="vader_lexicon.txt", custom_lexicon_file=None):
        from vaderSentiment.vaderSentiment import SentimentIntensityAnalyzer
        self.analyzer = SentimentIntensityAnalyzer()
        if custom_lexicon_file:
            self.load_custom_lexicon(custom_lexicon_file)

    def load_custom_lexicon(self, custom_lexicon_file):
        try:
            with open(custom_lexicon_file, "r", encoding="utf-8") as file:
                for line in file:
                    parts = line.strip().split("\t")
                    if len(parts) == 2:
                        word, measure = parts
                        try:
                            self.analyzer.lexicon[word] = float(measure)
                        except ValueError:
                            pass
        except Exception as e:
            QMessageBox.warning(None, "Warning", f"Failed to load custom lexicon: {e}")

    def polarity_scores(self, text):
        return self.analyzer.polarity_scores(text)

class CustomTextBlobSentimentAnalyzer:
    def __init__(self, custom_lexicon_file=None):
        self.lexicon = {}
        self.intensifiers = {}
        self.negations = set()
        if custom_lexicon_file:
            self.load_custom_lexicon(custom_lexicon_file)

    def load_custom_lexicon(self, custom_lexicon_file):
        try:
            with open(custom_lexicon_file, "r", encoding="utf-8") as file:
                for line in file:
                    parts = line.strip().split("\t")
                    if len(parts) == 2:
                        word, measure = parts
                        if measure == "negation":
                            self.negations.add(word)
                        elif measure.startswith("intensifier:"):
                            try:
                                self.intensifiers[word] = float(measure.split(":")[1])
                            except ValueError:
                                continue
                        else:
                            try:
                                self.lexicon[word] = float(measure)
                            except ValueError:
                                continue
        except Exception as e:
            QMessageBox.warning(None, "Warning", f"Failed to load custom lexicon: {e}")

    def analyze(self, text):
        from textblob import TextBlob
        blob = TextBlob(text)
        total_polarity = 0.0
        words_with_context = []

        for sentence in blob.sentences:
            words = sentence.words
            for i, word in enumerate(words):
                word_lower = word.lower()
                current_score = self.lexicon.get(word_lower, 0.0)
                if i > 0 and words[i-1].lower() in self.negations:
                    current_score *= -1
                if i > 0 and words[i-1].lower() in self.intensifiers:
                    multiplier = self.intensifiers[words[i-1].lower()]
                    current_score *= multiplier
                words_with_context.append(current_score)

        valid_scores = [s for s in words_with_context if s != 0]
        avg_polarity = sum(valid_scores) / len(valid_scores) if valid_scores else 0.0

        return {
            'polarity': avg_polarity,
            'subjectivity': blob.sentiment.subjectivity,
            'processed_words': len(valid_scores)
        }

class FlairModelLoaderThread(QThread):
    model_loaded = Signal(object)
    error_occurred = Signal(str)
    finished = Signal()
    _cached_model = None

    def __init__(self, model_path="sentiment"):
        super().__init__()
        self.model_path = model_path
        self._is_loading = False
        self._interrupt_requested = False

    def run(self):
        try:
            from flair.models import TextClassifier

            if self.isInterruptionRequested():
                return

            self._is_loading = True
            self._interrupt_requested = False

            if FlairModelLoaderThread._cached_model is None:
                try:
                    FlairModelLoaderThread._cached_model = TextClassifier.load(self.model_path)
                except Exception as e:
                    if not self.isInterruptionRequested():
                        self.error_occurred.emit(str(e))
                    return

            if self.isInterruptionRequested():
                return

            self.model_loaded.emit(FlairModelLoaderThread._cached_model)
        except Exception as e:
            if not self.isInterruptionRequested():
                self.error_occurred.emit(str(e))
        finally:
            self._is_loading = False
            self.finished.emit()

    def requestInterruption(self):
        logger.info("FlairModelLoader: Interruption requested")
        self._interrupt_requested = True
        self._is_loading = False
        super().requestInterruption()

class TopicAnalysisThread(QThread):
    """Thread for topic analysis to prevent UI freezing"""
    finished = Signal(list)
    error = Signal(str)
    
    def __init__(self, text, num_topics, method='lda'):
        super().__init__()
        self.text = text
        self.num_topics = num_topics
        self.method = method
        
    def run(self):
        try:
            if self.method == 'lda':
                result = self.parent().analyze_topics_lda(self.text, self.num_topics)
            else:
                result = self.parent().analyze_topics_nmf(self.text, self.num_topics)
            self.finished.emit(result)
        except Exception as e:
            self.error.emit(str(e))

class KeywordExtractionThread(QThread):
    """Thread for keyword extraction to prevent UI freezing"""
    finished = Signal(list)
    error = Signal(str)
    
    def __init__(self, text, method, num_keywords):
        super().__init__()
        self.text = text
        self.method = method
        self.num_keywords = num_keywords
        
    def run(self):
        try:
            result = self.parent().extract_keywords(self.text, self.method, self.num_keywords)
            self.finished.emit(result)
        except Exception as e:
            self.error.emit(str(e))

class TopicAnalysisTab(QWidget):
    def __init__(self, parent=None):
        super().__init__(parent)
        self.text = ""
        self.parent_widget = parent
        self.setFixedHeight(140)
        self.initUI()
        self.topic_thread = None
        self.keyword_thread = None
        
    def initUI(self):
        main_layout = QVBoxLayout()
        main_layout.setContentsMargins(1, 1, 1, 1)
        main_layout.setSpacing(5)

        group_layout = QHBoxLayout()
     
        topic_group = QGroupBox("Topic Modeling")
        topic_layout = QGridLayout()
        topic_layout.setContentsMargins(10, 5, 10, 5)
        topic_layout.setVerticalSpacing(2)
        
        topic_layout.addWidget(QLabel("Method:"), 0, 0)
        self.topic_method = QComboBox()
        self.topic_method.addItems(['LDA', 'NMF'])
        self.topic_method.setToolTip(
            "LDA (For discovering topics in large text corpora)\n"
            "NMF (For extracting hidden topics using non-negative features)"
        )
        topic_layout.addWidget(self.topic_method, 0, 1)
        
        topic_layout.addWidget(QLabel("Number of Topics:"), 1, 0)
        self.num_topics = QSpinBox()
        self.num_topics.setRange(1, 20)
        self.num_topics.setValue(5)
        topic_layout.addWidget(self.num_topics, 1, 1)
        
        self.analyze_topics_btn = QPushButton("Analyze Topics")
        topic_layout.addWidget(self.analyze_topics_btn, 2, 0, 1, 2)
        
        topic_group.setLayout(topic_layout)
        group_layout.addWidget(topic_group)

        keyword_group = QGroupBox("Keyword Extraction") 
        keyword_layout = QGridLayout()
        keyword_layout.setContentsMargins(10, 5, 10, 5)
        keyword_layout.setVerticalSpacing(2)
        
        keyword_layout.addWidget(QLabel("Method:"), 0, 0)
        self.keyword_method = QComboBox()
        self.keyword_method.addItems(['TF-IDF', 'YAKE', 'RAKE'])
        self.keyword_method.setToolTip(
            "TF-IDF (For identifying important words based on frequency)\n"
            "RAKE (For extracting keywords from short texts)\n"
            "YAKE (For extracting context-based keywords from documents)"            
        )
        keyword_layout.addWidget(self.keyword_method, 0, 1)
        
        keyword_layout.addWidget(QLabel("Number of Keywords:"), 1, 0)
        self.num_keywords = QSpinBox()
        self.num_keywords.setRange(5, 50)
        self.num_keywords.setValue(10)
        keyword_layout.addWidget(self.num_keywords, 1, 1)
        
        self.extract_keywords_btn = QPushButton("Extract Keywords")
        keyword_layout.addWidget(self.extract_keywords_btn, 2, 0, 1, 2)
        
        keyword_group.setLayout(keyword_layout)
        group_layout.addWidget(keyword_group)
        
        main_layout.addLayout(group_layout)
        self.setLayout(main_layout)

        self.analyze_topics_btn.setEnabled(False)
        self.extract_keywords_btn.setEnabled(False)
        
        self.analyze_topics_btn.clicked.connect(self.analyze_topics)
        self.extract_keywords_btn.clicked.connect(self.extract_keywords)

    def set_text(self, text):
        self.text = text
        has_text = bool(text.strip())
        self.analyze_topics_btn.setEnabled(has_text)
        self.extract_keywords_btn.setEnabled(has_text)
        
    def analyze_topics(self):
        if not self.text.strip():
            QMessageBox.warning(self, "Error", "Please load text first!")
            return
            
        try:
            if self.parent_widget:
                custom_stopwords = self.parent_widget.import_stopwords()
                self.parent_widget.update_stopwords_for_topic(custom_stopwords)

            method = self.topic_method.currentText()
            num_topics = self.num_topics.value()
            
            self.parent_widget.button_manager.disable_other_buttons('analyze_topics_btn')
            
            if self.parent_widget:
                self.parent_widget.set_progress('topics')
            
            self.topic_thread = TopicAnalysisThread(self.text, num_topics, method.lower())
            self.topic_thread.setParent(self.parent_widget)
            self.topic_thread.finished.connect(self.on_topic_analysis_complete)
            self.topic_thread.error.connect(self.on_analysis_error)
            self.topic_thread.start()
            
        except Exception as e:
            QMessageBox.critical(self, "Error", f"Topic analysis failed: {str(e)}")
            self.parent_widget.button_manager.restore_states()
            if self.parent_widget:
                self.parent_widget.set_progress('topics', False)
            
    def extract_keywords(self):
        if not hasattr(self, 'text'):
            QMessageBox.warning(self, "Error", "Please load text first!")
            return
            
        try:
            if self.parent_widget:
                custom_stopwords = self.parent_widget.import_stopwords()
                self.parent_widget.update_stopwords_for_topic(custom_stopwords)

            method = self.keyword_method.currentText().lower().replace('-', '')
            num_keywords = self.num_keywords.value()
            
            self.parent_widget.button_manager.disable_other_buttons('extract_keywords_btn')
            
            if self.parent_widget:
                self.parent_widget.set_progress('keywords')
            
            self.keyword_thread = KeywordExtractionThread(self.text, method, num_keywords)
            self.keyword_thread.setParent(self.parent_widget)
            self.keyword_thread.finished.connect(self.on_keyword_extraction_complete)
            self.keyword_thread.error.connect(self.on_analysis_error)
            self.keyword_thread.start()
            
        except Exception as e:
            QMessageBox.critical(self, "Error", f"Keyword extraction failed: {str(e)}")
            self.parent_widget.button_manager.restore_states()
            if self.parent_widget:
                self.parent_widget.set_progress('keywords', False)

    def on_topic_analysis_complete(self, topics):
        if self.parent_widget:
            self.parent_widget.set_progress('topics', False)
        self.parent_widget.button_manager.restore_states()
        self.show_results_dialog(self.topic_method.currentText(), topics)

    def on_keyword_extraction_complete(self, keywords):
        if self.parent_widget:
            self.parent_widget.set_progress('keywords', False)
        self.parent_widget.button_manager.restore_states()
        self.show_keyword_dialog(self.keyword_method.currentText(), keywords)

    def on_analysis_error(self, error_msg):
        if self.parent_widget:
            self.parent_widget.set_progress('topics', False)
            self.parent_widget.set_progress('keywords', False)
        self.parent_widget.button_manager.restore_states()
        QMessageBox.critical(self, "Error", error_msg)

    def show_results_dialog(self, method, topics):
        result_html = "<h3>Topic Analysis Results</h3>"
        for topic in topics:
            result_html += f"<p><b>{topic['topic']}</b> (weight: {topic['weight']:.2f})<br>"
            result_html += ", ".join(topic['terms']) + "</p>"

        dialog = QDialog(self)
        dialog.setWindowTitle(f"{method} Topic Results") 
        dialog.setSizeGripEnabled(True)
        self.setup_results_dialog(dialog, result_html)
        
    def show_keyword_dialog(self, method, keywords):
        result_html = "<h3 style='text-align: center;'>Keyword Extraction Results</h3>"
        result_html += "<table border='1' cellspacing='0' cellpadding='3' width='100%' style='margin-top: 20px;'>"
        result_html += "<tr style='background-color: #d3d3d3; color: black'><th>Keyword</th><th>Score</th></tr>"
        
        for kw in keywords:
            result_html += f"<tr><td>{kw['keyword']}</td><td>{kw['score']:.4f}</td></tr>"
            
        dialog = QDialog(self)
        dialog.setWindowTitle(f"{method.upper()} Keywords")
        dialog.setSizeGripEnabled(True)
        self.setup_results_dialog(dialog, result_html)

    def setup_results_dialog(self, dialog, html_content):
        dialog.setMinimumSize(500, 300)
        
        text_browser = QTextBrowser()
        text_browser.setHtml(html_content)
        
        close_btn = QPushButton("Close", dialog)
        close_btn.clicked.connect(dialog.close)
        
        layout = QVBoxLayout()
        layout.addWidget(text_browser)
        layout.addWidget(close_btn)
        dialog.setLayout(layout)
        
        dialog.setWindowModality(Qt.NonModal)
        dialog.show()

class WarningEmitter(QObject):
    token_warning = Signal(str)

class CustomWarningHandler:
    def __init__(self, emitter):
        self.emitter = emitter
        self.reported_tokens = set()

    def handle_warning(self, message, category, filename, lineno, file=None, line=None):
        if "Your stop_words may be inconsistent" in str(message):
            tokens = self._extract_tokens(str(message))
            if tokens:
                self.emitter.token_warning.emit(", ".join(tokens))
                self.reported_tokens.update(tokens)

    def _extract_tokens(self, message):
        start = message.find('[') + 1
        end = message.find(']')
        if start > 0 and end > start:
            return [token.strip("'") for token in message[start:end].split(', ')]
        return []

class PerformanceMonitor:
    """Monitor system resources and performance"""
    
    @staticmethod
    def check_resources():
        """Check available system resources"""
        try:
            mem = psutil.virtual_memory()
            mem_available = mem.available / 1024 / 1024
            cpu_available = 100 - psutil.cpu_percent()
            return mem_available, cpu_available
        except Exception as e:
            logger.error(f"Resource check error: {e}")
            return 1000, 50

class ButtonStateManager:
    """Manages button states during processing operations"""
    def __init__(self, main_window):
        self.main_window = main_window
        self.button_states = {}
        self.process_buttons = [
            'generate_wordcloud_button',
            'sentiment_button',
            'analyze_topics_btn',
            'extract_keywords_btn',
            'view_fulltext_button',
            'text_stats_button',
            'save_wc_button',
            'summarize_button',
            'load_file_button',
        ]
        
    def save_states(self):
        """Save current state of all process buttons"""
        for btn_name in self.process_buttons:
            if hasattr(self.main_window, btn_name):
                btn = getattr(self.main_window, btn_name)
                self.button_states[btn_name] = btn.isEnabled()
            elif hasattr(self.main_window.topic_tab, btn_name):
                btn = getattr(self.main_window.topic_tab, btn_name)
                self.button_states[btn_name] = btn.isEnabled()
                
    def disable_other_buttons(self, active_button):
        """Disable all process buttons except the active one"""
        self.save_states()
        for btn_name in self.process_buttons:
            if btn_name != active_button:
                if hasattr(self.main_window, btn_name):
                    getattr(self.main_window, btn_name).setEnabled(False)
                elif hasattr(self.main_window.topic_tab, btn_name):
                    getattr(self.main_window.topic_tab, btn_name).setEnabled(False)
                    
    def restore_states(self):
        """Restore buttons to their previous states"""
        for btn_name, was_enabled in self.button_states.items():
            if hasattr(self.main_window, btn_name):
                getattr(self.main_window, btn_name).setEnabled(was_enabled)
            elif hasattr(self.main_window.topic_tab, btn_name):
                getattr(self.main_window.topic_tab, btn_name).setEnabled(was_enabled)
        self.button_states.clear()

class ImportThread(QThread):
    """Dedicated thread for import initialization"""
    finished = Signal()
    
    def __init__(self):
        super().__init__()
        self._is_running = True
        
    def run(self):
        try:
            if not self._is_running:
                return
                
            import matplotlib.pyplot as plt
            if not self._is_running:
                return
                
            import nltk
            if not self._is_running:
                return
                
            from textblob import TextBlob
            if not self._is_running:
                return
                
            try:
                nltk.data.find('tokenizers/punkt')
            except LookupError:
                if self._is_running:
                    nltk.download('punkt', quiet=True)
        except Exception as e:
            logger.error(f"Import error: {e}")
        finally:
            if self._is_running:
                self.finished.emit()

    def stop(self):
        """Metode untuk menghentikan thread dengan aman"""
        self._is_running = False
        self.wait()

class FileLoaderThread(QThread):
    file_loaded = Signal(str, str)
    file_error = Signal(str)

    def __init__(self, file_path):
        super().__init__()
        self.file_path = file_path

    def run(self):
        try:
            if self.isInterruptionRequested():
                return

            if not os.path.exists(self.file_path):
                raise FileLoadError(f"File not found: {self.file_path}")

            file_size = os.path.getsize(self.file_path)
            if file_size > 100 * 1024 * 1024:
                self.process_large_file()
            else:
                self.process_normal_file()

        except Exception as e:
            self.file_error.emit(f"Error loading {os.path.basename(self.file_path)}: {str(e)}")

    def extract_text_from_pdf(self, pdf_path: str) -> str:
        """Extract text from PDF with proper error handling"""
        try:
            import pypdf
            text = ""
            
            try:
                with open(pdf_path, "rb") as file:
                    reader = pypdf.PdfReader(file)
                    if not reader.pages:
                        raise FileLoadError("PDF contains no readable pages")
                    
                    for page in reader.pages:
                        if self.isInterruptionRequested():
                            return ""
                        page_text = page.extract_text()
                        if page_text:
                            text += page_text + "\n"
                return text
                
            except pypdf.PdfReadError as e:
                logger.error(f"PDF read error: {e}")
                raise FileLoadError(f"Failed to read PDF: {e}")
            except PermissionError as e:
                logger.error(f"Permission denied: {e}")
                raise FileLoadError("Permission denied accessing PDF file")
                
        except ImportError as e:
            logger.error(f"PDF processing module not found: {e}")
            raise FileLoadError("PDF processing module not available")
        except Exception as e:
            logger.error(f"Unexpected error processing PDF: {e}")
            raise FileLoadError(f"PDF processing failed: {e}")

    def process_normal_file(self):
        """Process normal sized files with improved error handling"""
        try:
            if self.isInterruptionRequested():
                return

            text_data = ""
            file_ext = Path(self.file_path).suffix.lower()
            
            try:
                if file_ext == ".txt":
                    try:
                        with open(self.file_path, "r", encoding="utf-8") as file:
                            text_data = file.read()
                    except UnicodeDecodeError:
                        # Fallback to different encoding
                        with open(self.file_path, "r", encoding="latin-1") as file:
                            text_data = file.read()
                            
                elif file_ext == ".pdf":
                    text_data = self.extract_text_from_pdf(self.file_path)
                    
                elif file_ext in (".doc", ".docx"):
                    text_data = self.extract_text_from_word(self.file_path)
                    
                elif file_ext in (".xlsx", ".xls"):
                    text_data = self.extract_text_from_excel(self.file_path)
                    
                elif file_ext == ".csv":
                    text_data = self.extract_text_from_csv(self.file_path)
                    
                else:
                    raise FileLoadError(f"Unsupported file format: {file_ext}")

            except (PermissionError, OSError) as e:
                logger.error(f"File access error: {e}")
                raise FileLoadError(f"Cannot access file: {e}")
            except UnicodeError as e:
                logger.error(f"Text encoding error: {e}")
                raise FileLoadError("File contains invalid text encoding")

            if self.isInterruptionRequested():
                return

            if not text_data.strip():
                raise FileLoadError("File is empty or contains no extractable text")

            self.file_loaded.emit(self.file_path, text_data)

        except FileLoadError as e:
            logger.error(f"File load error: {e}")
            self.file_error.emit(str(e))
        except Exception as e:
            logger.error(f"Unexpected error: {e}", exc_info=True)
            self.file_error.emit(f"Unexpected error loading file: {e}")

    def extract_text_from_word(self, word_path):
        try:
            import docx
            doc = docx.Document(word_path)
            if not doc.paragraphs:
                raise FileLoadError("Word document contains no readable text")
            
            text = ""
            for p in doc.paragraphs:
                if self.isInterruptionRequested():
                    return ""
                if p.text.strip():
                    text += p.text.strip() + "\n"
                    
            for table in doc.tables:
                if self.isInterruptionRequested():
                    return ""
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            text += cell.text.strip() + "\n"
                            
            if not text.strip():
                raise FileLoadError("No readable text found in document")
                
            text = re.sub(r'\s+', ' ', text)
            text = text.replace('\n\n', '\n').strip()
            
            return text
            
        except Exception as e:
            logger.error(f"Word processing failed: {str(e)}")
            raise FileLoadError(f"Word processing failed: {str(e)}") from e

    def extract_text_from_excel(self, excel_path):
        try:
            import pandas as pd
            text_data = ""
            df = pd.read_excel(excel_path, sheet_name=None)
            if not df:
                raise FileLoadError("Excel file is empty or unreadable")
            
            for sheet_name, sheet_df in df.items():
                if self.isInterruptionRequested():
                    return ""
                text_data += f"\n--- {sheet_name} ---\n"
                text_data += sheet_df.to_string(index=False, header=True)
            return text_data
        except Exception as e:
            logger.error(f"Excel processing failed: {str(e)}")
            raise FileLoadError(f"Excel processing failed: {str(e)}") from e

    def extract_text_from_csv(self, csv_path):
        try:
            import pandas as pd
            df = pd.read_csv(csv_path)
            if df.empty:
                raise FileLoadError("CSV file is empty")
            
            if self.isInterruptionRequested():
                return ""
                
            return df.to_string(index=False, header=True)
        except Exception as e:
            logger.error(f"CSV processing failed: {str(e)}")
            raise FileLoadError(f"CSV processing failed: {str(e)}") from e

class SentimentAnalysisThread(QThread):
    sentiment_analyzed = Signal(dict)
    offline_warning = Signal(str)
    translation_failed = Signal(str)

    def __init__(self, text_data, sentiment_mode, vader_analyzer, flair_classifier, flair_classifier_cuslang, textblob_analyzer=None):
        super().__init__()
        self.text_data = text_data
        self.sentiment_mode = sentiment_mode
        self.vader_analyzer = vader_analyzer
        self.flair_classifier = flair_classifier
        self.flair_classifier_cuslang = flair_classifier_cuslang
        self.textblob_analyzer = textblob_analyzer
        self.cached_translation = None
        self.temp_dir = Path(tempfile.gettempdir()) / "textplora_cache"
        
        self.text_hash = hashlib.md5(text_data.encode()).hexdigest()
        self.temp_file = self.temp_dir / f"trans_{self.text_hash}.txt"
        
        self.temp_dir.mkdir(parents=True, exist_ok=True)

    async def _async_translate(self, text):
        try:
            translated = GoogleTranslator(source="auto", target="en").translate(text)
            return translated
        except Exception as e:
            logger.error(f"Translation error: {e}")
            return None

    def translate_text(self, text):
        try:
            loop = asyncio.get_event_loop()
        except RuntimeError:
            loop = asyncio.new_event_loop()
            asyncio.set_event_loop(loop)

        sentences = [s.strip() for s in text.split(".") if s.strip()]
        translated = []

        async def run_translations():
            for sentence in sentences:
                try:
                    result = await self._async_translate(sentence)
                    if result:
                        translated.append(result)
                except Exception as e:
                    logger.error(f"Translation error: {e}")
                    continue
            return ". ".join(translated)

        try:
            return loop.run_until_complete(run_translations())
        except Exception as e:
            logger.error(f"Translation error: {e}")
            return None
        finally:
            if loop.is_running():
                loop.close()

    def analyze_textblob(self, text, use_custom=False):
        if use_custom and self.textblob_analyzer:
            analysis = self.textblob_analyzer.analyze(text)
            polarity = analysis['polarity']
            subjectivity = analysis['subjectivity']
        else:
            from textblob import TextBlob
            blob = TextBlob(text)
            polarity = blob.sentiment.polarity
            subjectivity = blob.sentiment.subjectivity

        if abs(polarity) < 0.1:
            neutral_score = 1.0
            positive_score = negative_score = 0.0
        else:
            positive_score = (polarity + 1) / 2 if polarity > 0 else 0
            negative_score = (-polarity + 1) / 2 if polarity < 0 else 0
            neutral_score = 1 - (positive_score + negative_score)

        return {
            "positive_score": positive_score,
            "negative_score": negative_score,
            "neutral_score": neutral_score,
            "compound_score": polarity,
            "sentiment_label": "POSITIVE" if polarity > 0 else "NEGATIVE" if polarity < 0 else "NEUTRAL",
            "subjectivity": subjectivity
        }

    def get_cached_translation(self):
        """Get translation from cache if exists"""
        if self.temp_file.exists():
            try:
                with open(self.temp_file, "r", encoding="utf-8") as f:
                    return f.read()
            except Exception as e:
                logger.error(f"Error reading cached translation: {e}")
                return None
        return None
        
    def save_translation(self, translated_text):
        """Save translation to cache"""
        try:
            with open(self.temp_file, "w", encoding="utf-8") as f:
                f.write(translated_text)
        except Exception as e:
            logger.error(f"Error saving translation: {e}")

    def run(self):
        result = {
            "positive_score": 0,
            "neutral_score": 0,
            "negative_score": 0,
            "compound_score": 0,
            "sentiment_label": "N/A",
            "subjectivity": "N/A",
        }

        try:
            text_to_analyze = self.text_data
            needs_translation = False

            if self.sentiment_mode in ["VADER", "Flair", "TextBlob"]:
                try:
                    from langdetect import detect
                    detected_lang = detect(self.text_data)
                    needs_translation = detected_lang != "en"
                except Exception as e:
                    logger.error(f"Language detection error: {e}")
                    needs_translation = False

            if needs_translation:
                if not is_connected():
                    self.offline_warning.emit("Your text is not in English.\nTranslation required but no internet connection.\nPlease use the mode with custom lexicon/model instead.")
                    return

                cached = self.get_cached_translation()
                if cached:
                    text_to_analyze = cached
                else:
                    translated_text = self.translate_text(self.text_data)
                    if not translated_text:
                        self.translation_failed.emit("Translation failed - analysis aborted")
                        return
                        
                    self.save_translation(translated_text)
                    text_to_analyze = translated_text

            if self.sentiment_mode == "TextBlob":
                result.update(self.analyze_textblob(text_to_analyze))
            elif self.sentiment_mode == "TextBlob (Custom Lexicon)":
                result.update(self.analyze_textblob(self.text_data, True))

            elif self.sentiment_mode == "VADER":
                if self.vader_analyzer:
                    sentiment = self.vader_analyzer.polarity_scores(text_to_analyze)
                    total = sentiment["pos"] + sentiment["neg"] + sentiment["neu"]
                    positive_score = sentiment["pos"]
                    negative_score = sentiment["neg"] / total if total > 0 else 0
                    neutral_score = sentiment["neu"] / total if total > 0 else 0

                    result.update({
                        "positive_score": positive_score,
                        "negative_score": negative_score,
                        "neutral_score": neutral_score,
                        "compound_score": sentiment["compound"],
                        "sentiment_label": "POSITIVE" if sentiment["compound"] >= 0.05 else "NEGATIVE" if sentiment["compound"] <= -0.05 else "NEUTRAL",
                    })
                else:
                    result["sentiment_label"] = "Error: VADER not initialized"

            elif self.sentiment_mode == "VADER (Custom Lexicon)":
                if self.vader_analyzer:
                    sentiment = self.vader_analyzer.polarity_scores(self.text_data)
                    result.update({
                        "positive_score": sentiment["pos"],
                        "negative_score": sentiment["neg"],
                        "neutral_score": sentiment["neu"],
                        "compound_score": sentiment["compound"],
                        "sentiment_label": "POSITIVE" if sentiment["compound"] >= 0.05 else "NEGATIVE" if sentiment["compound"] <= -0.05 else "NEUTRAL",
                    })
                else:
                    result["sentiment_label"] = "Error: Custom VADER not initialized"

            elif self.sentiment_mode == "Flair":
                from flair.data import Sentence
                sentence = Sentence(text_to_analyze)
                self.flair_classifier.predict(sentence)
                sentiment = sentence.labels[0]
                confidence = sentiment.score
                sentiment_label = sentiment.value

                model_labels = self.flair_classifier.label_dictionary.get_items()
                
                if 'NEUTRAL' not in model_labels:
                    if abs(confidence - 0.5) < 0.05:
                        sentiment_label = "NEUTRAL"
                        positive_score = 0.0
                        negative_score = 0.0
                        neutral_score = 1.0
                    else:
                        positive_score = confidence if sentiment_label == "POSITIVE" else 0.0
                        negative_score = confidence if sentiment_label == "NEGATIVE" else 0.0
                        neutral_score = 0.0
                else:
                    if confidence < 0.55:
                        sentiment_label = "NEUTRAL"
                        positive_score = 0.0
                        negative_score = 0.0
                        neutral_score = 1.0
                    else:
                        positive_score = confidence if sentiment_label == "POSITIVE" else 0.0
                        negative_score = confidence if sentiment_label == "NEGATIVE" else 0.0
                        neutral_score = confidence if sentiment_label == "NEUTRAL" else 0.0

                result.update({
                    "compound_score": confidence,
                    "sentiment_label": sentiment_label,
                    "positive_score": positive_score,
                    "negative_score": negative_score,
                    "neutral_score": neutral_score,
                })            

            elif self.sentiment_mode == "Flair (Custom Model)":
                try:
                    from flair.data import Sentence
                    sentence = Sentence(text_to_analyze)
                    
                    if self.flair_classifier_cuslang is None:
                        raise ValueError("Custom model not loaded")
                        
                    self.flair_classifier_cuslang.predict(sentence)
                    if not sentence.labels:
                        raise ValueError("Model produced no labels")
                        
                    sentiment = sentence.labels[0]
                    confidence = sentiment.score
                    sentiment_label = sentiment.value

                    model_labels = self.flair_classifier.label_dictionary.get_items()
                    
                    if 'NEUTRAL' not in model_labels:
                        if abs(confidence - 0.5) < 0.05:
                            sentiment_label = "NEUTRAL"
                            positive_score = negative_score = 0.0
                            neutral_score = 1.0
                        else:
                            positive_score = confidence if sentiment_label == "POSITIVE" else 0
                            negative_score = confidence if sentiment_label == "NEGATIVE" else 0
                            neutral_score = 0
                    else:
                        if confidence < 0.55:
                            sentiment_label = "NEUTRAL"
                            positive_score = negative_score = 0.0
                            neutral_score = 1.0
                        else:
                            positive_score = confidence if sentiment_label == "POSITIVE" else 0
                            negative_score = confidence if sentiment_label == "NEGATIVE" else 0
                            neutral_score = 1.0 if sentiment_label == "NEUTRAL" else 0

                    result.update({
                        "compound_score": confidence,
                        "sentiment_label": sentiment_label,
                        "positive_score": positive_score,
                        "negative_score": negative_score, 
                        "neutral_score": neutral_score,
                    })            

                except ImportError as e:
                    result["sentiment_label"] = "Error: Flair not installed properly"
                except Exception as e:
                    result["sentiment_label"] = f"Error: {str(e)}"

            self.sentiment_analyzed.emit(result)

        except Exception as e:
            result["sentiment_label"] = f"Error: {str(e)}"
            self.sentiment_analyzed.emit(result)

class WordCloudGeneratorThread(QThread):
    """Thread for generating word clouds"""
    progress = Signal(int)
    finished = Signal(object)
    error = Signal(str)

    def __init__(self, text, wc_params):
        super().__init__()
        self.text = text
        self.wc_params = wc_params
        self._cancelled = False

    def run(self):
        try:
            if len(self.text) > 100000:
                chunk_size = 10000
                chunks = [self.text[i:i+chunk_size] 
                         for i in range(0, len(self.text), chunk_size)]
                frequencies = {}
                
                for i, chunk in enumerate(chunks):
                    if self._cancelled:
                        return
                        
                    chunk_freqs = MPWordCloud._process_chunk(chunk, self.wc_params.get('stopwords'))
                    
                    for word, freq in chunk_freqs.items():
                        frequencies[word] = frequencies.get(word, 0) + freq
                        
                    progress = int((i + 1) / len(chunks) * 100)
                    self.progress.emit(progress)
                    
                wordcloud = WordCloud(**self.wc_params)
                wordcloud.generate_from_frequencies(frequencies)
            else:
                wordcloud = WordCloud(**self.wc_params)
                wordcloud.generate(self.text)
                self.progress.emit(100)
                
            self.finished.emit(wordcloud)
            
        except Exception as e:
            self.error.emit(str(e))

    def cancel(self):
        self._cancelled = True
        self.requestInterruption()

class CustomFileLoaderThread(QThread):
    file_loaded = Signal(object, bool)

    def __init__(self, file_path, file_type):
        super().__init__()
        self.file_path = file_path
        self.file_type = file_type

    def run(self):
        try:
            if self.file_type == "lexicon":
                from vaderSentiment.vaderSentiment import SentimentIntensityAnalyzer
                self.vader_analyzer = CustomVaderSentimentIntensityAnalyzer(custom_lexicon_file=self.file_path)
                self.file_loaded.emit(self.file_path, True)

            elif self.file_type == "TextBlob (Custom Lexicon)":
                from textblob import TextBlob
                self.textblob_analyzer = CustomTextBlobSentimentAnalyzer(self.file_path)
                self.file_loaded.emit(self.file_path, True)

            elif self.file_type == "model":
                from flair.models import TextClassifier
                from flair.data import Sentence
                try:
                    model = TextClassifier.load(self.file_path)
                    test = Sentence("test")
                    model.predict(test)
                    if not test.labels or test.labels[0].value not in ['POSITIVE', 'NEGATIVE', 'NEUTRAL']:
                        raise ValueError("Model is not a valid sentiment classifier")
                    self.file_loaded.emit(model, True)
                except Exception as e:
                    self.file_loaded.emit(f"Invalid sentiment model: {str(e)}", False)

        except Exception as e:
            self.file_loaded.emit(str(e), False)

class MPWordCloud(WordCloud):
    """Thread-enabled WordCloud for large texts"""
    def __init__(self, *args, **kwargs):
        super().__init__(*args, **kwargs)
        self.generator_thread = None
        
    @staticmethod
    def _process_chunk(chunk, stopwords=None):
        """Static method to process text chunks with stopwords support"""
        temp_wc = WordCloud(stopwords=stopwords)
        return temp_wc.process_text(chunk)

    def generate_threaded(self, text, progress_callback=None, finished_callback=None, error_callback=None):
        """Generate word cloud using threading"""
        wc_params = {
            'width': self.width,
            'height': self.height,
            'background_color': self.background_color,
            'mask': self.mask,
            'max_words': self.max_words,
            'min_font_size': self.min_font_size,
            'font_path': self.font_path,
            'colormap': self.colormap,
            'stopwords': self.stopwords,
        }
        
        self.generator_thread = WordCloudGeneratorThread(text, wc_params)
        
        if progress_callback:
            self.generator_thread.progress.connect(progress_callback)
        if finished_callback:    
            self.generator_thread.finished.connect(finished_callback)
        if error_callback:
            self.generator_thread.error.connect(error_callback)
            
        self.generator_thread.start()
        return self

class SummarizeThread(QThread):
    summary_ready = Signal(str)
    error_occurred = Signal(str)

    def __init__(self, text):
        super().__init__()
        self.text = text

    def run(self):
        try:
            parser = PlaintextParser.from_string(self.text, Tokenizer("english"))
            summarizer = LsaSummarizer()
            summary = summarizer(parser.document, 5)
            summary_text = "\n".join(str(sentence) for sentence in summary)
            self.summary_ready.emit(summary_text)
        except Exception as e:
            self.error_occurred.emit(str(e))

def safe_open(file_path, mode='r', encoding=None, **kwargs):
    """
    Membuka file dengan aman setelah memvalidasi path
    
    Args:
        file_path (str): Path file yang akan dibuka
        mode (str): Mode untuk membuka file ('r', 'w', dll)
        encoding (str, optional): Encoding yang digunakan
        **kwargs: Argumen tambahan untuk fungsi open()
        
    Returns:
        file: File object yang telah dibuka
        
    Raises:
        ValueError: Jika path tidak valid
        IOError: Jika file tidak dapat dibuka
    """
    validated_path = sanitize_path(file_path)
    return open(validated_path, mode=mode, encoding=encoding, **kwargs)

def safe_read_file(file_path, encoding='utf-8'):
    """
    Membaca isi file dengan aman dengan penanganan kesalahan yang lebih baik
    
    Args:
        file_path (str): Path file yang akan dibaca
        encoding (str): Encoding yang digunakan
        
    Returns:
        str: Isi file
        
    Raises:
        FileNotFoundError: Jika file tidak ditemukan
        PermissionError: Jika tidak ada izin untuk membaca file
        UnicodeDecodeError: Jika file tidak dapat didekode dengan encoding yang diberikan
        IOError: Jika file tidak dapat dibaca karena alasan lain
    """
    try:
        validated_path = sanitize_path(file_path)
        
        # Periksa apakah file ada
        if not os.path.exists(validated_path):
            raise FileNotFoundError(f"File not found: {validated_path}")
        
        # Periksa apakah file dapat dibaca
        if not os.access(validated_path, os.R_OK):
            raise PermissionError(f"Permission denied: {validated_path}")
        
        # Periksa apakah file terlalu besar
        file_size = os.path.getsize(validated_path)
        if file_size > 100 * 1024 * 1024:  # 100MB
            raise IOError(f"File too large: {file_size} bytes")
        
        # Baca file
        with open(validated_path, 'r', encoding=encoding) as f:
            return f.read()
            
    except UnicodeDecodeError:
        # Re-raise untuk penanganan khusus di level yang lebih tinggi
        raise
    except Exception as e:
        logger.error(f"Error reading file {file_path}: {e}")
        raise FileLoadError(f"Tidak dapat membaca file: {e}")

def safe_write_file(file_path, content, encoding='utf-8'):
    """
    Menulis ke file dengan aman
    
    Args:
        file_path (str): Path file yang akan ditulis
        content (str): Konten yang akan ditulis
        encoding (str): Encoding yang digunakan
        
    Raises:
        ValueError: Jika path tidak valid
        IOError: Jika file tidak dapat ditulis
    """
    try:
        with safe_open(file_path, 'w', encoding=encoding) as f:
            f.write(content)
    except Exception as e:
        logger.error(f"Error writing to file {file_path}: {e}")
        raise IOError(f"Tidak dapat menulis ke file: {e}")

class LayoutManager:
    """Kelas untuk mengelola pembuatan layout yang umum digunakan"""
    
    @staticmethod
    def create_form_layout(parent, fields):
        """
        Membuat form layout dengan label dan input fields
        
        Args:
            parent: Widget parent
            fields: List of tuples (label_text, input_widget, [optional_tooltip])
            
        Returns:
            QGridLayout: Layout yang sudah dikonfigurasi
        """
        layout = QGridLayout()
        
        for row, field_info in enumerate(fields):
            label_text = field_info[0]
            input_widget = field_info[1]
            
            label = QLabel(label_text, parent)
            layout.addWidget(label, row, 0)
            layout.addWidget(input_widget, row, 1)
            
            # Tambahkan tooltip jika ada
            if len(field_info) > 2 and field_info[2]:
                input_widget.setToolTip(field_info[2])
                label.setToolTip(field_info[2])
        
        return layout
    
    @staticmethod
    def create_button_row(parent, buttons):
        """
        Membuat row dengan beberapa button
        
        Args:
            parent: Widget parent
            buttons: List of tuples (button_text, callback_function, [optional_tooltip])
            
        Returns:
            QHBoxLayout: Layout dengan button yang sudah dikonfigurasi
        """
        layout = QHBoxLayout()
        
        for btn_info in buttons:
            btn_text = btn_info[0]
            callback = btn_info[1]
            
            button = QPushButton(btn_text, parent)
            button.clicked.connect(callback)
            
            # Tambahkan tooltip jika ada
            if len(btn_info) > 2 and btn_info[2]:
                button.setToolTip(btn_info[2])
                
            layout.addWidget(button)
        
        return layout
    
    @staticmethod
    def create_group_box(parent, title, inner_layout):
        """
        Membuat group box dengan layout di dalamnya
        
        Args:
            parent: Widget parent
            title: Judul group box
            inner_layout: Layout yang akan dimasukkan ke dalam group box
            
        Returns:
            QGroupBox: Group box yang sudah dikonfigurasi
        """
        group_box = QGroupBox(title, parent)
        group_box.setLayout(inner_layout)
        return group_box

class DialogFactory:
    """Factory class untuk membuat dialog yang umum digunakan"""
    
    @staticmethod
    def create_info_dialog(parent, title, content, modal=True, min_size=(500, 400)):
        """
        Membuat dialog informasi dengan QTextBrowser
        
        Args:
            parent: Widget parent
            title: Judul dialog
            content: Konten HTML untuk ditampilkan
            modal: Apakah dialog bersifat modal
            min_size: Ukuran minimum dialog (width, height)
            
        Returns:
            QDialog: Dialog yang sudah dikonfigurasi
        """
        dialog = QDialog(parent)
        dialog.setWindowTitle(title)
        dialog.setMinimumSize(*min_size)
        
        if modal:
            dialog.setWindowModality(Qt.ApplicationModal)
        else:
            dialog.setWindowModality(Qt.NonModal)
            
        dialog.setSizeGripEnabled(True)
        
        layout = QVBoxLayout()
        
        text_browser = QTextBrowser()
        text_browser.setHtml(content)
        text_browser.setOpenExternalLinks(True)
        text_browser.setReadOnly(True)
        layout.addWidget(text_browser)
        
        close_button = QPushButton("Close")
        close_button.clicked.connect(dialog.accept)
        layout.addWidget(close_button)
        
        dialog.setLayout(layout)
        return dialog
    
    @staticmethod
    def create_input_dialog(parent, title, fields, on_accept, modal=True):
        """
        Membuat dialog input dengan multiple fields
        
        Args:
            parent: Widget parent
            title: Judul dialog
            fields: List of tuples (label_text, input_widget_type, default_value)
            on_accept: Callback function when dialog is accepted
            modal: Apakah dialog bersifat modal
            
        Returns:
            QDialog: Dialog yang sudah dikonfigurasi
        """
        dialog = QDialog(parent)
        dialog.setWindowTitle(title)
        
        if modal:
            dialog.setWindowModality(Qt.ApplicationModal)
        else:
            dialog.setWindowModality(Qt.NonModal)
            
        layout = QVBoxLayout()
        form_layout = QGridLayout()
        
        input_widgets = {}
        
        for row, (label_text, widget_type, default_value) in enumerate(fields):
            label = QLabel(label_text)
            form_layout.addWidget(label, row, 0)
            
            if widget_type == "line_edit":
                widget = QLineEdit(default_value)
            elif widget_type == "combo_box":
                widget = QComboBox()
                widget.addItems(default_value)
            elif widget_type == "spin_box":
                widget = QSpinBox()
                widget.setValue(default_value)
            else:
                widget = QLineEdit(str(default_value))
                
            form_layout.addWidget(widget, row, 1)
            input_widgets[label_text] = widget
            
        layout.addLayout(form_layout)
        
        button_layout = QHBoxLayout()
        ok_button = QPushButton("OK")
        cancel_button = QPushButton("Cancel")
        
        ok_button.clicked.connect(lambda: on_accept(dialog, input_widgets))
        cancel_button.clicked.connect(dialog.reject)
        
        button_layout.addWidget(ok_button)
        button_layout.addWidget(cancel_button)
        layout.addLayout(button_layout)
        
        dialog.setLayout(layout)
        return dialog

class ThreadFactory:
    """Factory class untuk membuat dan mengonfigurasi thread"""
    
    @staticmethod
    def create_worker_thread(worker, on_finished=None, on_progress=None, on_error=None, thread_name=None):
        """
        Membuat dan mengonfigurasi worker thread
        
        Args:
            worker: Worker object yang akan dijalankan di thread
            on_finished: Callback ketika thread selesai
            on_progress: Callback untuk update progress
            on_error: Callback ketika terjadi error
            thread_name: Nama thread untuk debugging
            
        Returns:
            QThread: Thread yang sudah dikonfigurasi
        """
        thread = QThread()
        if thread_name:
            thread.setObjectName(thread_name)
            
        worker.moveToThread(thread)
        thread.started.connect(worker.run)
        
        # Connect signals
        worker.finished.connect(thread.quit)
        worker.finished.connect(worker.deleteLater)
        thread.finished.connect(thread.deleteLater)
        
        # Connect custom callbacks
        if on_finished:
            worker.finished.connect(on_finished)
        if on_progress and hasattr(worker, 'progress'):
            worker.progress.connect(on_progress)
        if on_error and hasattr(worker, 'error'):
            worker.error.connect(on_error)
            
        return thread

class FileOperations:
    """Kelas untuk mengelola operasi file dengan validasi path"""
    
    @staticmethod
    def load_text_file(parent, file_path):
        """
        Load text from file with proper validation and enhanced error handling
        
        Args:
            parent: Parent widget for error messages
            file_path: Path to the file
            
        Returns:
            str: File content or empty string if error
        """
        try:
            # Validasi ekstensi file
            ext = Path(file_path).suffix.lower()
            supported_exts = ['.txt', '.csv', '.md', '.json', '.html', '.xml']
            
            if ext not in supported_exts:
                logger.warning(f"Unsupported file extension: {ext}")
                reply = QMessageBox.question(
                    parent,
                    "Unsupported File Type",
                    f"The file type '{ext}' may not be supported. Try to open anyway?",
                    QMessageBox.Yes | QMessageBox.No,
                    QMessageBox.No
                )
                if reply == QMessageBox.No:
                    return None
            
            # Coba baca file dengan encoding utf-8
            try:
                content = safe_read_file(file_path, encoding='utf-8')
            except UnicodeDecodeError:
                # Fallback ke encoding lain jika utf-8 gagal
                logger.info(f"UTF-8 decoding failed for {file_path}, trying other encodings")
                encodings = ['latin-1', 'cp1252', 'iso-8859-1']
                
                for encoding in encodings:
                    try:
                        content = safe_read_file(file_path, encoding=encoding)
                        logger.info(f"Successfully decoded with {encoding}")
                        break
                    except UnicodeDecodeError:
                        continue
                else:
                    # Jika semua encoding gagal
                    raise UnicodeDecodeError("All encodings failed", b"", 0, 1, "Cannot decode file")
            
            # Periksa apakah file kosong
            if not content.strip():
                logger.warning(f"Empty file: {file_path}")
                QMessageBox.warning(
                    parent, 
                    "Empty File", 
                    "The file appears to be empty. Please select a file with content."
                )
                return None
            
            # Periksa ukuran file
            if len(content) > 10 * 1024 * 1024:  # 10MB
                logger.warning(f"Large file detected: {file_path} ({len(content)} bytes)")
                reply = QMessageBox.question(
                    parent,
                    "Large File",
                    "This file is very large and may cause performance issues. Continue loading?",
                    QMessageBox.Yes | QMessageBox.No,
                    QMessageBox.No
                )
                if reply == QMessageBox.No:
                    return None
            
            parent.statusBar().showMessage(f"Loaded: {os.path.basename(file_path)}")
            return content
            
        except FileNotFoundError:
            QMessageBox.critical(
                parent, 
                "File Not Found", 
                f"The file '{os.path.basename(file_path)}' could not be found. It may have been moved or deleted."
            )
            logger.error(f"File not found: {file_path}")
            return None
            
        except PermissionError:
            QMessageBox.critical(
                parent, 
                "Permission Denied", 
                f"You don't have permission to access '{os.path.basename(file_path)}'. Check file permissions."
            )
            logger.error(f"Permission denied: {file_path}")
            return None
            
        except UnicodeDecodeError as e:
            QMessageBox.critical(
                parent, 
                "Encoding Error", 
                f"Could not decode the file. The file may be binary or use an unsupported encoding."
            )
            logger.error(f"Unicode decode error for {file_path}: {e}")
            return None
            
        except Exception as e:
            QMessageBox.critical(
                parent, 
                "Error", 
                f"Failed to load file: {str(e)}\n\nTry checking if the file is accessible and not corrupted."
            )
            logger.error(f"Error loading file {file_path}: {e}")
            return None
    
    @staticmethod
    def save_text_file(parent, file_path, content):
        """
        Save text to file with proper validation
        
        Args:
            parent: Parent widget for error messages
            file_path: Path to save the file
            content: Content to save
            
        Returns:
            bool: True if successful, False otherwise
        """
        try:
            safe_write_file(file_path, content)
            parent.statusBar().showMessage(f"Saved: {os.path.basename(file_path)}")
            return True
        except Exception as e:
            QMessageBox.critical(parent, "Error", f"Failed to save file: {str(e)}")
            logger.error(f"Error saving file {file_path}: {e}")
            return False
    
    @staticmethod
    def open_file_dialog(parent, title="Open File", file_filter=None):
        """
        Show open file dialog with path validation
        
        Args:
            parent: Parent widget
            title: Dialog title
            file_filter: File filter string
            
        Returns:
            str: Selected file path or None
        """
        options = QFileDialog.Options()
        
        if file_filter is None:
            file_filter = (
                "Text Files (*.txt);;CSV Files (*.csv);;Markdown (*.md);;"
                "JSON Files (*.json);;HTML Files (*.html);;XML Files (*.xml);;"
                "All Files (*)"
            )
        
        file_path, _ = QFileDialog.getOpenFileName(
            parent, title, "", file_filter, options=options
        )
        
        if file_path:
            try:
                return sanitize_path(file_path)
            except ValueError as e:
                QMessageBox.critical(parent, "Security Error", str(e))
                logger.warning(f"Blocked access to file: {file_path}")
                return None
        return None
    
    @staticmethod
    def save_file_dialog(parent, title="Save File", file_filter="All Files (*)"):
        """
        Show save file dialog with path validation
        
        Args:
            parent: Parent widget
            title: Dialog title
            file_filter: File filter string
            
        Returns:
            str: Selected file path or None
        """
        options = QFileDialog.Options()
        file_path, _ = QFileDialog.getSaveFileName(
            parent, title, "", file_filter, options=options
        )
        
        if file_path:
            try:
                return sanitize_path(file_path)
            except ValueError as e:
                QMessageBox.critical(parent, "Security Error", str(e))
                logger.warning(f"Blocked access to file: {file_path}")
                return None
        return None

def create_wordcloud_settings_ui(self):
    """Create wordcloud settings UI using LayoutManager"""
    group_box = QGroupBox("WordCloud Settings", self)
    layout = QVBoxLayout()
    
    # Create form fields
    self.width_input = QSpinBox(self)
    self.width_input.setRange(100, 3000)
    self.width_input.setValue(800)
    
    self.height_input = QSpinBox(self)
    self.height_input.setRange(100, 3000)
    self.height_input.setValue(400)
    
    self.max_words_input = QSpinBox(self)
    self.max_words_input.setRange(10, 1000)
    self.max_words_input.setValue(200)
    
    self.bg_color_button = QPushButton("Background Color", self)
    self.bg_color_button.clicked.connect(self.choose_bg_color)
    self.bg_color = "white"
    
    fields = [
        ("Width:", self.width_input, "Width of the word cloud image"),
        ("Height:", self.height_input, "Height of the word cloud image"),
        ("Max Words:", self.max_words_input, "Maximum number of words to include"),
        ("Background:", self.bg_color_button, "Background color of the word cloud")
    ]
    
    form_layout = LayoutManager.create_form_layout(self, fields)
    layout.addLayout(form_layout)
    
    # Create buttons
    buttons = [
        ("Generate", self.generate_wordcloud, "Generate word cloud from text"),
        ("Save Image", self.save_wordcloud, "Save word cloud as image file")
    ]
    
    button_layout = LayoutManager.create_button_row(self, buttons)
    layout.addLayout(button_layout)
    
    group_box.setLayout(layout)
    return group_box

class CacheManager:
    """
    Kelas untuk mengelola pembersihan cache secara manual
    berdasarkan penggunaan memori sistem
    """
    
    def __init__(self, threshold_percent=75, check_interval=60000):
        """
        Inisialisasi cache manager
        
        Args:
            threshold_percent (int): Persentase penggunaan memori yang memicu pembersihan cache
            check_interval (int): Interval pengecekan memori dalam milidetik
        """
        self.threshold_percent = threshold_percent
        self.check_interval = check_interval
        self.cached_functions = []
        self.timer = None
        
    def register_cached_function(self, func):
        """
        Mendaftarkan fungsi yang menggunakan cache
        
        Args:
            func: Fungsi dengan dekorator lru_cache
        """
        if hasattr(func, 'cache_clear'):
            self.cached_functions.append(func)
            logger.debug(f"Registered cached function: {func.__name__}")
        else:
            logger.warning(f"Function {func.__name__} has no cache_clear method")
    
    def register_cached_property(self, instance, property_name):
        """
        Mendaftarkan cached_property
        
        Args:
            instance: Instance objek yang memiliki cached_property
            property_name: Nama property yang di-cache
        """
        if hasattr(instance.__class__, property_name):
            # Simpan referensi ke instance dan nama property
            self.cached_functions.append((instance, property_name))
            logger.debug(f"Registered cached property: {property_name}")
        else:
            logger.warning(f"Property {property_name} is not found in class")
    
    def clear_all_caches(self):
        """Membersihkan semua cache yang terdaftar"""
        cleared_count = 0
        
        for item in self.cached_functions:
            try:
                if isinstance(item, tuple):
                    # Ini adalah cached_property
                    instance, prop_name = item
                    # Reset atribut cache untuk memaksa recalculation
                    if hasattr(instance, f"_{prop_name}"):
                        delattr(instance, f"_{prop_name}")
                        cleared_count += 1
                else:
                    # Ini adalah fungsi dengan lru_cache
                    item.cache_clear()
                    cleared_count += 1
            except Exception as e:
                logger.error(f"Error clearing cache: {e}")
        
        # Juga bersihkan cache LazyLoader
        if 'LazyLoader' in globals() and hasattr(LazyLoader, '_cache'):
            LazyLoader._cache.clear()
            cleared_count += 1
            
        logger.info(f"Cleared {cleared_count} caches due to high memory usage")
        
    def check_memory_usage(self):
        """
        Memeriksa penggunaan memori dan membersihkan cache jika melebihi threshold
        """
        memory_percent = psutil.virtual_memory().percent
        
        if memory_percent > self.threshold_percent:
            logger.warning(f"Memory usage high ({memory_percent}%), clearing caches")
            self.clear_all_caches()
            
            # Paksa garbage collection
            import gc
            gc.collect()
    
    def start_monitoring(self):
        """Mulai memantau penggunaan memori secara periodik"""
        if self.timer is None:
            self.timer = QTimer()
            self.timer.timeout.connect(self.check_memory_usage)
            self.timer.start(self.check_interval)
            logger.info(f"Started memory monitoring (threshold: {self.threshold_percent}%)")
    
    def stop_monitoring(self):
        """Hentikan pemantauan memori"""
        if self.timer is not None:
            self.timer.stop()
            self.timer = None
            logger.info("Stopped memory monitoring")

if __name__ == "__main__":
    app = QApplication(sys.argv)
    loop = QEventLoop(app)
    asyncio.set_event_loop(loop)
    
    ex = MainClass()
    ex.show()

    def cleanup():
        """Enhanced cleanup"""
        try:
            ex.cleanup()
            loop.stop()
        finally:
            LazyLoader._cache.clear()
            for func in [func for func in globals().values() if hasattr(func, 'cache_clear')]:
                func.cache_clear()

    app.aboutToQuit.connect(cleanup)

    with loop:
        sys.exit(loop.run_forever())
